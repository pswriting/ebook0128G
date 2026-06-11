# -*- coding: utf-8 -*-
import os
import logging

# 환경 변수
os.environ['PYTHONIOENCODING'] = 'utf-8'

# 로깅 비활성화
logging.getLogger('anthropic').setLevel(logging.ERROR)
logging.getLogger('httpx').setLevel(logging.ERROR)

import streamlit as st
import re
import json
import html
import base64
import urllib.parse
import uuid
import platform
import hashlib
import requests
from datetime import datetime, timedelta
from pathlib import Path

# Claude API
try:
    import anthropic
    CLAUDE_AVAILABLE = True
except ImportError:
    CLAUDE_AVAILABLE = False

# Gemini (이미지 생성용으로만 사용)
try:
    import google.generativeai as genai
    from google import genai as google_genai
    from google.genai import types as genai_types
    GEMINI_AVAILABLE = True
    IMAGEN_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    IMAGEN_AVAILABLE = False

# YouTube 자막 추출용
try:
    from youtube_transcript_api import YouTubeTranscriptApi
    YOUTUBE_TRANSCRIPT_AVAILABLE = True
except ImportError:
    YOUTUBE_TRANSCRIPT_AVAILABLE = False

# 브라우저 ID용 (클라우드 배포 시 필요)
try:
    from streamlit_javascript import st_javascript
    BROWSER_ID_AVAILABLE = True
except ImportError:
    BROWSER_ID_AVAILABLE = False

# 쿠키 매니저 (데이터 저장용)
try:
    import extra_streamlit_components as stx
    COOKIE_AVAILABLE = True
except ImportError:
    COOKIE_AVAILABLE = False

# Word 문서 생성용
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ==========================================
# 설정
# ==========================================
def get_config_path():
    return Path.home() / ".ebook_app_config.json"

def load_config():
    try:
        if get_config_path().exists():
            with open(get_config_path(), 'r') as f:
                return json.load(f)
    except:
        pass
    return {}

def save_config(data):
    try:
        config = load_config()
        config.update(data)
        with open(get_config_path(), 'w') as f:
            json.dump(config, f)
    except:
        pass

USED_WORDINGS_PATH = Path.home() / ".writey_used_wordings.json"

def load_used_wordings():
    """이전에 생성된 컨셉명 목록 (중복 방지용, 최근 40개)"""
    try:
        if USED_WORDINGS_PATH.exists():
            with open(USED_WORDINGS_PATH, 'r', encoding='utf-8') as f:
                return json.load(f)[-40:]
    except:
        pass
    return []

def record_used_wording(name):
    """생성된 컨셉명을 이력에 기록"""
    try:
        lst = load_used_wordings()
        if name and name not in lst:
            lst.append(name)
            with open(USED_WORDINGS_PATH, 'w', encoding='utf-8') as f:
                json.dump(lst[-40:], f, ensure_ascii=False)
    except:
        pass

def load_saved_api_key():
    return load_config().get('api_key', '')

def save_api_key(api_key):
    save_config({'api_key': api_key})

def is_authenticated():
    return load_config().get('authenticated', False)

def save_authenticated():
    save_config({'authenticated': True})

# ==========================================
# 간단 비밀번호 인증
# ==========================================
CORRECT_PASSWORD = "cashmaker2024"  # ← 비밀번호 변경하려면 여기만 수정


# 비디오 배경용 base64 인코딩
@st.cache_data
def get_video_base64(video_path):
    try:
        with open(video_path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    except:
        return None

@st.cache_data(show_spinner=False)
def get_video_base64_cached(video_path):
    """헤더 배경 영상 — 세션 간 캐시 (rerun마다 재인코딩 방지)"""
    return get_video_base64(video_path)

st.set_page_config(page_title="Writey", layout="wide", page_icon="✍")

# 쿠키 매니저 초기화 및 데이터 불러오기/저장
# 승인 유효기간: 365일 + 롤링 만료 (방문할 때마다 자동 갱신)
COOKIE_LIFETIME_DAYS = 365  # 한 번 로그인 후 365일 유지. 방문할 때마다 자동 연장됨

if COOKIE_AVAILABLE:
    import time
    cookie_manager = stx.CookieManager(key="writey_cookies")
    cookies = cookie_manager.get_all()

    # [중요] CookieManager는 새 세션의 '첫 실행'에서는 브라우저와의 통신이 끝나기 전이라
    # 쿠키가 실제로 있어도 빈 dict({})를 반환한다. 그 상태로 자동 로그인을 판정하면
    # 항상 로그인 화면이 떠서 "로그인 정보가 기억되지 않는" 증상이 발생한다.
    # → 첫 실행에서 쿠키가 비어 있으면 딱 한 번만 재실행해서 쿠키를 읽을 기회를 준다.
    if not cookies and not st.session_state.get('_cookie_loaded'):
        st.session_state['_cookie_loaded'] = True
        time.sleep(0.3)
        st.rerun()
    st.session_state['_cookie_loaded'] = True

    # 쿠키에서 비밀번호/API키 복원
    if cookies:
        if 'writey_password' in cookies and cookies['writey_password']:
            if 'saved_password' not in st.session_state:
                st.session_state['saved_password'] = cookies['writey_password']
        if 'writey_api_key' in cookies and cookies['writey_api_key']:
            if 'saved_api_key' not in st.session_state:
                st.session_state['saved_api_key'] = cookies['writey_api_key']

    # 롤링 만료: 매 세션마다 만료일 연장
    if 'cookie_rolling_refreshed' not in st.session_state:
        if st.session_state.get('saved_password'):
            st.session_state['pending_save_password'] = st.session_state['saved_password']
        if st.session_state.get('saved_api_key'):
            st.session_state['pending_save_api'] = st.session_state['saved_api_key']
        st.session_state['cookie_rolling_refreshed'] = True

    # pending 값을 쿠키에 저장
    if 'pending_save_password' in st.session_state:
        cookie_manager.set('writey_password', st.session_state['pending_save_password'], key='set_writey_pw', expires_at=datetime.now() + timedelta(days=COOKIE_LIFETIME_DAYS))
        del st.session_state['pending_save_password']
    if 'pending_save_api' in st.session_state:
        cookie_manager.set('writey_api_key', st.session_state['pending_save_api'], key='set_writey_api', expires_at=datetime.now() + timedelta(days=COOKIE_LIFETIME_DAYS))
        del st.session_state['pending_save_api']
else:
    cookie_manager = None

def save_password_to_browser(password):
    """비밀번호를 쿠키에 저장"""
    st.session_state['saved_password'] = password
    st.session_state['pending_save_password'] = password

def save_api_key_to_browser(api_key):
    """API 키를 쿠키에 저장"""
    st.session_state['saved_api_key'] = api_key
    st.session_state['pending_save_api'] = api_key

def get_saved_api_key():
    """저장된 API 키 반환"""
    return st.session_state.get('saved_api_key', None)

# ==========================================
# APPLE STYLE CSS
# ==========================================
st.markdown("""
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
@import url('https://fonts.googleapis.com/css2?family=Noto+Serif+KR:wght@400;500;600;700;900&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;500;600;700;800;900&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@400;500;600;700&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Cinzel:wght@400;500;600;700&display=swap');

/* S-Core Dream 폰트 */
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-1Thin.woff') format('woff');
    font-weight: 100;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-2ExtraLight.woff') format('woff');
    font-weight: 200;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-3Light.woff') format('woff');
    font-weight: 300;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-4Regular.woff') format('woff');
    font-weight: 400;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-5Medium.woff') format('woff');
    font-weight: 500;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-6Bold.woff') format('woff');
    font-weight: 600;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-7ExtraBold.woff') format('woff');
    font-weight: 700;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-8Heavy.woff') format('woff');
    font-weight: 800;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-9Black.woff') format('woff');
    font-weight: 900;
}

:root {
    --gold: #C9A24B;
    --gold-light: #E0C074;
    --gold-dark: #A8852F;
    --rose-gold: #b76e79;
    --cream: #FAF8F4;
    --charcoal: #141416;
    --dark: #0B0B0D;
    --card: rgba(255,255,255,0.025);
    --card2: rgba(255,255,255,0.05);
    --text: #F5F3EF;
    --text2: #8A8780;
    --text3: #7A776F;
    --line: rgba(201,162,75,0.18);
    --line2: rgba(255,255,255,0.06);
    --glow: rgba(201,162,75,0.32);
    --success: #6FA86F;
    --warning: #E0C074;
    --danger: #C97A6F;
}

/* 애니메이션 정의 */
@keyframes fadeInUp {
    from { opacity: 0; transform: translateY(30px); }
    to { opacity: 1; transform: translateY(0); }
}
@keyframes shimmer {
    0% { background-position: -200% 0; }
    100% { background-position: 200% 0; }
}
@keyframes pulse {
    0%, 100% { opacity: 1; transform: scale(1); }
    50% { opacity: 0.8; transform: scale(1.02); }
}
@keyframes borderGlow {
    0%, 100% { box-shadow: 0 0 5px var(--glow), inset 0 0 5px rgba(201,162,75,0.1); }
    50% { box-shadow: 0 0 20px var(--glow), inset 0 0 10px rgba(201,162,75,0.2); }
}
@keyframes float {
    0%, 100% { transform: translateY(0); }
    50% { transform: translateY(-5px); }
}
@keyframes goldShine {
    0% { background-position: -100% 0; }
    100% { background-position: 200% 0; }
}

*:not([data-testid*="Icon"]):not(.material-icons):not([class*="icon"]):not(span[aria-hidden="true"]) {
    font-family: 'S-CoreDream', 'Pretendard', -apple-system, sans-serif !important;
}
/* 아이콘 폰트 복원 */
[data-testid*="Icon"], .material-icons, span[aria-hidden="true"], button[kind="header"] span {
    font-family: 'Material Symbols Rounded', 'Material Icons', sans-serif !important;
}
.stDeployButton, footer, #MainMenu { display: none !important; }
/* 헤더 투명하게 (사이드바 버튼은 보임) */
header[data-testid="stHeader"] {
    background: transparent !important;
}

/* 럭셔리 배경 - 미세한 그라데이션 */
.stApp {
    background:
        radial-gradient(ellipse at 20% 0%, rgba(201,162,75,0.04) 0%, transparent 55%),
        radial-gradient(ellipse at 80% 100%, rgba(201,162,75,0.025) 0%, transparent 55%),
        linear-gradient(180deg, #0B0B0D 0%, #08080A 50%, #0B0B0D 100%) !important;
    background-attachment: fixed;
}

.main .block-container { max-width: 1000px; padding: 3rem 2rem; }

/* 사이드바 - 미니멀 */
[data-testid="stSidebar"] {
    background: var(--charcoal) !important;
    border-right: 1px solid var(--line);
}
[data-testid="stSidebar"] * { color: var(--text2) !important; }

/* 타이포그래피 - 가독성 향상 */
h1, h2, h3 { color: var(--text) !important; font-weight: 300 !important; letter-spacing: 0.5px; }
h1 { font-size: 34px !important; color: var(--cream) !important; font-weight: 300 !important; }
h2 { font-size: 26px !important; margin-bottom: 20px !important; font-weight: 300 !important; }
h3 { font-size: 21px !important; color: var(--gold) !important; font-weight: 400 !important; }
p, span, label, div { color: var(--text) !important; font-size: 16px !important; line-height: 1.7 !important; }
li { font-size: 16px !important; line-height: 1.8 !important; }

/* 버튼 - 채워진 골드 그라데이션 (첨부 디자인) */
.stButton > button {
    background: linear-gradient(135deg, #E0C074 0%, #C9A24B 100%) !important;
    color: #0B0B0D !important;
    -webkit-text-fill-color: #0B0B0D !important;
    border: none !important;
    border-radius: 12px;
    font-weight: 600;
    font-size: 15px !important;
    padding: 15px 36px;
    letter-spacing: 0.4px;
    text-transform: none;
    transition: all 0.35s cubic-bezier(0.4, 0, 0.2, 1);
    position: relative;
    overflow: hidden;
    box-shadow: 0 6px 20px rgba(201,162,75,0.22);
}
.stButton > button * {
    color: #0B0B0D !important;
    -webkit-text-fill-color: #0B0B0D !important;
}
.stButton > button::before {
    content: '';
    position: absolute;
    top: 0;
    left: -100%;
    width: 100%;
    height: 100%;
    background: linear-gradient(90deg, transparent, rgba(255,255,255,0.45), transparent);
    transition: left 0.6s ease;
}
.stButton > button:hover::before {
    left: 100%;
}
.stButton > button:hover {
    background: linear-gradient(135deg, #EBCE86 0%, #D4AC56 100%) !important;
    box-shadow: 0 10px 32px rgba(201,162,75,0.4);
    transform: translateY(-2px);
}
.stButton > button:active {
    transform: translateY(0);
    box-shadow: 0 4px 15px rgba(201,162,75,0.3);
}

/* 입력 필드 - 밝은 배경 + 검은 글씨 */
.stTextInput input, .stTextArea textarea, .stNumberInput input {
    background: #ffffff !important;
    background-color: #ffffff !important;
    border: 0.5px solid var(--line) !important;
    border-radius: 10px !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
    caret-color: #000000 !important;
    padding: 18px !important;
    font-size: 17px !important;
}
.stTextInput input:focus, .stTextArea textarea:focus, .stNumberInput input:focus {
    border-color: var(--gold) !important;
    box-shadow: 0 0 0 2px rgba(201,162,75,0.2) !important;
}

/* 셀렉트박스 컨테이너 */
.stSelectbox > div > div {
    background: var(--card) !important;
    border: 0.5px solid var(--line) !important;
    border-radius: 10px;
}
/* 셀렉트박스 선택된 값 - 흰색 */
.stSelectbox [data-baseweb="select"] > div {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

/* 스코어 카드 - 럭셔리 */
.score-card {
    background: linear-gradient(145deg, var(--card) 0%, rgba(30,30,30,0.95) 100%) !important;
    border: 0.5px solid var(--gold);
    border-radius: 20px;
    padding: 50px 40px;
    text-align: center;
    animation: fadeInUp 0.6s ease-out;
    transition: all 0.5s cubic-bezier(0.4, 0, 0.2, 1);
    position: relative;
    overflow: hidden;
    box-shadow: 0 10px 40px rgba(201,162,75,0.15);
}
.score-card::before {
    content: '';
    position: absolute;
    top: 0;
    left: 0;
    right: 0;
    height: 3px;
    background: linear-gradient(90deg, transparent, var(--gold), transparent);
    opacity: 1;
}
.score-card:hover {
    border-color: var(--gold);
    box-shadow: 0 20px 60px rgba(201,162,75,0.3), inset 0 1px 0 rgba(201,162,75,0.1);
    transform: translateY(-5px);
}
.score-card:hover::before {
    opacity: 1;
}
.score-number {
    font-size: 140px;
    font-weight: 300;
    background: linear-gradient(135deg, var(--gold-light) 0%, var(--gold) 50%, var(--gold-dark) 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    line-height: 1;
    letter-spacing: -4px;
    animation: fadeInUp 0.8s ease-out;
    filter: drop-shadow(0 2px 4px rgba(201,162,75,0.3));
}

/* 정보 카드 + 애니메이션 */
.info-card {
    background: transparent !important;
    border: none;
    border-left: 2px solid var(--gold);
    padding: 20px 24px;
    margin: 20px 0;
    animation: fadeInUp 0.5s ease-out;
    transition: all 0.3s ease;
}
.info-card:hover {
    background: rgba(201,169,98,0.05) !important;
    border-left-width: 4px;
    padding-left: 22px;
}

/* 스탯 박스 + 애니메이션 */
.stat-box {
    background: var(--card) !important;
    border: 0.5px solid var(--line2);
    border-radius: 14px;
    padding: 32px;
    text-align: center;
    animation: fadeInUp 0.5s ease-out;
    transition: all 0.4s ease;
}
.stat-box:hover {
    transform: translateY(-4px);
    box-shadow: 0 8px 25px rgba(0,0,0,0.3);
    border-color: var(--gold);
}
.stat-value {
    font-size: 42px;
    font-weight: 200;
    color: var(--gold) !important;
    letter-spacing: -2px;
    transition: transform 0.3s ease;
}
.stat-box:hover .stat-value {
    transform: scale(1.05);
}
.stat-label {
    font-size: 11px;
    color: var(--text2) !important;
    margin-top: 12px;
    text-transform: uppercase;
    letter-spacing: 3px;
}

/* 데이터 카드 + 애니메이션 */
.data-card {
    background: var(--card) !important;
    border-left: 2px solid var(--gold);
    padding: 20px 24px;
    margin: 16px 0;
    animation: fadeInUp 0.4s ease-out;
    transition: all 0.3s ease;
}
.data-card:hover {
    border-left-width: 4px;
    background: var(--card2) !important;
}

/* 서머리 허브 + 애니메이션 */
.summary-hub {
    background: var(--card) !important;
    border: 0.5px solid var(--line2);
    border-radius: 14px;
    padding: 40px;
    animation: fadeInUp 0.5s ease-out;
    transition: all 0.4s ease;
}
.summary-hub:hover {
    border-color: var(--gold);
}

/* 배지 - 미니멀 + 펄스 */
.verdict-go {
    background: transparent !important;
    color: var(--success) !important;
    border: 1px solid var(--success);
    padding: 12px 32px;
    border-radius: 20px;
    font-weight: 400;
    font-size: 12px;
    letter-spacing: 3px;
    text-transform: uppercase;
    animation: fadeInUp 0.6s ease-out;
}
.verdict-wait {
    background: transparent !important;
    color: var(--warning) !important;
    border: 1px solid var(--warning);
    padding: 12px 32px;
    border-radius: 20px;
    font-weight: 400;
    font-size: 12px;
    letter-spacing: 3px;
    animation: fadeInUp 0.6s ease-out;
}
.verdict-no {
    background: transparent !important;
    color: var(--danger) !important;
    border: 1px solid var(--danger);
    padding: 12px 32px;
    border-radius: 20px;
    font-weight: 400;
    font-size: 12px;
    letter-spacing: 3px;
    animation: fadeInUp 0.6s ease-out;
}

/* 네비게이션 */
.premium-nav-container {
    background: transparent;
    border-top: 1px solid var(--line);
    border-bottom: 1px solid var(--line);
    padding: 0;
    margin-bottom: 48px;
}
.nav-item {
    padding: 18px 12px;
    text-align: center;
    font-size: 14px;
    color: var(--text2);
    letter-spacing: 1px;
    transition: all 0.3s ease;
}
.nav-item.active {
    background: linear-gradient(135deg, rgba(201,162,75,0.2) 0%, rgba(201,162,75,0.1) 100%);
    color: var(--gold) !important;
    font-weight: 600;
    border-bottom: 3px solid var(--gold);
    box-shadow: 0 4px 15px rgba(201,162,75,0.2);
}

/* 섹션 타이틀 - 미니멀 (첨부 디자인) */
.section-title-box {
    background: rgba(255,255,255,0.025);
    border: 0.5px solid var(--line2);
    border-radius: 14px;
    padding: 30px 36px;
    margin-bottom: 35px;
    text-align: left;
    position: relative;
    overflow: hidden;
    animation: fadeInUp 0.5s ease-out;
}
.section-title-box h2 {
    font-size: 26px !important;
    color: #FAF8F4 !important;
    margin: 0 0 8px 0 !important;
    font-weight: 300 !important;
    letter-spacing: 0.5px;
}
.section-title-box p {
    color: var(--text2) !important;
    font-size: 14px !important;
    margin: 0 !important;
}
.section-step {
    display: inline-block;
    background: transparent;
    color: var(--gold) !important;
    font-size: 11px;
    font-weight: 500;
    padding: 0;
    margin-bottom: 12px;
    letter-spacing: 0.22em;
}

/* 제목 카드 + 애니메이션 */
.title-card {
    background: var(--card);
    border: 0.5px solid var(--line2);
    border-radius: 12px;
    padding: 28px;
    margin: 16px 0;
    transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
    animation: fadeInUp 0.5s ease-out;
}
.title-card:hover {
    border-color: var(--gold);
    background: rgba(201,169,98,0.05);
    transform: translateX(8px);
    box-shadow: -4px 0 20px rgba(201,169,98,0.15);
}
.title-main {
    font-size: 18px;
    font-weight: 400;
    color: var(--text) !important;
    letter-spacing: 1px;
    transition: color 0.3s ease;
}
.title-card:hover .title-main {
    color: var(--gold) !important;
}
.title-sub {
    font-size: 13px;
    color: var(--text2) !important;
    margin-top: 8px;
}

/* 로그인 - 럭셔리 */
.login-card {
    max-width: 420px;
    margin: 100px auto;
    padding: 70px 50px;
    background: linear-gradient(145deg, rgba(26,26,31,0.98) 0%, rgba(11,11,13,0.98) 100%);
    border: 0.5px solid var(--line);
    border-radius: 16px;
    text-align: center;
    animation: fadeInUp 0.8s ease-out;
    position: relative;
    box-shadow: 0 25px 80px rgba(0,0,0,0.5), 0 0 40px rgba(201,162,75,0.05);
}
.login-card::before {
    content: '';
    position: absolute;
    top: -1px;
    left: 20%;
    right: 20%;
    height: 2px;
    background: linear-gradient(90deg, transparent, var(--gold), transparent);
}
.login-card::after {
    content: '';
    position: absolute;
    bottom: -1px;
    left: 20%;
    right: 20%;
    height: 1px;
    background: linear-gradient(90deg, transparent, var(--gold-dark), transparent);
}
.login-title {
    font-size: 32px;
    font-weight: 300;
    background: linear-gradient(135deg, var(--cream) 0%, var(--gold-light) 50%, var(--gold) 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    letter-spacing: 7px;
    animation: fadeInUp 1s ease-out;
}
.login-subtitle {
    font-size: 11px;
    color: var(--gold) !important;
    margin-top: 20px;
    letter-spacing: 4px;
    text-transform: uppercase;
    animation: fadeInUp 1.2s ease-out;
    opacity: 0.8;
}

/* 헤더 - 럭셔리 */
.main-header {
    text-align: center;
    padding: 80px 20px 60px;
    margin-bottom: 50px;
    border-bottom: 1px solid var(--line);
    animation: fadeInUp 0.6s ease-out;
    position: relative;
    background: linear-gradient(180deg, rgba(201,162,75,0.02) 0%, transparent 100%);
}
.main-header::after {
    content: '';
    position: absolute;
    bottom: -1px;
    left: 10%;
    right: 10%;
    height: 1px;
    background: linear-gradient(90deg, transparent, var(--gold), transparent);
}
.main-header-brand {
    font-size: 11px;
    color: var(--gold) !important;
    letter-spacing: 10px;
    text-transform: uppercase;
    animation: fadeInUp 0.8s ease-out;
    text-shadow: 0 0 20px rgba(201,162,75,0.3);
}
.main-header-title {
    font-size: 42px;
    font-weight: 200;
    background: linear-gradient(135deg, var(--cream) 0%, var(--gold-light) 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    letter-spacing: 8px;
    margin-top: 24px;
    animation: fadeInUp 1s ease-out;
}
.header-tagline {
    font-size: 13px;
    color: var(--text2) !important;
    margin-top: 24px;
    letter-spacing: 3px;
    animation: fadeInUp 1.2s ease-out;
}

/* Expander + 애니메이션 */
.stExpander {
    background: var(--card) !important;
    border: 0.5px solid var(--line2) !important;
    border-radius: 12px !important;
    animation: fadeInUp 0.4s ease-out;
    transition: border-color 0.3s ease;
}
.stExpander:hover {
    border-color: var(--gold) !important;
}
.stProgress > div > div > div {
    background: linear-gradient(90deg, var(--gold-dark), var(--gold), var(--gold-light), var(--gold), var(--gold-dark)) !important;
    background-size: 300% 100%;
    animation: goldShine 3s ease infinite;
    border-radius: 4px;
    box-shadow: 0 0 15px rgba(201,162,75,0.4);
}
.stProgress > div > div {
    background: rgba(20,20,20,0.8);
    border-radius: 4px;
    border: 1px solid var(--line);
}

/* 라디오 & 탭 */
.stRadio > div { background: transparent; border: 1px solid var(--line); padding: 16px; }
.stTabs [data-baseweb="tab-list"] { background: transparent; border-bottom: 1px solid var(--line); }
.stTabs [aria-selected="true"] {
    background: transparent !important;
    color: var(--gold) !important;
    border-bottom: 2px solid var(--gold) !important;
}

/* 알림 */
.stSuccess > div { background: rgba(111,168,111,0.1) !important; border: 0.5px solid rgba(111,168,111,0.3) !important; border-radius: 10px; }
.stWarning > div { background: rgba(224,192,116,0.1) !important; border: 0.5px solid rgba(224,192,116,0.3) !important; border-radius: 10px; }
.stError > div { background: rgba(201,122,111,0.1) !important; border: 0.5px solid rgba(201,122,111,0.3) !important; border-radius: 10px; }
.stInfo > div { background: rgba(201,162,75,0.08) !important; border: 0.5px solid var(--line) !important; border-radius: 10px; }

/* 스크롤바 */
::-webkit-scrollbar { width: 6px; }
::-webkit-scrollbar-track { background: var(--dark); }
::-webkit-scrollbar-thumb { background: var(--gold-dark); }

/* 다운로드 버튼 - 럭셔리 골드 */
.stDownloadButton button {
    background: linear-gradient(135deg, var(--gold) 0%, var(--gold-dark) 50%, var(--gold) 100%) !important;
    background-size: 200% 100%;
    color: var(--dark) !important;
    border: none !important;
    border-radius: 12px;
    font-weight: 600;
    letter-spacing: 1.5px;
    box-shadow: 0 4px 20px rgba(201,162,75,0.3);
    transition: all 0.4s ease;
    text-shadow: 0 1px 1px rgba(255,255,255,0.2);
}
.stDownloadButton button:hover {
    background-position: 100% 0 !important;
    box-shadow: 0 8px 35px rgba(201,162,75,0.5);
    transform: translateY(-2px);
}

/* 구분선 */
hr { border: none; height: 1px; background: var(--line); margin: 40px 0; }

/* 표지 미리보기 - 실제 책처럼 */
.book-wrapper {
    perspective: 1000px;
    display: flex;
    justify-content: center;
    padding: 30px;
    background: linear-gradient(135deg, #1a1a1a 0%, #2d2d2d 100%);
    border-radius: 8px;
}
.ebook-cover {
    font-family: 'Pretendard', sans-serif !important;
    box-shadow:
        0 0 5px rgba(0,0,0,0.3),
        5px 5px 15px rgba(0,0,0,0.4),
        10px 10px 30px rgba(0,0,0,0.3),
        15px 15px 50px rgba(0,0,0,0.2),
        inset -3px 0 10px rgba(0,0,0,0.2);
    transform: rotateY(-3deg);
    border-radius: 0 3px 3px 0;
    position: relative;
}
.ebook-cover::before {
    content: '';
    position: absolute;
    left: 0;
    top: 0;
    bottom: 0;
    width: 25px;
    background: linear-gradient(90deg,
        rgba(0,0,0,0.4) 0%,
        rgba(0,0,0,0.1) 30%,
        rgba(255,255,255,0.05) 50%,
        rgba(0,0,0,0.1) 70%,
        rgba(0,0,0,0.3) 100%);
    border-radius: 3px 0 0 3px;
}
.ebook-cover::after {
    content: '';
    position: absolute;
    right: 0;
    top: 2px;
    bottom: 2px;
    width: 8px;
    background: linear-gradient(90deg,
        rgba(255,255,255,0.03) 0%,
        rgba(255,255,255,0.08) 50%,
        rgba(0,0,0,0.1) 100%);
}
.ebook-cover * {
    color: inherit !important;
    -webkit-text-fill-color: inherit !important;
}

/* ============================================
   입력 필드 텍스트 색상 - 최우선 적용
   ============================================ */

/* 모든 입력 필드 - 흰 배경 + 검은 글씨 */
.stTextInput input,
.stTextArea textarea,
.stNumberInput input,
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea,
[data-testid="stNumberInput"] input {
    background: #ffffff !important;
    background-color: #ffffff !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
    caret-color: #000000 !important;
}

/* Placeholder 색상 */
input::placeholder,
textarea::placeholder {
    color: #888888 !important;
    -webkit-text-fill-color: #888888 !important;
}

/* 셀렉트박스 - 선택된 값 (어두운 배경에 흰 글씨) */
.stSelectbox [data-baseweb="select"] > div,
.stSelectbox [data-baseweb="select"] span,
.stSelectbox > div > div > div {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

/* ============================================
   드롭다운/팝오버 - 검은 글씨 (흰 배경)
   ============================================ */
[data-baseweb="popover"],
[data-baseweb="popover"] *,
[data-baseweb="menu"],
[data-baseweb="menu"] *,
[data-baseweb="list"],
[data-baseweb="list"] *,
[role="listbox"],
[role="listbox"] *,
[role="option"],
[role="option"] *,
.stSelectbox ul,
.stSelectbox ul *,
.stSelectbox li,
.stSelectbox li * {
    background: #ffffff !important;
    background-color: #ffffff !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
}

/* 드롭다운 옵션 호버 */
[role="option"]:hover,
[data-baseweb="menu"] li:hover,
.stSelectbox li:hover {
    background: #f0f0f0 !important;
    background-color: #f0f0f0 !important;
}

/* select 요소 */
select,
select option {
    color: #000000 !important;
    background: #ffffff !important;
}

/* Expander 스타일 정리 */
.stExpander details summary {
    background: var(--card) !important;
    overflow: hidden !important;
}
/* 모든 텍스트 숨기기 (keyboard_arrow 등 영어 텍스트 포함) */
.stExpander details summary * {
    font-size: 0 !important;
    color: transparent !important;
    -webkit-text-fill-color: transparent !important;
}
/* 한국어 제목만 보이게 */
.stExpander details summary p {
    font-size: 15px !important;
    color: var(--text) !important;
    -webkit-text-fill-color: var(--text) !important;
}
/* 화살표 아이콘만 보이게 */
.stExpander details summary svg {
    width: 20px !important;
    height: 20px !important;
    color: var(--gold) !important;
    fill: var(--gold) !important;
}

/* 버튼 앞 불필요한 라벨 숨기기 */
.stButton > div:not([data-testid="baseButton-secondary"]):not([data-testid="baseButton-primary"]) > p,
.stButton > div > div > p:first-child:not(:last-child),
.stButton label,
.stExpander .stButton > div:first-child > p {
    display: none !important;
}
/* 링크버튼 라벨 숨기기 */
.stLinkButton > div:first-child > p {
    display: none !important;
}

/* ═══════════════════════════════════════════
   LUXURY REFINEMENT — 최종 오버라이드
   (아트 갤러리 톤: 다크 글래스 + 골드 헤어라인 + 세리프)
   ═══════════════════════════════════════════ */

/* 1. Primary 버튼 — 스트림릿 기본 빨강 제거, 샴페인 골드 */
.stButton > button[kind="primary"],
.stButton > button[kind="primaryFormSubmit"],
.stButton > button[data-testid="stBaseButton-primary"],
.stButton > button[data-testid="baseButton-primary"],
button[kind="primary"] {
    background: linear-gradient(135deg, #E8CC85 0%, #C9A24B 55%, #B08C36 100%) !important;
    color: #15120A !important;
    -webkit-text-fill-color: #15120A !important;
    border: 1px solid rgba(232,204,133,0.6) !important;
    border-radius: 12px !important;
    font-weight: 600 !important;
    letter-spacing: 0.12em !important;
    box-shadow: 0 10px 30px rgba(201,162,75,0.3), inset 0 1px 0 rgba(255,255,255,0.4) !important;
}
.stButton > button[kind="primary"] *,
.stButton > button[data-testid="stBaseButton-primary"] *,
button[kind="primary"] * {
    color: #15120A !important;
    -webkit-text-fill-color: #15120A !important;
}
.stButton > button[kind="primary"]:hover,
button[kind="primary"]:hover {
    filter: brightness(1.08);
    transform: translateY(-1px);
}

/* 2. 입력 필드 — 흰 종이 대신 다크 글래스 + 골드 헤어라인 */
.stTextInput input, .stTextArea textarea, .stNumberInput input,
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea,
[data-testid="stNumberInput"] input {
    background: rgba(17,16,13,0.78) !important;
    background-color: rgba(17,16,13,0.78) !important;
    border: 1px solid rgba(201,162,75,0.22) !important;
    border-radius: 10px !important;
    color: #F5F3EF !important;
    -webkit-text-fill-color: #F5F3EF !important;
    caret-color: #E0C074 !important;
    box-shadow: inset 0 1px 8px rgba(0,0,0,0.4) !important;
    transition: border-color .3s ease, box-shadow .3s ease !important;
}
.stTextInput input:focus, .stTextArea textarea:focus, .stNumberInput input:focus {
    border-color: rgba(224,192,116,0.75) !important;
    box-shadow: inset 0 1px 8px rgba(0,0,0,0.4), 0 0 0 3px rgba(201,162,75,0.14), 0 0 18px rgba(201,162,75,0.18) !important;
}
input::placeholder, textarea::placeholder {
    color: #6E6A60 !important;
    -webkit-text-fill-color: #6E6A60 !important;
}

/* 3. 제목 타이포 — 세리프 (전시 도록 느낌) */
h1, h2, h3,
.section-title-box h2 {
    font-family: 'Playfair Display', 'Noto Serif KR', serif !important;
    font-weight: 400 !important;
    letter-spacing: 0.04em !important;
}

/* 4. 사이드바 — 패널 정돈 + 고스트 버튼 */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #0C0B09 0%, #121009 140%) !important;
    border-right: 1px solid rgba(201,162,75,0.14) !important;
}
[data-testid="stSidebar"] .stButton > button {
    background: rgba(201,162,75,0.05) !important;
    border: 1px solid rgba(201,162,75,0.35) !important;
    border-radius: 9px !important;
    box-shadow: none !important;
    letter-spacing: 0.08em !important;
    font-weight: 500 !important;
    color: #D8C9A4 !important;
    -webkit-text-fill-color: #D8C9A4 !important;
}
[data-testid="stSidebar"] .stButton > button * {
    color: #D8C9A4 !important;
    -webkit-text-fill-color: #D8C9A4 !important;
}
[data-testid="stSidebar"] .stButton > button:hover {
    background: linear-gradient(135deg, #E0C074 0%, #C9A24B 100%) !important;
    transform: none !important;
}
[data-testid="stSidebar"] .stButton > button:hover * {
    color: #15120A !important;
    -webkit-text-fill-color: #15120A !important;
}

/* 5. 진행바 — 가는 골드 라인 + 은은한 발광 */
.stProgress > div > div {
    background: rgba(255,255,255,0.06) !important;
    height: 3px !important;
    border-radius: 2px !important;
}
.stProgress > div > div > div {
    background: linear-gradient(90deg, #B08C36, #E8CC85, #C9A24B) !important;
    box-shadow: 0 0 12px rgba(201,162,75,0.55) !important;
    height: 3px !important;
}

/* 6. 알림 박스 — 다크 글래스 통일 */
[data-testid="stAlert"], .stAlert {
    background: rgba(19,18,15,0.88) !important;
    border: 1px solid rgba(201,162,75,0.25) !important;
    border-radius: 10px !important;
}
[data-testid="stAlert"] p, [data-testid="stAlert"] div {
    color: #EDE8DC !important;
}

/* 7. 셀렉트박스 — 입력창과 같은 다크 글래스 */
.stSelectbox [data-baseweb="select"] > div {
    background: rgba(17,16,13,0.78) !important;
    border: 1px solid rgba(201,162,75,0.22) !important;
    border-radius: 10px !important;
}

/* 8. 파일 업로더 — 갤러리 드롭존 */
[data-testid="stFileUploader"] section {
    background: rgba(17,16,13,0.55) !important;
    border: 1px dashed rgba(201,162,75,0.35) !important;
    border-radius: 10px !important;
}
[data-testid="stFileUploader"] section span,
[data-testid="stFileUploader"] section small {
    color: #8A8780 !important;
}
[data-testid="stFileUploader"] section button {
    background: rgba(201,162,75,0.08) !important;
    border: 1px solid rgba(201,162,75,0.4) !important;
    color: #D8C9A4 !important;
    border-radius: 8px !important;
}

/* 9. 익스팬더 — 헤어라인 카드 */
.stExpander, [data-testid="stExpander"] {
    background: rgba(17,16,13,0.5) !important;
    border: 1px solid rgba(201,162,75,0.16) !important;
    border-radius: 10px !important;
}

/* 10. 다운로드 버튼 — primary와 동일한 골드 */
.stDownloadButton > button {
    background: linear-gradient(135deg, #E8CC85 0%, #C9A24B 60%, #B08C36 100%) !important;
    color: #15120A !important;
    -webkit-text-fill-color: #15120A !important;
    border: 1px solid rgba(232,204,133,0.6) !important;
    border-radius: 10px !important;
    font-weight: 600 !important;
    letter-spacing: 0.1em !important;
    box-shadow: 0 6px 20px rgba(201,162,75,0.25) !important;
}
.stDownloadButton > button * {
    color: #15120A !important;
    -webkit-text-fill-color: #15120A !important;
}

/* 11. 체크박스/라디오 — 골드 액센트 */
input[type="checkbox"], input[type="radio"] {
    accent-color: #C9A24B !important;
}

/* 12. 스크롤바 — 가늘고 어둡게 */
::-webkit-scrollbar { width: 8px; height: 8px; }
::-webkit-scrollbar-track { background: #0B0B0D; }
::-webkit-scrollbar-thumb {
    background: rgba(201,162,75,0.3);
    border-radius: 4px;
}
::-webkit-scrollbar-thumb:hover { background: rgba(201,162,75,0.55); }

/* 13. 본문 미리보기(책 페이지)는 종이 흰색 유지 */
.content-preview-box, .content-preview-box * {
    background-color: #ffffff;
}

/* ═══════════════════════════════════════════
   DESIGN v2 — 버튼 위계 + 타이포 + 라벨
   ═══════════════════════════════════════════ */

/* 일반(보조) 버튼: 고스트 — 골드는 핵심 CTA(primary/다운로드)에만 */
.stButton > button:not([kind="primary"]):not([kind="primaryFormSubmit"]):not([data-testid="stBaseButton-primary"]):not([data-testid="baseButton-primary"]) {
    background: rgba(201,162,75,0.04) !important;
    border: 1px solid rgba(201,162,75,0.32) !important;
    color: #D8C9A4 !important;
    -webkit-text-fill-color: #D8C9A4 !important;
    box-shadow: none !important;
    border-radius: 10px !important;
    font-weight: 500 !important;
    letter-spacing: 0.06em !important;
}
.stButton > button:not([kind="primary"]):not([kind="primaryFormSubmit"]):not([data-testid="stBaseButton-primary"]):not([data-testid="baseButton-primary"]) * {
    color: #D8C9A4 !important;
    -webkit-text-fill-color: #D8C9A4 !important;
}
.stButton > button:not([kind="primary"]):not([kind="primaryFormSubmit"]):not([data-testid="stBaseButton-primary"]):not([data-testid="baseButton-primary"]):hover {
    background: linear-gradient(135deg, #E0C074 0%, #C9A24B 100%) !important;
    border-color: rgba(232,204,133,0.7) !important;
}
.stButton > button:not([kind="primary"]):not([kind="primaryFormSubmit"]):not([data-testid="stBaseButton-primary"]):not([data-testid="baseButton-primary"]):hover * {
    color: #15120A !important;
    -webkit-text-fill-color: #15120A !important;
}

/* 헤딩: 골드 남용 제거 — h3는 크림 세리프, 골드는 마이크로 라벨 전용 */
h3 {
    color: var(--cream) !important;
    font-weight: 500 !important;
    letter-spacing: 0.02em !important;
}

/* 사이드바 마이크로 라벨 */
.sb-label {
    font-size: 10px !important;
    font-weight: 600 !important;
    letter-spacing: 0.38em !important;
    color: var(--gold) !important;
    opacity: 0.92;
    margin: 6px 0 10px 2px;
    font-family: 'Pretendard', sans-serif !important;
}

/* 입력 라벨: 조용한 회색 + 자간 */
.stTextInput label p, .stTextArea label p, .stNumberInput label p,
.stSelectbox label p, .stFileUploader label p {
    font-size: 13px !important;
    color: #A39E92 !important;
    letter-spacing: 0.05em !important;
}

/* 캡션 */
[data-testid="stCaptionContainer"] p {
    color: #6E6A60 !important;
    font-size: 12.5px !important;
    letter-spacing: 0.03em !important;
}

/* 본문 영역 여백 — 갤러리처럼 여유 있게 */
.main .block-container {
    max-width: 960px;
    padding-top: 2.2rem;
}

/* 구분선: 더 옅게 */
hr {
    background: linear-gradient(90deg, transparent, rgba(201,162,75,0.25), transparent) !important;
    margin: 34px 0 !important;
}
</style>
""", unsafe_allow_html=True)



# ==========================================
# 비밀번호 인증 (단순)
# ==========================================
if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False

# 저장된 비밀번호로 자동 로그인 (매 실행마다 재확인)
# 쿠키는 첫 실행 뒤 재실행 시점에 늦게 로드될 수 있으므로, 미인증 상태에서는
# 매번 saved_password를 다시 확인해야 쿠키가 늦게 들어와도 자동 로그인이 된다.
if not st.session_state['authenticated']:
    if st.session_state.get('saved_password', '') == CORRECT_PASSWORD:
        st.session_state['authenticated'] = True

if not st.session_state['authenticated']:
    st.markdown("""
    <div class="login-card">
        <div class="login-subtitle">CASHMAKER</div>
        <div class="login-title">WRITEY</div>
        <div class="login-subtitle">Premium E-Book Studio</div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        pw = st.text_input("비밀번호", type="password", key="pw_login", placeholder="Enter password...")
        remember_pw = st.checkbox("비밀번호 저장 (다음 접속 시 자동 로그인)", value=True, key="remember_pw")
        if st.button("입장", key="btn_login", use_container_width=True):
            if pw == CORRECT_PASSWORD:
                st.session_state['authenticated'] = True
                if remember_pw:
                    save_password_to_browser(pw)
                st.rerun()
            else:
                st.error("비밀번호가 올바르지 않습니다")
    st.stop()


# 세션 초기화
defaults = {
    'topic': '', 'target_persona': '', 'pain_points': '',
    'outline': [], 'chapters': {}, 'book_title': '', 'subtitle': '',
    'score_details': None, 'generated_titles': None, 'suggested_targets': None,
    'analyzed_pains': None, 'review_analysis': None, 'market_gaps': None,
    'knowledge_hub': [], 'study_summary': None, 'current_page': 0,
    'recommended_refs': None, 'generated_ideas': None,
    # 인터뷰 관련 변수
    'interview_completed': False,
    'interview_data': {},
    'author_name': '',
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# 사이드바
with st.sidebar:
    # API 키 섹션 (접기/펼치기 가능)
    if 'show_api_section' not in st.session_state:
        st.session_state['show_api_section'] = True

    if 'api_key' not in st.session_state:
        # 쿠키에서 API 키 불러오기
        st.session_state['api_key'] = st.session_state.get('saved_api_key', '') or ''
    # api_key가 비어 있는데 쿠키에 저장된 키가 (늦게) 들어왔다면 복원
    elif not st.session_state['api_key'] and st.session_state.get('saved_api_key'):
        st.session_state['api_key'] = st.session_state['saved_api_key']

    # API 키가 입력되어 있으면 기본적으로 접힌 상태로
    api_key_exists = bool(st.session_state['api_key'])

    col_title, col_toggle = st.columns([4, 1])
    with col_title:
        st.markdown('<div class="sb-label">CLAUDE API KEY</div>', unsafe_allow_html=True)
    with col_toggle:
        toggle_label = "▼" if st.session_state['show_api_section'] else "▶"
        if st.button(toggle_label, key="toggle_api_section", help="접기/펼치기"):
            st.session_state['show_api_section'] = not st.session_state['show_api_section']
            st.rerun()

    if st.session_state['show_api_section']:
        api_key = st.text_input("키 입력", value=st.session_state['api_key'], type="password", key="api_sidebar", label_visibility="collapsed", placeholder="sk-ant-api03-... 형식")
        if api_key != st.session_state['api_key']:
            st.session_state['api_key'] = api_key
            # 쿠키에 저장
            if api_key:
                save_api_key_to_browser(api_key)
                # 비밀번호 흐름과 동일하게 즉시 재실행해 pending 값을 쿠키에 바로 기록
                # (재실행이 없으면 다음 상호작용 전까지 쿠키 기록이 미뤄져 저장이 누락될 수 있음)
                st.rerun()

        if api_key:
            st.success("✅ Claude 키 입력 완료!")
        else:
            st.error("⚠️ Claude API 키를 입력하세요")
    else:
        # 접힌 상태에서 간단한 상태 표시
        if st.session_state['api_key']:
            st.caption("✅ API 키 설정됨")
        else:
            st.caption("⚠️ API 키 필요")

    # 모델 선택
    st.markdown('<div class="sb-label">MODEL</div>', unsafe_allow_html=True)
    if 'claude_model' not in st.session_state:
        st.session_state['claude_model'] = "claude-sonnet-4-5"

    model_options = {
        "Claude Sonnet 4.5 (추천)": "claude-sonnet-4-5",
        "Claude Fable 5 (최신·최고 지능)": "claude-fable-5",
        "Claude Opus 4.5 (최고 품질)": "claude-opus-4-5",
        "Claude Haiku 4.5 (저렴)": "claude-haiku-4-5"
    }
    selected_model = st.selectbox(
        "모델 선택",
        options=list(model_options.keys()),
        index=0,
        label_visibility="collapsed"
    )
    st.session_state['claude_model'] = model_options[selected_model]

    if "Haiku" in selected_model:
        st.info("💰 가장 저렴하고 빠름.\n📌 단, 목차·본문·프롤로그·에필로그·컨셉·제목 생성은 품질 보장을 위해 자동으로 Sonnet 4.5 사용")
    elif "Fable" in selected_model:
        st.info("🌟 Anthropic 최신 모델 — 가장 높은 지능. 글 품질이 가장 중요할 때 추천 (단가는 Opus급)\n⏱️ 추론형 모델이라 목차·본문 생성에 단계당 1~3분 걸릴 수 있어요. 멈춘 게 아니니 기다려주세요.")
    elif "Opus" in selected_model:
        st.info("💎 최고 품질, 단가가 가장 높음")
    else:
        st.info("⚡ 균형잡힌 품질/가격, 일반적으로 가장 추천")

    # API 키 발급 방법 안내
    with st.expander("📖 Claude API 키 발급 방법 (상세)", expanded=False):
        st.markdown("""
        ### 🟣 1단계: Anthropic 회원가입

        1. 아래 버튼을 클릭하세요
        2. **"Sign up"** 클릭
        3. Google 계정 또는 이메일로 가입
        """)
        st.link_button("🔗 Anthropic 가입 페이지", "https://console.anthropic.com/", use_container_width=True)

        st.markdown("""
        ---
        ### 💳 2단계: 결제 수단 등록

        1. 로그인 후 왼쪽 메뉴에서 **"Settings"** 클릭
        2. **"Billing"** 클릭
        3. **"Add payment method"** 클릭
        4. 카드 정보 입력 후 저장
        5. **"Add credits"**로 크레딧 충전 ($5~10 추천)
        """)
        st.link_button("🔗 Billing 페이지 바로가기", "https://console.anthropic.com/settings/billing", use_container_width=True)

        st.markdown("""
        ---
        ### 🔑 3단계: API 키 발급

        1. 왼쪽 메뉴에서 **"API Keys"** 클릭
        2. **"Create Key"** 버튼 클릭
        3. 이름 입력 (예: ebook)
        4. **"Create Key"** 클릭
        5. 생성된 키 **복사** (sk-ant-api03-... 형식)
        6. 위 입력창에 **붙여넣기**
        """)
        st.link_button("🔗 API Keys 페이지 바로가기", "https://console.anthropic.com/settings/keys", use_container_width=True)

        st.markdown("---")
        st.warning("⚠️ API 키는 한 번만 보여줍니다. 복사해두세요!")
        st.success("💰 예상 비용: 전자책 1권당 약 200~500원")

    st.markdown("---")
    st.markdown('<div class="sb-label">PROGRESS</div>', unsafe_allow_html=True)
    # 실제 제작 단계 기준: 인터뷰(0~50%) → 목차 완성(50%) → 본문(50~100%)
    _ch_map = st.session_state.get('chapters', {})
    _total_subs = sum(len(c.get('subtopics', [])) for c in _ch_map.values())
    _done_subs = sum(1 for c in _ch_map.values() for sd in c.get('subtopic_data', {}).values() if sd.get('content'))
    if _total_subs > 0 and _done_subs >= _total_subs:
        _prog, _stage = 1.0, "✓ 완성 — 다운로드 가능"
    elif _total_subs > 0 and _done_subs > 0:
        _prog = 0.5 + 0.5 * (_done_subs / _total_subs)
        _stage = f"본문 작성 {_done_subs}/{_total_subs}"
    elif st.session_state.get('outline'):
        _prog, _stage = 0.5, "목차 완성 — 본문 생성 대기"
    elif not st.session_state.get('interview_completed', False):
        _istep = st.session_state.get('interview_step', 1)
        _prog = max(0.04, (_istep - 1) / 6 * 0.5)
        _stage = f"인터뷰 {_istep}/6"
    else:
        _prog, _stage = 0.04, "주제 입력 대기"
    st.progress(_prog)
    st.caption(_stage)

    st.markdown("---")
    st.markdown('<div class="sb-label">MENU</div>', unsafe_allow_html=True)
    sidebar_pages = ["① 주제", "② 목차", "③ 본문", "④ 완성"]
    sidebar_mapping = [0, 4, 5, 7]
    for i, p in enumerate(sidebar_pages):
        if st.button(p, key=f"sidebar_nav_{i}", use_container_width=True):
            st.session_state['current_page'] = sidebar_mapping[i]
            st.rerun()

    # ── 프로젝트 저장/복원 ──
    # Streamlit Cloud는 새로고침/절전 시 세션이 사라지므로 수시 저장 권장
    st.markdown("---")
    st.markdown('<div class="sb-label">BACKUP</div>', unsafe_allow_html=True)
    st.caption("새로고침하면 작업이 사라질 수 있어요. 수시로 저장하세요!")

    _SAVE_KEYS = ['topic', 'target_persona', 'pain_points', 'outline', 'chapters',
                  'book_title', 'subtitle', 'book_concept', 'author_name',
                  'interview_data', 'temp_interview', 'interview_completed',
                  'interview_step', 'current_page',
                  '_prologue_cache', '_prologue_cache_key',
                  '_epilogue_cache', '_epilogue_cache_key']
    try:
        _proj = {k: st.session_state.get(k) for k in _SAVE_KEYS if k in st.session_state}
        _proj_json = json.dumps(_proj, ensure_ascii=False, indent=2, default=str)
        _fname = (st.session_state.get('book_title') or st.session_state.get('topic') or 'writey_project').strip()[:30]
        st.download_button("작업 저장 (JSON)", _proj_json,
                           file_name=f"{_fname}.json", mime="application/json",
                           use_container_width=True, key="proj_save")
    except Exception:
        st.caption("⚠️ 저장 파일 생성 실패")

    _uploaded_proj = st.file_uploader("저장 파일 불러오기", type=['json'], key="proj_load")
    if _uploaded_proj is not None and not st.session_state.get('_proj_restored'):
        try:
            _data = json.loads(_uploaded_proj.read().decode('utf-8'))
            if not isinstance(_data, dict) or 'chapters' not in _data:
                st.error("올바른 저장 파일이 아닙니다")
            else:
                for k in _SAVE_KEYS:
                    if k in _data:
                        st.session_state[k] = _data[k]
                st.session_state['_proj_restored'] = True
                st.success("✅ 프로젝트 복원 완료!")
                st.rerun()
        except Exception as e:
            st.error(f"불러오기 실패: {str(e)[:50]}")
    if _uploaded_proj is None and st.session_state.get('_proj_restored'):
        st.session_state['_proj_restored'] = False

    # 사이드바 하단 제작자 정보
    st.markdown("---")
    st.markdown("""
    <div style="text-align:center; padding:10px 0; color:#C9A24B !important; font-size:12px;">
        <strong>CASHMAKER</strong><br>
        <span style="color:#ffffff !important;">제작: 남현우 작가</span>
    </div>
    """, unsafe_allow_html=True)

# ==========================================
# 헬퍼 함수
# ==========================================
def get_api_key():
    return st.session_state.get('api_key', '')

def clean_text(text):
    if not text:
        return ""
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*([^*]+)\*\*', r'「\1」', text)
    text = text.replace('**', '').replace('*', '').replace('###', '').replace('##', '').replace('#', '')
    return text.strip()

def _html_table_to_pipes(match):
    """HTML 표를 마크다운 파이프 표로 변환 (표를 버리지 않고 살려서 DOCX 표로 렌더링)"""
    rows = re.findall(r'<tr[^>]*>(.*?)</tr>', match.group(0), flags=re.DOTALL | re.IGNORECASE)
    out = []
    for r in rows:
        cells = re.findall(r'<t[hd][^>]*>(.*?)</t[hd]>', r, flags=re.DOTALL | re.IGNORECASE)
        cells = [re.sub(r'<[^>]+>', '', c).replace('\n', ' ').strip() for c in cells]
        cells = [c for c in cells if c != '']
        if cells:
            out.append('| ' + ' | '.join(cells) + ' |')
    return ('\n' + '\n'.join(out) + '\n') if out else ''

def _strip_inline_summary_section(text):
    """AI가 본문 끝에 붙인 '핵심정리/요약' 섹션 제거 (파트 끝 핵심 정리 박스와 중복 방지)"""
    lines = text.split('\n')
    n = len(lines)
    for i in range(n - 1, -1, -1):
        bare = re.sub(r'[#\s━─=*◆📌:.\-]+', '', lines[i])
        if bare in ('핵심정리', '핵심요약', '요약정리', '정리하면', '오늘의핵심', '핵심포인트', '요약'):
            if i >= n * 0.5:  # 본문 후반부에 있을 때만 (그 지점부터 끝까지 삭제)
                return '\n'.join(lines[:i]).strip()
            break
    return text

def _split_long_paragraphs(text, max_sentences=3, max_chars=320):
    """긴 문단을 2~3문장 단위로 분할 — 가독성 (표·★ 인용구 블록은 보존)"""
    out = []
    for para in text.split('\n\n'):
        p = para.strip()
        if not p or '|' in p or p.startswith('★'):
            out.append(para)
            continue
        sentences = [s for s in re.split(r'(?<=[.!?…])\s+', p) if s.strip()]
        if len(sentences) <= max_sentences and len(p) <= max_chars:
            out.append(p)
            continue
        chunks, cur, cur_len = [], [], 0
        for s in sentences:
            cur.append(s)
            cur_len += len(s)
            if len(cur) >= max_sentences or cur_len > max_chars:
                chunks.append(' '.join(cur))
                cur, cur_len = [], 0
        if cur:
            chunks.append(' '.join(cur))
        out.append('\n\n'.join(chunks))
    return '\n\n'.join(out)

def clean_content(text, subtopic=None):
    if not text:
        return ""
    # HTML 테이블은 마크다운 표로 변환해 보존, 나머지 HTML 태그만 제거
    text = re.sub(r'<table[^>]*>.*?</table>', _html_table_to_pipes, text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    # 마크다운 헤더 제거
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    # **굵게**는 보존 (DOCX에서 볼드로 렌더링), *기울임*만 제거
    text = re.sub(r'(?<!\*)\*([^*\n]+)\*(?!\*)', r'\1', text)
    text = re.sub(r'\*\*\s*\*\*', '', text)  # 빈 볼드 정리
    text = text.replace('###', '').replace('##', '')
    text = re.sub(r'^#\s*', '', text, flags=re.MULTILINE)
    # 연속 줄바꿈 정리
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()

    # 본문 첫 줄이 소제목과 동일하면 제거 (AI가 소제목을 본문 시작에 또 박는 경우)
    if subtopic:
        sub_clean = re.sub(r'[「」"\'\s\.\?!]+', '', subtopic).strip()
        lines = text.split('\n')
        if lines:
            first_line_clean = re.sub(r'[「」"\'\s\.\?!]+', '', lines[0]).strip()
            # 정확 일치 또는 거의 같은 경우 (90% 이상 매칭)
            if first_line_clean == sub_clean or (sub_clean and sub_clean in first_line_clean and len(first_line_clean) <= len(sub_clean) * 1.2):
                text = '\n'.join(lines[1:]).lstrip('\n').strip()

    # 본문 끝 '핵심정리' 중복 섹션 제거 + 긴 문단 분할 (가독성)
    text = _strip_inline_summary_section(text)
    text = _split_long_paragraphs(text)
    return text

def parse_json(response):
    """JSON 파싱 - 개선된 에러 처리"""
    if not response:
        return None
    try:
        # 먼저 전체 응답에서 JSON 블록 찾기
        json_match = re.search(r'```json\s*([\s\S]*?)\s*```', response)
        if json_match:
            return json.loads(json_match.group(1))

        # JSON 블록이 없으면 중괄호로 시작하는 객체 찾기
        match = re.search(r'\{[\s\S]*\}', response)
        if match:
            json_str = match.group()
            # 불완전한 JSON 수정 시도
            json_str = re.sub(r',\s*}', '}', json_str)  # 마지막 쉼표 제거
            json_str = re.sub(r',\s*]', ']', json_str)  # 배열 마지막 쉼표 제거
            return json.loads(json_str)
    except json.JSONDecodeError as e:
        st.warning(f"JSON 파싱 경고: {str(e)[:50]}")
    except Exception as e:
        st.warning(f"파싱 오류: {str(e)[:50]}")
    return None

def _build_msg_kwargs(model, temp, max_tokens, prompt):
    """모델별 파라미터 제약 처리
    - Fable 5: temperature 미지정 필수 + 적응형 추론이 항상 켜져 있어 effort로 깊이 조절
      (effort 없이는 모든 호출에 깊은 추론이 돌아 목차/본문 생성이 매우 느려짐)
    - 그 외 모델: 기존처럼 작업별 temperature 사용"""
    kwargs = dict(model=model, max_tokens=max_tokens,
                  messages=[{"role": "user", "content": prompt}])
    if 'fable' in model.lower():
        # temp는 작업 성격의 힌트로만 사용: 검수·요약(저온) → low, 창작(목차·본문) → medium
        # 속도가 충분하고 더 깊은 품질을 원하면 'medium' → 'high'로 바꿀 것
        effort = 'low' if temp <= 0.4 else 'medium'
        # effort는 output_config 안에 넣어야 함 (최상위에 넣으면 400 invalid_request_error)
        kwargs['extra_body'] = {'output_config': {'effort': effort}}  # extra_body: SDK 버전 무관하게 전달
        # 추론(thinking) 토큰이 max_tokens에 포함되므로 한도를 넉넉히 —
        # 한도가 작으면 추론만 하다 잘려서 텍스트 없이 끝나고, 재시도가 반복되며 멈춘 것처럼 보임
        kwargs['max_tokens'] = max(max_tokens, 16000)
    else:
        kwargs['temperature'] = temp
    return kwargs

def _extract_text(message):
    """응답에서 텍스트 블록만 추출 — Fable 5는 thinking 블록이 앞에 올 수 있어 content[0].text가 깨짐"""
    parts = []
    for block in getattr(message, 'content', []) or []:
        if getattr(block, 'type', '') == 'text':
            t = getattr(block, 'text', None)
            if t:
                parts.append(t)
    if parts:
        return '\n'.join(parts)
    # 타입 정보가 없는 구버전 SDK 대비: text 속성을 가진 첫 블록
    for block in getattr(message, 'content', []) or []:
        t = getattr(block, 'text', None)
        if t:
            return t
    return None

def claude_call_threadsafe(prompt, api_key, model, temp=0.7, max_tokens=8000):
    """병렬 본문 생성용 — Streamlit 비의존, 일시 오류 자동 재시도"""
    import time as _t
    try:
        client = anthropic.Anthropic(api_key=api_key)
    except Exception as e:
        return None, str(e)[:120]
    last_err = None
    for attempt in range(3):
        try:
            msg = client.messages.create(**_build_msg_kwargs(model, temp, max_tokens, prompt))
            _txt = _extract_text(msg)
            if _txt is None:
                return None, "응답에 텍스트가 없습니다"
            return _txt, None
        except Exception as e:
            last_err = e
            m = str(e).lower()
            if 'rate' in m and attempt < 2:
                _t.sleep(15 * (attempt + 1))  # 레이트리밋: 길게 대기
                continue
            if any(k in m for k in ['overloaded', '529', '500', '503', 'internal server', 'connection', 'timeout', 'timed out']) and attempt < 2:
                _t.sleep(4 * (attempt + 1))
                continue
            return None, str(e)[:120]
    return None, str(last_err)[:120]

def ask_ai(prompt, temp=0.7, ensure_quality=False):
    """Claude API 호출

    ensure_quality=True 시 Haiku 선택해도 자동으로 Sonnet 4.5로 업그레이드.
    목차/본문/프롤로그/에필로그/컨셉 생성처럼 정교한 프롬프트를 따라야 하는 작업에 사용.
    """
    api_key = get_api_key()
    if not api_key:
        st.error("Claude API 키를 입력해주세요")
        return None

    if not CLAUDE_AVAILABLE:
        st.error("anthropic 패키지가 설치되지 않았습니다. pip install anthropic")
        return None

    # 선택된 모델 가져오기 (기본값: Sonnet 4.5)
    user_model = st.session_state.get('claude_model', 'claude-sonnet-4-5')

    # 핵심 생성 작업은 Haiku 자동 업그레이드 (품질 일관성 보장)
    if ensure_quality and 'haiku' in user_model.lower():
        model = 'claude-sonnet-4-5'
    else:
        model = user_model

    try:
        client = anthropic.Anthropic(api_key=api_key)
        # 일시적 오류(과부하 529, 서버 오류, 연결 끊김)는 최대 2회 자동 재시도
        for _attempt in range(3):
            try:
                message = client.messages.create(**_build_msg_kwargs(model, temp, 8000, prompt))
                return _extract_text(message)
            except (anthropic.AuthenticationError, anthropic.RateLimitError, anthropic.BadRequestError):
                raise  # 재시도 무의미한 오류는 즉시 처리
            except Exception as _e:
                _msg = str(_e).lower()
                _retryable = any(k in _msg for k in ['overloaded', '529', '500', '503', 'internal server', 'connection', 'timeout', 'timed out'])
                if _retryable and _attempt < 2:
                    import time as _time
                    _time.sleep(4 * (_attempt + 1))  # 4초, 8초 대기 후 재시도
                    continue
                raise
    except anthropic.AuthenticationError:
        st.error("API 키가 유효하지 않습니다. Claude API 키를 확인해주세요.")
        return None
    except anthropic.RateLimitError:
        st.error("API 할당량이 초과되었습니다. 잠시 후 다시 시도해주세요.")
        return None
    except anthropic.BadRequestError as e:
        try:
            err_msg = str(e).encode('utf-8', errors='ignore').decode('utf-8')[:100]
        except:
            err_msg = "요청 형식 오류"
        st.error(f"요청 오류: {err_msg}")
        return None
    except Exception as e:
        try:
            err_msg = str(e).encode('utf-8', errors='ignore').decode('utf-8')[:100]
        except:
            err_msg = "알 수 없는 오류"
        st.error(f"AI 오류: {err_msg}")
        return None

def generate_cover_image_gemini(title, subtitle, theme_keywords):
    """Google Gemini로 표지 배경 이미지 생성"""

    api_key = get_api_key()
    if not api_key:
        return None, "Gemini API 키가 필요합니다."

    if not IMAGEN_AVAILABLE:
        return None, "google-genai 패키지가 필요합니다: pip install google-genai"

    try:
        client = google_genai.Client(api_key=api_key)

        # 베스트셀러급 고급 표지 프롬프트 - 텍스트 절대 금지 강조
        prompt = f"""Create an ABSTRACT background image for a book cover.

Theme keywords: {theme_keywords}

STYLE: Dark, moody, cinematic atmosphere. Abstract shapes, gradients, smoke, light rays, or geometric patterns. Luxury aesthetic with gold/amber accent lighting on deep black background.

CRITICAL RULES:
- ONLY abstract visuals: smoke, light, shadows, gradients, textures
- NO objects, NO people, NO faces, NO hands
- NO text, NO letters, NO words, NO numbers, NO symbols, NO characters of ANY language
- NO Korean, NO English, NO Chinese, NO Japanese characters
- Pure abstract art only

OUTPUT: Dark dramatic background with subtle golden light accents, suitable for text overlay."""

        # Gemini 이미지 생성
        response = client.models.generate_content(
            model='gemini-2.0-flash-exp-image-generation',
            contents=prompt,
            config=genai_types.GenerateContentConfig(
                response_modalities=['IMAGE', 'TEXT']
            )
        )

        if response.candidates and response.candidates[0].content.parts:
            for part in response.candidates[0].content.parts:
                if hasattr(part, 'inline_data') and part.inline_data:
                    image_base64 = base64.b64encode(part.inline_data.data).decode('utf-8')
                    return image_base64, None

        return None, "이미지가 생성되지 않았습니다."

    except Exception as e:
        error_msg = str(e)
        if "quota" in error_msg.lower() or "limit" in error_msg.lower():
            return None, "API 할당량 초과. 잠시 후 다시 시도해주세요."
        elif "safety" in error_msg.lower():
            return None, "안전 필터에 의해 차단되었습니다. 다른 키워드로 시도해주세요."
        return None, f"이미지 생성 오류: {error_msg[:80]}"

def generate_cover_prompt_ai(title, subtitle, topic):
    """AI가 표지 디자인 컨셉과 이미지 프롬프트 생성"""
    prompt = f"""당신은 베스트셀러 책 표지 디자이너입니다.

책 제목: {title}
부제: {subtitle}
주제: {topic}

이 책의 표지 이미지를 위한 영문 프롬프트를 만들어주세요.

[요구사항]
1. 실제 베스트셀러 표지 스타일 분석 기반
2. 제목의 핵심 메시지를 시각적으로 표현
3. 고급스럽고 전문적인 느낌
4. 텍스트 오버레이를 위한 여백 고려
5. 추상적이거나 상징적인 이미지

[출력 형식]
IMAGE_PROMPT: (영문 이미지 생성 프롬프트, 50단어 이내)
COLOR_SCHEME: (추천 컬러 팔레트, 예: dark, gold, minimal)
STYLE: (디자인 스타일, 예: editorial, bold, elegant)

영문 프롬프트만 출력하세요. 한국어 설명 불필요."""

    result = ask_ai(prompt, temp=0.7)
    if result:
        # 파싱
        image_prompt = ""
        color_scheme = "dark"
        style = "editorial"

        for line in result.split('\n'):
            if 'IMAGE_PROMPT:' in line:
                image_prompt = line.split('IMAGE_PROMPT:')[-1].strip()
            elif 'COLOR_SCHEME:' in line:
                color_scheme = line.split('COLOR_SCHEME:')[-1].strip().lower()
            elif 'STYLE:' in line:
                style = line.split('STYLE:')[-1].strip().lower()

        return image_prompt, color_scheme, style
    return None, "dark", "editorial"


# ==========================================
# 고급 표지 렌더러 (외부 API 불필요, 벡터 SVG)
# ==========================================
COVER_TEMPLATES = {
    "two_tone_frame": "투톤 프레임 — 블랙 블록+초대형 제목+포인트 프레임 (베스트셀러 1군)",
    "diagonal_block": "대각 블록 — 비비드 단색+대각 분할 (베스트셀러 1군)",
    "solid_impact": "솔리드 임팩트 — 밝은 단색+검정 타이포 (베스트셀러 매대 표준)",
    "noir_gold": "느와르 골드 — 블랙+금박 (부·재테크)",
    "deep_navy": "딥 네이비 — 프레스티지 (투자·금융)",
    "editorial": "에디토리얼 — 아이보리 클래식 (인문·마인드)",
    "minimal_type": "미니멀 타이포 — 딥그린 (에세이·철학)",
    "color_block": "컬러 블록 — 코발트 임팩트 (동기부여)",
    "big_number": "빅 넘버 — 숫자 강조 (기간·챌린지)",
    "gradient_modern": "모던 그라데이션 (IT·트렌드·SNS)",
    "warm_essay": "웜 에세이 — 살구빛 (습관·건강)",
    "geometric": "스위스 기하 — 화이트 (비즈니스·마케팅)",
    "red_punch": "레드 펀치 — 블랙+레드 (도발·머니)",
}

def pick_cover_template(topic, title=""):
    """주제 키워드로 표지 아키타입 자동 매칭 (책마다 다른 디자인)"""
    text = f"{topic} {title}"
    candidates = None
    rules = [
        (["주식", "투자", "배당", "ETF", "연금", "경제"], ["deep_navy", "noir_gold"]),
        (["부동산", "재테크", "자산", "부자", "억"], ["noir_gold", "deep_navy"]),
        (["부업", "수익", "창업", "매출", "월급", "돈", "전자책", "크몽", "PDF", "클래스101", "강의", "지식창업"], ["two_tone_frame", "diagonal_block", "solid_impact", "noir_gold"]),
        (["마케팅", "브랜딩", "전략", "사업", "영업", "세일즈"], ["geometric", "color_block"]),
        (["AI", "챗GPT", "GPT", "디지털", "트렌드", "유튜브", "SNS", "블로그", "인스타", "스마트스토어", "쓰레드"], ["gradient_modern", "geometric"]),
        (["습관", "운동", "다이어트", "건강", "아침", "새벽", "루틴", "미라클"], ["solid_impact", "warm_essay", "big_number"]),
        (["글쓰기", "독서", "공부", "영어", "마인드", "심리", "멘탈", "철학"], ["minimal_type", "editorial"]),
        (["성공", "동기", "자기계발", "인생"], ["two_tone_frame", "diagonal_block", "solid_impact", "red_punch"]),
    ]
    for keywords, tmpls in rules:
        if any(k.lower() in text.lower() for k in keywords):
            candidates = tmpls
            break
    if re.search(r'\d+(일|주|개월|시간|분)', text):
        candidates = (candidates or []) + ["big_number"]
    if not candidates:
        candidates = list(COVER_TEMPLATES.keys())
    # 같은 주제라도 제목이 다르면 다른 디자인이 나오도록 해시 분산
    return candidates[abs(hash(text)) % len(candidates)]


def _wrap_title_lines(title, max_chars=7, max_lines=3):
    """제목을 표지용으로 줄바꿈 (한글 글자수 기준)"""
    title = (title or "").strip()
    if not title:
        return ["제목"]
    words = title.split()
    lines, cur = [], ""
    for w in words:
        cand = (cur + " " + w).strip()
        if not cur or len(cand) <= max_chars:
            cur = cand
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    # 띄어쓰기 없는 긴 단어는 강제 줄바꿈
    if len(lines) == 1 and len(lines[0]) > max_chars:
        s = lines[0]
        lines = [s[i:i + max_chars] for i in range(0, len(s), max_chars)]
    return lines[:max_lines]


def _cover_motif(context_text, accent='#C9A24B', soft='#E8CD8B'):
    """책 주제(제목+부제 키워드)에 맞는 표지 모티프 — 주제와 무관한 범용 그림 방지"""
    def has(*kws):
        return any(k in context_text for k in kws)

    if has("전자책", "글쓰기", "블로그", "콘텐츠", "PDF", "크몽", "출판", "책 쓰"):
        pass  # 아래 책 모티프로 (전자책이 월급·수익 키워드보다 우선)
    elif has("배당", "현금흐름", "월급", "적금", "연금", "입금", "적립"):
        # 월 12칸 입금 캘린더 — 매달 들어오는 돈 (자신감 있는 크기)
        cells = []
        filled_idx = (0, 3, 5, 8, 10, 11)
        for r in range(3):
            for c in range(4):
                i = r * 4 + c
                x = 390 + c * 220
                y = 1560 + r * 190
                if i in filled_idx:
                    cells.append(f'<rect x="{x}" y="{y}" width="160" height="130" rx="20" fill="{accent}" opacity="0.85"/>')
                else:
                    cells.append(f'<rect x="{x}" y="{y}" width="160" height="130" rx="20" fill="none" stroke="{accent}" stroke-width="2.5" opacity="0.45"/>')
        return ''.join(cells)
    if has("전자책", "글쓰기", "블로그", "콘텐츠", "PDF", "크몽", "출판", "책 쓰"):
        # 펼친 책 (자신감 있는 크기)
        return (f'<g stroke="{accent}" stroke-width="3" opacity="0.8" fill="none">'
                f'<path d="M 800 2080 C 620 1980 450 1980 320 2030 L 320 1640 C 450 1590 620 1590 800 1690 Z"/>'
                f'<path d="M 800 2080 C 980 1980 1150 1980 1280 2030 L 1280 1640 C 1150 1590 980 1590 800 1690 Z"/>'
                f'<line x1="800" y1="1690" x2="800" y2="2080"/>'
                f'<line x1="410" y1="1745" x2="710" y2="1700" opacity="0.5"/>'
                f'<line x1="410" y1="1840" x2="710" y2="1795" opacity="0.5"/>'
                f'<line x1="410" y1="1935" x2="710" y2="1890" opacity="0.5"/>'
                f'<line x1="890" y1="1700" x2="1190" y2="1745" opacity="0.5"/>'
                f'<line x1="890" y1="1795" x2="1190" y2="1840" opacity="0.5"/>'
                f'<line x1="890" y1="1890" x2="1190" y2="1935" opacity="0.5"/></g>')
    if has("습관", "루틴", "아침", "새벽", "운동", "다이어트", "미라클"):
        # 체크 동그라미 연속 — 쌓이는 실행 (자신감 있는 크기)
        out = []
        for i in range(5):
            cx = 400 + i * 200
            done = i < 3
            out.append(f'<circle cx="{cx}" cy="1820" r="62" fill="none" stroke="{accent}" stroke-width="3" opacity="{0.9 if done else 0.35}"/>')
            if done:
                out.append(f'<path d="M {cx - 24} 1820 L {cx - 5} 1842 L {cx + 27} 1796" stroke="{soft}" stroke-width="7" fill="none"/>')
        return ''.join(out)
    if has("마케팅", "사업", "매출", "브랜딩", "세일즈", "스토어"):
        # 상승 막대 (자신감 있는 크기)
        bars = []
        for i, h in enumerate((120, 185, 160, 265, 350)):
            bars.append(f'<rect x="{420 + i * 160}" y="{2060 - h}" width="95" height="{h}" fill="{accent}" opacity="{0.32 + i * 0.13}"/>')
        return ''.join(bars)
    # 기본 (주식·투자 등): 상승 라인
    return (f'<path d="M 180 2030 L 470 1890 L 740 1950 L 1050 1750 L 1420 1610" fill="none" stroke="{accent}" stroke-width="3.5" opacity="0.65"/>'
            f'<g fill="{soft}" opacity="0.8"><circle cx="1050" cy="1750" r="9"/><circle cx="1420" cy="1610" r="12"/></g>')


def build_cover_svg(template, title, subtitle, author):
    """베스트셀러 표지 v4 — 절제된 장식 + 강한 타이포 위계 (1600x2560)
    원칙: 큰 제목, 넉넉한 여백, 색은 2가지 이내, 장식은 주제 맞춤 모티프 1개만"""
    sub_raw = (subtitle or "").strip()
    sub = html.escape(sub_raw)
    auth_raw = (author or "").strip()
    auth = html.escape(auth_raw)
    auth_jieum = (html.escape(auth_raw) + " 지음") if auth_raw else ""
    serif = "'Noto Serif KR','Nanum Myeongjo',serif"
    songmyung = "'Song Myung','Noto Serif KR',serif"
    gowun = "'Gowun Batang','Noto Serif KR',serif"
    gmarket = "'GmarketSans','Noto Sans KR',sans-serif"
    blackhan = "'Black Han Sans','Noto Sans KR',sans-serif"
    pretendard = "'Pretendard','Noto Sans KR',sans-serif"
    sans = "'Noto Sans KR','Malgun Gothic',sans-serif"
    fonts = (
        '<style><![CDATA['
        "@import url('https://fonts.googleapis.com/css2?family=Noto+Serif+KR:wght@600;700;900&family=Song+Myung&family=Gowun+Batang:wght@400;700&family=Black+Han+Sans&family=Noto+Sans+KR:wght@400;500;700;900&display=swap');"
        "@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');"
        "@font-face{font-family:'GmarketSans';src:url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_2001@1.1/GmarketSansMedium.woff') format('woff');font-weight:500;}"
        "@font-face{font-family:'GmarketSans';src:url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_2001@1.1/GmarketSansBold.woff') format('woff');font-weight:700;}"
        ']]></style>'
    )

    # 질감 필터: 필름 그레인(다크용) + 종이결(라이트용) — 인쇄물 질감
    texture_defs = (
        '<defs>'
        '<filter id="grainD" x="0" y="0" width="100%" height="100%">'
        '<feTurbulence type="fractalNoise" baseFrequency="0.8" numOctaves="3" stitchTiles="stitch"/>'
        '<feColorMatrix type="matrix" values="0 0 0 0 1  0 0 0 0 1  0 0 0 0 0.9  0 0 0 0.04 0"/>'
        '</filter>'
        '<filter id="grainL" x="0" y="0" width="100%" height="100%">'
        '<feTurbulence type="fractalNoise" baseFrequency="0.75" numOctaves="3" stitchTiles="stitch"/>'
        '<feColorMatrix type="matrix" values="0 0 0 0 0.2  0 0 0 0 0.18  0 0 0 0 0.15  0 0 0 0.03 0"/>'
        '</filter>'
        '</defs>'
    )
    grain_dark = '<rect width="1600" height="2560" filter="url(#grainD)"/>'
    grain_light = '<rect width="1600" height="2560" filter="url(#grainL)"/>'

    # 강조어: 제목에서 가장 긴 단어 (2글자 이상)
    _words = [w for w in title.split() if len(w) >= 2]
    key_word = max(_words, key=len) if _words else ""

    def wrap(maxc):
        return [l for l in _wrap_title_lines(title, max_chars=maxc, max_lines=3)]

    def em_line(line, accent):
        """줄 안의 강조어에만 색 적용"""
        esc = html.escape(line)
        if key_word and key_word in line:
            k = html.escape(key_word)
            return esc.replace(k, f'<tspan fill="{accent}">{k}</tspan>', 1)
        return esc

    def title_block(x, y_center, lines, fs, family, weight, fill, accent=None, anchor="middle", ls=2, lh=1.26):
        line_h = int(fs * lh)
        top = y_center - int((len(lines) - 1) * line_h / 2)
        spans = "".join(
            f'<tspan x="{x}" dy="{0 if i == 0 else line_h}">{em_line(l, accent) if accent else html.escape(l)}</tspan>'
            for i, l in enumerate(lines))
        return (f'<text x="{x}" y="{top}" text-anchor="{anchor}" font-family="{family}" '
                f'font-size="{fs}" font-weight="{weight}" fill="{fill}" letter-spacing="{ls}">{spans}</text>'), top, line_h

    def fit(lines, base, mid, small):
        lg = max((len(l) for l in lines), default=4)
        return base if lg <= 4 else (mid if lg <= 6 else small)

    def eyebrow(x, y, text, color, family=None, fs=28, ls=12, anchor="middle"):
        fam = family or pretendard
        return f'<text x="{x}" y="{y}" text-anchor="{anchor}" font-family="{fam}" font-size="{fs}" font-weight="500" fill="{color}" letter-spacing="{ls}">{text}</text>'

    def hook(x, y, fill, font=None, anchor="middle", weight=600):
        """제목 위/아래에 놓는 부제 한 줄 (길이에 따라 자동 축소)"""
        if not sub:
            return ""
        f = font or serif
        hfs = 50 if len(sub_raw) <= 18 else (44 if len(sub_raw) <= 26 else 37)
        return (f'<text x="{x}" y="{y}" text-anchor="{anchor}" font-family="{f}" '
                f'font-size="{hfs}" font-weight="{weight}" fill="{fill}" letter-spacing="1">{sub}</text>')

    def copy_two_lines(x, y, text, fill, font=None, fs=None, weight=700):
        """카피 한 덩어리 — 길이에 따라 자동 1~2줄 (어떤 길이든 반드시 표시됨)"""
        if not text:
            return ''
        tf = font or sans
        n = len(text)
        f = fs or (64 if n <= 18 else (56 if n <= 26 else 46))
        esc = html.escape(text)
        if n <= 20:
            return f'<text x="{x}" y="{y}" text-anchor="middle" font-family="{tf}" font-size="{f}" font-weight="{weight}" fill="{fill}">{esc}</text>'
        mid_i = n // 2
        spaces = [i for i, c in enumerate(text) if c == ' ' and 3 < i < n - 3]
        cut = min(spaces, key=lambda i: abs(i - mid_i)) if spaces else mid_i
        l1, l2 = html.escape(text[:cut].strip()), html.escape(text[cut:].strip())
        g = int(f * 1.45)
        return (f'<text x="{x}" y="{y}" text-anchor="middle" font-family="{tf}" font-size="{f}" font-weight="{weight}" fill="{fill}">{l1}</text>'
                f'<text x="{x}" y="{y + g}" text-anchor="middle" font-family="{tf}" font-size="{f}" font-weight="{weight}" fill="{fill}">{l2}</text>')

    def belt(content, band_bg, text_fill, hairline=None, text_font=None, y=2240, h=320):
        """하단 띠지 — 실물 단행본의 카피 밴드 (슬림, 자동 줄바꿈)"""
        if not content:
            return f'<rect x="0" y="{y}" width="1600" height="{h}" fill="{band_bg}"/>'
        tf = text_font or serif
        qfs = 56 if len(content) <= 18 else (48 if len(content) <= 26 else 41)
        line = f'<line x1="0" y1="{y}" x2="1600" y2="{y}" stroke="{hairline}" stroke-width="3"/>' if hairline else ''
        if len(content) > 24:
            mid_i = len(content) // 2
            spaces = [i for i, c in enumerate(content) if c == ' ' and 3 < i < len(content) - 3]
            def _cut_score(i):
                s = abs(i - mid_i)
                if i + 1 < len(content) and content[i + 1].isdigit():
                    s += 10  # 숫자 앞에서 자르지 않기 ("월 / 100" 방지)
                prev_tok = content[:i].split()[-1] if content[:i].split() else ''
                if len(prev_tok) <= 1:
                    s += 6   # 한 글자 단어 뒤에서 자르지 않기
                return s
            cut = min(spaces, key=_cut_score) if spaces else mid_i
            l1, l2 = content[:cut].strip(), content[cut:].strip()
            txt = (f'<text x="800" y="{y + 130}" text-anchor="middle" font-family="{tf}" font-size="{qfs}" font-weight="600" fill="{text_fill}">{l1}</text>'
                   f'<text x="800" y="{y + 212}" text-anchor="middle" font-family="{tf}" font-size="{qfs}" font-weight="600" fill="{text_fill}">{l2}</text>')
        else:
            txt = f'<text x="800" y="{y + h // 2 + 18}" text-anchor="middle" font-family="{tf}" font-size="{qfs}" font-weight="600" fill="{text_fill}">{content}</text>'
        return f'<rect x="0" y="{y}" width="1600" height="{h}" fill="{band_bg}"/>{line}{txt}'

    quote = f'“{sub_raw}”' if sub_raw else ''
    q_esc = html.escape(quote)

    # AI가 제목과 함께 생성한 표지 카피 (눈썹 카피 + 띠지 약속 카피) — 카피 밀도가 구매 욕구를 만든다
    try:
        eyebrow_copy = html.escape((st.session_state.get('cover_eyebrow') or '').strip())
        belt_copy_raw = (st.session_state.get('cover_belt_copy') or '').strip()
    except Exception:
        eyebrow_copy, belt_copy_raw = '', ''

    if template == "two_tone_frame":
        # 베스트셀러 1군 공식 ①: 블랙 블록 + 블록을 가득 채우는 초대형 백색 제목
        # + 포인트 컬러 프레임 + 하단 화이트 영역에 별점·카피 (참고: 매대 상위권 구도)
        acc_palette = ['#5B2D8E', '#B0322A', '#0F6E5C', '#1A3FA8']
        acc = acc_palette[abs(hash(title or '')) % len(acc_palette)]
        lines = wrap(6); fs = fit(lines, 300, 240, 185)
        t, top, lh = title_block(800, 840, lines, fs, gmarket, 700, '#FFFFFF', accent=None, ls=-3, lh=1.18)
        bottom = top + (len(lines) - 1) * lh
        stars = ''.join(f'<text x="{668 + i * 66}" y="2300" text-anchor="middle" font-family="{sans}" font-size="40" fill="{acc}" opacity="0.85">★</text>' for i in range(5))
        body = (
            '<rect width="1600" height="2560" fill="#FFFFFF"/>'
            '<rect x="0" y="0" width="1600" height="1560" fill="#101010"/>'
            # 포인트 컬러 프레임 — 가는 ㄱ자 브래킷 + 절제된 말풍선 꼬리 (두껍지 않게)
            f'<path d="M 110 380 L 110 110 L 440 110" fill="none" stroke="{acc}" stroke-width="13"/>'
            f'<path d="M 1490 1290 L 1490 1640 L 1160 1640" fill="none" stroke="{acc}" stroke-width="13"/>'
            f'<path d="M 1030 1560 L 1130 1560 L 1030 1690 Z" fill="{acc}"/>'
            + (eyebrow(1450, 196, auth_jieum, '#C9C9C9', sans, 34, 2, "end") if auth_jieum else "")
            # 눈썹 카피 — 제목 위 관형구 ("품격 있는 대화를 위한" 위계, 화면 밖 잘림 방지)
            + (f'<text x="800" y="{max(300, top - int(fs * 0.95) - 80)}" text-anchor="middle" font-family="{gmarket}" font-size="58" font-weight="500" fill="#E8E8E8">{eyebrow_copy}</text>' if eyebrow_copy else "")
            + t
            # 핵심 카피 — 띠지 약속(있으면) 또는 부제, 포인트 컬러 대형
            + copy_two_lines(800, 1880, belt_copy_raw or sub_raw, acc, gmarket)
            + f'<line x1="720" y1="2120" x2="880" y2="2120" stroke="{acc}" stroke-width="5"/>'
            + stars
            # 띠지 카피를 썼으면 부제도 작게 한 줄 더 (카피 밀도)
            + (copy_two_lines(800, 2420, sub_raw, '#3A3A3A', sans, fs=42, weight=500) if belt_copy_raw and sub_raw else "")
            + grain_light
        )
    elif template == "diagonal_block":
        # 베스트셀러 1군 공식 ②: 비비드 단색 + 초대형 검정 제목 + 부제 2줄 + 대각 블랙 블록(주제 모티프)
        dg_palette = ['#FF7A1A', '#FFD43B', '#3ECF9A', '#4DA8E8']
        bg = dg_palette[abs(hash(title or '')) % len(dg_palette)]
        lines = wrap(6); fs = fit(lines, 310, 245, 190)
        # 3줄 제목이면 중심을 내려 눈썹 카피 공간 확보 (카피가 위로 잘리는 것 방지)
        t, top, lh = title_block(800, 600 if len(lines) <= 2 else 760, lines, fs, gmarket, 700, '#111111', accent=None, ls=-3, lh=1.14)
        bottom = top + (len(lines) - 1) * lh
        _eb_y = max(155, top - int(fs * 0.95) - 70)  # 화면 밖으로 못 나가게 고정
        body = (
            f'<rect width="1600" height="2560" fill="{bg}"/>'
            # 눈썹 카피 — 제목 위 관형구
            + (f'<text x="800" y="{_eb_y}" text-anchor="middle" font-family="{gmarket}" font-size="56" font-weight="500" fill="#111111">{eyebrow_copy}</text>' if eyebrow_copy else "")
            + t
            # 부제 — 길이와 무관하게 항상 1~2줄, 제목보다 가벼운 굵기 (타이포 위계)
            + copy_two_lines(800, bottom + 190, sub_raw, '#111111', gmarket, weight=500)
            + f'<g stroke="#111111" stroke-width="3.5" fill="{bg}">'
              f'<line x1="420" y1="{bottom + 410}" x2="1180" y2="{bottom + 410}"/>'
              f'<circle cx="420" cy="{bottom + 410}" r="12" fill="#111111"/>'
              f'<circle cx="1180" cy="{bottom + 410}" r="12" fill="#111111"/></g>'
            + (eyebrow(800, bottom + 525, auth_jieum, '#111111', sans, 38, 2) if auth_jieum else "")
            + '<polygon points="0,1900 1600,1620 1600,2560 0,2560" fill="#111111"/>'
            # 블랙 블록: 띠지 약속 카피(있으면) — 없으면 주제 모티프
            + (copy_two_lines(800, 2150, belt_copy_raw, bg, gmarket, fs=72)
               + ''.join(f'<text x="{700 + i * 100}" y="2400" text-anchor="middle" font-family="{sans}" font-size="44" fill="{bg}" opacity="0.85">★</text>' for i in range(3))
               if belt_copy_raw else
               f'<g transform="translate(0,310)">{_cover_motif(f"{title} {sub_raw}", accent=bg, soft=bg)}</g>')
            + grain_light
        )
    elif template == "solid_impact":
        # 실제 베스트셀러 매대 표준 공식 — 밝은 단색 배경 + 묵직한 검정 타이포 + 채움형 오브제 + 검정 띠지
        palettes = [('#FFD93B', '#171511'), ('#6FD6B9', '#0E2A21'), ('#7EC8F2', '#0D2233'), ('#FF8A4C', '#2A1305')]
        bg, ink_c = palettes[abs(hash(title or '')) % len(palettes)]
        lines = wrap(6); fs = fit(lines, 275, 220, 170)
        t, top, lh = title_block(800, 880, lines, fs, gmarket, 700, ink_c, accent='#FFFFFF', ls=-3, lh=1.2)
        bottom = top + (len(lines) - 1) * lh
        body = (
            f'<rect width="1600" height="2560" fill="{bg}"/>'
            + (eyebrow(800, 300, auth_jieum, ink_c, sans, 36, 2) if auth_jieum else "")
            + f'<rect x="730" y="{top - int(fs * 0.95) - 60}" width="140" height="20" fill="{ink_c}"/>'
            + t
            + _cover_motif(f"{title} {sub_raw}", accent=ink_c, soft=bg)
            + belt(q_esc or auth_jieum, ink_c, bg, text_font=sans)
            + grain_light
        )
    elif template == "noir_gold":
        # 블랙 + 금박 — 종이책 베스트셀러 위계 (제목 → 부제 → 저자), 띠지 없이 여백으로 승부
        lines = wrap(7); fs = fit(lines, 320, 230, 178)
        t, top, lh = title_block(800, 1000, lines, fs, serif, 900, '#EFE5C9', accent='#C9A24B', ls=0, lh=1.22)
        bottom = top + (len(lines) - 1) * lh
        body = (
            '<defs><radialGradient id="ng" cx="50%" cy="32%" r="90%"><stop offset="0%" stop-color="#1C1913"/><stop offset="58%" stop-color="#0E0D0A"/><stop offset="100%" stop-color="#040404"/></radialGradient></defs>'
            '<rect width="1600" height="2560" fill="url(#ng)"/>'
            + f'<line x1="744" y1="{top - int(fs * 0.95) - 70}" x2="856" y2="{top - int(fs * 0.95) - 70}" stroke="#C9A24B" stroke-width="3"/>'
            + t
            + (eyebrow(800, bottom + 185, auth_jieum, '#9A8C68', serif, 36, 1) if auth_jieum else "")
            + _cover_motif(f"{title} {sub_raw}")
            + belt(q_esc or auth_jieum, '#EFE5C9', '#241F15', hairline='#C9A24B')
            + grain_dark
        )
    elif template == "deep_navy":
        # 딥 네이비 — 프레스티지 금융 (굵은 명조 대형 제목 + 은은한 상승 라인)
        lines = wrap(7); fs = fit(lines, 320, 230, 178)
        t, top, lh = title_block(800, 980, lines, fs, serif, 900, '#F3ECD9', accent='#D9B45E', ls=0, lh=1.22)
        bottom = top + (len(lines) - 1) * lh
        body = (
            '<defs><radialGradient id="dn" cx="50%" cy="30%" r="92%"><stop offset="0%" stop-color="#1D3A5F"/><stop offset="100%" stop-color="#091320"/></radialGradient></defs>'
            '<rect width="1600" height="2560" fill="url(#dn)"/>'
            + _cover_motif(f"{title} {sub_raw}")
            + t
            + (eyebrow(800, bottom + 250, auth_jieum, '#93A9C2', serif, 38, 1) if auth_jieum else "")
            + belt(q_esc or auth_jieum, '#F3ECD9', '#13253C')
            + grain_dark
        )
    elif template == "editorial":
        # 아이보리 클래식 — 더블 룰 + 명조 (서점 인문 베스트셀러)
        lines = wrap(7); fs = fit(lines, 290, 215, 168)
        t, top, lh = title_block(800, 1060, lines, fs, gowun, 700, '#1C1A17', accent='#8C3A2E', ls=0, lh=1.22)
        bottom = top + (len(lines) - 1) * lh
        body = (
            '<rect width="1600" height="2560" fill="#F5F1E8"/>'
            '<line x1="120" y1="170" x2="1480" y2="170" stroke="#1C1A17" stroke-width="5"/>'
            '<line x1="120" y1="186" x2="1480" y2="186" stroke="#1C1A17" stroke-width="1.2"/>'
            '<line x1="120" y1="2380" x2="1480" y2="2380" stroke="#1C1A17" stroke-width="1.2"/>'
            '<line x1="120" y1="2396" x2="1480" y2="2396" stroke="#1C1A17" stroke-width="5"/>'
            + '<g transform="translate(800 520)"><rect x="-13" y="-13" width="26" height="26" transform="rotate(45)" fill="none" stroke="#8C3A2E" stroke-width="2"/>'
              '<rect x="-6" y="-6" width="12" height="12" transform="rotate(45)" fill="#8C3A2E"/>'
              '<line x1="-110" y1="0" x2="-34" y2="0" stroke="#8A7A55" stroke-width="1.2"/><line x1="34" y1="0" x2="110" y2="0" stroke="#8A7A55" stroke-width="1.2"/></g>'
            + hook(800, top - int(fs * 0.95) - 80, '#6E6450')
            + t
            + f'<circle cx="800" cy="{bottom + 110}" r="6" fill="#8C3A2E"/>'
            + (eyebrow(800, bottom + 240, auth_jieum, '#55503F', serif, 38, 1) if auth_jieum else "")
            + grain_light
        )
    elif template == "minimal_type":
        # 딥그린 미니멀 — 거대한 배경 글자 + 좌측 정렬 타이포
        lines = wrap(6); fs = fit(lines, 270, 210, 162)
        t, top, lh = title_block(150, 1000, lines, fs, gowun, 700, '#F2EDE3', accent='#D9BC7A', anchor="start", ls=0, lh=1.22)
        bottom = top + (len(lines) - 1) * lh
        body = (
            '<rect width="1600" height="2560" fill="#0F3D2E"/>'
            + f'<text x="1640" y="2330" text-anchor="end" font-family="{gowun}" font-size="1450" font-weight="700" fill="#155039" opacity="0.95">{html.escape((title or "글")[0])}</text>'
            + f'<line x1="150" y1="{top - int(fs * 0.95) - 60}" x2="390" y2="{top - int(fs * 0.95) - 60}" stroke="#C8B27C" stroke-width="4"/>'
            + t
            + hook(150, bottom + 175, '#A9C0B4', anchor="start")
            + (eyebrow(150, 2200, auth_jieum, '#C8B27C', serif, 38, 1, "start") if auth_jieum else "")
            + '<line x1="150" y1="2300" x2="1450" y2="2300" stroke="#C8B27C" stroke-width="1" opacity="0.45"/>'
            + grain_dark
        )
    elif template == "color_block":
        # 코발트 임팩트 — 색면 + 헤비 산세리프 (동기부여 베스트셀러)
        lines = wrap(6); fs = fit(lines, 270, 215, 168)
        t, top, lh = title_block(150, 1000, lines, fs, gmarket, 700, '#FFFFFF', accent='#FFD43B', anchor="start", ls=-3, lh=1.2)
        bottom = top + (len(lines) - 1) * lh
        body = (
            '<rect width="1600" height="2560" fill="#1D4ED8"/>'
            '<circle cx="1420" cy="330" r="430" fill="#1A45BE"/>'
            '<circle cx="1505" cy="245" r="150" fill="#FFD43B"/>'
            + eyebrow(150, 300, "CHANGE STARTS TODAY", '#9DB6F5', sans, 30, 9, "start")
            + f'<rect x="150" y="{top - int(fs * 0.92) - 60}" width="170" height="22" fill="#FFD43B"/>'
            + t
            + (eyebrow(150, bottom + 250, auth_jieum, '#FFFFFF', sans, 38, 4, "start") if auth_jieum else "")
            + belt(q_esc or auth_jieum, '#0B0F1E', '#FFFFFF', hairline='#FFD43B', text_font=sans)
        )
    elif template == "big_number":
        # 빅 넘버 — 숫자가 주인공 (기간·챌린지)
        m = re.search(r'\d[\d,\.]*', f"{title} {sub_raw}")
        bignum = html.escape(m.group(0)) if m else html.escape((title or "1")[:1])
        nfs = 920 if len(bignum) <= 2 else (680 if len(bignum) <= 3 else 480)
        lines = wrap(7); fs = fit(lines, 215, 175, 140)
        t, top, lh = title_block(800, 1530, lines, fs, gmarket, 700, '#15130F', accent='#D8483A', ls=-2, lh=1.22)
        bottom = top + (len(lines) - 1) * lh
        body = (
            '<rect width="1600" height="2560" fill="#F5F2EC"/>'
            f'<text x="800" y="1020" text-anchor="middle" font-family="{gmarket}" font-size="{nfs}" font-weight="700" fill="#ECE2D0">{bignum}</text>'
            f'<text x="800" y="1020" text-anchor="middle" font-family="{gmarket}" font-size="{nfs}" font-weight="700" fill="none" stroke="#D8483A" stroke-width="5">{bignum}</text>'
            + f'<rect x="730" y="{top - int(fs * 0.95) - 55}" width="140" height="18" fill="#D8483A"/>'
            + t
            + (eyebrow(800, bottom + 270, auth_jieum, '#8E8675', sans, 36, 4) if auth_jieum else "")
            + belt(q_esc or auth_jieum, '#D8483A', '#FFF6EE', text_font=sans)
            + grain_light
        )
    elif template == "gradient_modern":
        # 모던 그라데이션 — IT·트렌드 (클린, 글로우 포인트 1개)
        lines = wrap(6); fs = fit(lines, 280, 220, 170)
        t, top, lh = title_block(800, 1000, lines, fs, pretendard, 800, '#FFFFFF', accent='#99F6E4', ls=-2, lh=1.2)
        bottom = top + (len(lines) - 1) * lh
        body = (
            '<defs><linearGradient id="gm" x1="0%" y1="0%" x2="100%" y2="100%">'
            '<stop offset="0%" stop-color="#3730A3"/><stop offset="55%" stop-color="#7C3AED"/><stop offset="100%" stop-color="#0D9488"/></linearGradient>'
            '<radialGradient id="bk1" cx="50%" cy="50%" r="50%"><stop offset="0%" stop-color="#FFFFFF" stop-opacity="0.32"/><stop offset="100%" stop-color="#FFFFFF" stop-opacity="0"/></radialGradient></defs>'
            '<rect width="1600" height="2560" fill="url(#gm)"/>'
            '<circle cx="1300" cy="400" r="360" fill="url(#bk1)"/>'
            '<circle cx="1300" cy="400" r="300" fill="none" stroke="#FFFFFF" stroke-width="1.4" opacity="0.35"/>'
            '<circle cx="270" cy="1900" r="430" fill="url(#bk1)" opacity="0.5"/>'
            + t
            + (eyebrow(800, bottom + 250, auth_jieum, '#FFFFFF', sans, 36, 1) if auth_jieum else "")
            + belt(q_esc or auth_jieum, '#0B0B14', '#FFFFFF', text_font=sans).replace('fill="#0B0B14"/>', 'fill="#0B0B14" opacity="0.55"/>')
            + grain_dark
        )
    elif template == "warm_essay":
        # 웜 에세이 — 종이 위 노을 (원 + 수평선의 미니멀 풍경)
        lines = wrap(7); fs = fit(lines, 240, 190, 150)
        t, top, lh = title_block(800, 1240, lines, fs, gowun, 700, '#5B4636', accent='#C96F4A', ls=0, lh=1.22)
        bottom = top + (len(lines) - 1) * lh
        body = (
            '<rect width="1600" height="2560" fill="#F7EFE6"/>'
            '<circle cx="800" cy="640" r="240" fill="#EFBE96" opacity="0.9"/>'
            '<circle cx="800" cy="640" r="240" fill="none" stroke="#E0A87C" stroke-width="1.5" opacity="0.6"/>'
            '<rect x="0" y="760" width="1600" height="6" fill="#F7EFE6"/>'
            '<line x1="330" y1="762" x2="1270" y2="762" stroke="#B98C68" stroke-width="2.5" opacity="0.7"/>'
            '<line x1="450" y1="826" x2="1150" y2="826" stroke="#C8A287" stroke-width="1.4" opacity="0.55"/>'
            + t
            + (eyebrow(800, bottom + 270, auth_jieum, '#8A7361', serif, 38, 1) if auth_jieum else "")
            + belt(q_esc or auth_jieum, '#EFD9C3', '#6B523D')
            + grain_light
        )
    elif template == "geometric":
        # 스위스 기하 — 화이트 + 블랙 대각 + 레드 서클 (비즈니스 전략서)
        lines = wrap(6); fs = fit(lines, 255, 205, 160)
        t, top, lh = title_block(150, 970, lines, fs, gmarket, 700, '#141414', accent='#E03E2F', anchor="start", ls=-3, lh=1.2)
        bottom = top + (len(lines) - 1) * lh
        body = (
            '<rect width="1600" height="2560" fill="#FAFAF7"/>'
            '<polygon points="1600,0 1600,760 880,0" fill="#141414"/>'
            '<circle cx="1230" cy="330" r="150" fill="#E03E2F"/>'
            '<circle cx="1230" cy="330" r="150" fill="none" stroke="#FAFAF7" stroke-width="3"/>'
            + eyebrow(150, 300, "STRATEGY PLAYBOOK", '#9A968A', sans, 29, 8, "start")
            + f'<rect x="150" y="{top - int(fs * 0.92) - 60}" width="200" height="18" fill="#E03E2F"/>'
            + t
            + (eyebrow(150, bottom + 250, auth_jieum, '#141414', sans, 36, 4, "start") if auth_jieum else "")
            + belt(q_esc or auth_jieum, '#141414', '#FAFAF7', hairline='#E03E2F', text_font=sans)
            + grain_light
        )
    elif template == "red_punch":
        # 레드 펀치 — 블랙 + 레드 도발 (머니·각성)
        lines = wrap(6); fs = fit(lines, 285, 228, 178)
        t, top, lh = title_block(800, 1000, lines, fs, blackhan, 400, '#E8362B', accent='#FFFFFF', ls=0, lh=1.18)
        bottom = top + (len(lines) - 1) * lh
        body = (
            '<rect width="1600" height="2560" fill="#0C0C0C"/>'
            + eyebrow(800, 320, "WAKE UP CALL", '#5A5A5A', sans, 28, 13)
            + t
            + f'<line x1="560" y1="{bottom + 90}" x2="1040" y2="{bottom + 90}" stroke="#E8362B" stroke-width="10"/>'
            + (eyebrow(800, bottom + 330, auth_jieum, '#9A9A9A', sans, 36, 4) if auth_jieum else "")
            + belt(q_esc or auth_jieum, '#E8362B', '#1A0503', text_font=sans)
            + grain_dark
        )
    else:  # 알 수 없는 템플릿 → noir_gold 재귀
        return build_cover_svg("noir_gold", title, subtitle, author)

    return ('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1600 2560" width="100%" preserveAspectRatio="xMidYMid meet">'
            + fonts + texture_defs + body + '</svg>')


def _estimate_text_width(s, fs, ls=0.0):
    """한글/영문/숫자 폭 추정 (cairosvg의 폰트 폴백 폭 계산 오류를 우회하기 위한 자체 계산)"""
    w = 0.0
    for ch in s:
        o = ord(ch)
        if 0x1100 <= o <= 0x11FF or 0x2E80 <= o <= 0xD7FF or 0xF900 <= o <= 0xFAFF or 0xFF01 <= o <= 0xFF60:
            w += 1.0      # 한글/CJK 전각
        elif ch == ' ':
            w += 0.34
        elif ch.isdigit():
            w += 0.60
        elif ch.isupper():
            w += 0.72
        else:
            w += 0.52
    return w * fs + ls * max(len(s) - 1, 0)


def _fix_text_anchors_for_png(svg):
    """cairosvg는 폰트 폴백 시 글자 폭을 잘못 재서 middle/end 정렬이 깨진다(제목이 잘려 보임).
    → 앵커를 전부 start로 바꾸고 좌표를 직접 계산해 어떤 PC에서도 정렬이 유지되게 한다."""
    def fix_text(m):
        tag = m.group(0)
        open_tag = m.group(1)
        inner = m.group(2)
        am = re.search(r'text-anchor="(middle|end)"', open_tag)
        if not am:
            return tag
        anchor = am.group(1)
        fs_m = re.search(r'font-size="([\d.]+)"', open_tag)
        ls_m = re.search(r'letter-spacing="(-?[\d.]+)"', open_tag)
        fs = float(fs_m.group(1)) if fs_m else 16.0
        ls = float(ls_m.group(1)) if ls_m else 0.0

        def shift(x_val, content):
            text_only = re.sub(r'<[^>]+>', '', content)
            text_only = html.unescape(text_only)
            w = _estimate_text_width(text_only, fs, ls)
            return x_val - (w / 2 if anchor == 'middle' else w)

        new_open = open_tag.replace(f'text-anchor="{anchor}"', 'text-anchor="start"')

        if '<tspan' in inner:
            # x 속성을 가진 '줄 tspan' 단위로 분할 (강조용 중첩 tspan은 줄 내용에 포함해 폭 계산)
            parts = re.split(r'(<tspan\b[^>]*\bx="[^"]*"[^>]*>)', inner)
            out = parts[0]
            for i in range(1, len(parts), 2):
                open_t = parts[i]
                rest = parts[i + 1] if i + 1 < len(parts) else ''
                xm = re.search(r'x="(-?[\d.]+)"', open_t)
                if xm:
                    content = re.sub(r'<[^>]+>', '', rest)
                    new_x = shift(float(xm.group(1)), html.unescape(content))
                    open_t = open_t.replace(f'x="{xm.group(1)}"', f'x="{new_x:.1f}"', 1)
                out += open_t + rest
            return new_open + out + '</text>'
        else:
            xm = re.search(r'\bx="(-?[\d.]+)"', new_open)
            if not xm:
                return new_open + inner + '</text>'
            new_x = shift(float(xm.group(1)), inner)
            new_open = new_open.replace(f'x="{xm.group(1)}"', f'x="{new_x:.1f}"', 1)
            return new_open + inner + '</text>'

    return re.sub(r'(<text\b[^>]*>)(.*?)(</text>)',
                  lambda m: fix_text(m), svg, flags=re.DOTALL)


def prepare_svg_for_png(svg):
    """cairosvg 변환용 전처리: 미지원 필터(그레인) 제거 + 서버 한글 폰트로 매핑 + 정렬 좌표 직접 계산"""
    # 웹폰트 import 제거
    svg = re.sub(r'<style><!\[CDATA\[.*?\]\]></style>', '', svg, flags=re.DOTALL)
    # 그레인 필터 정의 + 적용 rect 제거 (cairosvg가 필터를 무시하고 검은 사각형으로 그리는 문제)
    svg = re.sub(r'<filter id="grain[DL]".*?</filter>', '', svg, flags=re.DOTALL)
    svg = re.sub(r'<rect width="1600" height="2560" filter="url\(#grain[DL]\)"/>', '', svg)
    # ⚠️ cairosvg는 font-family 목록 중 '첫 번째 폰트만' 사용하고 폴백을 무시한다
    # (cairosvg/text.py: .split(',')[0]) → 첫 폰트가 없는 PC에서는 글자가 아예 안 그려져
    # 표지가 검은 배경만 나올 수 있다. OS별로 반드시 설치돼 있는 한글 폰트를 1순위로 지정.
    _os_name = platform.system()
    if _os_name == 'Windows':
        _kfont = "'Malgun Gothic'"          # 윈도우 기본 한글 폰트
    elif _os_name == 'Darwin':
        _kfont = "'Apple SD Gothic Neo'"    # macOS 기본 한글 폰트
    else:
        _kfont = "'Noto Sans CJK KR'"       # 리눅스/Streamlit Cloud (packages.txt: fonts-noto-cjk)
    _serif_to = _kfont + ",sans-serif"
    _sans_to = _kfont + ",sans-serif"
    for fam, repl in [
        ("'Noto Serif KR','Nanum Myeongjo',serif", _serif_to),
        ("'Song Myung','Noto Serif KR',serif", _serif_to),
        ("'Gowun Batang','Noto Serif KR',serif", _serif_to),
        ("'GmarketSans','Noto Sans KR',sans-serif", _sans_to),
        ("'Black Han Sans','Noto Sans KR',sans-serif", _sans_to),
        ("'Pretendard','Noto Sans KR',sans-serif", _sans_to),
        ("'Noto Sans KR','Malgun Gothic',sans-serif", _sans_to),
        ("Georgia,serif", _serif_to),
    ]:
        svg = svg.replace(fam, repl)
    # 가운데/끝 정렬을 좌표 직접 계산으로 변환 (폰트 폴백 환경에서 제목 잘림 방지)
    svg = _fix_text_anchors_for_png(svg)
    return svg

def extract_video_id(url):
    """YouTube URL에서 video ID 추출"""
    patterns = [
        r'(?:youtube\.com\/watch\?v=|youtu\.be\/|youtube\.com\/embed\/)([^&\n?#]+)',
        r'youtube\.com\/watch\?.*v=([^&\n?#]+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def get_youtube_transcript(video_id):
    """YouTube 자막 가져오기"""
    if not YOUTUBE_TRANSCRIPT_AVAILABLE:
        return None, "youtube-transcript-api가 설치되지 않았습니다. pip install youtube-transcript-api"

    try:
        # 한국어 자막 우선, 없으면 영어, 없으면 자동생성
        transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)

        transcript = None
        # 수동 자막 먼저 시도
        for lang in ['ko', 'en']:
            try:
                transcript = transcript_list.find_transcript([lang])
                break
            except:
                continue

        # 수동 자막 없으면 자동 생성 자막
        if not transcript:
            try:
                transcript = transcript_list.find_generated_transcript(['ko', 'en'])
            except:
                # 아무 자막이나 가져오기
                for t in transcript_list:
                    transcript = t
                    break

        if transcript:
            fetched = transcript.fetch()
            full_text = ' '.join([item['text'] for item in fetched])
            return full_text, None
        else:
            return None, "자막을 찾을 수 없습니다"

    except Exception as e:
        return None, f"자막 추출 오류: {str(e)[:100]}"

def analyze_youtube_video_direct(url):
    """YouTube 영상 자막 기반 분석 (빠르고 정확)"""
    api_key = get_api_key()
    if not api_key:
        st.error("API 키를 입력해주세요")
        return None

    # 1. Video ID 추출
    video_id = extract_video_id(url)
    if not video_id:
        st.error("올바른 YouTube URL이 아닙니다")
        return None

    # 2. 자막 가져오기
    transcript, error = get_youtube_transcript(video_id)
    if error:
        st.warning(f"자막 추출 실패: {error}")
        st.info("자막이 없는 영상입니다. 텍스트 입력으로 직접 내용을 입력해주세요.")
        return None

    if not transcript or len(transcript) < 50:
        st.warning("자막 내용이 너무 짧습니다")
        return None

    # 3. 자막 기반 분석
    prompt = f"""다음은 YouTube 영상의 자막입니다. 이 내용을 분석해주세요.

[자막 내용]
{transcript[:15000]}

[분석 요청]
위 자막 내용을 바탕으로 분석해주세요. 자막에 없는 내용은 추측하지 마세요.

JSON 형식으로 응답:
{{
    "title": "영상의 핵심 주제 (자막 기반 추론)",
    "creator": "알 수 없음",
    "main_topic": "메인 주제 한 줄 요약",
    "key_points": ["핵심 포인트 1", "핵심 포인트 2", "핵심 포인트 3", "핵심 포인트 4", "핵심 포인트 5"],
    "detailed_notes": ["상세 내용 1", "상세 내용 2", "상세 내용 3"],
    "actionable_tips": ["실천 팁 1", "실천 팁 2", "실천 팁 3"],
    "quotes": ["인상적인 문장 1", "인상적인 문장 2"],
    "vocabulary": [{{"term": "용어", "definition": "설명"}}],
    "study_questions": ["학습 질문 1", "학습 질문 2"],
    "summary": "전체 내용 5-7문장 요약"
}}"""

    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash')
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"분석 오류: {str(e)[:150]}")
        return None

def get_full_content():
    full = ""
    for ch in st.session_state.get('outline', []):
        if ch in st.session_state.get('chapters', {}):
            ch_data = st.session_state['chapters'][ch]
            ch_content = ""
            for s in ch_data.get('subtopics', []):
                c = ch_data.get('subtopic_data', {}).get(s, {}).get('content', '')
                if c:
                    ch_content += f"\n\n【{s}】\n\n{clean_content(c)}"
            if ch_content:
                full += f"\n\n{'='*50}\n{ch}\n{'='*50}{ch_content}"
    return full.strip()

def create_ebook_docx(title, subtitle, author, chapters_data, outline, interview_data=None):
    """베스트셀러 스타일의 전문적인 워드 문서 생성"""
    if not DOCX_AVAILABLE:
        return None, "python-docx 패키지가 필요합니다: pip install python-docx"

    try:
        doc = Document()

        # 페이지 설정 (A5 크기 - 전자책에 적합)
        section = doc.sections[0]
        section.page_width = Cm(14.8)
        section.page_height = Cm(21)
        # 좌우 여백 2.4cm → 글줄 길이 10cm (단행본 최적 글줄 길이 8~10cm)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        def set_font(run, size, bold=False, color=None, italic=False):
            run.font.size = Pt(size)
            run.font.name = 'Malgun Gothic'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = RGBColor(*color)

        def serif_run(run):
            """명조(바탕) 폰트 적용 — 단행본 타이포"""
            run.font.name = 'Batang'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '바탕')

        def add_hrule(indent_cm=4.2, color='A4803D', size=6, space=8):
            """골드 헤어라인 괘선 (문단 하단 보더)"""
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(indent_cm)
            pf.right_indent = Cm(indent_cm)
            pf.space_before = Pt(space)
            pf.space_after = Pt(space)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), str(size))
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), color)
            pBdr.append(bottom)
            pPr.append(pBdr)
            return p

        def add_bookmark(paragraph, bookmark_name):
            """문단에 북마크 추가"""
            # 북마크 이름에서 특수문자 제거 (Word 북마크 규칙)
            clean_name = re.sub(r'[^\w가-힣]', '_', bookmark_name)[:40]

            bookmark_start = OxmlElement('w:bookmarkStart')
            bookmark_start.set(qn('w:id'), str(hash(clean_name) % 10000))
            bookmark_start.set(qn('w:name'), clean_name)

            bookmark_end = OxmlElement('w:bookmarkEnd')
            bookmark_end.set(qn('w:id'), str(hash(clean_name) % 10000))

            paragraph._p.insert(0, bookmark_start)
            paragraph._p.append(bookmark_end)
            return clean_name

        def add_hyperlink(paragraph, text, bookmark_name, font_size=10, bold=False, color=(70, 70, 70), serif=False):
            """북마크로 연결되는 하이퍼링크 추가"""
            # 북마크 이름 정리
            clean_name = re.sub(r'[^\w가-힣]', '_', bookmark_name)[:40]

            # 하이퍼링크 요소 생성
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('w:anchor'), clean_name)

            # 텍스트 실행 요소
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')

            # 폰트 설정
            rFonts = OxmlElement('w:rFonts')
            _fam, _fam_e = ('Batang', '바탕') if serif else ('Malgun Gothic', '맑은 고딕')
            rFonts.set(qn('w:ascii'), _fam)
            rFonts.set(qn('w:eastAsia'), _fam_e)
            rFonts.set(qn('w:hAnsi'), _fam)
            rPr.append(rFonts)

            # 폰트 크기 (half-point 단위, 정수 필수)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_size * 2)))
            rPr.append(szCs)

            # 볼드
            if bold:
                b = OxmlElement('w:b')
                rPr.append(b)

            # 색상
            if color:
                c = OxmlElement('w:color')
                c.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*color))
                rPr.append(c)

            new_run.append(rPr)

            # 텍스트
            text_elem = OxmlElement('w:t')
            text_elem.text = text
            new_run.append(text_elem)

            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)

            return hyperlink

        def set_char_spacing(run, pt_value):
            """자간 설정 (1/20pt 단위) — 단행본 타이포 디테일"""
            rPr = run._element.get_or_add_rPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:val'), str(int(pt_value * 20)))
            rPr.append(spacing)

        BODY_INK = (45, 42, 38)
        BOLD_INK = (24, 21, 17)
        GOLD = (164, 128, 61)
        INK = (28, 26, 22)

        def add_styled_body_paragraph(text, is_first=False):
            """본문·프롤로그·에필로그 공용 문단 — 실제 단행본 조판
            (명조 10.5pt · 행간 185% · 자간 -0.2pt · 들여쓰기 1글자 · **볼드** 렌더링)"""
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            segments = re.split(r'(\*\*[^*\n]+\*\*)', text)
            for seg in segments:
                if not seg:
                    continue
                is_bold = seg.startswith('**') and seg.endswith('**') and len(seg) > 4
                content = seg[2:-2] if is_bold else seg
                if content:
                    run = para.add_run(content)
                    set_font(run, 10.5, bold=is_bold, color=BOLD_INK if is_bold else BODY_INK)
                    serif_run(run)
                    set_char_spacing(run, -0.2)  # 단행본 자간 (살짝 좁게)
            pf = para.paragraph_format
            pf.line_spacing = 1.85
            pf.space_after = Pt(18)  # 문단 사이 한 줄 띄기 (가독성)
            if not is_first:
                pf.first_line_indent = Cm(0.38)  # 1글자 들여쓰기 (단행본 관례)
            return para

        def add_pull_quote(text):
            """★ 핵심 문장 — 상하 골드 괘선의 인용 블록 (단행본 풀쿼트)"""
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            set_font(run, 12, bold=True, color=(60, 48, 26))
            serif_run(run)
            pf = para.paragraph_format
            pf.line_spacing = 1.6
            pf.space_before = Pt(20)
            pf.space_after = Pt(20)
            pf.left_indent = Cm(1.1)
            pf.right_indent = Cm(1.1)
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            for side in ('top', 'bottom'):
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '6')
                el.set(qn('w:space'), '10')
                el.set(qn('w:color'), 'A4803D')
                pBdr.append(el)
            pPr.append(pBdr)
            return para

        def gothic_run(run):
            """제목용 고딕 — 명조 본문과 대비되는 헤드라인 폰트 (단행본 표제 체계)"""
            run.font.name = 'Malgun Gothic'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        def add_page_number_footer(section):
            """하단 가운데 페이지 번호 (실제 단행본 하시라)"""
            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fp.add_run()
            run.font.size = Pt(8.5)
            run.font.name = 'Batang'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '바탕')
            run.font.color.rgb = RGBColor(120, 114, 102)
            f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
            it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
            f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
            run._r.append(f1); run._r.append(it); run._r.append(f2)

        # ══════════════════════════════════════════════════════════════
        # 1페이지: 그래픽 표지 (선택된 표지 디자인을 PNG로 변환해 삽입)
        # ══════════════════════════════════════════════════════════════
        cover_added = False
        try:
            import cairosvg
            _tmpl = None
            try:
                _label = st.session_state.get('cover_template_choice')
                _l2i = {v: k for k, v in COVER_TEMPLATES.items()}
                _tmpl = _l2i.get(_label)
            except Exception:
                _tmpl = None
            if not _tmpl:
                _topic_for_cover = st.session_state.get('topic', '') or (interview_data or {}).get('topic', '')
                _tmpl = pick_cover_template(_topic_for_cover, title)
            _svg = prepare_svg_for_png(build_cover_svg(_tmpl, title, subtitle, author))
            _png_bytes = cairosvg.svg2png(bytestring=_svg.encode('utf-8'), output_width=1240, output_height=1984)

            # 표지 전용 섹션: 여백을 최소화해 풀페이지 표지가 정확히 1페이지에 들어가게
            from docx.enum.section import WD_SECTION
            _sec0 = doc.sections[0]
            _sec0.left_margin = Cm(0.4)
            _sec0.right_margin = Cm(0.4)
            _sec0.top_margin = Cm(0.4)
            _sec0.bottom_margin = Cm(0.4)
            _avail_w = _sec0.page_width - _sec0.left_margin - _sec0.right_margin
            _avail_h = _sec0.page_height - _sec0.top_margin - _sec0.bottom_margin
            # 비율(1600:2560) 유지하며 페이지 높이에 거의 꽉 차게 — 절대 넘치지 않도록 0.98 마진
            _pic_h = int(int(_avail_h) * 0.98)
            _pic_w = int(_pic_h * 1600 / 2560)
            if _pic_w > int(_avail_w):
                _pic_w = int(_avail_w)
                _pic_h = int(_pic_w * 2560 / 1600)
            _cover_p = doc.add_paragraph()
            _cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _cover_p.paragraph_format.space_before = Pt(0)
            _cover_p.paragraph_format.space_after = Pt(0)
            _cover_p.paragraph_format.line_spacing = 1.0
            _cover_p.add_run().add_picture(io.BytesIO(_png_bytes), width=_pic_w, height=_pic_h)

            # 본문은 새 섹션에서 시작 (페이지 나눔 없이 섹션 경계로 분리 → 표지는 딱 1페이지)
            _body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
            _body_sec.page_width = Cm(14.8)
            _body_sec.page_height = Cm(21)
            _body_sec.left_margin = Cm(2.4)
            _body_sec.right_margin = Cm(2.4)
            _body_sec.top_margin = Cm(2.5)
            _body_sec.bottom_margin = Cm(2.5)
            add_page_number_footer(_body_sec)  # 본문부터 페이지 번호 (표지는 없음)
            cover_added = True
        except Exception:
            # cairosvg 미설치/변환 실패 시 텍스트 속표지 사용 — 여백 원복
            _sec0 = doc.sections[0]
            _sec0.left_margin = Cm(2.4)
            _sec0.right_margin = Cm(2.4)
            _sec0.top_margin = Cm(2.5)
            _sec0.bottom_margin = Cm(2.5)
            add_page_number_footer(_sec0)

        # 속표지(텍스트)는 이미지 표지 실패 시에만 — 표지는 한 번만 나온다
        if not cover_added:
            # 속표지 (단행본 에디토리얼 스타일: 골드 괘선 + 명조 + 지음)
            # ══════════════════════════════════════════════════════════════
            GOLD = (164, 128, 61)
            INK = (28, 26, 22)

            for _ in range(5):
                doc.add_paragraph()

            # 상단 골드 헤어라인
            add_hrule(indent_cm=4.2, size=8)

            # 작은 다이아 오너먼트
            orn = doc.add_paragraph()
            orn.alignment = WD_ALIGN_PARAGRAPH.CENTER
            orn_run = orn.add_run("◆")
            set_font(orn_run, 8, color=GOLD)
            orn.paragraph_format.space_before = Pt(14)
            orn.paragraph_format.space_after = Pt(20)

            # 메인 타이틀 — 명조 대형
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title)
            set_font(title_run, 30, bold=True, color=INK)
            serif_run(title_run)
            title_para.paragraph_format.space_after = Pt(18)
            title_para.paragraph_format.line_spacing = 1.25

            # 부제 — 명조 회색
            if subtitle:
                subtitle_para = doc.add_paragraph()
                subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub_run = subtitle_para.add_run(subtitle)
                set_font(sub_run, 12, color=(95, 90, 80))
                serif_run(sub_run)
                subtitle_para.paragraph_format.space_before = Pt(4)
                subtitle_para.paragraph_format.space_after = Pt(20)

            # 하단 골드 헤어라인
            add_hrule(indent_cm=4.2, size=8)

            # 중간 여백
            for _ in range(7):
                doc.add_paragraph()

            # 저자 — "○○○ 지음"
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"{author if author else '저자'}  지음")
            set_font(author_run, 13, color=(60, 56, 48))
            serif_run(author_run)
            author_para.paragraph_format.space_after = Pt(30)

            # 하단 마크
            mark = doc.add_paragraph()
            mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mark_run = mark.add_run("· · ·")
            set_font(mark_run, 9, color=GOLD)

            doc.add_page_break()

        # ══════════════════════════════════════════════════════════════
        # 프롤로그 (단행본 에디토리얼 스타일 — 에필로그와 동일 양식)
        # ══════════════════════════════════════════════════════════════
        for _ in range(4):
            doc.add_paragraph()

        # 프롤로그 라벨 (영문 소제목)
        pl_label = doc.add_paragraph()
        pl_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pl_label_run = pl_label.add_run("P R O L O G U E")
        set_font(pl_label_run, 9, color=GOLD)
        pl_label.paragraph_format.space_after = Pt(14)

        # 프롤로그 제목 — 명조
        prologue_title = doc.add_paragraph()
        prologue_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt_run = prologue_title.add_run("들어가며")
        set_font(pt_run, 16, bold=True, color=INK)
        gothic_run(pt_run)
        set_char_spacing(pt_run, 1.5)
        prologue_title.paragraph_format.space_after = Pt(8)

        # 제목 아래 골드 헤어라인
        add_hrule(indent_cm=5.2, size=6)
        _pl_sp = doc.add_paragraph()
        _pl_sp.paragraph_format.space_after = Pt(20)

        # 프롤로그 내용 - AI가 인터뷰 내용을 참고해서 자연스럽게 작성
        # [캐시] 페이지가 갱신될 때마다 재생성하지 않도록 세션에 1회만 생성
        _pe_cache_key = f"{title}|{interview_data.get('topic', '') if interview_data else ''}"
        if st.session_state.get('_prologue_cache_key') == _pe_cache_key and st.session_state.get('_prologue_cache'):
            prologue_text = st.session_state['_prologue_cache']
            interview_data_for_prologue = None  # 생성 건너뛰기
        else:
            prologue_text = None
            interview_data_for_prologue = interview_data
        if prologue_text is None and interview_data_for_prologue:
            prologue_prompt = f"""당신은 한국 자기계발 베스트셀러 작가입니다. 독자가 첫 문장에서 "이거 내 얘기야"라고 무릎 치고, 마지막 문장에서 "다음 페이지가 너무 궁금해"라며 책장을 넘기게 만드는 프롤로그를 작성하세요.

[저자 정보 - 참고용, 그대로 복사하지 말 것]
- 분야: {interview_data.get('field', '')}
- 경력: {interview_data.get('experience_years', '')}
- 저자의 현재 직업/상황: {interview_data.get('author_job', '')}
- 책 주제: {interview_data.get('topic', '')}
- 타겟 독자: {interview_data.get('target_reader', '')}
- 독자의 고민: {interview_data.get('target_problem', '')}
- 집필 동기: {interview_data.get('why_write', '')}

🚨 저자 정체성 규칙 (최우선 — 어기면 전체 무효)
- 1인칭('저')의 경험담은 [저자의 현재 직업/상황]에 적힌 것만 사용하라.
- '타겟 독자'의 직업·상황을 저자의 이야기로 바꿔 쓰지 마라. 저자가 직장인이라는 정보가 없으면
  출근·퇴근·월급·회사·사무실을 1인칭 경험으로 쓰는 것 절대 금지.
- 저자 직업 정보가 비어 있으면 직업을 특정하지 말고, 누구에게나 있는 보편적 장면(잔고를 확인하던 밤,
  남들과 비교하던 순간)으로만 1인칭을 써라. 없는 사실을 지어내지 마라.
- 독자 공감 장면(1막)은 "이런 분들이 많습니다" 식으로 독자의 이야기로 쓰고, 저자 고백(2막)과 섞지 마라.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 프롤로그의 단 하나의 미션
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

독자가 첫 문장부터 마지막 문장까지 한숨에 읽고, 자기 이야기처럼 공감하면서, 본문이 미치도록 궁금해서 1장으로 넘어갈 수밖에 없게 만들 것.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ 톤: 공감 후킹 + 스토리텔링 + 호기심 갭
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[5막 구조 - 베스트셀러 프롤로그 공식]

1막) 누구나 겪는 장면 (3~4문장) - 공감 끌어올리기
   - 독자가 "이거 내 얘기야"라고 즉시 떠올리는 구체적 장면
   - 시간·장소·감정의 디테일 (예: "수요일 밤 11시. 또 인스타를 켰다. 친구 OO이 부자가 돼 있었다.")
   - 일반론 절대 금지. 손에 잡히는 장면 하나

2막) 저자의 고백 (3~4문장) - "저도 그랬습니다"
   - 구체적 실패담 + 그때의 감정 (수치심, 막막함, 분노 등)
   - 진심 어린 톤 (위에서 내려다보지 말 것)
   - 약점 노출이 신뢰를 만든다

3막) 결정적 전환 (2~3문장) - 그러던 어느 날
   - 발견 / 사건 / 깨달음의 순간
   - "그날 발견한 한 가지가 모든 걸 바꿨다"
   - 그게 뭔지 다 말하지 말고 살짝 가리기
   - ⚠️ 이 책의 방법은 저자가 직접 만든 것 — 남에게 들었거나 배운 것처럼 쓰지 마라 (스스로 발견·정리한 시점으로)

4막) 약속 + 변화 (2문장) - 이 책이 줄 것
   - 본문이 다룰 핵심 변화를 구체적으로 (숫자/기간 포함)
   - 추상 X, 구체 O (예: "이 책은 그 90초가 어떻게 작동하는지 단계별로 풀어냅니다.")

5막) 본문 미끼 (1~2문장) - 페이지 넘기게 하기
   - 본문 1장이 다룰 가장 강한 장면 또는 통찰의 일부만 흘리기
   - "그런데 그 출발점은 의외의 한 가지였습니다." 같은 호기심 갭
   - "다음 페이지부터 시작됩니다" 같은 직접 안내는 절대 X
   - 🚨 미끼도 한 번에 이해되는 문장이어야 한다. '1장'과 행동을 한 문장에 뒤섞지 마라:
     ❌ "1장에서 가장 먼저 여는 건, 엑셀 파일 하나였거든요" — '장을 연다'와 '파일을 연다'가 섞인 비문
     ✅ "그 시작은 차트가 아니라 엑셀 파일 하나였습니다. 그 파일에 무엇이 적혀 있었는지는 곧 보시게 됩니다"

[문체 - 가장 중요한 규칙]
⚠️ 모든 문장은 반드시 존댓말로 끝나야 합니다. 단 한 문장도 예외 없이.
- 허용 종결: ~습니다, ~입니다, ~까요?, ~거든요, ~더라고요, ~잖아요, ~죠, ~네요
- 금지 종결(반말): ~했다, ~이다, ~있다, ~것이다, ~한다, ~해라, ~다.
- 장면 묘사도 존댓말로: "또 인스타를 켰다" ❌ → "또 인스타를 켜고 있었습니다" ⭕
- 작성 후 모든 문장의 끝을 검토해서 반말이 하나라도 있으면 존댓말로 고칠 것
- 짧은 문장 위주, 가끔 긴 문장으로 호흡 변화
- 디테일이 살아 있는 묘사 (시간, 숫자, 표정, 사물, 장소)
- 한 문단은 2~3문장만 (절대 4문장 초과 금지). 문단 사이는 반드시 빈 줄. 긴 이야기는 잘게 쪼개라
- 쉼표 최소화: 한 문장에 쉼표 최대 1개, 가능하면 0개. 쉼표로 절을 이어붙이지 말고 문장을 끊어라
  ("하루 47분만 투자하고, 월 1,127만원의 수익을…" ❌ → "하루 47분을 씁니다. 월 1,127만원이 들어옵니다." ⭕)
- 목적어-동사 결합이 자연스러운지 모든 문장 검토: "루틴을 담았습니다" ❌ → "방법을 담았습니다" ⭕,
  "시선을 만났습니다" ❌ → "생각이 바뀌었습니다" ⭕. 어색한 명사+동사 조합은 즉시 고칠 것

[분량] 600~800자

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
❌ 마크다운 문법 일체 (# ## ### **굵게** > 인용 - 글머리) — 제목/라벨 출력 금지, 본문 텍스트만
❌ "프롤로그", "Prologue", "Prologue.", "들어가며" 같은 제목/라벨 출력 금지 (이미 본문 위에 'Prologue' 표시됨)
❌ 위 저자 정보를 그대로 복사 붙여넣기
❌ 자청, 역행자, 자의식 해체, 유전자 역행, 원시인, 추월차선 (특정 작가 고유 표현)
❌ 교과서 표현: "여정", "발걸음", "함께 하시길 바랍니다", "진정한", "이 책의 여정"
❌ AI 어휘: "~의 중요성", "다양한", "효과적인", "~를 통해", "이 책을 통해"
❌ 과장: "놀라운", "혁신적인", "충격적인", "어마어마한", "기적의"
❌ 뻔한 말: "포기하지 마세요", "꾸준히 하세요", "당신도 할 수 있습니다", "함께 가요"
❌ 직접 호칭: 여러분, 당신, 독자님
❌ 위에서 내려다보는 어조 (당신은 이래서 안 됩니다 X)

본문 텍스트만 출력하세요. 어떤 마크다운 헤더(#)나 '프롤로그' 라벨도 출력하지 마세요. 첫 문장부터 바로 본문이 시작되어야 합니다."""

            generated_prologue = ask_ai(prologue_prompt, 0.7, ensure_quality=True)
            if generated_prologue:
                prologue_text = generated_prologue
                st.session_state['_prologue_cache'] = prologue_text
                st.session_state['_prologue_cache_key'] = _pe_cache_key

        if not prologue_text:
            prologue_text = """밤 11시, 불 꺼진 방에서 휴대폰 화면만 바라보던 날이 있었습니다. 남들은 다 잘 풀리는 것 같은데 나만 제자리인 것 같은 기분. 아마 한 번쯤은 느껴보셨을 겁니다.

저도 오래 그 자리에 있었습니다. 책을 사 모으고, 강의를 결제하고, 새해마다 계획을 세웠지만 석 달을 넘기지 못했거든요. 문제는 의지가 아니었습니다. 방법을 몰랐던 것뿐이었습니다.

그러던 어느 날, 우연히 알게 된 한 가지 원칙이 모든 것을 바꿔놓았습니다. 거창한 것이 아니었습니다. 오히려 너무 단순해서 그동안 지나쳤던 것이었죠.

이 책에는 그 원칙을 현장에서 부딪히며 다듬어온 과정과, 누구나 따라 할 수 있게 정리한 단계들이 담겨 있습니다. 이론이 아니라 직접 해보고 실패하고 다시 고친 기록입니다.

그 출발점이 무엇이었는지는, 바로 다음 장에서 시작됩니다."""

        # 마크다운 헤더(#, ##, ###) + 굵게(**) + '프롤로그/Prologue' 라벨 모두 제거
        # AI가 어떤 형태로 라벨을 박든 다 잡아냄
        prologue_text = re.sub(r'^\s*#+\s*(프롤로그|Prologue|들어가며|머리말)\s*\.?\s*$', '', prologue_text, flags=re.MULTILINE | re.IGNORECASE)
        prologue_text = re.sub(r'^\s*\*+\s*(프롤로그|Prologue|들어가며|머리말)\s*\*+\s*$', '', prologue_text, flags=re.MULTILINE | re.IGNORECASE)
        prologue_text = re.sub(r'^\s*(프롤로그|Prologue|들어가며|머리말)\s*\.?\s*\n', '', prologue_text, flags=re.IGNORECASE)
        prologue_text = prologue_text.replace('**프롤로그**', '').replace('**Prologue**', '')
        # 시작 부분의 빈 줄 제거 + 긴 문단 자동 분할 (가독성 — 본문과 동일)
        prologue_text = prologue_text.lstrip('\n').strip()
        prologue_text = _split_long_paragraphs(prologue_text)

        # 프롤로그 본문 - 본문과 완전히 동일한 양식 (명조, 10.5pt, 행간 1.9, 드롭캡, 볼드 렌더링)
        _pl_first = True
        for para_text in prologue_text.split('\n\n'):
            if para_text.strip():
                add_styled_body_paragraph(para_text.strip(), is_first=_pl_first)
                _pl_first = False

        # (페이지 시작은 각 섹션의 첫 내용 단락에 page_break_before 속성으로 처리 — 빈 단락 없음)

        # ══════════════════════════════════════════════════════════════
        # 목차 (프리미엄 미니멀 디자인)
        # ══════════════════════════════════════════════════════════════

        # 목차 제목 — 종이책 '차례' (단락 속성으로 새 페이지 + 상단 여백, 빈 단락 없음)
        toc_eyebrow = doc.add_paragraph()
        toc_eyebrow.paragraph_format.page_break_before = True
        toc_eyebrow.paragraph_format.space_before = Pt(86)
        toc_eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_eb_run = toc_eyebrow.add_run("CONTENTS")
        set_font(toc_eb_run, 9, color=GOLD)
        set_char_spacing(toc_eb_run, 4)
        toc_eyebrow.paragraph_format.space_after = Pt(8)

        toc_title = doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_run = toc_title.add_run("차례")
        set_font(toc_run, 20, bold=True, color=INK)
        gothic_run(toc_run)
        set_char_spacing(toc_run, 6)
        toc_title.paragraph_format.space_after = Pt(10)

        add_hrule(indent_cm=4.6, size=8)
        _sp = doc.add_paragraph()
        _sp.paragraph_format.space_after = Pt(16)

        for idx, chapter in enumerate(outline):
            # 챕터 제목 정리 (PART X. 등 접두사 제거)
            clean_chapter = chapter
            for prefix in [f"PART {idx + 1}.", f"PART{idx + 1}.", f"PART {idx + 1} ", f"PART{idx + 1} ", f"{idx + 1}.", f"{idx + 1})"]:
                clean_chapter = clean_chapter.replace(prefix, "").strip()

            # ─────────────────────────────────────
            # PART 라벨 — 골드 넘버 + 가는 헤어라인 (단행본 차례 양식)
            # ─────────────────────────────────────
            ch_num_para = doc.add_paragraph()
            ch_num_para.paragraph_format.space_before = Pt(26)
            ch_num_para.paragraph_format.space_after = Pt(0)
            ch_num_run = ch_num_para.add_run(f"PART {idx + 1:02d}")
            set_font(ch_num_run, 9.5, bold=True, color=GOLD)
            set_char_spacing(ch_num_run, 2.5)
            # 라벨 아래 가는 골드 헤어라인 (오른쪽 절반)
            _tPr = ch_num_para._p.get_or_add_pPr()
            _tBdr = OxmlElement('w:pBdr')
            _tb = OxmlElement('w:bottom')
            _tb.set(qn('w:val'), 'single')
            _tb.set(qn('w:sz'), '4')
            _tb.set(qn('w:space'), '4')
            _tb.set(qn('w:color'), 'D8C9A6')
            _tBdr.append(_tb)
            _tPr.append(_tBdr)

            # ─────────────────────────────────────
            # 챕터 제목 (하이퍼링크 · 명조 볼드)
            # ─────────────────────────────────────
            ch_title_para = doc.add_paragraph()
            ch_title_para.paragraph_format.space_before = Pt(8)
            ch_title_para.paragraph_format.space_after = Pt(10)
            ch_title_para.paragraph_format.line_spacing = 1.3
            chapter_bookmark_name = f"chapter_{idx + 1}"
            add_hyperlink(ch_title_para, clean_chapter, chapter_bookmark_name, font_size=13, bold=True, color=(28, 26, 22), serif=False)

            # ─────────────────────────────────────
            # 소제목들 — 골드 두 자리 번호 + 명조 (점 불릿 대신)
            # ─────────────────────────────────────
            if chapter in chapters_data:
                ch_data = chapters_data[chapter]
                subtopics = ch_data.get('subtopics', [])

                for sub_idx, sub in enumerate(subtopics):
                    toc_sub = doc.add_paragraph()
                    toc_sub.paragraph_format.left_indent = Cm(0.45)
                    toc_sub.paragraph_format.space_after = Pt(7)
                    toc_sub.paragraph_format.line_spacing = 1.25

                    # 골드 두 자리 번호
                    num_run = toc_sub.add_run(f"{sub_idx + 1:02d}")
                    set_font(num_run, 9, bold=True, color=(186, 156, 98))
                    sp_run = toc_sub.add_run("   ")
                    set_font(sp_run, 9)

                    # 소제목 텍스트 (하이퍼링크로 연결)
                    subtopic_bookmark_name = f"subtopic_{idx + 1}_{sub_idx + 1}"
                    add_hyperlink(toc_sub, sub, subtopic_bookmark_name, font_size=10.5, bold=False, color=(72, 67, 58), serif=True)

        # ══════════════════════════════════════════════════════════════
        # 본문 (프리미엄 에디토리얼 스타일)
        # ══════════════════════════════════════════════════════════════

        def add_horizontal_line(doc, width_cm=3, color=(220, 220, 220)):
            """가로 구분선 추가"""
            line_para = doc.add_paragraph()
            line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            line_run = line_para.add_run("─" * 12)
            set_font(line_run, 10, color=color)
            line_para.paragraph_format.space_before = Pt(20)
            line_para.paragraph_format.space_after = Pt(20)
            return line_para

        def add_chapter_opener(doc, idx, chapter_title):
            """챕터 시작 페이지 - 프리미엄 에디토리얼 스타일 (새 페이지에서, 빈 단락 없이)
            PART/번호/괘선/제목이 절대 쪼개지지 않도록 keep_with_next로 묶음"""
            # PART 라벨 (작은 대문자 · 골드) — 단락 속성으로 새 페이지 + 상단 여백
            part_label = doc.add_paragraph()
            part_label.paragraph_format.page_break_before = True
            part_label.paragraph_format.space_before = Pt(190)
            part_label.paragraph_format.keep_with_next = True
            part_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            part_run = part_label.add_run("PART")
            set_font(part_run, 9, color=GOLD)
            set_char_spacing(part_run, 5)
            part_label.paragraph_format.space_after = Pt(8)

            # 챕터 번호 (매우 큰 숫자 · 명조)
            ch_num_para = doc.add_paragraph()
            ch_num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ch_num_para.paragraph_format.keep_with_next = True
            ch_num_run = ch_num_para.add_run(f"{idx + 1}")
            set_font(ch_num_run, 52, bold=False, color=(38, 34, 28))
            serif_run(ch_num_run)
            ch_num_para.paragraph_format.space_after = Pt(14)

            # 골드 헤어라인 구분선
            _co_rule = add_hrule(indent_cm=5.6, size=6, space=4)
            _co_rule.paragraph_format.keep_with_next = True
            _co_sp = doc.add_paragraph()
            _co_sp.paragraph_format.space_after = Pt(10)
            _co_sp.paragraph_format.keep_with_next = True

            # 챕터 제목 — 고딕 볼드 (명조 본문과 대비)
            ch_name = doc.add_paragraph()
            ch_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            chn_run = ch_name.add_run(chapter_title)
            set_font(chn_run, 16, bold=True, color=INK)
            gothic_run(chn_run)
            ch_name.paragraph_format.line_spacing = 1.35
            ch_name.paragraph_format.space_after = Pt(60)

            return ch_name

        def add_subtopic_header(doc, subtopic_text, sub_idx):
            """소제목 - 단행본 스타일 (골드 번호 + 고딕 + 헤어라인) — 항상 새 페이지, 빈 단락 없음"""
            # 골드 번호 라벨 (이 단락 자체에 '앞에서 페이지 나눔' — 빈 단락·검은 점 없음)
            num_para = doc.add_paragraph()
            num_para.paragraph_format.page_break_before = True
            num_para.paragraph_format.space_before = Pt(46)
            num_para.paragraph_format.keep_with_next = True
            num_run = num_para.add_run(f"{sub_idx + 1:02d}")
            set_font(num_run, 11, bold=True, color=(164, 128, 61))
            num_para.paragraph_format.space_after = Pt(2)

            # 소제목 — 고딕 볼드 (단행본 표제 체계)
            sub_title = doc.add_paragraph()
            sub_run = sub_title.add_run(subtopic_text)
            set_font(sub_run, 14, bold=True, color=(25, 23, 20))
            gothic_run(sub_run)
            set_char_spacing(sub_run, -0.3)
            sub_title.paragraph_format.space_after = Pt(6)
            sub_title.paragraph_format.line_spacing = 1.3

            # 제목 아래 가는 골드 헤어라인
            _rl = doc.add_paragraph()
            _rlf = _rl.paragraph_format
            _rlf.right_indent = Cm(5.5)
            _rlf.space_before = Pt(0)
            _rlf.space_after = Pt(22)
            _pPr = _rl._p.get_or_add_pPr()
            _pBdr = OxmlElement('w:pBdr')
            _bt = OxmlElement('w:bottom')
            _bt.set(qn('w:val'), 'single')
            _bt.set(qn('w:sz'), '6')
            _bt.set(qn('w:space'), '1')
            _bt.set(qn('w:color'), 'A4803D')
            _pBdr.append(_bt)
            _pPr.append(_pBdr)

            return sub_title

        def add_part_summary_box(doc, part_no, ch_data):
            """파트 마지막 '핵심 정리' 박스 — 크림 배경 + 골드 보더"""
            lines = ch_data.get('part_summary') or []
            if isinstance(lines, str):
                lines = [l.strip() for l in lines.split('\n') if l.strip()]
            fallback = False
            if not lines:
                lines = list(ch_data.get('subtopics', []))[:4]
                fallback = True
            if not lines:
                return

            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(26)

            # 단행본 '장 끝 정리' — 상하단 가는 괘선만 두른 절제된 박스
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.cell(0, 0)
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            for side, sz in (('top', '12'), ('bottom', '12'), ('left', None), ('right', None)):
                el = OxmlElement(f'w:{side}')
                if sz:
                    el.set(qn('w:val'), 'single')
                    el.set(qn('w:sz'), sz)
                    el.set(qn('w:color'), '8A7A55')
                else:
                    el.set(qn('w:val'), 'nil')
                borders.append(el)
            tcPr.append(borders)

            # 라벨 (박스 안 첫 줄, 고딕)
            lp = cell.paragraphs[0]
            lp_run = lp.add_run(f"PART {part_no} 핵심 정리" if not fallback else f"PART {part_no} 에서 다룬 것")
            set_font(lp_run, 9.5, bold=True, color=(122, 96, 46))
            gothic_run(lp_run)
            set_char_spacing(lp_run, 1)
            lp.paragraph_format.space_before = Pt(10)
            lp.paragraph_format.space_after = Pt(8)

            for line in lines:
                p = cell.add_paragraph()
                is_action = line.startswith('오늘 할 일')
                b_run = p.add_run('—  ')
                set_font(b_run, 9, color=(150, 132, 96))
                t_run = p.add_run(line)
                set_font(t_run, 10, bold=is_action, color=(45, 42, 38))
                serif_run(t_run)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.5
            cell.paragraphs[-1].paragraph_format.space_after = Pt(10)

            after = doc.add_paragraph()
            after.paragraph_format.space_after = Pt(10)

        def set_serif(run):
            """본문용 명조 계열 폰트 (실제 단행본 본문 타이포)"""
            run.font.name = 'Batang'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '바탕')

        def format_body_paragraph(doc, text, is_first=False):
            """본문 문단 - 베스트셀러 단행본 스타일 (명조 + 여유 행간 + **볼드** 렌더링)
            ★로 시작하는 문단은 골드 괘선 인용 블록으로"""
            stripped = text.lstrip()
            if stripped.startswith('★'):
                quote_text = stripped.lstrip('★').strip().strip('"“”').strip()
                quote_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', quote_text)
                if quote_text:
                    return add_pull_quote(quote_text)
                return None
            return add_styled_body_paragraph(text, is_first=is_first)

        def is_table_text(text):
            """텍스트가 표 형식인지 감지"""
            lines = text.strip().split('\n')
            if len(lines) < 2:
                return False

            # 마크다운 테이블 감지: | 로 시작하고 | 로 끝남
            pipe_lines = sum(1 for line in lines if line.strip().startswith('|') and line.strip().endswith('|'))
            if pipe_lines >= 2:
                return True

            # 파이프로 구분된 테이블 (| 가 있지만 시작/끝이 아닐 수 있음)
            pipe_content_lines = sum(1 for line in lines if '|' in line and len(line.split('|')) >= 2)
            if pipe_content_lines >= 2:
                return True

            # 탭 구분 테이블
            tab_lines = sum(1 for line in lines if '\t' in line)
            if tab_lines >= 2:
                return True

            # 콜론 기반 비교 테이블 감지 (Before: xxx / After: xxx)
            colon_lines = sum(1 for line in lines
                            if ':' in line
                            and not line.strip().startswith('http')
                            and len(line.split(':')[0]) < 30)
            if colon_lines >= 2 and colon_lines >= len(lines) * 0.6:
                return True

            return False

        def parse_table_data(text):
            """표 텍스트를 파싱하여 2D 배열로 변환"""
            lines = text.strip().split('\n')
            table_data = []

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # 순수 구분선만 스킵 (실제 내용이 없는 라인)
                # 마크다운 구분선: |---|---| 또는 |:---:|:---:|
                if re.match(r'^\|[\s\-:]+\|$', line):
                    continue
                # 박스 그리기 문자만 있는 라인
                if re.match(r'^[─━┌┬┐├┼┤└┴┘│┃]+$', line):
                    continue
                # 하이픈만 있는 라인 (--- 또는 - - -)
                if re.match(r'^[\s\-]+$', line) and len(line.replace(' ', '').replace('-', '')) == 0:
                    continue

                # 마크다운 테이블 파싱 (| cell | cell |)
                if line.startswith('|') and line.endswith('|'):
                    cells = [cell.strip() for cell in line.split('|')]
                    cells = [c for c in cells if c]  # 빈 셀 제거
                    if cells:
                        table_data.append(cells)
                # 일반 파이프 구분 (cell | cell)
                elif '|' in line and not line.startswith('|'):
                    cells = [cell.strip() for cell in line.split('|')]
                    cells = [c for c in cells if c]
                    if cells:
                        table_data.append(cells)
                # 탭 구분 테이블
                elif '\t' in line:
                    cells = [cell.strip() for cell in line.split('\t')]
                    cells = [c for c in cells if c]
                    if len(cells) >= 2:
                        table_data.append(cells)
                # 콜론 기반 파싱 (Before: xxx) - 단, URL이 아닌 경우
                elif ':' in line and not line.startswith('http'):
                    # 첫 번째 콜론으로만 분리
                    parts = line.split(':', 1)
                    if len(parts) == 2 and len(parts[0]) < 30:  # 키가 너무 길면 제외
                        table_data.append([parts[0].strip(), parts[1].strip()])

            # 구분선 행 제거 — 모든 셀이 -, :, = 만으로 된 행 (마크다운 구분선이 셀로 새는 것 방지)
            table_data = [r for r in table_data
                          if not all(re.match(r'^[\s\-:=·–—]+$', str(c)) for c in r)]
            return table_data

        def add_premium_table(doc, table_data):
            """인포그래픽 스타일 테이블 - 시각적 이해도 향상"""
            if not table_data or len(table_data) < 1:
                return None

            rows = len(table_data)
            cols = max(len(row) for row in table_data)

            # 2열 비교 테이블인 경우 (Before/After, 항목/설명 등)
            is_comparison = cols == 2 and rows >= 2

            # 테이블 생성
            table = doc.add_table(rows=rows, cols=cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 테이블 전체 너비 설정
            table.autofit = False
            for col_idx in range(cols):
                for row in table.rows:
                    if col_idx < len(row.cells):
                        if is_comparison:
                            # 2열 비교: 첫 열 좁게, 둘째 열 넓게
                            width = Cm(3) if col_idx == 0 else Cm(7)
                        else:
                            width = Cm(10 / cols)
                        row.cells[col_idx].width = width

            # 각 셀 스타일링 — 실제 단행본 표 조판 (촘촘한 행, 고딕 글자, 절제된 색)
            for i, row_data in enumerate(table_data):
                row = table.rows[i]
                # 행 높이는 내용에 맞게 자동 (고정 높이 제거 — 책 표는 촘촘해야 한다)

                for j, cell_text in enumerate(row_data):
                    if j < cols:
                        cell = row.cells[j]
                        cell.text = ''

                        para = cell.paragraphs[0]
                        para.paragraph_format.space_before = Pt(2)
                        para.paragraph_format.space_after = Pt(2)
                        para.paragraph_format.line_spacing = 1.3

                        # 첫 번째 행(헤더) - 잉크 블랙 배경 + 크림 텍스트
                        if i == 0:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = para.add_run(str(cell_text))
                            set_font(run, 9.5, bold=True, color=(238, 228, 200))
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:fill'), '26231E')
                            cell._tc.get_or_add_tcPr().append(shading)

                        # 첫 번째 열 (라벨/항목) - 볼드 고딕, 배경 없이 깔끔하게
                        elif j == 0:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = para.add_run(str(cell_text))
                            set_font(run, 9.5, bold=True, color=(60, 52, 38))

                        # 일반 내용 셀
                        else:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = para.add_run(str(cell_text))
                            set_font(run, 9.5, color=(55, 50, 42))

                        # 셀 여백 (촘촘하게)
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcMar = OxmlElement('w:tcMar')
                        for margin_name, margin_val in [('top', '50'), ('left', '100'), ('bottom', '50'), ('right', '100')]:
                            margin = OxmlElement(f'w:{margin_name}')
                            margin.set(qn('w:w'), margin_val)
                            margin.set(qn('w:type'), 'dxa')
                            tcMar.append(margin)
                        tcPr.append(tcMar)

                        # 셀 수직 정렬 (가운데)
                        vAlign = OxmlElement('w:vAlign')
                        vAlign.set(qn('w:val'), 'center')
                        tcPr.append(vAlign)

            # 테이블 테두리 스타일 (깔끔한 라인)
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            tblBorders = OxmlElement('w:tblBorders')

            # 상하단 골드 굵은 선 + 내부는 옅은 크림 헤어라인 (좌우 개방형 — 고급 단행본 표)
            for border_name in ['top', 'bottom']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '12')
                border.set(qn('w:color'), 'A4803D')
                tblBorders.append(border)

            for border_name in ['left', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')  # 좌우 테두리 없음
                tblBorders.append(border)

            for border_name in ['insideH']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), 'E5DCC8')
                tblBorders.append(border)

            for border_name in ['insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), 'EFE8D8')
                tblBorders.append(border)

            tblPr.append(tblBorders)

            # 테이블 후 여백
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(15)

            return table

        def process_content_with_tables(doc, text):
            """본문 텍스트에서 표를 감지하고 처리"""
            blocks = []
            current_block = []
            table_lines = []

            lines = text.split('\n')
            i = 0

            def is_table_start(line, next_line=None):
                """표 시작 라인인지 확인"""
                stripped = line.strip()
                # 마크다운 테이블 (| cell | cell |)
                if stripped.startswith('|') and stripped.endswith('|') and stripped.count('|') >= 2:
                    return True
                # 파이프로 구분된 내용 (cell | cell)
                if '|' in stripped and len(stripped.split('|')) >= 2:
                    parts = stripped.split('|')
                    if any(p.strip() and not re.match(r'^[\s\-:]+$', p) for p in parts):
                        return True
                # 콜론 기반 테이블 (키: 값) - 연속된 경우
                if next_line and ':' in stripped and ':' in next_line:
                    if len(stripped.split(':')[0].strip()) < 30 and len(next_line.split(':')[0].strip()) < 30:
                        return True
                return False

            def is_table_continue(line):
                """표 계속 라인인지 확인"""
                stripped = line.strip()
                # 빈 줄은 표 종료
                if not stripped:
                    return False
                # 마크다운 테이블
                if stripped.startswith('|') and stripped.endswith('|'):
                    return True
                # 마크다운 구분선
                if re.match(r'^\|[\s\-:]+\|$', stripped):
                    return True
                # 파이프로 구분된 내용
                if '|' in stripped:
                    return True
                # 콜론 기반 (키: 값)
                if ':' in stripped and len(stripped.split(':')[0].strip()) < 30:
                    return True
                return False

            while i < len(lines):
                line = lines[i]
                next_line = lines[i + 1] if i + 1 < len(lines) else None

                if is_table_start(line, next_line):
                    # 이전 일반 텍스트 저장
                    if current_block:
                        blocks.append(('text', '\n'.join(current_block)))
                        current_block = []

                    # 표 라인 수집
                    table_lines = [line]
                    i += 1
                    while i < len(lines) and is_table_continue(lines[i]):
                        table_lines.append(lines[i])
                        i += 1

                    if len(table_lines) >= 2:
                        blocks.append(('table', '\n'.join(table_lines)))
                    else:
                        current_block.extend(table_lines)
                    table_lines = []
                else:
                    current_block.append(line)
                    i += 1

            # 마지막 블록 저장
            if current_block:
                blocks.append(('text', '\n'.join(current_block)))

            return blocks

        for idx, chapter in enumerate(outline):
            if chapter in chapters_data:
                # ─────────────────────────────────────
                # 챕터 시작 페이지 (프리미엄 오프너)
                # ─────────────────────────────────────
                clean_chapter = chapter
                for prefix in [f"PART {idx + 1}.", f"PART{idx + 1}.", f"PART {idx + 1} ", f"PART{idx + 1} ", f"{idx + 1}.", f"{idx + 1})"]:
                    clean_chapter = clean_chapter.replace(prefix, "").strip()

                ch_name = add_chapter_opener(doc, idx, clean_chapter)
                add_bookmark(ch_name, f"chapter_{idx + 1}")

                # ─────────────────────────────────────
                # 본문 시작
                # ─────────────────────────────────────
                ch_data = chapters_data[chapter]
                subtopics = ch_data.get('subtopics', [])

                for sub_idx, sub in enumerate(subtopics):
                    content = ch_data.get('subtopic_data', {}).get(sub, {}).get('content', '')
                    if content:
                        # 소제목마다 새 페이지에서 시작 (헤더 단락의 page_break_before로 처리 — 빈 페이지 없음)
                        sub_title = add_subtopic_header(doc, sub, sub_idx)
                        add_bookmark(sub_title, f"subtopic_{idx + 1}_{sub_idx + 1}")

                        # 본문 내용 (표 감지 및 처리 포함, 소제목 중복도 제거)
                        cleaned = clean_content(content, subtopic=sub)

                        # 표가 포함된 콘텐츠 처리
                        content_blocks = process_content_with_tables(doc, cleaned)

                        is_first_para = True
                        for block_type, block_content in content_blocks:
                            if block_type == 'table':
                                # 표 데이터 파싱 및 프리미엄 테이블 생성
                                table_data = parse_table_data(block_content)
                                if table_data and len(table_data) >= 2:
                                    # 표 전 여백
                                    spacer = doc.add_paragraph()
                                    spacer.paragraph_format.space_after = Pt(10)
                                    add_premium_table(doc, table_data)
                                    is_first_para = False
                            else:
                                # 일반 텍스트 처리
                                paragraphs = block_content.split('\n\n')
                                if not paragraphs or not any(p.strip() for p in paragraphs):
                                    paragraphs = block_content.split('\n')

                                for para_text in paragraphs:
                                    if para_text.strip():
                                        format_body_paragraph(doc, para_text.strip(), is_first=is_first_para)
                                        is_first_para = False

                        # 소제목은 각자 새 페이지에서 시작하므로 별도 구분 장식 없음 (절제)

                # 파트 마무리: 핵심 정리 박스
                add_part_summary_box(doc, idx + 1, ch_data)

        # ══════════════════════════════════════════════════════════════
        # 에필로그 (프리미엄 에디토리얼 스타일)
        # ══════════════════════════════════════════════════════════════

        # 에필로그 라벨 (프롤로그와 동일 양식) — 단락 속성으로 새 페이지 + 상단 여백
        ep_label = doc.add_paragraph()
        ep_label.paragraph_format.page_break_before = True
        ep_label.paragraph_format.space_before = Pt(190)
        ep_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ep_label_run = ep_label.add_run("E P I L O G U E")
        set_font(ep_label_run, 9, color=GOLD)
        ep_label.paragraph_format.space_after = Pt(14)

        # 에필로그 제목 — 명조
        epilogue_title = doc.add_paragraph()
        epilogue_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ep_run = epilogue_title.add_run("마치며")
        set_font(ep_run, 16, bold=True, color=INK)
        gothic_run(ep_run)
        set_char_spacing(ep_run, 1.5)
        epilogue_title.paragraph_format.space_after = Pt(8)

        # 제목 아래 골드 헤어라인
        add_hrule(indent_cm=5.2, size=6)
        _ep_sp = doc.add_paragraph()
        _ep_sp.paragraph_format.space_after = Pt(20)

        # 에필로그 내용 - AI가 인터뷰 내용을 참고해서 자연스럽게 작성
        # [캐시] 프롤로그와 동일하게 1회만 생성
        if st.session_state.get('_epilogue_cache_key') == _pe_cache_key and st.session_state.get('_epilogue_cache'):
            epilogue_text = st.session_state['_epilogue_cache']
            interview_data_for_epilogue = None
        else:
            epilogue_text = None
            interview_data_for_epilogue = interview_data
        if epilogue_text is None and interview_data_for_epilogue:
            epilogue_prompt = f"""당신은 한국 자기계발 베스트셀러 작가입니다. 마지막 페이지를 덮은 독자가 한 번 더 처음으로 돌아가게 만드는 에필로그를 작성하세요.

[저자 정보 - 참고용, 그대로 복사하지 말 것]
- 분야: {interview_data.get('field', '')}
- 경력 기간: {interview_data.get('experience_years', '')}
- 저자의 현재 직업/상황: {interview_data.get('author_job', '')}
- 책 주제: {interview_data.get('topic', '')}
- 타겟 독자: {interview_data.get('target_reader', '')}
- 독자에게 전하고 싶은 말: {interview_data.get('final_message', '')}
- 작가 경력/경험: {interview_data.get('author_career', '')}
- 어려움/실패 경험: {interview_data.get('struggle_story', '')}
- 극복 스토리: {interview_data.get('breakthrough', '')}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

## 에필로그 작성 원칙 (스토리텔링)

### 1. 나의 스토리로 시작 (3-4문장)
- 작가 경력/경험을 자연스럽게 녹여서
- "저는 ~했습니다" 형식으로 간결하게
- 구체적 숫자나 사실 포함
- 🚨 1인칭 경험은 [저자의 현재 직업/상황]에 적힌 것만. 타겟 독자의 직업(직장인 등)을 저자 이야기로 지어내지 마라

### 2. 왜 이 책을 썼는지 (2-3문장)
- 내가 겪은 어려움 + 극복 과정 힌트
- 독자를 위해 책을 쓴 진심

### 3. 독자에게 한마디 (2-3문장)
- 지금 당장 할 수 있는 구체적 행동 하나
- 진심 어린 마무리 (근데 뻔하지 않게)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[문체]
- 짧은 문장, 짧은 문단 (2-4문장)
- 구어체 + 합쇼체 ("~거든요", "~잖아요" OK)
- ⚠️ 모든 문장은 존댓말로 종결 (반말 종결 ~했다/~이다/~한다 절대 금지, 단 한 문장도 예외 없이)

[분량] 400-500자

[금지 - 절대 쓰지 말 것]
- 저자 정보를 그대로 복사 붙여넣기
- 자청, 역행자, 자의식 해체, 유전자 역행, 원시인, 추월차선 (특정 작가 고유 표현)
- 교과서 표현: "여정", "발걸음", "함께 하시길 바랍니다", "진정한"
- AI 표현: "~의 중요성", "다양한", "효과적인", "~를 통해"
- 과장: "놀라운", "혁신적인", "충격적인"
- 뻔한 말: "포기하지 마세요", "꾸준히 하세요", "화이팅"
- 직접 호칭: 여러분, 당신, 독자님
- 마크다운 문법

에필로그만 출력하세요."""

            generated_epilogue = ask_ai(epilogue_prompt, 0.7, ensure_quality=True)
            if generated_epilogue:
                epilogue_text = generated_epilogue
                st.session_state['_epilogue_cache'] = epilogue_text
                st.session_state['_epilogue_cache_key'] = _pe_cache_key

        if not epilogue_text:
            epilogue_text = """여기까지 읽어주셔서 감사합니다.

이 책에 담긴 내용이 당신의 삶에 작은 변화라도 만들어낸다면 그것으로 충분합니다.

완벽할 필요 없습니다. 지금 당장 할 수 있는 것 하나만 시작해보세요.

작은 시작이 큰 결과를 만듭니다.

항상 응원합니다."""

        # 마크다운 헤더(#, ##) + 굵게(**) + '에필로그/Epilogue' 라벨 모두 제거
        epilogue_text = re.sub(r'^\s*#+\s*(에필로그|Epilogue|마치며|맺음말)\s*\.?\s*$', '', epilogue_text, flags=re.MULTILINE | re.IGNORECASE)
        epilogue_text = re.sub(r'^\s*\*+\s*(에필로그|Epilogue|마치며|맺음말)\s*\*+\s*$', '', epilogue_text, flags=re.MULTILINE | re.IGNORECASE)
        epilogue_text = re.sub(r'^\s*(에필로그|Epilogue|마치며|맺음말)\s*\.?\s*\n', '', epilogue_text, flags=re.IGNORECASE)
        epilogue_text = epilogue_text.replace('**에필로그**', '').replace('**Epilogue**', '')
        epilogue_text = epilogue_text.lstrip('\n').strip()
        epilogue_text = _split_long_paragraphs(epilogue_text)

        # 에필로그 본문 - 본문과 완전히 동일한 양식 (명조 + 볼드 렌더링)
        _ep_first = True
        for para_text in epilogue_text.split('\n\n'):
            if para_text.strip():
                add_styled_body_paragraph(para_text.strip(), is_first=_ep_first)
                _ep_first = False

        # 저자 서명 (프리미엄 스타일)
        for _ in range(3):
            doc.add_paragraph()

        # 서명 라인
        sign_line = doc.add_paragraph()
        sign_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_line_run = sign_line.add_run("─────")
        set_font(sign_line_run, 10, color=(200, 200, 200))
        sign_line.paragraph_format.space_after = Pt(10)

        sign_para = doc.add_paragraph()
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_run = sign_para.add_run(f"{author if author else '저자'}")
        set_font(sign_run, 11, italic=True, color=(80, 80, 80))

        # ══════════════════════════════════════════════════════════════
        # 저자 소개 페이지 (프리미엄 에디토리얼 스타일)
        # ══════════════════════════════════════════════════════════════

        # 저자 소개 라벨 (프롤로그·에필로그와 동일 양식) — 단락 속성으로 새 페이지 + 상단 여백
        about_label = doc.add_paragraph()
        about_label.paragraph_format.page_break_before = True
        about_label.paragraph_format.space_before = Pt(190)
        about_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        about_label_run = about_label.add_run("A B O U T   T H E   A U T H O R")
        set_font(about_label_run, 9, color=GOLD)
        about_label.paragraph_format.space_after = Pt(14)

        # 저자명 (크게 · 명조)
        author_name_para = doc.add_paragraph()
        author_name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_name_run = author_name_para.add_run(author if author else "저자")
        set_font(author_name_run, 16, bold=True, color=INK)
        gothic_run(author_name_run)
        set_char_spacing(author_name_run, 1.5)
        author_name_para.paragraph_format.space_after = Pt(8)

        # 이름 아래 골드 헤어라인
        add_hrule(indent_cm=5.2, size=6)
        _ab_sp = doc.add_paragraph()
        _ab_sp.paragraph_format.space_after = Pt(20)

        # 저자 소개 내용
        if interview_data:
            career_text = interview_data.get('author_career', '')
            field = interview_data.get('field', '')
            exp = interview_data.get('experience_years', '')
            method = interview_data.get('core_method', '')

            if career_text:
                author_bio = f"""{field} 분야에서 {exp}간 활동해온 실전가.

{career_text}

{method[:100] if method else ''}"""
            else:
                author_bio = f"""{field} 분야에서 {exp}간 활동해온 실전가.

{method}"""
        else:
            author_bio = """실전에서 직접 부딪히며 쌓은 노하우를 독자들과 나누고자 이 책을 썼다."""

        for para_text in author_bio.split('\n\n'):
            if para_text.strip():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para_run = para.add_run(para_text.strip())
                set_font(para_run, 10, color=(72, 67, 58))
                serif_run(para_run)
                para_format = para.paragraph_format
                para_format.line_spacing = 1.7
                para_format.space_after = Pt(14)

        # 하단 장식
        for _ in range(2):
            doc.add_paragraph()

        end_mark = doc.add_paragraph()
        end_mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
        end_run = end_mark.add_run("◆")
        set_font(end_run, 11, color=GOLD)

        # ══════════════════════════════════════════════════════════════
        # 판권 (책 맨 뒤 — 종이책 관례)
        # ══════════════════════════════════════════════════════════════
        cp_title = doc.add_paragraph()
        cp_title.paragraph_format.page_break_before = True
        cp_title.paragraph_format.space_before = Pt(320)
        cp_title.paragraph_format.keep_with_next = True
        cp_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp_title_run = cp_title.add_run(title)
        set_font(cp_title_run, 11, bold=True, color=(70, 64, 54))
        serif_run(cp_title_run)
        cp_title.paragraph_format.space_after = Pt(14)

        for line in [f"지은이  {author if author else '저자'}",
                     "",
                     "이 책의 저작권은 저자에게 있습니다.",
                     "무단 전재와 복제를 금합니다."]:
            cp_para = doc.add_paragraph()
            cp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if line:
                cp_run = cp_para.add_run(line)
                set_font(cp_run, 8.5, color=(130, 124, 112))
                serif_run(cp_run)
            cp_para.paragraph_format.space_after = Pt(3)

        # 메모리에 저장
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue(), None

    except Exception as e:
        return None, f"문서 생성 오류: {str(e)}"

def polish_outline(chapters, subtopics, topic):
    """목차 비문 검수 패스 — 멀쩡한 줄은 그대로, '문장이 깨진 줄만' 최소 수정 (저온 교열)"""
    lines = []
    for ci, ch in enumerate(chapters):
        lines.append(f"C{ci+1}: {ch}")
        for si, s in enumerate(subtopics.get(ch, [])):
            lines.append(f"S{ci+1}-{si+1}: {s}")
    numbered = "\n".join(lines)

    prompt = f"""당신은 한국 대형 출판사의 교열 전문가다. 아래 목차에서 '문장이 깨진 줄'만 찾아 최소한으로 고친다.

[절대 원칙]
1. 멀쩡한 줄은 글자 하나도 바꾸지 마라. 입력의 80% 이상이 그대로 나가는 게 정상이다.
2. 고칠 때도 후킹(호기심 갭, 숫자, 도발, 떡밥)은 그대로 유지하고 문장 구조만 바로잡아라.
3. 새 내용·새 숫자를 지어내지 마라. 호기심 갭을 풀어서 설명하지 마라.
4. 어미는 평어체 유지. 콜론(:) 금지. 글자 수는 원래와 비슷하게.

[깨진 줄 판정 기준 — 아래 11가지에 해당할 때만 수정]
① 소리 내어 읽어 3초 안에 뜻이 안 들어오는 줄 (두 번 읽어야 이해되면 깨진 것)
② 수식어 겹침 — 관형절 2개 이상이 겹쳐 핵심이 안 보임:
   "첫 구매자가 석 달 뒤 187만원을 쓴 그 흐름도 한 장" → "첫 구매자 한 명이 석 달 뒤 187만원을 더 썼다"
③ 어색한 수량·동사 결합:
   "카톡 문의가 세 건 쌓이기 시작했다" → "카톡 문의 세 건이 연달아 도착했다"
④ 사물+사람동사 오결합: "통장이 울렸다" → "입금 알림이 울렸다"
⑤ 방향이 모순되는 동사 연결: "꺾여 올라간" → 한 방향 동사 하나로
⑥ 비교+명사 꼬리표 비문: "~보다 빠른 증명" → 완결 문장으로
⑦ 추상명사 의인화: "전략이 깨어난다", "공식이 멈춘 날" 류
⑧ 꼬인 부정형: "~이 안 만들어지는" → 단순 부정(없다/모른다)으로
⑨ 허위 단정: "절대/무조건/100% ~안 된다" → 단정을 빼고 사실이 되게
⑩ 어색한 합성어: "~세팅법", "전략을 돌리다"(→ 적용하다/따르다), 컨셉을 "들은/배운"(→ 발견한)
⑪ 존댓말 혼입("~습니다") → 평어체로
⑫ 타사 브랜드·인물 이름(프드프, 탈잉, 자청, 신사임당 등)과 지어낸 리서치 주장("상위 50개를 분석") → 브랜드명은 일반 표현(전자책 시장)으로, 리서치 주장은 저자 경험 장면으로 교체
⑬ 물성 충돌 비유: 디지털 대상에 물리 동사·장소 ("전자책이 창고에서 썩는다") → 디지털의 실패 장면("조회수 4에 멈춰 있는")으로 교체
⑭ 꼬리 강조어: 명사 종결 뒤에 덧붙은 "전부/모두/다" 삭제 ("퇴근 후 한 일 전부" → "퇴근 후 한 일")
⑮ 길이 초과·숫자 과적: 챕터 18자, 소제목 24자를 넘거나 숫자가 2개 이상인 줄 → 조건절을 버리고 가장 충격적인 숫자 하나만 남겨 압축
⑯ 범주 불일치 비교: "시간이 연봉을 추월했다" 류 → 같은 범주로 ("2시간의 수익이 연봉을 넘어섰다")
⑰ 인용+서술 오결합·부정 프레임: "질문을 따지다" → "질문을 던지다/물어오다"로, 따지다·항의하다 프레임은 놀라움·반가움으로
⑱ 업계 은어 직역: "팔리는 사람", "나를 판다" → 실제 내용을 일반어로 ("독자는 글보다 글쓴이를 먼저 본다")

[주제] {topic}

[목차]
{numbered}

[출력 규칙 — 어기면 전체 무효]
- 입력과 정확히 같은 줄 수, 같은 라벨(C1, S1-1 형식)로만 출력
- 각 줄: 라벨 + 콜론 + (수정된 또는 그대로인) 제목
- 챕터 줄(C1~C5)에 'PART 1.' 같은 접두사가 있으면 그대로 유지
- 설명, 주석, 빈 줄, 다른 텍스트 일절 금지"""

    result = ask_ai(prompt, 0.15, ensure_quality=True)
    if not result:
        return chapters, subtopics

    fixed = {}
    for line in result.strip().split('\n'):
        m = re.match(r'^\s*([CS][\d\-]+)\s*[:.]\s*(.+)$', line.strip())
        if m:
            val = m.group(2).strip()
            # 양끝이 '짝으로' 감싸진 경우에만 따옴표 제거 — 문장 첫머리의 인용 따옴표를
            # 일방적으로 벗기면 ("더 열심히 써라"는 → 더 열심히 써라"는) 짝이 깨진다
            if len(val) >= 2 and val[0] == val[-1] and val[0] in '"\'':
                val = val[1:-1].strip()
            if len(val) >= 5:
                fixed[m.group(1)] = val

    new_chapters = []
    new_subtopics = {}
    for ci, ch in enumerate(chapters):
        nch = fixed.get(f"C{ci+1}", ch)
        if nch in new_chapters:  # 교정 결과가 중복되면 원본 유지
            nch = ch
        new_chapters.append(nch)
        subs = []
        for si, s in enumerate(subtopics.get(ch, [])):
            ns = fixed.get(f"S{ci+1}-{si+1}", s)
            subs.append(ns if ns not in subs else s)
        new_subtopics[nch] = subs
    return new_chapters, new_subtopics

def generate_outline_only(interview_data, progress_placeholder):
    """인터뷰 데이터를 기반으로 목차까지만 생성 (본문 제외)"""
    try:
        topic = interview_data.get('topic', '')
        if not topic:
            return False

        # 1. 타겟 자동 설정
        progress_placeholder.info("🎯 1/4 타겟 독자 분석 중...")
        target = f"{interview_data.get('target_reader', '')} - {interview_data.get('target_problem', '')}"
        st.session_state['target_persona'] = target

        # 2. 책 고유 컨셉 생성 (가장 중요!)
        progress_placeholder.info("💡 2/4 책 고유 컨셉 설계 중...")
        concept_prompt = f"""당신은 크몽/클래스101 베스트셀러 전자책 기획자입니다.
이 책만의 '고유한 시스템/공식'을 만들어야 합니다.

[저자 정보]
주제: {topic}
핵심 방법: {interview_data.get('core_method', '')}
저자만의 차별점: {interview_data.get('unique_point', '')}
타겟의 고민: {interview_data.get('target_problem', '')}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 실제 잘 팔리는 전자책의 고유 시스템/공식 예시
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[작명 원칙 - 고급스러운 영문 약어 스타일]
• 형태: 영문 대문자 약어(2~4자) + 공식/시스템/전략 — 예시 형태: "DIP 공식 (Dip-Invest-Profit)"처럼 약어 뒤에 풀네임 병기
• 약어는 반드시 이 주제의 핵심 영단어들의 이니셜에서 도출할 것 (주제가 다르면 약어도 완전히 달라져야 정상)
• 발음 가능하고 한 번에 읽히는 약어일 것
• 이미 널리 쓰이는 용어(ROI, KPI, PER, MBTI 등) 및 실존 기관·방송사·기업 약어(EBS, KBS, MBC, SBS, IMF, GDP, ETF 등)와 겹치면 폐기 후 재작명 — 유명 브랜드와 겹치면 책 전체가 우스워진다
• 사물 비유 단어(수도관, 물탱크, 파이프, 엔진, 눈덩이 등)는 이름에도 설명에도 쓰지 마라
• 건축 비유도 금지: 층, 적층, 쌓기(Stacking), 탑, 계단, 벽돌, 블록 — 이름·약어 풀네임·설명 어디에도 쓰지 마라
  (예: "DFS (Dividend Flow Stacking)" 같은 작명 금지 — 본문 전체가 "층을 쌓는" 식의 어색한 비유로 도배된다)

[이름 접미사 규칙]
• 허용: 공식, 시스템, 전략, 법칙, 원칙 (예: "○○ 공식", "○○ 전략")
• 금지 접미사: 루프, 사이클, 엔진, 부스터, 매트릭스, 퍼널, 트리거, 스택, 자석, 파이프라인, 스노우볼
  — 외래어 접미사는 한국어 문장에 넣었을 때 어색하다 ("○○ 루프를 발견한 날" ❌)
• 판별법: "○○○를 처음 발견한 날"이라는 문장에 넣어 소리 내 읽었을 때 자연스러워야 통과

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✅ 만들어야 할 것
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. 고유 시스템/공식 이름 (가장 중요!)
   - 영어 약자 + 한글 설명 (예: "CPM 전략")
   - 또는 직관적인 한글 조어 (예: "단단한 매출 구조")
   - 목차 전체에서 이 용어가 반복되어야 함

2. 핵심 관점
   - 이 주제를 어떤 새로운 시각으로 보는가?
   - 남들과 다른 접근법

3. 핵심 메시지
   - "[시스템명]만 알면 ~할 수 있다" 형식

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ 부자연스러운 과장:
- "제국을 건설", "왕좌에 오르다", "언더그라운드"
- "전설의", "역사를 바꾼", "세계 최초"

❌ 유치한 단어:
- 황금, 보물, 비밀, 마법, 연금술

❌ 모든 분야에 같은 비유 사용:
- 부동산 비유만 반복하지 말 것
- 주제에 맞는 다양한 비유 사용

❌ 이미 유명한 이름:
- 역행자, 추월차선, 아토믹 해빗 등 그대로 사용

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 출력 형식 — 어기면 전체 무효
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- 전체 500자 이내의 줄글. 마크다운 기호(#, *, -, |, 표, 굵게), 이모지, 체크표시(✅) 절대 금지.
- 아래 4개 라벨과 그 내용만 출력. 목차 예시, 장 구성 제안, 검증 결과 따위를 덧붙이지 마라.
  (특히 "B(Brand) — ..." 식의 낱글자 장 제목 예시는 목차 규칙 위반을 유발하므로 절대 금지)

[고유 시스템/공식 이름]
(영어 약자+한글 또는 참신한 한글 조어 — 한 줄)

[핵심 관점]
(2~3문장의 줄글 — 이 책이 주제를 보는 새로운 시각)

[핵심 메시지]
(한 문장)

[목차에서 반복할 키워드]
(시스템 이름 또는 핵심 단어 1~2개)"""

        # 이전에 쓴 컨셉명 금지 목록 주입 (사용자 간 워딩 중복 방지)
        _used = load_used_wordings()
        if _used:
            concept_prompt += "\n\n[절대 사용 금지 - 이미 다른 책에 쓰인 이름들]\n" + ", ".join(_used) + "\n위 이름들과 같거나 비슷한(한 글자만 다른) 이름도 금지."

        book_concept = ask_ai(concept_prompt, 0.9, ensure_quality=True)
        st.session_state['book_concept'] = book_concept

        # 생성된 컨셉명을 이력에 기록
        if book_concept:
            _m = re.search(r'\[고유 시스템/공식 이름\]\s*\n+\s*([^\n]+)', book_concept)
            if _m:
                _nm = _m.group(1).strip().strip('"\'`*()[]').strip()
                if 2 <= len(_nm) <= 25:
                    record_used_wording(_nm)

        # 3. 제목 생성
        progress_placeholder.info("📝 3/4 제목 생성 중...")
        title_prompt = f"""당신은 크몽/클래스101 베스트셀러 전자책 제목을 만드는 전문가입니다.
결제 버튼을 누르게 만드는 제목을 써주세요.

⚠️ 이 지시문에 나오는 실존 책 제목과 예시들은 패턴 참고용입니다. 그대로 또는 단어만 바꿔 쓰면 무효입니다.

[이 책의 컨셉]
{book_concept}

[주제]
{topic}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 실제 잘 팔리는 전자책 제목 분석
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[컨셉 중심형 - 짧은 신조어 + 부제]
• 단단한 돈 - 잃지 않는 사람의 7가지 원칙
• 돈의 속성 - 최소한 이것만은 알아야 할
• 1억 모으는 통장 - 30대 직장인의 5단계 공식

[신사임당/클래스101 스타일 - 결과 중심]
• 퇴사 후 월 1000만원 버는 글쓰기
• 블로그로 월 300 만드는 현실적인 방법
• 투잡러의 시간관리 비법

[크몽 베스트셀러 - 구체적 약속]
• 30일 만에 첫 수익 내는 스마트스토어
• 3개월 안에 월 100 만드는 전자책 공식
• 회사 다니면서 월 200 추가 수입 만들기

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✅ 좋은 제목의 공식
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

실제 교보문고·예스24 베스트셀러 제목 중 아래 4가지 패턴으로만 만들어라.

🚨 "X의 Y" 소유격 명사구 패턴은 사용 금지.
  ("한 권의 월급", "하락장의 태도", "퇴사 전의 월급"처럼 비문·밋밋한 조합이 양산되는 패턴이다.
   '의'로 두 명사를 잇는 제목은 어떤 형태든 만들지 마라.)

[패턴 1] 완결된 선언 문장 — 도발/통념 뒤집기
  실례: 가난은 습관이다 / 행동하지 않으면 인생은 바뀌지 않는다 / 미움받을 용기
  조건: 주어와 서술어가 갖춰진, 친구에게 말해도 자연스러운 문장
  🚨 수긍 테스트: "A는 B다" 단정은 충격적이되 5초 안에 "아, 그럴 수도 있겠다"로 수긍되어야 한다.
  "가난은 습관이다"가 작동하는 건 '나쁜 행동의 반복 → 가난'이라는 인과가 바로 보이기 때문.
  ❌ "수익은 성격이다" — B가 근거 없는 비약이라 도발이 아니라 억지가 된다. 인과의 다리가 안 보이면 폐기.
  🚨 방향 테스트: 선언은 '독자'를 향해야 한다 — 독자의 통념을 때리거나 욕망을 건드려야 제목이다.
  책 방법론의 작동 원리를 요약한 문장은 제목이 아니라 본문 소제목감이다.
  ❌ "한 권이 다음 권을 판다" — 방법의 메커니즘 설명일 뿐, 읽는 사람의 삶과 충돌하는 게 없어 평범한 문장으로 들린다.
  판별: 이 문장을 읽고 '내 이야기'로 느껴지는가, '그 방법에 대한 설명'으로 느껴지는가. 설명이면 폐기.
  🚨 기시감 테스트: 어디서 들어본 듯한 일반 문장을 그대로 가져온 느낌이면 탈락 — 단어 하나라도 의외의 자리에 있어야 신선하다.
  🚨 비유 단정 금지: "A는 [비유]다" 형태는 쓰지 마라. B 자리는 문자 그대로 말이 되는 단어여야 한다.
  자연물·사물 비유 전부 포함: 씨앗, 나무, 뿌리, 열매, 농사, 그릇, 다리, 지도 — "전자책은 씨앗이다" 같은 후보는
  점수가 높아도 출력 자체가 금지다. 후보 3개를 내보내기 전에 전원 이 검사를 다시 통과시켜라.
  ❌ "배당은 다시 심는 것이다" — 농사 비유가 단정에 들어가 문자 그대로는 뜻이 통하지 않는다 (배당을 심을 수 없다).
  ⭕ "가난은 습관이다" — '습관'은 비유가 아니라 문자 그대로의 진단이라 바로 이해된다.
  비유로 말하고 싶은 내용이면 비유를 풀어서 사실로 써라 ("배당은 쓰지 않으면 불어난다" 방향).
  🚨 빈 부정 선언 금지: "○○는 필요 없다", "○○는 중요하지 않다", "○○는 답이 아니다"처럼
  부정만 있고 대안이 없는 선언은 광고에서 닳은 프레임이라 기시감과 공허함으로 즉시 폐기.
  ❌ "글재주는 필요 없다" — 그래서 뭐가 필요한지가 없어서 김이 빠진다.
  부정을 쓰려면 버린 것과 택한 것이 한 문장에: ⭕ "나는 주식 대신 달러를 산다"

[패턴 2] 1인칭 행동 선언
  실례: 나는 4시간만 일한다 / 나는 주식 대신 달러를 산다
  조건: 반드시 "나는"으로 시작하는 완결문. 행동이 구체적이고 의외성이 있을 것.
  부정은 문어체로만: "보지 않는다/사지 않는다" ⭕, "안 본다/안 산다" ❌
  (❌ "주가를 안 본다" — 주어 없는 구어체 부정은 전단지 카피처럼 가볍다
   ✅ "나는 주가를 보지 않는다" — 같은 내용도 문어체 완결문이면 선언의 무게가 생긴다)
  🚨 동사 정밀성: 동사의 사전적 의미가 문장의 의도와 정확히 일치해야 한다.
  ❌ "나는 퇴사 전에 월급을 넘긴다" — '넘기다'는 건네다/페이지를 넘기다라는 뜻이라 의미가 어긋난 비문.
  의도가 '월급보다 많이 번다'면 동사도 그 뜻이어야: ✅ "나는 퇴사 전에 월급을 넘어섰다"
  작성 후 모든 동사를 사전 뜻으로 검증하고, 어긋나면 그 후보는 폐기하라.

[패턴 3] 한 단어 신조어 (2~4자)
  실례: 역행자 / 웰씽킹 / 그릿
  조건: 발음이 쉽고, 듣는 순간 뜻이 짐작되거나 궁금해질 것. 아무 단어나 줄여 붙이면 실패

[패턴 4] 시간/장면 + 보상
  실례: 퇴근 후 두 번째 월급 / 마흔에 읽는 쇼펜하우어
  조건: 구체적 시간·장면이 들어가고 읽는 사람의 삶이 연상될 것

🚨 제목 비문 검증 (가장 중요 — 이걸 통과 못 하면 전부 무효)
- 관형형+명사 압축 금지: "자는 전자책", "버는 글쓰기", "잠드는 통장"처럼
  동사 관형형을 명사에 그대로 붙인 압축은 듣는 순간 뜻이 안 통한다 — 즉시 폐기
- 소리 내어 읽었을 때 완결된 자연스러운 한국어여야 한다. 서점 직원에게
  "○○○ 있나요?"라고 물었을 때 어색하지 않아야 통과
- 의미 없는 명사 나열 금지: "수면 매출 설계도" 류

🚨 품격 검증 (전문 작가의 문장인가)
- 구어체 부정·축약 금지: "안 본다", "안 산다", "못 번다", "안 망한다" → "보지 않는다/사지 않는다"로.
  구어체 부정이 들어간 제목은 광고 전단지처럼 가벼워진다 — 즉시 폐기
- "관형형(~는/~한) + 명사" 두 단어 제목 전면 금지: ❌ "팔리는 순서" / "이기는 습관" / "버는 구조"
  — 문법은 맞아도 강의 슬라이드 제목처럼 가볍다. 특히 순서/방법/단계/구조 같은 실용 어휘로 끝나면 최악.
  같은 내용도 개념어로 승격시켜라: 속성, 문법, 감각, 기술, 태도, 수업, 심리학 같은 단어가 격을 만든다
  (예: '파는 행위'를 다룬 책이면 "팔리는 순서" ❌ → "파는 사람의 문법" 방향 ⭕ — 단어만 바꿔 재사용 금지)
- 줄임말, 유행어, 인터넷 입말 금지. 10년 뒤 서점 매대에 있어도 촌스럽지 않을 문장만
- 판별법: 이 제목을 하드커버 양장본 표지에 박았을 때 어울리는가. 어울리지 않으면 폐기

🚨 평이함 탈락 기준 (비문 검증과 동급으로 중요)
- 주제 키워드 두 개를 그대로 붙인 제목은 컨셉이 아니라 검색어다 — 즉시 폐기
  ❌ "전자책 복리" / "블로그 부자" / "배당 월급" — 광고 배너처럼 평이하고 딱딱함
- 핵심 키워드(전자책, 배당, 블로그 등)는 부제로 보내라. 메인 제목은 그 키워드 없이도
  핵심 컨셉이 전달될 때 가장 고급스럽다 ("돈의 속성"에는 '재테크'가 없다)
- 메인 제목에는 '비틀기 한 끗'이 반드시 있어야 한다 (다음 중 하나):
  · 역설 결합: "죽고 싶지만 떡볶이는 먹고 싶어" (모순이 호기심을 만든다)
  · 통념 비틀기: "가난은 습관이다" (습관=좋은 것이라는 통념을 뒤집음)
  · 당연한 것의 재발견: "돈의 속성" (돈을 관찰 대상으로 보는 시선 자체가 새로움)
  · 장면의 승격: "퇴근 후 두 번째 월급" (구체적 장면이 곧 컨셉이 됨)

[시적 품격 — 카피라이터가 아니라 시인이 쓴 문장처럼]
좋은 제목은 광고 카피가 아니라 시의 한 행처럼 읽힌다. 다음 장치 중 1개 이상을 넣어라:
· 반어의 여운: "하마터면 열심히 살 뻔했다" — 당연한 가치를 뒤집어 놓고 설명하지 않는다
· 일상 사물의 승격: 통장, 새벽, 퇴근길, 알림 같은 평범한 사물이 상징이 되게
· 병치의 긴장: 안 어울리는 두 단어를 나란히 두되, 뜻은 한 번에 통해야 한다
· 절제: 광고 어휘(지금, 바로, 무조건, 최고)와 느낌표 0개 — 조용한 문장이 더 오래 남는다

[어휘의 격 — 단어 선택이 품격을 만든다]
· 격조 레퍼런스 (톤만 흡수, 복제 금지): "하마터면 열심히 살 뻔했다" / "나는 나로 살기로 했다" /
  "내가 틀릴 수도 있습니다" — 에세이 베스트셀러의 절제된 품격. 자극적 단어 없이 사람을 멈춰 세운다
· 흔한 동사(하다, 되다, 만들다, 벌다)보다 정확하고 단정한 동사를 골라라.
  "돈을 번다" → "월급을 넘어선다", "돈이 들어온다" → "월급이 두 번 도착한다" — 같은 뜻도 동사가 격을 바꾼다
· 리듬: 7~12자 안에서 두 호흡으로 자연스럽게 나뉘면 우아하다 ("나는 나로 / 살기로 했다")
· 단어 배치: 한 단어만 뜻밖의 자리에 두고, 나머지는 단정하게 — 전부 튀면 소란스럽고 전부 얌전하면 평범하다
⚠️ 시적이라는 핑계로 뜻이 흐려지면 실패다. 아래 모든 검증을 통과한 문장에만 여운을 얹어라.

[내부 검증 절차 — 출력하지 말 것]
위 4가지 패턴으로 후보 7개를 만들고 각각 채점하라 ("X의 Y" 형태가 하나라도 섞이면 그 후보는 무효):
① 말이 되는가 ② 갖고 싶은가 ③ 무슨 책인지 짐작되는가 ④ 비틀기 한 끗이 있는가 ⑤ 어디서 본 듯한 평이함이 없는가
⑥ 품격 — 구어체·축약 없이 전문 작가의 문장인가 (하드커버 양장본에 어울리는가)
⑦ 여운 — 읽고 3초 뒤에도 머릿속에 남는가 (시인의 문장 테스트)
하나라도 미달인 후보는 버려라. 전부 미달이면 7개를 다시 만들어라.
통과한 후보 중 상위 3개를 점수 순으로 출력하라 (1번 = 최고작).

[후보 순서 규칙 — 어기면 전체 무효]
- 1번 후보(자동 선택됨)는 반드시 '1인칭 행동 선언'("나는 글 대신 구조를 판다" 결) 또는
  '시간/장면+보상'("퇴근 후 두 번째 월급" 결) 패턴이어야 한다.
- "A는 B다 / A는 B가 아니다" 선언문 패턴은 비문 위험이 가장 높으므로,
  모든 검증을 통과했더라도 1번 후보로는 금지 — 2~3번 슬롯에만 배치하라.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ 유치한 비유:
나침반, 지도, 열쇠, 보물, 황금, 마법, 연금술

❌ AI스러운 제목:
"~의 이해", "~가이드", "~완벽 정복"
"효과적인 ~", "성공적인 ~"

❌ 너무 추상적:
의미를 알 수 없는 신조어
무슨 내용인지 전혀 감이 안 오는 제목

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 출력 — 베스트셀러 제목의 절대 규칙
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[메인 제목] 2~12자. 위 4가지 패턴 중 하나로, 서점 매대의 실제 베스트셀러처럼.
  ✅ 좋은 형태: "가난은 습관이다" / "나는 4시간만 일한다" / "역행자" / "퇴근 후 두 번째 월급"
  ❌ "X의 Y" 소유격 형태는 전면 금지 ("돈의 속성"류 포함 — 이 책에서는 쓰지 않는다)
  ❌ 즉시 폐기: 관형형+명사 압축 ("자는 전자책", "버는 글쓰기"), 대시/하이픈 구조 ("배당 달력 - 12칸으로 완성하는 월급 시스템"), 기능 설명형 ("~으로 ~하는 ~"), 콜론, 15자 초과, 소리 내어 읽어 어색한 모든 것

[부제] 15~25자. 호기심 + 구체적 약속. 결과·기간·수단 중 2개 이상.
  ✅ "잃지 않는 사람들의 일곱 가지 원칙" / "퇴근 후 1시간이 만드는 두 번째 월급"
  ❌ 즉시 폐기: "~만들기", "~시스템", "~방법", "~하는 법"으로 끝나는 설명형 / 컨셉명 약어를 부제에 그대로 노출 ("DTS 전략으로...")

[공통] 평어체. 메인 제목만 읽고도 갖고 싶어야 하고, 부제를 읽으면 사야 하는 이유가 생겨야 한다.

[표지 카피 2종 — 표지의 구매 욕구는 카피 밀도에서 나온다]
- cover_eyebrow: 제목 '위'에 얹는 관형 카피 (8~14자). 제목과 이어 읽으면 한 문장이 되는 자연스러운 구절.
  형태 예: "품격 있는 대화를 위한" + 제목 / "퇴근 후 2시간으로 완성하는" + 제목. 광고체·과장 금지.
- belt_copy: 표지 하단 띠지 카피 (12~20자). 이 책의 가장 강한 약속 한 줄 — 구체적 숫자·기간 포함.
  부제와 겹치지 않는 다른 내용으로. 형태 예: "첫 달, 첫 입금이 시작된다"

JSON만 출력 — 모든 검증을 통과한 최고 후보 3개를 점수 순으로 (1번이 최고작):
{{
    "candidates": [
        {{
            "title": "메인 제목 (2~12자, 명사구/선언)",
            "subtitle": "부제 (15~25자, 호기심+약속)",
            "cover_eyebrow": "제목 위 관형 카피 (8~14자)",
            "belt_copy": "띠지 약속 카피 (12~20자, 숫자 포함)"
        }},
        {{ "title": "...", "subtitle": "...", "cover_eyebrow": "...", "belt_copy": "..." }},
        {{ "title": "...", "subtitle": "...", "cover_eyebrow": "...", "belt_copy": "..." }}
    ]
}}
(3개는 서로 다른 패턴이어야 한다 — 같은 패턴 변주 3개 금지)"""

        title_result = ask_ai(title_prompt, 0.4)
        title_data = parse_json(title_result)
        if title_data:
            _cands = title_data.get('candidates') or []
            if not _cands and title_data.get('title'):  # 단일 형식 호환
                _cands = [title_data]
            _cands = [c for c in _cands if isinstance(c, dict) and c.get('title')][:5]
            if _cands:
                st.session_state['title_candidates'] = _cands
                _best = _cands[0]
                st.session_state['book_title'] = _best.get('title', topic)
                st.session_state['subtitle'] = _best.get('subtitle', '')
                st.session_state['cover_eyebrow'] = (_best.get('cover_eyebrow') or '').strip()
                st.session_state['cover_belt_copy'] = (_best.get('belt_copy') or '').strip()

        # 4. 목차 생성 (책 컨셉 기반)
        if 'fable' in st.session_state.get('claude_model', '').lower():
            progress_placeholder.info("📋 4/4 목차 설계 중... (Fable 5는 깊은 추론 모델이라 1~3분 걸릴 수 있어요)")
        else:
            progress_placeholder.info("📋 4/4 목차 설계 중...")
        outline_prompt = f"""당신은 한국 PDF 전자책 시장 베스트셀러를 만드는 기획자다. 목차만 보고 결제 버튼을 누르게 만드는 5장짜리 목차를 쓴다.

⚠️ 예시 복제 금지 — 전역 규칙 (다른 모든 규칙보다 우선)
이 지시문에 등장하는 모든 예시 문장(✅ 좋은 예 포함)은 '구조'를 보여주는 견본일 뿐이다.
예시 문장을 그대로, 또는 단어 한두 개만 바꿔 출력에 쓰면 그 목차 전체가 무효다.
출력 전에 25줄 각각을 이 지시문의 예시들과 대조해서, 비슷한 줄이 있으면 완전히 새로 써라.
모든 줄은 이 책의 주제·실제 수단·새로운 숫자에서 만들어야 한다.

[이 책의 시그니처 컨셉]
{book_concept}

[주제] {topic}

━━━━━━━━━━━━━━━━━━━━━━━━━━
1. 뼈대 — 한 사람의 변화 드라마
━━━━━━━━━━━━━━━━━━━━━━━━━━
- PART 5개 × 소제목 4개 = 25줄.
- 5개 챕터 제목을 이어 읽으면 "막막했던 사람이 무엇을 만나서 어떻게 달라졌는가"의 줄거리가 보여야 한다.
  흐름: 충격적 현실 → 통념이 틀린 이유 → 첫 사건(전환점) → 시스템이 굴러감 → 달라진 삶
- 각 챕터의 마지막 소제목(PART 2~4)은 다음 챕터를 열어보게 만드는 떡밥. 단, 떡밥에는 반드시 '이 책의 주제에서 나온 실물 하나'(화면 캡처, 메시지, 알림, 숫자 한 줄 등 — 이 예시 단어를 그대로 쓰지 말고 주제의 실물로)가 박혀 있어야 한다. "여기서 끝이 아니었다" 같은 정보 0짜리 문장은 폐기.

━━━━━━━━━━━━━━━━━━━━━━━━━━
2. 한 줄 작법 — 실제 베스트셀러 패턴 9개 (이 패턴들로만 써라, 각 1~3회)
━━━━━━━━━━━━━━━━━━━━━━━━━━
① 지시어 숨김: "나는 '이것' 하나로 상대의 미래를 예측한다" / "인생에 개사기 스킬이 딱 두 가지 있다"
② 1인칭 고백(평어체): "내가 철학을 버리고 100억을 택한 이유" / "평생 후회 없던 내가 딱 2개를 후회한다"
③ 통념 도발: "돈이 없다면 운이 억세게 좋은 것이다" / "노력할수록 가난해지는 구조가 있다"
④ 숫자·기간·금액 공개 + 방법 은닉: "일주일 만에 월 4천 버는 자동수익의 구조" (수단은 보여주고 만드는 법만 본문에)
⑤ 괄호 장치: "돈을 버는 짧은 공식 (어려움주의)"
⑥ 장면 한 컷: "새벽 3시 결제 알림이 공포에서 설렘으로 바뀐 날"
⑦ 질문 도발: "처음부터 1등하는 사람은 타고난 걸까"
⑧ 컨셉 단정: "'○○○'만 알면 인생이 바뀐다" / "결국 '○○○'를 아는 사람만 살아남는다"
⑨ 컨셉 이유형: "'○○○'가 삶을 바꾸는 이유" / "'○○○'를 아는 사람과 모르는 사람의 1년 뒤"
   (⑧⑨는 컨셉명 허용 3회 슬롯 안에서만 사용 — 컨셉을 책의 운명을 가르는 분기점으로 보이게 한다)

핵심 원칙 — WHAT은 보여주고 HOW만 숨겨라:
독자가 "내가 뭘 하게 되는지"(블로그 글, 전자책, 배당 ETF, 위임, 검색 키워드 등 이 주제의 실제 수단)는 알 수 있어야 한다. 25줄 중 8줄 이상에 구체적 수단이 보이게. 수단이 하나도 안 보이는 목차는 사기처럼 보여서 안 팔린다.

🧠 뇌과학 마케팅 트리거 — 25줄 중 최소 14줄에 아래 중 하나가 박혀야 한다 (없는 줄은 '그냥 정보'다):
① 손실회피: "모르면 평생 ~한다" + 충격 통계(99%, 92%, 단 3%)
② 통념 박살(배신감): "여태 들은 조언이 정반대였다"
③ 호기심 갭: 결과·사건만 보여주고 이유·방법은 숨김 ("가장 먼저 끊은 것 하나", "정확히 둘째 달에 멈추는 지점")
④ 권위: 데이터, 수치, 연차, 실계좌
⑤ 사회증명: "~를 익힌 사람들의 5년 뒤", "상위 3%의 공통점"
⑥ 임박감: "이번 달 월급날부터", "읽은 그 주에"
⑦ 이중 보상: 돈 + 시간/마음이 같이 좋아지는 줄
온도 판별: 읽고 "아 그렇구나"로 끝나면 폐기. "뭔데? 왜? 나도?"가 터져야 통과. 밋밋한 줄 하나가 목차 전체를 죽인다.

[자극 배치 규칙 — 결제 버튼은 여기서 눌린다]
- PART 1 제목과 첫 소제목: 독자가 '지금 이 순간' 겪고 있는 고통을 정면으로 때려라. 읽는 순간 뜨끔해서 자기 얘기인 줄 알아야 한다.
- 각 PART의 첫 소제목 = 그 파트에서 가장 자극적인 줄 (제일 센 걸 맨 앞에 — 독자는 각 파트 첫 줄만 훑는다)
- 챕터 제목 5개 중 3개 이상은 손실회피·배신감 계열 (지금 잃고 있는 것, 여태 속고 있던 것)
- 다 쓴 뒤 '가장 약한 줄 5개'를 스스로 골라 한 단계 더 세게 다시 써라 — 단, 문법 규칙은 그대로 지키면서

[언박싱 테스트 — 밀봉된 카드팩을 까는 심리]
좋은 목차의 한 줄은 밀봉된 카드팩이다. 읽는 순간 "안에 뭐가 들었는지 당장 까보고 싶다"는
충동이 일어야 한다. 선물 포장을 뜯기 직전의 설렘을 25줄 전부에 심어라.
- 실루엣 원칙: 내용물이 '확실히 있다'는 건 보여주고, '정체'만 가려라.
  봉인된 보상 = 구체물의 윤곽으로 ("그날 통장에 찍힌 숫자", "정산서의 마지막 줄", "그가 보낸 한 문장").
  윤곽조차 없이 가리기만 하면 빈 팩처럼 보여서 아무도 안 깐다.
- 레어 카드 3장: 25줄 중 3줄은 '저자가 가장 아끼는 패'가 들어 있음이 문장에서 느껴지게 써라.
  ("끝까지 망설이다 처음 공개하는", "가장 많이 받았던 질문에 이제야 답하는" — 사실의 톤으로, 과장 금지)
- 개봉 직전 컷: 문장을 '까기 직전'에서 끊어라. 결과를 보여주지 말고,
  결과가 공개되는 순간의 문턱(봉투를 여는 손, 화면이 켜지는 순간)에서 멈춰라.
- 빈 팩 판별: 각 줄을 읽고 "지금 당장 이 장부터 펼치고 싶은가?" — 아니라면 그 줄은 빈 팩이다. 다시 써라.

━━━━━━━━━━━━━━━━━━━━━━━━━━
2-1. 차별화 — "이 책은 다르다"가 목차에서 보여야 한다
━━━━━━━━━━━━━━━━━━━━━━━━━━
독자는 이 주제의 책을 이미 한두 권 읽었거나 유튜브로 봤다. 그 책들과 똑같아 보이는 목차는 0원이다.
- 25줄 중 최소 4줄에 '기존 방법 부정 + 이 책만의 방식 암시'를 박아라:
  · 시중의 흔한 조언이 왜 실패하는지 정면으로 공격하는 줄 (PART 1~2에 배치)
  · 대부분이 쓰는 방법과 이 책의 방법이 갈라지는 결정적 지점을 보여주는 줄
  · 시그니처 컨셉('○○○')이 기존 방식과 정반대로 가는 부분을 암시하는 줄
- 차별화는 구체적으로: "남들과 다르다" 같은 빈 선언이 아니라, 무엇을 '안 하는지' 또는
  '반대로 하는지'가 보여야 한다. 형태 예: "○○부터 하라는 조언을 버리는 것이 첫 단계다"
  (형태만 참고 — 이 책의 주제에서 실제 통념과 실제 차이점을 찾아 새로 써라)
- 참신함(형식 파괴): 25줄 중 2~3줄은 시중 전자책 목차에서 본 적 없는 형식으로 써서
  "이 책은 만듦새부터 다르다"는 인상을 줘라. 허용 장치:
  · 실제 대사 인용으로 시작: "이거 진짜 되네요" — 수강생에게 가장 많이 받은 메시지
    (인용 뒤 서술은 자연스러운 결합만: 질문은 '던지다/물어오다', 메시지는 '받다'. "질문을 따지다" 같은 오결합 금지.
     부정적 프레임(따지다, 항의하다)보다 놀라움·반가움 프레임이 책의 격을 지킨다)
  · 괄호 증거 장치: 첫 달 정산의 전 과정 (계좌 캡처 그대로 실음)
  · 숫자가 주어인 장면: 3,840원이 내 계획을 전부 바꿨다
  단, 비문 금지·문외한 테스트는 그대로 적용 — 신선하되 한 번에 읽혀야 하고, 위 예시 문장 재사용은 금지.

━━━━━━━━━━━━━━━━━━━━━━━━━━
3. 어미와 표현
━━━━━━━━━━━━━━━━━━━━━━━━━━
- 전부 평어체(반말). "~습니다/~해요/제가/저는" 금지. 1인칭은 "내가/나는".
- 주어-서술어 호응: "전략이 돌아간다", "공식이 깨어난다", "구조가 숨쉰다" 같은 추상명사 의인화 절대 금지.
- 사물에는 실제로 가능한 동사만: 울리는 건 알림·휴대폰뿐이다 — "통장이 울렸다" 즉시 폐기. 통장은 "찍히다/불어나다/비어 있다", 돈은 "들어오다/빠져나가다"만 가능. 사물에 사람·기계 동사를 붙인 줄은 폐기 후 다시 쓴다.
- 명사구 쉼표 나열 금지: ❌ "1억을 모아야 월 30만원, 나는 그 계산에 질렸다" — 명사구와 문장을 쉼표로 이은 비문. ✅ "두 달을 뛰어도 몸무게가 그대로라는 사실에 나는 질렸다" (운동 주제의 예 — 한 줄은 완결된 한 문장. 이 예시를 단어만 바꿔 재사용 금지).
- 한 줄 = 한 생각: 명사구 두 개를 쉼표로 잇지 마라.
  ❌ "연봉 5,200만원, 60세 통장 잔고 시뮬레이션을 돌린 날" — 정보 두 개가 충돌해서 두 번 읽어야 한다
  ✅ "60세의 내 통장을 미리 열어본 날" — 한 장면, 한 번에 이해
  쉼표는 한 줄에 최대 1개, 자연스러운 호흡일 때만.
- 3초 이해 테스트: 소리 내어 읽었을 때 3초 안에 장면이 그려지지 않는 줄은 폐기.
- 자연어 테스트 (챕터 제목 필수): 친구에게 말로 해도 어색하지 않아야 한다. 컨셉을 압축하려고 명사구를 서술어 자리에 박지 마라.
  ❌ "운동은 시간이 아니라 회복 타이밍의 설계다" — 명사구가 술어 자리에 박힌 비문
  ❌ "루틴을 모르면 헬스장도 운동이 안 된다" — "~이 안 된다" 꼬인 부정, 두 번 읽어야 이해됨
  ✅ "같은 한 시간을 쓰고도 몸이 먼저 달라지는 사람들" — 장면이 그려지고 한 번에 읽힘
  ⚠️ 위 예시(이 프롬프트의 모든 예시 포함)를 단어만 바꿔 재사용하지 마라. 이 책의 주제에서 새 문장을 써라.
  규칙: "A는 B가 아니라 C다" 구조를 쓸 거면 C는 한 단어로. 서술어는 살아있는 동사로 끝내라. "~이/가 안 된다", "~이 아니라 ~이다"에 3어절 이상 명사구 금지.
- 챕터 제목: 공백 포함 18자 이내 — 초과하면 즉시 폐기 후 압축. 제목 5개 중 3개 이상에 호기심 장치(숨긴 대상 '이것', 구체 숫자, 뜻밖의 조합)를 넣어라. "~하게 된다/~할 수 있다" 금지.
  ❌ "월급 300만원으로 노후를 준비하면 60세에 통장 잔고가 얼마인지 아는가" (31자 — 읽다가 지친다)
  ✅ 같은 충격을 짧게: "60세의 내 통장을 미리 열어봤다" (14자 — 한눈에 들어오고 까보고 싶다)
  ❌ 단조로움: "배당 투자를 시작하다", "시스템이 완성되다"
  ✅ "월급날 아침 10분이 노후를 결정한다" / "은행이 말해주지 않는 두 번째 월급" / "첫 배당 3,840원이 바꾼 것"
- 소제목: 공백 포함 24자 이내 — 초과하면 즉시 폐기 후 압축. 콜론(:) 금지.
- 한 줄에 숫자는 1개까지: 숫자가 2개 이상 들어가면 정보가 충돌해 두 번 읽어야 한다.
  ❌ "연봉 5천만원을 25년 모아도 월 30만원 배당이 한계인 계산법의 함정" (숫자 3개 + 27자 — 무슨 말인지 한 번에 안 들어옴)
  ✅ 가장 충격적인 숫자 하나만 남겨라: "성실하게 모은 노후가 월 30만원인 이유" (조건절은 본문이 설명한다)
- 압축 원칙: 조건절("~하면", "~해도")을 버리고 결과의 충격만 남겨라. 자청 목차가 강한 건 짧아서다.
- 비유 금지: 물탱크, 수도관, 파이프, 엔진, 눈덩이, 나침반, 열쇠, 마법, 정원, 항해 등 사물 비유로 줄을 만들지 마라. 특히 목차 전체를 하나의 비유 테마로 도배하는 것 절대 금지. 비유 없이 사실·장면·숫자로만 써도 충분히 강하다.
- 금지 어휘: 자청/역행자/자의식 해체/유전자/원시인/추월차선(특정 작가 고유어), ~의 이해/~하는 방법/~의 중요성/효과적인/다양한(설명체), 비법/노하우로 끝나는 줄.
- 특정 플랫폼·인물 이름 금지: 프드프, 탈잉, 클래스101, 자청, 신사임당 등 타사 브랜드와 유명 인물 이름을 목차에 절대 쓰지 마라. (책 주제 자체가 특정 플랫폼일 때 그 플랫폼명만 '수단'으로 허용)
- 지어낸 리서치 주장 금지: "베스트셀러 상위 50개를 분석하고", "1만 명을 조사한" 같은, 저자가 실제로 했는지 알 수 없는 조사·분석 주장을 만들어내지 마라. 저자 정보에 있는 사실만 쓸 것.
- 컨셉명은 소제목에만 3회 (각 슬롯에 ⑧⑨ 패턴 사용 가능 — "'○○○'만 알면 인생이 바뀐다" / "'○○○'가 삶을 바꾸는 이유" 식):
  ① PART 1 마지막 — 도입 ("결국 모든 답은 '○○○' 안에 있었다" 식)
  ② PART 3 또는 4 — 적용 장면. 단, 컨셉명을 주어로 의인화하지 마라 ("○○○ 전략이 돌아가기 시작했다" ❌ → "'○○○'를 처음 적용한 날 생긴 일" ✅ — 행위 주체는 사람)
     전략·공식·시스템에 "돌리다/굴리다" 금지: ❌ "○○○를 돌린 지 18개월" — 돌리는 건 기계·시뮬레이션뿐.
     전략은 "쓰다/적용하다/따르다/시작하다"로: ✅ "○○○대로 투자한 지 18개월" / "○○○를 따른 첫 달"
  ③ PART 5 마지막 — "'○○○'를 익힌 사람들의 5년 뒤가 다른 이유"
  🚨 컨셉의 주인은 저자다: '○○○'는 저자가 시행착오 끝에 직접 만들어 이름 붙인 방법이다.
     남에게 전수받은 듯한 표현 절대 금지: ❌ "○○○를 처음 들은 날" / "○○○를 배운 날" / "○○○를 알게 된 날" / "○○○를 접한 순간"
     만든 사람의 시점으로만: ✅ "○○○를 처음 발견한 날" / "○○○라는 이름을 붙인 날" / "○○○를 완성한 날" / "○○○가 만들어진 새벽"
- 약어 해부: 컨셉이 영문 약어라면, 시스템을 소개하는 챕터(PART 3~4)의 소제목들이 약어를 구성하는 단어들의 '한국어 의미'를 한 줄씩 자연스럽게 보여줘야 한다. 목차를 다 읽으면 약어가 대략 무슨 뜻인지 짐작되게.
  🚨 알파벳 한 글자를 목차에 쓰는 것 전면 금지 — 어떤 형태로든:
     ❌ "D, ~" / "T는 ~" / "S가 ~" (글자 라벨링)
     ❌ "WPS 공식의 W, 쓰지 말고 채우는 5단계 템플릿" ("○○ 공식의 X" 분해형 — 가장 흔한 위반)
     ❌ "WPS 공식의 S, 복리 수익 구조를 만드는 전략"
     약어 전체("WPS 공식")를 통째로 쓰는 건 지정된 3회만 가능. 낱글자(W, P, S)는 0회.
  ✅ 글자가 뜻하는 내용을 '간접적으로' 장면에 녹여라 — 독자는 본문에서 "아, 그게 W였구나"를 깨닫게 된다:
     ❌ "WPS 공식의 W, 쓰지 말고 채우는 템플릿" → ✅ "쓰는 게 아니라 채우기만 하면 되는 5단계 템플릿"
     (구성 단어가 배당/자동매수/복리라면)
     "월급보다 먼저 들어오는 돈을 만드는 순서" (배당)
     "내가 자는 동안 매수 버튼이 눌리는 구조" (자동매수)
     "셋째 해부터 잔고가 불어나는 속도가 달라지는 이유" (복리)
     ⚠️ 이 세 줄은 형태 참고용 — 그대로/단어만 바꿔 쓰지 말고 이 책의 주제에서 새로 써라

━━━━━━━━━━━━━━━━━━━━━━━━━━
4. 1000만원 가치 테스트 — 구매 후의 삶이 보여야 산다
━━━━━━━━━━━━━━━━━━━━━━━━━━
독자가 목차만 읽고 "따라 하면 N주 안에 내 삶에 [구체적 변화]가 생긴다"를 그릴 수 있어야 1000만원에도 산다. "있으면 좋겠다"는 0원이다.
- 각 PART에 최소 1줄은 '약속 줄': 결과 + 시점 + 수단이 한 줄에.
  ✅ "읽은 그 주에 끝나는 첫 자동이체 세팅" / "석 달 뒤 월급날이 한 달에 두 번이 된다"
- 약속은 허풍 대신 검증 가능한 디테일로: 금액은 소박하게 구체적으로(3,840원, 27만원이 1,000만원보다 믿긴다), 기간은 현실적으로(첫 주, 둘째 달, 18개월).
- 마지막 PART에 '도착점' 줄 1개: 이 책을 끝낸 독자의 1년 뒤 하루가 그려지는 한 줄.
- 약속 줄은 목차 전체에 7개 이상. "막연히 좋아진다"는 0원이다 — "언제까지(첫 주, 둘째 달), 얼마가(27만원, 월 10만원), 무엇으로(자동이체, ETF 1개)"가 보여야 100만원짜리 기대 효과다.

━━━━━━━━━━━━━━━━━━━━━━━━━━
5. 온도 기준 — 모범 목차 (톤만 모방, 문장 구조·단어·숫자 복제 금지)
━━━━━━━━━━━━━━━━━━━━━━━━━━
(주제: 퇴근 후 부업 월 500 / '○○○' 자리에는 그 책만의 컨셉명)

PART 1. 성실한 사람부터 가난해진다
- 입사 7년 차, 통장에 247만원이 전부였다
- 야근이 늘수록 잔고가 줄어드는 구조의 정체
- 부업 30일 차에 92%가 조용히 접는 이유
- 결국 모든 답은 '○○○' 안에 있었다

PART 2. 노력보다 자리가 먼저 돈을 번다
- 하루 14시간 일한 내가 4시간 일하는 사람에게 졌다
- "더 열심히"가 가장 위험한 조언인 이유
- 월 500 찍은 사람들이 가장 먼저 끊은 것 하나
- 블로그 글 한 편이 그날 밤 모든 걸 뒤집었다

PART 3. 첫 27만원이 통장에 찍힌 날
- 전자책 한 권의 첫 정산이 월급보다 무거웠다
- 둘째 달에 90%가 멈추는 정확한 지점
- '○○○'가 작동을 시작하는 4가지 신호
- 그날 새벽, 자는 동안 두 번째 알림이 울렸다

PART 4. 잠든 사이에도 매출이 들어온다
- 손을 떼는 순간 매출이 커지는 위임의 역설
- 월 100과 월 500을 가르는 단 한 줄의 차이
- 나는 '이것' 하나만 매일 아침 확인한다
- 마지막 관문은 통장이 아니라 키보드 앞에 있었다

PART 5. 월급이 용돈으로 보이기 시작한다
- 퇴사를 통보하던 날 사장이 꺼낸 말
- 통장과 저녁 시간이 함께 불어나는 이중 보상
- 1년 뒤 가장 크게 달라지는 단 한 가지
- '○○○'를 익힌 사람들의 5년 뒤가 다른 이유

⚠️ 위 예시에서 가져갈 것은 '온도'뿐이다. 문장 구조를 줄 단위로 본뜨면(예: "○○보다 먼저 바뀐 건 △△였다" 구조 재활용) 어색한 문장이 나온다. 이 책의 주제에서 나온 실제 수단·장면·숫자로 새로 써라.

━━━━━━━━━━━━━━━━━━━━━━━━━━
5-1. 사실 검증 — 말이 안 되는 줄은 한 줄도 용납 안 된다 (최우선 필터)
━━━━━━━━━━━━━━━━━━━━━━━━━━
- 허위 단정 금지: "절대", "무조건", "100%", "평생 못 한다"로 사실이 아닌 일반화를 만들지 마라.
  ❌ "월급만으로는 절대 월 100이 안 나온다" — 거짓 단정(그런 사람은 흔하다). 독자가 읽는 순간 책 전체의 신뢰가 무너진다
- 꼬인 부정형 금지: "~이 안 나온다", "~이 안 만들어지는", "~가 안 되는" 같은 부정 서술 금지.
  부정이 필요하면 "없다 / 모른다 / 틀렸다"처럼 단순하게 끝내라.
- 건축·사물 비유 금지(확장): 층, 탑, 계단, 사다리, 벽돌, 블록, 퍼즐, 설계도 — "첫 번째 층을 쌓는 순서" 같은 줄 금지.
  비유 대신 실물(돈, 통장, 입금일, 결제 알림, 정산서)로 써라.
- 사물+사람동사 오결합 금지: ❌ "월급 없는 주에 통장이 울렸다" (통장은 울리지 않는다 — 울리는 건 알림) ❌ "계좌가 깨어났다"
- 전략·공식에 "돌리다/굴리다" 금지: ❌ "DCP 전략을 돌린 지 18개월" — 돌리는 건 기계·시뮬레이션뿐. 전략은 쓰다/적용하다/따르다 ("이 방식대로 투자한 지 18개월")
- 비교 구문을 명사 꼬리표에 구겨 넣지 마라: ❌ "배당금 100% 재투자가 원금 추가보다 빠른 증명" — "~보다 빠른 증명"은 비문.
  비교가 필요하면 단순하게: ✅ "돈을 더 넣는 것보다 빨랐던 한 가지" — 비교 대상 하나는 숨겨야 자연스럽고 궁금해진다
- 명사구 쉼표 나열 금지: ❌ "1억을 모아야 월 30만원, 나는 그 계산에 질렸다" — 명사구와 문장을 쉼표로 잇지 마라
- 모순 동사 연결 금지: ❌ "잔고 곡선이 꺾여 올라간 이유" — '꺾이다'와 '올라가다'처럼 방향이 충돌하는 동사를 잇지 마라.
  움직임은 한 방향 동사 하나로 ("잔고가 불어나는 속도가 달라진 이유")
- 비교·추월 구문은 같은 범주끼리만: 주어와 목적어가 돈↔돈, 시간↔시간으로 맞아야 한다.
  ❌ "퇴근 후 2시간이 연봉을 추월했다" — 시간이 돈을 추월할 수 없다.
  ✅ "퇴근 후 2시간의 수익이 연봉을 넘어섰다" — 수익↔연봉, 같은 범주라 말이 된다.
- 업계 은어를 문자 그대로 쓰지 마라: "사람이 팔리다", "나를 판다", "자신을 세일즈한다" 같은
  마케팅 은어는 일반 독자에겐 비문이거나 거부감을 준다.
  ❌ "팔리는 사람이 먼저다" — 사람은 팔리는 대상이 아니다.
  은어가 가리키는 실제 내용을 풀어서: ✅ "독자는 글보다 글쓴이를 먼저 본다"
- 수식어 겹쳐 쌓기 금지: 한 줄에 관형절(꾸미는 말)은 1개까지.
  ❌ "5천만원으로 월 200만원 시작한 사람의 정산서 한 장" — 수식이 세 겹이라 무슨 말인지 모름
- "~법/~세팅" 꼬리표 금지: ❌ "루틴 세팅법", "마인드 셋업" — 외래어+법 합성이나 "~법"으로 끝나는 줄 금지
- 물성 충돌 비유 금지: 디지털 대상(전자책, 블로그 글, 파일, 계좌)에 물리 세계의 장소·동사(창고, 썩다, 먼지, 녹슬다, 쌓인 재고)를 붙이지 마라.
  ❌ "브랜딩 없이 올린 전자책이 창고에서 썩는 이유" — 전자책은 창고에 없고 썩지 않는다.
  디지털엔 디지털의 실패 장면을: ✅ "올린 지 석 달, 조회수 4에 멈춰 있는 전자책" / "아무도 클릭하지 않는 판매 페이지"
- 꼬리 강조어 금지: 명사로 끝난 줄 뒤에 "전부", "모두", "다", "전체"를 덧붙이지 마라 — 군더더기다.
  ❌ "월 478만원 찍은 달에 내가 퇴근 후 한 일 전부" → ✅ "월 478만원 찍은 달에 내가 퇴근 후 한 일"
  명사에서 끝났으면 거기서 끝내라. 강조가 필요하면 "단 하나의 ~"처럼 문장 안에 녹여라.
- 막연한 전부-역전 표현 금지: "모든 게 반대였다", "모든 게 달라졌다", "전부 뒤집혔다"는
  무엇이 어떻게 변했는지 단 하나도 안 보이는 0정보 문장이다 — 즉시 폐기.
  반전은 가장 중요한 '하나'를 집어 구체적으로: ✅ "가장 먼저 버린 건 가격을 낮추라는 조언이었다"
- 따옴표는 반드시 짝으로: 여는 따옴표 없이 닫는 따옴표만 쓰지 마라 ("더 열심히 써라"는 ❌ → "더 열심히 써라"는 ⭕ 처럼 앞뒤 모두)
- 문외한 테스트: 이 주제를 처음 접한 사람이 읽어도 뜻이 한 번에 들어와야 한다. 두 번 읽어야 이해되는 줄은 폐기 후 다시 쓴다.
- ⚠️ 이 프롬프트의 모든 ✅ 예시는 '구조'만 배우는 용도다. 예시 문장을 그대로 또는 단어만 바꿔 목차에 쓰면 즉시 무효.
- 🔥 단, 이 검증 규칙들은 문장을 '바르게' 쓰라는 것이지 '약하게' 쓰라는 게 아니다.
  문법을 지키면서 도발·호기심·구체적 숫자의 수위는 베스트셀러 최상급으로 유지하라.
  안전하고 밋밋한 25줄보다, 문법이 정확하면서 심장을 건드리는 25줄이 정답이다.
  "배당 투자의 기본 원칙" 같은 교과서 줄이 하나라도 있으면 그 목차는 실패다.

━━━━━━━━━━━━━━━━━━━━━━━━━━
6. 출력 형식 (이 형식 외 어떤 텍스트도 출력 금지)
━━━━━━━━━━━━━━━━━━━━━━━━━━
PART 1. [챕터 제목]
- [소제목]
- [소제목]
- [소제목]
- [소제목]
(PART 5까지 동일)

출력 전 마지막 점검: ① 챕터 5개만 이어 읽어도 줄거리가 보이는가 ② 존댓말·추상명사 의인화("전략이 돌아간다" 류) 0건인가 ③ 수단이 보이는 줄 8개 이상인가 ④ 사물 비유 테마가 없는가 ⑤ 각 PART에 약속 줄(결과+시점+수단)이 있는가 ⑥ 약어 컨셉이면 글자별 의미가 어렴풋이 보이는가 ⑦ 허위 단정("절대 ~안 된다")·꼬인 부정형·사물+사람동사 오결합("통장이 울렸다" 류)·명사구 쉼표 나열·뜻이 안 들어오는 줄이 0건인가 ⑧ "언제까지 얼마가 무엇으로"가 보이는 약속 줄이 7개 이상이고 호기심 트리거가 14줄 이상인가 ⑨ 약어 낱글자("○○ 공식의 W" 류)가 25줄 어디에도 없는가 ⑩ '기존 방법과 다르게 간다'가 보이는 차별화 줄이 4개 이상인가 ⑪ 교과서처럼 밋밋한 줄("~의 기본", "~ 이해하기" 류)이 0건인가 ⑫ 마지막으로 독자가 되어 처음부터 읽어보라 — 이 목차에 지금 당장 12,900원을 결제하고 싶지 않다면, 가장 약한 줄 5개를 골라 더 세게 다시 쓴 뒤 출력하라. 점검 결과는 출력하지 마라."""

        chapters = []
        subtopics = {}
        for _outline_attempt in range(2):
            outline_result = ask_ai(outline_prompt, 0.85, ensure_quality=True)
            if not outline_result:
                continue

            chapters = []
            subtopics = {}
            current_ch = None

            lines = outline_result.split('\n')
            for i, orig_line in enumerate(lines):
                line = orig_line.strip()
                if not line:
                    continue

                # 챕터 감지 (PART, 파트, Chapter, 1., 2. 등 다양한 형식)
                is_chapter = False
                ch_name = None

                # 마크다운 강조/헤더 기호를 벗겨낸 감지용 라인
                # (신형 Claude 모델이 **PART 1. ...**, ## PART 1, > 등으로 헤더를 감싸 출력해도 인식)
                detect_line = re.sub(r'^[\s>#\*_`~]+', '', line)
                detect_line = re.sub(r'[\*_`~]+$', '', detect_line).strip()

                # PART 1. 제목 형식
                if re.match(r'^(PART|파트|Part)\s*\d+[\.\s]', detect_line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = detect_line
                # Chapter 1. 제목 형식
                elif re.match(r'^(Chapter|챕터)\s*\d+[\.\s]', detect_line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = detect_line
                # 마크다운 헤더 형식
                elif re.match(r'^#+\s*(PART|파트|Chapter|챕터|\d+)', line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = re.sub(r'^#+\s*', '', line)
                # 1. 제목 형식 (숫자로 시작, 들여쓰기 없음)
                elif re.match(r'^\d+[\.\)]\s', detect_line) and not orig_line.lstrip('*_`~ ').startswith(' '):
                    is_chapter = True
                    ch_name = detect_line
                # 【1부】 형식
                elif re.match(r'^[【\[]?\s*\d+\s*(부|장|편)[】\]]?', detect_line):
                    is_chapter = True
                    ch_name = detect_line

                if is_chapter and ch_name:
                    ch_name = re.sub(r'^[#\*\-\s]+', '', ch_name)
                    ch_name = ch_name.replace('**', '').replace('*', '').replace('#', '').strip()
                    if ch_name and len(ch_name) > 3:
                        current_ch = ch_name
                        if current_ch not in chapters:
                            chapters.append(current_ch)
                            subtopics[current_ch] = []
                elif current_ch:
                    # 소제목 감지
                    is_subtopic = False
                    st_name = None

                    # - 소제목 형식
                    if re.match(r'^[\-\•\·\*\→\▶]\s*', line):
                        is_subtopic = True
                        st_name = re.sub(r'^[\-\•\·\*\→\▶]\s*', '', line)
                    # 1) 소제목, a) 소제목 형식
                    elif re.match(r'^[a-z\d][\)\.\:]\s', line, re.IGNORECASE):
                        is_subtopic = True
                        st_name = re.sub(r'^[a-z\d][\)\.\:]\s*', '', line, flags=re.IGNORECASE)
                    # 들여쓰기된 라인
                    elif orig_line.startswith('  ') or orig_line.startswith('\t'):
                        is_subtopic = True
                        st_name = line.lstrip('- •·*→▶0123456789.):\t ')
                    # 챕터가 아닌 일반 텍스트 (이전이 챕터였고, 현재가 짧은 문장이면 소제목으로 간주)
                    elif len(chapters) > 0 and not re.match(r'^(PART|파트|Part|Chapter|챕터|\d+[\.\)])', line, re.IGNORECASE):
                        if len(line) > 5 and len(line) < 100:
                            is_subtopic = True
                            st_name = line.lstrip('- •·*→▶0123456789.):\t ')

                    if is_subtopic and st_name:
                        st_name = st_name.replace('**', '').replace('*', '').replace('#', '').strip()
                        st_name = re.sub(r'^\d+[\.\)\:]\s*', '', st_name)  # 앞 숫자 제거
                        if st_name and len(st_name) > 3 and len(subtopics[current_ch]) < 4:
                            # 챕터 이름과 동일하면 스킵
                            if st_name.lower() != current_ch.lower() and st_name not in subtopics[current_ch]:
                                subtopics[current_ch].append(st_name)

            # 챕터가 충분히 파싱됐으면 재시도 중단, 부족하면 한 번 더 생성 시도
            if len(chapters) >= 5:
                break

        if chapters:
            # 비문 검수 패스: 멀쩡한 줄은 글자 하나 안 바꾸고 '깨진 줄만' 최소 수정 (저온 0.15 교열)
            # 생성 1회만으로는 비문이 확률적으로 새어 나와 전담 검수가 필요함
            try:
                progress_placeholder.info("🔍 목차 비문 검수 중... (깨진 문장만 최소 수정)")
                chapters, subtopics = polish_outline(chapters, subtopics, topic)
            except Exception:
                pass
            st.session_state['outline'] = chapters
            st.session_state['chapters'] = {}
            for ch in chapters:
                st.session_state['chapters'][ch] = {
                    'subtopics': subtopics.get(ch, []),
                    'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in subtopics.get(ch, [])}
                }

        # 목차가 생성되지 않았으면 기본 목차 생성
        if not st.session_state.get('outline'):
            progress_placeholder.warning("⚠️ AI 목차 생성/파싱에 실패해 기본 목차를 사용합니다. (API 키·모델 설정을 확인하거나 다시 시도해 주세요)")
            default_chapters = [
                "PART 1. 왜 지금인가",
                "PART 2. 진짜 비밀",
                "PART 3. 실전 공식",
                "PART 4. 수익화",
                "PART 5. 다음 단계"
            ]
            default_subtopics = {
                default_chapters[0]: [f"90%가 {topic}에 실패하는 이유", "아무도 말해주지 않는 진실", "지금 시작해야 하는 3가지 이유"],
                default_chapters[1]: ["전문가들이 숨기는 핵심 원칙", f"{topic}의 본질을 꿰뚫는 법", "이것만 알면 절반은 성공"],
                default_chapters[2]: ["바로 써먹는 5단계 공식", "실패 없이 시작하는 체크리스트", "첫 달에 결과 내는 비법"],
                default_chapters[3]: ["월 100만원 만드는 구조", "자동화로 시간 벌기", "확장 전략 A to Z"],
                default_chapters[4]: ["1년 후 당신의 모습", "다음 레벨로 가는 로드맵", "지금 바로 해야 할 첫 번째 행동"]
            }
            st.session_state['outline'] = default_chapters
            st.session_state['chapters'] = {}
            for ch in default_chapters:
                st.session_state['chapters'][ch] = {
                    'subtopics': default_subtopics.get(ch, []),
                    'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in default_subtopics.get(ch, [])}
                }

        # 저자명 및 인터뷰 데이터 저장
        st.session_state['author_name'] = interview_data.get('author_name', '')
        st.session_state['interview_data'] = interview_data
        st.session_state['topic'] = topic

        progress_placeholder.success("✅ 목차 생성 완료! 목차를 확인하고 수정할 수 있습니다.")
        return True

    except Exception as e:
        progress_placeholder.error(f"오류 발생: {str(e)}")
        return False

def regenerate_single_subtopic(chapter_name, subtopic_index, existing_subtopics):
    """개별 소제목 AI 재생성 - 자기계발 베스트셀러 톤"""
    topic = st.session_state.get('topic', '')
    book_concept = st.session_state.get('book_concept', '')

    # 기존 소제목들 (중복 방지용)
    other_subtopics = [s for i, s in enumerate(existing_subtopics) if i != subtopic_index]

    prompt = f"""당신은 한국 자기계발 베스트셀러 편집자입니다. 목차만 보고 결제하게 만드는 소제목 하나를 써주세요.

[책 컨셉]
{book_concept}

[챕터]: {chapter_name}
[주제]: {topic}

[기존 소제목들 - 이것들과 완전히 다르게]
{chr(10).join(f'- {s}' for s in other_subtopics)}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 한국 자기계발 베스트셀러 톤 예시 (참고만, 그대로 베끼지 말 것)
- "5분만에 돈 되는 사업 아이템 찾는 비법"
- "월급 280만원이 월 1,000만원으로 바뀌기까지 47일"
- "회사 몰래 부업하다 잘리는 직장인의 진짜 이유"
- "노력이 결과로 안 바뀌는 결정적 함정"
- "첫 정산 받고 인생이 달라진 그날의 기록"

✅ 형식 (하나 선택, 매번 다른 패턴):
- 비법/법형: "5분만에 ~하는 비법", "~하는 정확한 방법"
- 충격형: "~는 거짓말이다", "~하면 오히려 망한다"
- 간증형: "~받고 인생이 달라졌다", "~를 깨닫기까지"
- 도발형: "~은 필요 없다", "~만 있으면 된다"
- 질문형: "왜 ~은 실패하는가"
- 숫자형: "정확히 47일 만에 일어난 일"
- 호기심 갭형(가장 강력): "월 500을 찍은 사람들이 가장 먼저 끊은 습관 한 가지" (결과만 보이고 방법은 숨김)

🧲 호기심 갭 우선 — 결과/사건/디테일은 보이고, 원리/방법/순서는 숨기면 결제 전환률 폭증

⚠️ 컨셉명을 쓰는 경우, 글자 수 단정 표현 절대 금지 ("세 글자", "네 글자", "단 한 단어" 등 — 컨셉명 길이와 안 맞으면 즉시 폐기)

❌ 금지:
- 자청, 역행자, 유전자, 원시인, 추월차선 등 특정 작가 고유 표현
- 설명서 표현: "~의 이해", "~하는 방법"
- 유치한 비유: 나침반, 열쇠, 마법, 톱니바퀴, 파이프라인
- AI 어휘: 효과적인, 다양한, ~를 통해
- 기존 소제목과 비슷한 패턴
- 컨셉명 글자 수 단정 ("세 글자에서 시작한다" 류)

소제목 하나만 (15~30자, 기호 없이):"""

    result = ask_ai(prompt, 0.9)
    if result:
        return result.strip().strip('"').strip("'").strip('-').strip()
    return None

def regenerate_chapter_subtopics(chapter_name, chapter_index):
    """챕터의 모든 소제목 AI 재생성 - 자기계발 베스트셀러 톤"""
    topic = st.session_state.get('topic', '')
    book_concept = st.session_state.get('book_concept', '')
    outline = st.session_state.get('outline', [])

    # 다른 챕터들의 소제목들 (중복 방지)
    other_chapter_subtopics = []
    for ch in outline:
        if ch != chapter_name:
            ch_data = st.session_state['chapters'].get(ch, {})
            other_chapter_subtopics.extend(ch_data.get('subtopics', []))

    # 챕터별 역할 정의
    chapter_roles = {
        0: "착각/각성 - 독자가 몰랐던 불편한 진실을 폭로",
        1: "해체 - 기존 상식과 믿음을 완전히 무너뜨림",
        2: "구조/재구축 - 저자만의 새로운 방법론 제시",
        3: "실전 - 구체적이고 따라할 수 있는 방법",
        4: "도약 - 변화된 미래와 행동 촉구"
    }
    current_role = chapter_roles.get(chapter_index, "핵심 내용 전달")

    prompt = f"""당신은 한국 자기계발 베스트셀러 편집자입니다. 이 챕터의 소제목 4개를 결제하고 싶게 써주세요.

[책 컨셉]
{book_concept}

[주제]: {topic}
[챕터]: {chapter_name}
[이 챕터의 역할]: {current_role}

[다른 챕터 소제목들 - 완전히 다르게 써야 함]
{chr(10).join(f'- {s}' for s in other_chapter_subtopics[:8])}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 한국 자기계발 베스트셀러 톤 예시 (참고만, 그대로 베끼지 말 것)
- "5분만에 돈 되는 사업 아이템 찾는 비법"
- "월급 280만원이 월 1,000으로 바뀌기까지 47일"
- "회사 몰래 부업하다 잘리는 직장인의 진짜 이유"
- "첫 정산 받고 인생이 달라진 그날의 기록"
- "노력이 결과로 안 바뀌는 결정적 함정"
- "통장 잔고 23만원에서 시작한 한 가지 시도"

✅ 4개 소제목 모두 다른 형식으로 (4개 다 같은 패턴 금지):
1번: 통계 충격/숫자형 (예: "정확히 47일째에 달라진 한 가지", "1년 안에 99%가 다시 무너지는 결정적 이유")
2번: 간증/사건형 (예: "27만원짜리 첫 정산서가 알려준 한 가지", "~받고 인생이 달라진 그날")
3번: 충격/통념 박살형 (예: "노력할수록 가난해진다", "~만 있으면 된다")
4번: 호기심 갭형 (결과만 보이고 방법은 숨김 — 예: "월 500을 찍은 사람들이 가장 먼저 끊은 습관 한 가지", "성공한 부업러가 매일 밤 11시에 반드시 끄는 것")

🧲 호기심 갭 규칙 (반드시 4개 중 1개 이상 포함)
   - 결과/사건/디테일은 보이고, 원리/방법/순서는 숨겨라
   - "방법", "비법", "노하우"로 끝내면 갭이 닫혀 결제 안 함
   - 한 줄에 "묘하게 구체적인 디테일 + 의문"이 같이 있어야 작동

⚠️ 컨셉명을 쓰는 경우, 글자 수 단정 표현 절대 금지 ("세 글자", "네 글자", "단 한 단어" 등)

❌ 금지:
- 자청, 역행자, 유전자, 원시인, 추월차선 등 특정 작가 고유 표현
- 설명서 표현: "~의 이해", "~하는 방법", "효과적인", "다양한"
- 유치한 비유: 나침반, 열쇠, 마법, 톱니바퀴, 파이프라인
- 같은 패턴 반복

소제목 정확히 4개만 출력 (줄바꿈으로 구분, 기호/번호 없이, 각 줄 15~30자):"""

    result = ask_ai(prompt, 0.8)
    if result:
        lines = [line.strip().strip('"').strip("'").strip('-').strip() for line in result.strip().split('\n') if line.strip() and len(line.strip()) > 5]
        return lines[:4] if lines else None
    return None

def generate_body_from_outline(interview_data, progress_placeholder):
    """생성된 목차를 기반으로 본문만 생성"""
    try:
        topic = interview_data.get('topic', '')
        book_concept = st.session_state.get('book_concept', '')

        if not st.session_state.get('outline') or not st.session_state.get('chapters'):
            progress_placeholder.error("먼저 목차를 생성해주세요.")
            return False

        # API 사전 점검 — 키/패키지/연결 문제를 시작 전에 표면화
        if not get_api_key():
            progress_placeholder.error("❌ Claude API 키가 없습니다. 사이드바에서 API 키를 입력해주세요.")
            return False
        progress_placeholder.info("🔌 Claude API 연결 확인 중...")
        if ask_ai("OK라고 한 단어로만 답하세요.", 0.0) is None:
            progress_placeholder.error("❌ Claude API 호출 실패. API 키/크레딧/모델 설정을 확인해주세요. (위에 표시된 오류 메시지 참고)")
            return False

        # 본문 생성 — 병렬 처리 (동시 4개, 순차 대비 약 4배 빠름)
        # 중복 방지: 이전 '본문 요약' 대신 '전체 소제목 목록'을 컨텍스트로 사용 → 병렬화 가능
        outline_list = list(st.session_state['outline'])
        chapters_map = st.session_state['chapters']
        total_subtopics = sum(len(chapters_map[ch]['subtopics']) for ch in outline_list)

        all_pairs = []
        for ch in outline_list:
            ch_data = chapters_map[ch]
            if 'subtopic_data' not in ch_data:
                ch_data['subtopic_data'] = {}
            for sub in ch_data['subtopics']:
                if sub not in ch_data['subtopic_data']:
                    ch_data['subtopic_data'][sub] = {'questions': [], 'answers': [], 'content': ''}
                all_pairs.append((ch, sub))

        # 이미 생성된 본문은 건너뛰기 (실패/중단 후 재시도 시 이어쓰기)
        pending = [(i, c, s) for i, (c, s) in enumerate(all_pairs)
                   if not chapters_map[c]['subtopic_data'][s].get('content')]
        done_count = total_subtopics - len(pending)
        failed = []

        if pending:
            hook_styles = [
                "도발적 질문 (예: '왜 99%는 이걸 모를까요?')",
                "충격적 고백 (예: '저도 3년간 완전히 잘못하고 있었습니다.')",
                "반전 사실 (예: '사실 정반대였습니다.')",
                "구체적 숫자 (예: '정확히 47일 만에 달라졌습니다.')",
                "생생한 에피소드 (예: '그날 카페에서 노트북을 열었을 때였습니다.')",
                "단호한 선언 (예: '결론부터 말씀드리겠습니다.')",
                "대화체 시작 (예: '\"이게 진짜 돼요?\" 처음 들었을 때 저도 그랬습니다.')",
                "before/after (예: '6개월 전만 해도 저는 완전히 다른 사람이었습니다.')",
                "상식 뒤집기 (예: '노력하면 된다? 완전히 틀렸습니다.')",
                "비유로 시작 (예: '이건 마치 고장난 네비게이션을 따라가는 것과 같습니다.')",
                "독자 공감 (예: '혹시 이런 경험 있으신가요?')",
                "미래 제시 (예: '3개월 후, 완전히 다른 결과를 보게 될 겁니다.')",
                "실패담 (예: '처음엔 완전히 망했습니다.')",
                "발견의 순간 (예: '그때 깨달았습니다. 방법이 틀렸던 거였죠.')",
                "핵심 한 줄 (예: '핵심은 딱 하나입니다.')",
            ]
            author_name = interview_data.get('author_name', '') or st.session_state.get('author_name', '') or '저자'
            api_key = get_api_key()
            user_model = st.session_state.get('claude_model', 'claude-sonnet-4-5')
            model = 'claude-sonnet-4-5' if 'haiku' in user_model.lower() else user_model

            def build_content_prompt(ch, sub, idx):
                other_subs = "\n".join(f"- {s}" for _c, s in all_pairs if s != sub)
                current_hook_style = hook_styles[(idx + 1) % len(hook_styles)]
                # 약 1/3 소제목에만 표 포함 (결정적 분배 — 책마다 적절히 섞임)
                want_table = (idx % 3 == 1)
                table_block = ("""
[표 - 이 글에는 필수]
본문 흐름상 자연스러운 위치에 마크다운 표 1개를 반드시 포함:
- 비교(기존 vs 새로운), 단계별 정리, 수치 데이터 중 내용에 맞는 것
- 형식 (표 앞뒤로 반드시 빈 줄):

| 구분 | 기존 방식 | 새로운 방식 |
| 시간 | 하루 3시간 | 하루 40분 |
| 결과 | 월 50만원 | 월 340만원 |

- 셀은 15자 이내로 간결하게, 3~5행
""" if want_table else """
[표 - 이 글에는 넣지 말 것]
이 소제목은 순수 텍스트로만 작성 (표는 다른 장에 들어감)
""")
                return f"""당신은 한국 자기계발 베스트셀러 작가입니다. 첫 문장으로 독자를 멈춰 세우고, 권석천 기자 칼럼처럼 정확한 디테일과 통찰로 끝까지 끌고 가는 본문을 씁니다.

🚨🚨🚨 최우선 규칙 🚨🚨🚨
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. 첫 문장이 가장 중요! 반드시 이 스타일로 시작:
   👉 {current_hook_style}

2. 이 책의 다른 소제목들과 절대 중복 금지!
   아래 소제목들이 각자 다룰 주제이므로, 이 글('{sub}')에서는 건드리지 마라.
   같은 사례·같은 숫자·같은 결론 반복도 금지:
{other_subs}

3. 독자 직접 호칭 금지
   ❌ "여러분", "당신", "독자님", "~하시는 분들"
   ✅ "저는", "우리는", "제가"

4. 특정 작가 고유 표현 절대 금지 (저작권/표절 위험)
   ❌ "자청", "역행자", "유전자 역행", "자의식 해체", "원시인", "추월차선", "아토믹 해빗"
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[집필 정보]
주제: {topic}
챕터: {ch}
소제목: {sub}
핵심 방법론: {interview_data.get('core_method', '')}
저자: {author_name}

[책 컨셉]
{book_concept}
⚠️ 위 컨셉/방법론은 저자가 직접 개발해 이름 붙인 것이다. "처음 들었을 때", "배웠을 때", "알게 됐을 때"처럼
남에게 전수받은 듯이 쓰지 마라. "발견했다", "정리하다 보니 만들어졌다", "이름을 붙였다"로 써라.

[문체]
- 합쇼체 기본 ("~입니다", "~합니다") + 가끔 구어체 ("~거든요", "~더라고요")
- ⚠️ 모든 문장은 존댓말로 종결. 반말(~했다, ~이다, ~한다) 혼용 절대 금지
- 현재 시제로 장면을 그리듯, 짧은 문장과 긴 문장 교차
- 추상보다 구체: "많은 사람" → "월급 280만원짜리 7년차 회사원"
- 결론부터, 분석은 그 다음. 사실 → 분석 → 통찰 순서로 인과 추적
- 본문 중간에 작은 반전 1회, 마지막 문장은 발견의 결과로서의 통찰 한 줄

[문단 - 가독성 필수 규칙]
- 한 문단은 2~3문장만. 절대 4문장을 넘기지 마라
- 문단과 문단 사이는 반드시 빈 줄로 구분
- 긴 설명은 여러 개의 짧은 문단으로 쪼갤 것

[쉽게 쓰기 - 전문성과 이해도를 동시에 (필수)]
- 내용은 논문 수준으로 정확하게, 문장은 중학생이 한 번에 이해하게
- 어려운 개념은 일상 비유 1개로 풀어라 (월급, 장보기, 통장, 운동 같은 생활 소재.
  단, 마법·황금열쇠 같은 유치한 비유와 층·탑·계단을 "쌓는" 건축 비유는 금지)
- 전문 용어는 처음 등장할 때 바로 한 줄로 뜻풀이 (예: "배당락일, 그러니까 이날 이후에 사면 이번 배당을 못 받는 날짜입니다")
- 비문 금지: 모든 문장을 소리 내어 읽었을 때 자연스러운 한국어여야 한다
  ❌ "1장에서 다룰 이야기는 종목도 차트도 아니었습니다. 통장을 바라보는 그 시선부터였습니다" — 꼬인 구문, 무슨 말인지 모름
  ✅ "1장에서는 종목 이야기를 꺼내지 않습니다. 그 전에 바로잡아야 할 게 있거든요. 통장 잔고를 바라보는 방식입니다"
- 한 문장에 추상 명사 2개 이상 금지 ("흐름의 적층이 만드는 구조" 같은 문장 즉시 폐기)

[강조 표시 - 필수]
- 독자가 밑줄 칠 핵심 문장·구절 2~3곳을 **굵게** 표시 (마크다운 ** 사용)
- 남발 금지: 문단당 최대 1곳. 통찰·반전·핵심 숫자에만
- 이 글의 단 하나의 핵심 통찰을 골라, 줄 시작에 ★ 를 붙여 단독 문단으로 1회 배치
  (책에서 괘선 인용구로 디자인됨. 짧고 강한 한 문장. 예: ★ 시간은 만드는 것이 아니라 발견하는 것입니다)
{table_block}
[구체성]
✅ 숫자: "많이" X → "월 340만원, 정확히 47일" O
✅ 도구/플랫폼 실명: 네이버, 카카오, 노션, 카톡, 구글 시트 등
✅ 실행 순서: "무엇을 → 어디서 → 어떻게" 명시

🚫 절대 금지
❌ 같은 이름 반복(민준, 지수 등), "김씨" 같은 성씨 호칭
❌ 유치한 표현: 후다닥, 짜잔, 대박, ㅋㅋ / 유치한 비유: 마법, 황금열쇠, 나침반, 로켓, 눈덩이
❌ 억지 메타포: 순환법, 엔진, 고리, 파이프라인, 톱니바퀴
❌ AI스러운: 중요합니다, 따라서, 결론적으로, ~를 통해, 다양한, 효과적인, 진정한
❌ 형식: 1. 2. 첫째, 둘째, 글머리 기호, 이모지, HTML 태그 (단, **굵게**·★·지시된 마크다운 표는 허용)
❌ 본문 끝 '핵심 정리'/'요약'/'마무리' 같은 제목·섹션 금지 — 파트 끝에 핵심 정리 박스가 자동으로 들어가니 절대 직접 쓰지 마라. 마지막도 일반 문단으로 끝낼 것

📏 분량: 1800~2200자

'{sub}' 본문 작성.
⛔ 본문 첫 줄에 소제목('{sub}')을 다시 쓰지 마라. 곧장 첫 후킹 문장으로 시작한다."""

            from concurrent.futures import ThreadPoolExecutor, as_completed
            progress_placeholder.info(f"✍️ 본문 {len(pending)}개 동시 작성 시작... (병렬 생성, 약 3~5분)")

            with ThreadPoolExecutor(max_workers=4) as _ex:
                _futures = {
                    _ex.submit(claude_call_threadsafe, build_content_prompt(c, s, i), api_key, model, 0.7): (c, s)
                    for i, c, s in pending
                }
                for _fut in as_completed(_futures):
                    _c, _s = _futures[_fut]
                    try:
                        _content, _err = _fut.result()
                    except Exception as _e:
                        _content, _err = None, str(_e)[:80]
                    if _content:
                        chapters_map[_c]['subtopic_data'][_s]['content'] = clean_content(_content, subtopic=_s)
                        done_count += 1
                        progress_placeholder.info(f"✍️ 본문 작성 중... ({done_count}/{total_subtopics}) 완료 — {_s[:18]}")
                    else:
                        failed.append(_s)

        # ── 파트별 핵심 정리 생성 (병렬, 본문 완료된 파트만) ──
        try:
            _api_key2 = get_api_key()
            _um2 = st.session_state.get('claude_model', 'claude-sonnet-4-5')
            _model2 = 'claude-sonnet-4-5' if 'haiku' in _um2.lower() else _um2
            _sum_jobs = []
            for _ch in outline_list:
                _cd = chapters_map[_ch]
                _subs = _cd.get('subtopics', [])
                if not _subs or _cd.get('part_summary'):
                    continue
                if any(not _cd.get('subtopic_data', {}).get(_s, {}).get('content') for _s in _subs):
                    continue
                _digest = "\n\n".join(f"[{_s}]\n" + _cd['subtopic_data'][_s]['content'][:350] for _s in _subs)
                _sum_prompt = f"""아래는 전자책 한 파트의 소제목과 본문 도입부다. 이 파트를 마무리하는 '핵심 정리'를 써라.

{_digest}

[규칙]
- 정확히 4줄 + 마지막에 실행 과제 1줄 (총 5줄)
- 각 줄 15~30자, 평어체 단언("~한다", "~이다")
- 번호/불릿/이모지 없이 줄바꿈으로만 구분
- 실행 과제 줄은 "오늘 할 일:"로 시작
- 본문에 없는 내용을 지어내지 말 것"""
                _sum_jobs.append((_ch, _sum_prompt))
            if _sum_jobs:
                from concurrent.futures import ThreadPoolExecutor as _TPE2, as_completed as _ac2
                progress_placeholder.info("📌 파트별 핵심 정리 작성 중...")
                with _TPE2(max_workers=5) as _ex2:
                    _fut2 = {_ex2.submit(claude_call_threadsafe, _p, _api_key2, _model2, 0.5, 1200): _c for _c, _p in _sum_jobs}
                    for _f2 in _ac2(_fut2):
                        _c2 = _fut2[_f2]
                        try:
                            _txt2, _e2 = _f2.result()
                        except Exception:
                            _txt2 = None
                        if _txt2:
                            chapters_map[_c2]['part_summary'] = [l.strip() for l in _txt2.strip().split('\n') if l.strip()][:5]
        except Exception:
            pass  # 요약 실패는 본문 완성에 영향 없음 (박스는 소제목 리스트로 대체됨)

        # 완료 처리 — 실패가 있으면 성공으로 위장하지 않음
        if failed:
            names = ', '.join(failed[:5]) + ('...' if len(failed) > 5 else '')
            progress_placeholder.error(
                f"⚠️ {len(failed)}개 소제목 본문 생성 실패: {names}\n\n"
                f"이미 완성된 본문은 저장되어 있습니다. '본문 생성하기'를 다시 누르면 실패한 부분만 재생성합니다."
            )
            return False

        st.session_state['interview_completed'] = True
        progress_placeholder.success("✅ 본문 생성 완료!")
        return True

    except Exception as e:
        progress_placeholder.error(f"오류 발생: {str(e)}")
        return False

def go_next():
    if st.session_state['current_page'] < 6:
        st.session_state['current_page'] += 1

def go_prev():
    if st.session_state['current_page'] > 0:
        st.session_state['current_page'] -= 1

def auto_generate_all(topic, progress_placeholder):
    """주제만 입력하면 목차+본문까지 자동 생성"""
    try:
        # 1. 타겟 자동 생성
        progress_placeholder.info("🎯 1/4 타겟 분석 중...")
        target_result = suggest_targets(topic)
        targets = parse_json(target_result)
        if targets and targets.get('targets'):
            first_target = targets['targets'][0]
            persona = f"{first_target.get('name', '')} - {first_target.get('description', '')}"
            st.session_state['target_persona'] = persona

            # 페인포인트 분석
            pain_result = analyze_pains_deep(topic, persona)
            pain_data = parse_json(pain_result)
            if pain_data:
                st.session_state['pains'] = pain_data.get('pains', [])

        # 2. 목차 자동 생성
        progress_placeholder.info("📋 2/4 목차 생성 중...")
        outline_result = generate_outline(
            topic,
            st.session_state.get('target_persona', ''),
            st.session_state.get('pains', [])
        )

        # 목차 텍스트 파싱 (PAGE 4와 동일한 방식)
        if outline_result:
            chapters = []
            subtopics = {}
            current_ch = None

            for line in outline_result.split('\n'):
                orig_line = line
                line = line.strip()
                if not line:
                    continue

                # 챕터 감지 (PART, 1., 2. 등)
                is_chapter = False
                ch_name = None

                if re.match(r'^(PART|파트)\s*\d+', line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = line
                elif re.match(r'^\d+[\.\)]\s', line):
                    is_chapter = True
                    ch_name = line
                elif re.match(r'^#+\s*(PART|파트|\d+)', line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = re.sub(r'^#+\s*', '', line)

                if is_chapter and ch_name:
                    ch_name = ch_name.replace('**', '').replace('*', '').replace('#', '').strip()
                    if ch_name and len(ch_name) > 3:
                        current_ch = ch_name
                        if current_ch not in chapters:
                            chapters.append(current_ch)
                            subtopics[current_ch] = []
                elif current_ch:
                    # 소제목 감지
                    is_subtopic = False
                    st_name = None

                    if line.startswith('-') or line.startswith('•') or line.startswith('·'):
                        is_subtopic = True
                        st_name = line.strip().lstrip('-•· ')
                    elif re.match(r'^\s+[\da-z][\)\.]', orig_line):
                        is_subtopic = True
                        st_name = re.sub(r'^[\s\da-z\)\.\-]+', '', line).strip()

                    if is_subtopic and st_name:
                        st_name = st_name.replace('**', '').replace('*', '').replace('#', '').strip()
                        if st_name and len(st_name) > 3 and not re.match(r'^(PART|파트|Chapter|챕터)', st_name, re.IGNORECASE):
                            subtopics[current_ch].append(st_name)

            if chapters:
                st.session_state['outline'] = chapters
                st.session_state['chapters'] = {}
                for ch in chapters:
                    st.session_state['chapters'][ch] = {
                        'subtopics': subtopics.get(ch, []),
                        'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in subtopics.get(ch, [])}
                    }

        # 3. 본문 자동 생성
        progress_placeholder.info("✍️ 3/4 본문 작성 중...")
        if st.session_state.get('outline') and st.session_state.get('chapters'):
            total_subtopics = sum(len(st.session_state['chapters'][ch]['subtopics']) for ch in st.session_state['outline'])
            done = 0

            for ch in st.session_state['outline']:
                ch_data = st.session_state['chapters'][ch]
                for sub in ch_data['subtopics']:
                    done += 1
                    progress_placeholder.info(f"✍️ 본문 작성 중... ({done}/{total_subtopics})")

                    content = generate_content_premium(sub, ch, [], [], topic, st.session_state.get('target_persona', ''))
                    if content:
                        ch_data['subtopic_data'][sub]['content'] = content
                        ch_data['subtopic_data'][sub]['formatted'] = format_content_html(content)

        # 4. 완료
        progress_placeholder.success("✅ 완료! 본문 페이지로 이동합니다...")
        return True

    except Exception as e:
        progress_placeholder.error(f"오류 발생: {str(e)}")
        return False


# ==========================================
# AI 함수들
# ==========================================
def analyze_market_deep(topic):
    prompt = f"""주제: {topic}

이 주제로 전자책 시장을 분석해주세요.

[중요] 모든 답변은 반드시 한국어로만 작성하세요.

JSON:
{{
    "verdict": "강력 추천/추천/보류/비추천 중 하나",
    "verdict_reason": "판정 이유 한국어로",
    "total_score": 85,
    "search_data": {{
        "naver_monthly": "네이버 월간 검색량 예시: 12,000회",
        "google_monthly": "구글 월간 검색량 예시: 8,500회",
        "naver_blog_posts": "블로그 게시물 수",
        "youtube_videos": "유튜브 영상 수",
        "search_trend": "상승 또는 유지 또는 하락"
    }},
    "market_size": {{
        "score": 85,
        "level": "매우 큼/큼/보통/작음 중 하나",
        "analysis": "분석 2문장 한국어로"
    }},
    "competition": {{
        "score": 70,
        "level": "치열함/보통/낮음 중 하나",
        "your_opportunity": "차별화 기회 한국어로"
    }},
    "profit": {{
        "score": 80,
        "price_range": "권장 가격대",
        "monthly_revenue": "예상 월 수익"
    }},
    "popular_ebooks": [
        {{
            "title": "이 주제 관련 인기 전자책 제목",
            "platform": "크몽/탈잉/클래스101/리디북스/yes24 중 하나",
            "url": "해당 전자책 실제 URL (예: https://kmong.com/xxx)",
            "price": "가격"
        }},
        {{
            "title": "두번째 인기 전자책",
            "platform": "플랫폼명",
            "url": "URL",
            "price": "가격"
        }},
        {{
            "title": "세번째 인기 전자책",
            "platform": "플랫폼명",
            "url": "URL",
            "price": "가격"
        }}
    ],
    "recommendation": "최종 권장 2문장 한국어로"
}}"""
    return ask_ai(prompt, 0.5)


def suggest_targets(topic):
    prompt = f"""주제: {topic}

이 주제의 전자책을 가장 많이 구매할 것 같은 핵심 타겟 3개만 추천해주세요.
가장 적합하고 구매 가능성이 높은 타겟만 엄선해서 3개만 알려주세요.

[중요] 모든 답변은 반드시 한국어로만 작성하세요.

JSON:
{{
    "personas": [
        {{
            "name": "타겟 이름 (구체적으로)",
            "demographics": "연령대, 직업",
            "needs": "이 타겟이 이 책을 사는 이유",
            "pain_points": ["핵심 고민1", "고민2", "고민3", "고민4", "고민5"]
        }}
    ]
}}"""
    return ask_ai(prompt, 0.7)


def analyze_pains_deep(topic, persona):
    prompt = f"""주제: {topic}
타겟: {persona}

이 타겟의 고민을 아주 깊이 분석해주세요.

[중요] 모든 답변은 반드시 한국어로만 작성하세요. 외국어 사용 금지.

JSON:
{{
    "surface_pains": {{
        "pains": ["표면적 고민1", "고민2", "고민3", "고민4", "고민5"],
        "description": "표면적 고민 설명 3문장"
    }},
    "hidden_pains": {{
        "pains": ["숨겨진 진짜 고민1", "고민2", "고민3", "고민4"],
        "description": "숨겨진 고민 설명 3문장"
    }},
    "emotional_pains": {{
        "pains": ["감정적 고통1", "고통2", "고통3"],
        "description": "감정적 고통 설명 2문장"
    }},
    "failed_attempts": {{
        "attempts": ["시도했지만 실패한 것1", "것2", "것3"],
        "why_failed": "실패 이유 2문장"
    }},
    "dream_outcome": {{
        "ideal_result": "이상적인 결과",
        "timeline": "원하는 기간",
        "what_changes": "달라지는 것 2문장"
    }},
    "buying_triggers": {{
        "triggers": ["구매 요인1", "요인2", "요인3"],
        "objections": ["망설임 이유1", "이유2"]
    }},
    "marketing_hook": "마케팅 훅 한 문장"
}}"""
    return ask_ai(prompt, 0.6)


def analyze_competitor_reviews(topic):
    prompt = f"""주제: {topic}

이 주제 관련 전자책/도서의 부정적 리뷰를 분석해주세요.

[매우 중요]
- 모든 답변은 반드시 한국어로만 작성하세요.
- 영어, 러시아어 등 외국어 절대 사용 금지
- 한글과 숫자만 사용하세요.

JSON:
{{
    "analysis_scope": {{
        "books_analyzed": "287권",
        "reviews_analyzed": "3,842개",
        "negative_reviews": "892개 (23%)",
        "platforms": ["크몽", "예스24", "알라딘", "교보문고"]
    }},
    "negative_patterns": [
        {{
            "pattern": "불만 패턴 한국어로",
            "frequency": "67%",
            "example_reviews": ["실제 리뷰 예시 한국어로", "리뷰2"],
            "reader_emotion": "독자 감정 한국어로",
            "hidden_need": "숨겨진 니즈 한국어로",
            "solution": "해결책 한국어로"
        }},
        {{
            "pattern": "두 번째 불만",
            "frequency": "54%",
            "example_reviews": ["리뷰1", "리뷰2"],
            "reader_emotion": "감정",
            "hidden_need": "니즈",
            "solution": "해결책"
        }},
        {{
            "pattern": "세 번째 불만",
            "frequency": "41%",
            "example_reviews": ["리뷰1", "리뷰2"],
            "reader_emotion": "감정",
            "hidden_need": "니즈",
            "solution": "해결책"
        }}
    ],
    "hidden_needs_summary": {{
        "needs": ["숨겨진 니즈1", "니즈2", "니즈3"],
        "insight": "핵심 인사이트 2문장"
    }},
    "concept_suggestions": [
        {{
            "concept": "차별화 컨셉1 한국어로",
            "why_works": "이유 한국어로",
            "unique_point": "차별점 한국어로"
        }},
        {{
            "concept": "컨셉2",
            "why_works": "이유",
            "unique_point": "차별점"
        }}
    ],
    "success_formula": {{
        "must_have": ["필수1", "필수2", "필수3"],
        "must_avoid": ["금지1", "금지2"],
        "differentiation": "차별화 전략 한국어로 2문장"
    }}
}}"""
    return ask_ai(prompt, 0.6)


def generate_titles_bestseller(topic, persona, pains):
    prompt = f"""당신은 교보문고 종이책 베스트셀러와 크몽·클래스101 전자책 베스트셀러를 동시에 분석하는 제목 카피라이터입니다.

주제: {topic}
독자: {persona}
독자의 고민: {pains}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
가장 중요한 규칙 (이걸 어기면 전부 실패)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
제목은 반드시 "실제로 말이 되는 자연스러운 한국어 한 구절"이어야 합니다.
주제에서 뽑은 단어들을 기계적으로 이어 붙이지 마세요.

❌ 단어 짜깁기 실패 (절대 이렇게 만들지 말 것):
  "수면 매출 설계도"  → 수면+매출+설계도, 서로 관계없는 명사를 그냥 붙인 말
  "다이어트 부자 공식"  → 의미가 안 통하는 조합
  "관계 성장 엔진"      → 추상명사만 나열
  이런 제목은 소리내어 읽으면 "이게 무슨 말이지?" 싶고, 어떤 서점에도 존재하지 않습니다.

✅ 자연스러운 제목 = 사람이 실제로 쓰는 어순과 의미가 살아있는 구절
  (수면 주제 예) "잠든 사이 일어나는 일" / "초저녁의 기술" / "다시 잠드는 법"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
두 시장의 제목 문법
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
[종이책 베스트셀러 — 품격·여운]
  돈의 속성 / 불변의 법칙 / 마흔에 읽는 쇼펜하우어 / 역행자 / 세이노의 가르침
  → 명사+의+명사, 한 단어 임팩트, 약간의 문학성. 광고 냄새가 없다.

[전자책 베스트셀러 — 구체적 약속·호기심]
  크몽·클래스101 상위권은 "독자가 무엇을 얻는지"가 살짝 드러난다.
  단, 유치하지 않게. "~하는 법/방법/노하우/비법/공식" 같은 흔한 꼬리표는 피하고,
  의외의 단어나 시선 전환을 한 번 넣어 호기심을 만든다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
피해야 할 것
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- 의미 없는 명사 나열 (가장 큰 실패 원인)
- 주제 키워드 두 개를 그대로 붙인 평이한 조합 ("전자책 복리", "블로그 부자") — 검색어지 제목이 아니다
- 비밀/비법/공식/바이블/마법/머니/시스템/파이프라인/연금술
- 완벽한·궁극의·최고의·기적의 같은 과장 형용사
- "월 1000만원" 류 숫자 과시, "직장인을 위한" 류 타겟 명시

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
출력
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
'{topic}' 주제로 서로 다른 결의 제목을 정확히 5개.
(3개는 종이책 품격 톤, 2개는 전자책 호기심 톤)
길이는 자유 — 보통 2~7어절. 짧게 만드는 것보다 자연스러움이 우선이다.

각 제목마다 self_check를 채워, 소리내어 읽었을 때 말이 되는지 스스로 검증할 것.
self_check가 "어색하다/말이 안 된다"면 그 제목은 버리고 다시 만들 것.

JSON만 출력:
{{
    "titles": [
        {{"title": "제목", "subtitle": "호기심을 더하는 부제 한 줄 (20자 이내)", "concept": "이 제목이 매력적인 이유 한 줄", "self_check": "소리내어 읽으면 자연스러운가에 대한 한 문장 자기검증"}}
    ]
}}"""
    return ask_ai(prompt, 0.6, ensure_quality=True)


def analyze_text_content(text, source=""):
    prompt = f"""출처: {source}
내용: {text[:5000]}

분석:

JSON:
{{
    "title": "주제",
    "key_points": ["핵심1", "핵심2", "핵심3", "핵심4", "핵심5"],
    "insights": ["인사이트1", "인사이트2", "인사이트3"],
    "action_items": ["실행1", "실행2", "실행3"],
    "ebook_ideas": ["아이디어1", "아이디어2"],
    "summary": "요약 3문장"
}}"""
    return ask_ai(prompt, 0.5)


def summarize_all_knowledge(items, topic):
    """전체 학습 내용 통합 요약"""
    all_points = []
    all_tips = []
    all_ideas = []

    for item in items:
        if isinstance(item, dict):
            all_points.extend(item.get('key_points', []))
            all_tips.extend(item.get('actionable_tips', item.get('action_items', [])))
            all_ideas.extend(item.get('ebook_applications', item.get('ebook_ideas', [])))

    prompt = f"""전자책 주제: {topic}

학습한 모든 정보를 통합 분석해주세요.

수집된 핵심 포인트들:
{chr(10).join([f"• {p}" for p in all_points[:25]])}

실행 팁들:
{chr(10).join([f"• {t}" for t in all_tips[:15]])}

전자책 활용 아이디어:
{chr(10).join([f"• {i}" for i in all_ideas[:10]])}

JSON:
{{
    "integrated_summary": "전체 학습 내용 통합 요약 5문장",
    "core_insights": [
        "핵심 인사이트 1",
        "인사이트 2",
        "인사이트 3",
        "인사이트 4",
        "인사이트 5"
    ],
    "action_plan": [
        "즉시 실행할 것 1",
        "실행 2",
        "실행 3"
    ],
    "ebook_structure": [
        "추천 목차 1장",
        "2장",
        "3장",
        "4장"
    ],
    "unique_angle": "이 전자책만의 차별화된 관점",
    "study_plan": {{
        "week1": "1주차: 무엇을 할지",
        "week2": "2주차: 무엇을 할지",
        "week3": "3주차: 무엇을 할지",
        "week4": "4주차: 무엇을 할지"
    }},
    "expert_tips": [
        "전문가 팁 1",
        "팁 2",
        "팁 3"
    ]
}}"""
    return ask_ai(prompt, 0.6)


def generate_outline(topic, persona, pains, gaps=None):
    """한국 자기계발 베스트셀러 톤: 결제 버튼을 누르게 하는 자극형 목차"""

    # 페르소나/고통/시장 빈틈 정리 (있으면 활용, 없어도 작동)
    persona_block = f"[타겟 독자]\n{persona}\n" if persona else ""
    if isinstance(pains, list):
        pains_text = "\n".join(f"- {p}" for p in pains if p)
    else:
        pains_text = str(pains) if pains else ""
    pains_block = f"[독자가 지금 느끼는 통증]\n{pains_text}\n" if pains_text else ""
    if gaps:
        gaps_text = "\n".join(f"- {g}" for g in gaps) if isinstance(gaps, list) else str(gaps)
        gaps_block = f"[시장의 빈틈 - 경쟁자가 안 다루는 것]\n{gaps_text}\n"
    else:
        gaps_block = ""

    prompt = f"""당신은 한국 자기계발 분야 톱 0.1% 기획자입니다. 서점에서 단 5초간 목차만 본 사람이 책을 손에서 못 놓게 만드는 5장짜리 목차를 씁니다.

⚠️ 예시 복제 금지 — 전역 규칙 (다른 모든 규칙보다 우선)
이 지시문의 모든 예시 문장(✅ 좋은 예 포함)은 구조 견본일 뿐이다. 예시를 그대로 또는
단어 한두 개만 바꿔 쓰면 전체 무효. 출력 전 모든 줄을 예시와 대조해 비슷하면 새로 써라.

목차의 단 하나의 목적: 독자가 "이 책을 안 읽으면 평생 손해"라고 느끼게 만드는 것.
정보 전달은 본문이 한다. 목차는 100% 구매심리만 다룬다.

[주제]: {topic}
{persona_block}{pains_block}{gaps_block}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🛒 구매 결정 5초 룰 (모든 규칙 중 1순위)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

독자는 목차를 5초만 본다. 그 5초 안에 다음 3가지가 동시에 작동해야 결제한다.

[1] 정체성 변화 발견 — "이걸 읽으면 나는 OO한 사람이 된다"
   → 5개 챕터 제목을 이어 읽으면 한 사람의 결정적 변화가 보여야 한다.
   ✅ "단단해진 멘탈은 인생을 통째로 바꾼다" (변화 서사 O)
   ❌ "멘탈 관리의 다양한 기법" (정체성 변화 X — 즉시 폐기)

[2] 손실회피 작동 — "이걸 모르면 평생 OO한다"
   → 통념 박살(인지부조화) 챕터/소제목이 최소 3개 들어가야 한다.
   ✅ "의지로 버틴 사람일수록 더 크게 무너진다"
   ❌ "멘탈 관리의 중요성" (잃을 게 안 보임)

[3] 구체성 — 추상 명사 1개당 구체적 숫자/장면 1개
   → "많은 사람" 금지, "월급 280만원짜리 7년차 회사원" 가능
   → 시간(90초, 47일), 금액(34만원, 1억), 비율(99%, 8할) 적극 사용

25줄(챕터 5 + 소제목 20) 중 한 줄이라도 "그냥 정보"가 섞이면 그 목차는 평이해진다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🧠 마케팅 뇌과학 8대 트리거 (목차 전체에 골고루 박을 것)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. **손해회피 + 충격 통계** — "1년 안에 99%가 다시 무너진다", "1억치 강의 들어도 망한다"
2. **인지부조화/통념 박살** — "의지로 버틴 사람일수록 더 크게 무너진다", "노력할수록 가난해진다"
3. **권위 어휘 (과학/임상)** — 뇌, 신경회로, N주 후, 임상, 데이터, 알고리즘
4. **임박감 + 절대성** — "이 90초를 놓치면 며칠 걸린다", "다시는 ~하지 않는다"
5. **정체성 전환 약속** — "회복한 뇌는 다시 무너지지 않는다", "단단해진 사람은 ~한다"
6. **인그룹 사회증명** — "○○를 익힌 사람들의 5년 뒤", "상위 1%만 도달하는"
7. **이중/삼중 보상** — "통장과 인간관계가 함께 변한다", "돈도 사람도 따라온다"
8. **호기심 갭** — "두 달 안에 가장 먼저 끊은 한 가지", "정확히 어디부터 멈추는가"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🧲 호기심 갭(Information Gap) 강화 — 결제를 부르는 가장 강력한 무기
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

조지 로웬스타인의 정보격차 이론: 사람은 자기가 모르는 것이 '뭔지'는 알지만 '내용'은 모를 때 가장 강하게 끌린다. 25줄 중 최소 14줄에 이 갭을 박아야 목차만 보고 결제한다. 갭이 없는 줄은 '그냥 정보'다.

[호기심 갭 6대 공식 - 결과만 보이고 방법/이유는 본문으로 미루기]

1. **숫자 + 미공개 결과** — '정확히/딱'으로 시작해 결과만 보여주고 내용은 숨기기
   ✅ "정확히 47일째에 통장이 처음 뒤집힌 그 순간"
   ✅ "27만원짜리 첫 정산서가 알려준 단 한 가지"
   ❌ "47일 만에 돈을 버는 방법" (방법을 다 보여줘버림 → 결제 안 함)

2. **이미 벌어진 사건 + 원인 숨김** — '왜?'를 유발하는 결과만
   ✅ "3년 차 베테랑이 신입에게 6개월 만에 따라잡힌 단 하나의 이유"
   ✅ "월 1,000을 찍은 사람들이 가장 먼저 끊은 습관 한 가지"

3. **묘하게 구체적인 행동/대상 + 이유 숨김** — 디테일이 호기심을 폭발시킨다
   ✅ "성공한 부업러가 매일 밤 11시에 반드시 끄는 것"
   ✅ "1년 만에 1억 모은 사람들이 절대 안 쓰는 5단어"
   ✅ "월 500 넘긴 사람들 카톡 프로필에서 사라진 한 단어"

4. **반대 결과 미스터리** — 통념과 정반대 결과만 던지고 메커니즘은 본문
   ✅ "더 열심히 할수록 더 가난해진 7년의 비밀"
   ✅ "잠을 늘렸더니 매출이 2배가 된 이상한 메커니즘"

5. **'딱 하나' 절대성** — 수많은 변수 중 단 하나만 보여주기
   ✅ "월 100 / 월 500을 가르는 단 한 줄의 차이"
   ✅ "결국 모든 게 무너지는 사람들의 공통점 단 하나"

6. **시간 압축 미스터리** — 짧은 시간에 큰 일이 일어났는데 그 사이를 숨기기
   ✅ "퇴근 후 90분이 1년 뒤 인생을 갈라놓는다"
   ✅ "주말 4시간이 5년치 월급을 바꾼 그 과정"

[호기심 갭 만들 때 절대 어기지 말 것]
• 답을 같은 줄에 다 보여주지 마라. "왜 ~한가" "어떻게 ~하는가"로 끝나면 본문을 사야 알 수 있게.
• "방법" "비법" "노하우" 같은 말로 끝내면 갭이 닫힘 → 결제 안 함.
• 결과/사건/디테일은 보이고, 원리/메커니즘/순서는 숨겨라.
• 한 줄 안에 "장면 + 의문"이 같이 있어야 호기심이 작동한다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 단 하나의 미션
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

5개 챕터 제목만 빠르게 이어 읽었을 때 한 사람의 변화 이야기가 보이고, 한 줄 한 줄이 다음 챕터를 못 참게 만들어야 한다.

[스토리 아크 - 목차가 한 편의 드라마처럼 읽혀야 한다]
- 5개 챕터 제목을 "그래서 → 하지만 → 그 순간 → 마침내"로 이어 읽었을 때 자연스러운 한 문단(한 사람의 줄거리)이 되어야 한다.
  예: "99%는 1년 안에 다시 무너진다 → (하지만) 의지로 잡으려는 순간 뇌는 반대로 움직인다 → (그 순간) 회복의 8할이 결판나는 90초가 있다 → (그래서) 한 번 회복한 뇌는 다시는 무너지지 않는다 → (마침내) 단단해진 멘탈은 인생을 통째로 바꾼다"
- 각 챕터는 앞 챕터가 남긴 질문에 답하면서, 동시에 새로운 질문 하나를 다음 챕터로 넘긴다. 인과가 끊긴 챕터는 즉시 다시 써라.
- 각 챕터 안의 소제목 4개도 미니 아크: 사건/충격 → 원인 추적 → 전환 → 다음으로 넘어가는 미끼.

[클리프행어 규칙 - 챕터 사이를 끊지 못하게]
- PART 2~4의 마지막 소제목은 "다음 PART를 안 읽고는 못 배기는 떡밥"이어야 한다.
- 🚨 떡밥에는 반드시 '이 책의 주제에서 나온 실물 하나'가 박혀야 한다 (화면 캡처, 메시지, 알림, 숫자 한 줄 등 — 이 예시 단어를 그대로 옮기지 말고 주제의 실물로 새로 쓸 것).
  ✅ "블로그 글 한 편이 그날 밤 모든 걸 뒤집었다" (구체물 + 미해결)
  ✅ "정산서의 마지막 줄에서 이상한 숫자를 발견했다"
  ❌ "여기서 끝이 아니었다", "진짜 문제는 따로 있었다", "이 책을 쓴 이유가 여기 있다" — 아무 정보 없는 빈 떡밥, 즉시 폐기
- ⚠️ 세 챕터가 같은 패턴 반복 금지 — 각각 다른 구체물, 다른 방식으로
- PART 1 마지막은 컨셉 도입 풀에서 선택 (이것 자체가 떡밥 역할)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🗣️ 어미 규칙 — 목차는 전부 평어체(반말)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- 존댓말 절대 금지: "~습니다", "~해요", "제가", "저는" → 발견 즉시 그 줄 폐기
- 1인칭은 "나는/내가"로: "내가 철학을 버리고 100억을 택한 이유" (O) / "제가 ~했습니다" (X)
- 챕터 제목에서 "~하게 된다", "~할 수 있다" 같은 결과 설명형 어미 금지
  ❌ "주가 폭락에도 매도 버튼을 안 누르게 된다" (남 얘기 설명)
  ✅ "폭락장에서 웃는 사람들의 계좌를 열어봤다" (장면), "버티는 게 아니라 설계가 버틴다" (선언)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔧 수단 노출 원칙 — WHAT은 보여주고 HOW만 숨겨라
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

독자는 "내가 뭘 하게 되는지"도 모르면 사지 않는다. 25줄 중 최소 8줄에는
구체적 수단/레버(예: 블로그 글, 전자책, 배당 ETF, 위임, 검색 키워드, 새벽 루틴, 정산 구조 — 주제에 맞는 것)가 보여야 한다.
- ✅ "일주일 만에 월 4천 버는 자동수익의 구조" — 수단(자동수익)·기간·금액 공개, 만드는 법만 본문에
- ✅ "경제적 자유에 이르는 '책 고르는 법'을 공개한다" — 수단 공개, 기준은 본문에
- ❌ "성공의 진짜 원리" — 뭘 하라는 건지 0% 노출, 사기처럼 보임

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎪 국내 1위 자기계발 전자책들의 실전 기법 (실제 베스트셀러 목차 분석)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. 지시어 숨김 — "나는 '이것' 하나로 상대의 미래를 예측한다", "인생에 개사기 스킬이 딱 두 가지 있다"
2. 1인칭 고백 — "내가 ○○을 버리고 ○○을 택한 이유", "평생 후회 없던 내가 딱 2개를 후회한다"
3. 통념 도발 — "돈이 없다면 운이 억세게 좋은 것이다", "흙수저 월세 22만원이면 행운인 줄 알아"
4. 괄호 장치 — "돈을 버는 짧은 공식 (어려움주의)", "~하는 법 (불편할 수 있음)"
5. 숫자 대구 — "3분의 기록으로 3배 빠르게 크는 법"
6. 강렬한 은유 명사형 — "돈 버는 것을 가로막는 심리적 기생충"
각 기법은 목차 전체에서 1~2회만. 같은 기법 반복은 티가 난다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🌡️ 감정 온도 테스트 — 설명문 박멸 (가장 중요한 필터)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

판별법: 줄을 읽고 "아, 그렇구나"로 끝나면 설명문이다. "뭔데? 왜? 그래서?"가 터져야 목차다.
모든 줄은 다음 다섯 감정 중 하나를 반드시 건드려야 한다. 하나도 못 건드리면 즉시 폐기:
① 불안 (나도 저렇게 될까봐) ② 배신감 (여태 속고 있었다) ③ 욕망 (나도 저거 갖고 싶다) ④ 안도 (아직 늦지 않았다) ⑤ 통쾌함 (반전의 짜릿함)

❌ "수익 구조 이해하기" → ✅ "통장에 마이너스가 찍힌 날 알게 된 구조"
❌ "자동화 시스템 만들기" → ✅ "잠든 새벽 3시에 도착한 첫 매출 알림"
❌ "꾸준함의 힘" → ✅ "92%가 정확히 둘째 달에 조용히 접는 이유"

규칙: 개념 명사(수익화, 자동화, 시스템, 마인드, 습관)로 시작하는 줄 금지. 모든 줄은 '장면(구체적 순간)'이나 '미스터리(왜?)'로 시작한다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📕 완성형 모범 목차 — 이 온도를 그대로 모방하라 (내용 베끼기 금지, 톤만)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

(주제: 퇴근 후 부업 월 500 / 컨셉명은 '○○○'로 표기 — 실제로는 그 책만의 이름을 지어 넣어라)

PART 1. 성실한 사람부터 가난해진다
- 입사 7년 차, 통장에 247만원이 전부였다
- 야근이 늘수록 잔고가 줄어드는 구조가 있다
- 부업 30일 차에 92%가 조용히 접는 이유
- 결국 모든 답은 '○○○' 안에 있었다

PART 2. 노력보다 자리가 먼저 돈을 번다
- 하루 14시간 일한 내가 4시간 일하는 사람에게 졌다
- "더 열심히"가 사실은 가장 위험한 조언이었다
- 월 500을 찍은 사람들이 가장 먼저 끊은 습관 한 가지
- 블로그 글 한 편이 그날 밤 모든 걸 뒤집었다

PART 3. 첫 27만원이 통장에 찍힌 날
- 전자책 한 권의 첫 정산 27만원이 월급보다 무거웠다
- 둘째 달에 90%가 멈추는 정확한 지점이 있다
- '○○○'가 작동을 시작하는 4가지 신호
- 정산보다 먼저 늘어난 건 새벽의 댓글이었다

PART 4. 잠든 사이에도 매출이 들어온다
- 새벽 3시 알림이 공포에서 설렘으로 바뀐 날
- 손을 떼는 순간 매출이 커지는 역설
- 월 100과 월 500을 가르는 단 한 줄의 차이
- 마지막 관문은 통장이 아니라 키보드 앞에 있었다

PART 5. 월급이 용돈으로 보이기 시작한다
- 퇴사를 통보하던 날 사장님이 꺼낸 말
- 통장과 저녁 시간이 함께 불어나는 이중 보상
- 1년 뒤 가장 크게 달라지는 단 한 가지
- '○○○'를 익힌 사람들의 5년 뒤가 다른 이유

위 예시처럼: 모든 줄이 장면이거나 미스터리다. 단 한 줄도 "정보 요약"이 없다. 챕터 제목만 이어 읽어도 한 사람의 이야기다. 이 온도에 못 미치면 다시 써라.\n🚨 단, 예시의 문장·숫자·단어(247만원, 92%, 27만원, 새벽 3시 등)를 그대로 옮기면 즉시 폐기 — 톤만 모방하고 내용은 이 책의 주제에서 꺼내라.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📐 챕터 제목 형식 (가장 중요)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[형식]
- 13~18자의 자연스러운 한국어 한 문장
- 단어 한두 개짜리 라벨 금지, 라벨 분리(— ㅣ :) 금지
- 평서문 또는 단언형. "~이유"로 끝나는 설명체는 한 PART에만 사용
- 명사 엔딩과 동사 엔딩을 챕터별로 섞어라 (5개 모두 명사 엔딩 금지)
- 5개를 이으면 [좌절 → 통념 박살 → 첫 사건 → 안정화 → 도약]의 5막

[좋은 예 - 성공적인 멘탈 관리 비결]
PART 1. 1년 안에 99%가 다시 무너지는 결정적 이유
PART 2. 의지로 잡으려는 순간 뇌는 반대로 움직인다
PART 3. 회복의 8할이 결판나는 폭발 직후 90초
PART 4. 한 번 회복한 뇌는 다시는 무너지지 않는다
PART 5. 단단해진 멘탈은 인생을 통째로 바꾼다

[좋은 예 - 30대 직장인 N잡 월 500]
PART 1. 직장인 99%가 부업 30일을 못 버티는 이유
PART 2. 노력보다 자리가 먼저다
PART 3. 첫 30만원이 통장에 찍힌 그날
PART 4. 새벽 3시에도 매출이 들어온다
PART 5. 월 500 다음, 억대 수익으로 가는 길

[나쁜 예 - 즉시 폐기]
- "발굴", "폭로", "전환" 같은 한두 단어짜리
- "발굴 — 부업의 90%는 첫 단추에서 망한다" (라벨 + 대시)
- "DPS의 첫 관문, 노력 없이도 돈이 따라오는 자리를 찾는 법" (시스템명 라벨화 + 너무 김)
- "이제 무너지는 게 더 이상 사건이 아니다" (사건이 아니다 ← 말이 안 됨)
- "한 번 흔들려도 다음 날엔 흔적도 없다" (AI식 과장)
- "월수도 시스템의 첫 설계" (의미 불명 + 설계라는 설명체 어휘)
- "MDS 파이프라인" (파이프라인이라는 영어 외래어를 시스템 접미사로)
- "주가 -12% 떨어져도 매도 안 하는 뇌 회로가 박혔다" (뇌 회로가 박히다 = 어법 어색, 비유 남용)
- "월급만으로는 절대 월 100이 안 나온다" (허위 단정 — "절대/무조건/100%"로 틀린 일반화 금지)
- "고배당 ETF를 사고도 월급이 안 만들어지는 사람들" ("~이 안 만들어지는" 꼬인 부정 — 뜻이 안 들어옴)
- "월 10만원으로 첫 번째 층을 쌓는 정확한 순서" (층/탑/계단/벽돌 등 건축 비유 금지 — 실물로 쓸 것)
- "여섯 번째 입금일, 월급 없는 주에 통장이 울렸다" (통장은 울리지 않는다 — 울리는 건 알림. 사물+사람동사 오결합 금지)
- "1억을 모아야 월 30만원, 나는 그 계산에 질렸다" (명사구와 문장을 쉼표로 이은 비문 — 한 줄은 완결된 한 문장으로)
- "WPS 공식의 W, 쓰지 말고 채우는 5단계 템플릿" (약어 낱글자 분해 금지 — 글자가 뜻하는 내용을 장면으로 간접 전달: "쓰는 게 아니라 채우기만 하면 되는 5단계 템플릿")
- "1년 뒤 손을 뗀 사이 잔고 곡선이 꺾여 올라간 이유" ('꺾이다'와 '올라가다'는 방향이 모순 — 움직임은 한 방향 동사 하나로)
- "5천만원으로 월 200만원 시작한 사람의 정산서 한 장" (관형절 세 겹 — 한 줄에 꾸미는 말은 1개까지, 두 번 읽어야 이해되면 폐기)
- "퇴근 후 30분 루틴 세팅법" ("~세팅법" 외래어+법 합성 금지, "~법" 꼬리표로 끝나는 줄 금지)
- "DCP 전략을 돌린 지 18개월, 주가를 보지 않게 된 날" (전략은 돌리는 게 아니다 — 쓰다/적용하다/따르다. 돌리는 건 기계·시뮬레이션뿐)
- "DCP 전략을 처음 들은 날 달력을 다시 펼쳤다" (저자가 직접 만든 방법을 남에게 들은 것처럼 쓰지 마라 — "처음 발견한 날"로)
- "배당금 100% 재투자가 원금 추가보다 빠른 증명" ("~보다 빠른 증명"은 비문 — 비교는 단순한 완결 문장으로, 비교 대상 하나는 숨겨라)
- "프드프 베스트셀러 상위 50개를 분석하고 발견한 공통점 3개" (타사 브랜드명 + 저자가 했는지 알 수 없는 지어낸 리서치 주장 — 둘 다 금지)
- "브랜딩 없이 올린 전자책이 창고에서 썩는 이유" (물성 충돌 — 전자책은 창고에 없고 썩지 않는다. 디지털엔 디지털의 실패 장면: "조회수 4에 멈춰 있는")
- "월 478만원 찍은 달에 내가 퇴근 후 한 일 전부" (꼬리 강조어 — '전부'를 빼고 "퇴근 후 한 일"로 끝내야 자연스럽다)
- "신경회로가 새로 깔리고 있다는 신호" 류 (한 번까진 OK, 같은 비유 두 번 X)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
💎 책의 시그니처 컨셉 (필수) - 작명 규칙 엄격
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

먼저 책의 시그니처 컨셉명을 하나 만든다.

[좋은 컨셉명 형식]
✅ 영문 3~4자 약어 (반드시 발음 가능 + 의미 있어야 함):
   - DPS (Discover-Plug-Scale), ARC (Awareness-Reset-Control)
   - PAM (Prompt-Automate-Monetize), RFM (Recency-Frequency-Monetary)
   - BTS, TRE, KFC, ZAP 식의 강한 자음 조합

✅ 깔끔한 한글 합성어 (3~5자):
   - 복리 자산 공식, 단단한 매출 구조, 안전 마진 법칙

[즉시 폐기 - 어색한 작명]
❌ 의미·발음이 어색한 한글:
   - "월수도", "수익도", "월500", "월천만", "부자도" (의미가 바로 안 잡히는 한글)
❌ 시스템 접미사로 영어 외래어 사용:
   - 파이프라인, 모듈, 엔진, 회로, 시너지, 매트릭스, 어레이, 클러스터, 프레임워크, 인프라
   → "MDS 파이프라인" "ABC 모듈" "XYZ 엔진" 같이 박는 즉시 탈락
   ✅ 반드시 "○○○ 시스템 / 공식 / 구조 / 법칙" 중 하나로 끝낼 것
❌ 단어가 길거나 너무 직설적인 한글 합성:
   - "잠자는 동안 매출 굴러가는 시스템" (설명체)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 컨셉명 등장 규칙 - 3번, 자연스럽게 분산
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

이 컨셉명을 5개 PART 안에 정확히 3곳에 등장시킨다:

⚠️ 컨셉명 글자 수/단어 수를 단정하는 표현 절대 금지 (가장 자주 어기는 실수!)
   ❌ "○○○ 세 글자에서 시작한다" — 컨셉명이 3글자 아니면 거짓말이 됨
   ❌ "○○○ 네 글자만 기억하라" — 글자 수 단정 금지
   ❌ "단 한 단어가 모든 걸 바꾼다" — 단어 수 단정 금지
   ✅ 컨셉명이 'DPS'(3자)든 '단단한 매출 구조'(8자)든 '복리 자산 공식'(7자)이든 모두 자연스럽게 작동하는 문장만 사용

1. **PART 1 마지막 소제목** (도입) — 아래 풀에서 1개 선택, 컨셉명 글자 수와 무관하게 자연스러운 것
   - 예: "결국 모든 답은 '○○○' 안에 있었다"
   - 예: "이 책의 모든 페이지는 '○○○' 하나를 향해 간다"
   - 예: "지금부터 '○○○' 단 하나만 기억하면 된다"
   - 예: "여기서부터 진짜 이야기, '○○○'가 시작된다"
   - 예: "마지막에 도달하는 곳은 결국 '○○○'다"
   - 예: "'○○○'를 만나기 전과 후는 완전히 다른 게임이다"

2. **PART 3 또는 PART 4 안 (1곳)** (작동·전환점)
   - 예: "○○○가 본격 작동하기 시작하는 4가지 신호"
   - 예: "○○○를 처음 적용한 사람들이 가장 먼저 느낀 변화"
   - 예: "○○○가 통장에 처음 흔적을 남기는 순간"
   - 예: "○○○ 한 달 차에 가장 먼저 무너지는 한 가지"

3. **PART 5 마지막 소제목** (확장/사회증명)
   - 예: "○○○를 익힌 사람들의 5년 뒤가 완전히 다른 이유"
   - 예: "○○○로 자리잡은 사람들이 다시는 돌아가지 않는 이유"
   - 예: "○○○ 이후, 1년 만에 가장 크게 달라지는 단 한 가지"

❌ 챕터 제목에는 컨셉명 절대 등장 X (5개 챕터 제목엔 안 들어감)
❌ 매 PART에 박지 말 것. 정확히 3곳.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ 소제목 톤: 한국 자기계발 베스트셀러 + 마케팅 뇌과학
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

각 PART당 소제목 4개. 모두 다른 패턴 + 명사·동사 엔딩 섞기 (4개 모두 같은 엔딩 금지).

[좋은 패턴 풀 - 매번 다르게]
1. 통계 충격형 — "1년 안에 99%가 다시 무너진다", "한 달 차에 90%가 다시 무너지는 정확한 이유"
2. 인지부조화형 — "의지로 버틴 사람일수록 더 크게 무너진다"
3. 뇌과학 권위형 — "감정이 폭발할 때 뇌는 정확히 어디부터 멈추는가", "신경회로가 새로 깔리고 있다는 4가지 신호"
4. 임박감/손해회피형 — "이 90초를 놓치면 다시 일어서는 데 며칠이 걸린다"
5. 정체성 전환형 — "한 번 회복한 뇌는 다시는 무너지지 않는다"
6. 호기심 갭형 — "단번에 회복한 사람들이 모두 무의식적으로 하는 행동"
7. 이중/삼중 보상형 — "회복 후 6개월 만에 통장과 인간관계가 함께 변한다"
8. 인그룹 사회증명형 — "○○를 익힌 사람들의 5년 뒤가 완전히 다른 이유"

[자연스러움 원칙 - 절대 어기지 말 것]
- 어법 검사: 주어와 동사가 자연스럽게 연결되는가? ("회복이 굴러간다" X — 회복은 안 굴러감)
- 추상 명사 의인화 금지: "○○가 멈춘다", "○○가 굴러간다", "○○ 위에 얹는다" (시스템/회복 같은 추상 명사를 사람/물건처럼)
- 과장 형용사 금지: "흔적도 없다", "통째로", "완전히" (꼭 필요할 때만)
- 추상 X, 구체 O: "많은 사람" → "월급 280만원짜리 7년차 회사원"
- 도구/플랫폼명 적극: 네이버, 카카오, 노션, 카톡, 캘린더, 구글 시트

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 즉시 폐기 표현
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

특정 작가 고유어 (절대 금지): 자청 / 역행자 / 자의식 해체 / 유전자 역행 / 원시인 / 추월차선 / 아토믹 해빗 / 언카피어블
AI 클리셰: 졸업 / 정체 / 마지막 한 수 / 다른 차원 / 결정적 시그널 / 진짜 게임 / 흔적도 없다 / 사건이 아니다 / 회로가 박혔다
시스템 의인화: "○○가 멈춘 날", "○○ 위에 얹다", "○○를 졸업한", "회복이 굴러가다"
뇌과학 비유 남용: "뇌 회로가 박혔다", "뇌 회로가 새로 깔린다" (전체 목차에 뇌·신경회로는 사실 진술로 1~2회만, 비유 남용 X)
밍밍: 효과적인 / 성공적인 / ~의 모든 것 / ~하는 방법 / 알아야 할 / 의 중요성
유치 비유: 나침반 / 열쇠 / 보물 / 황금 / 마법 / 파이프라인 / 엔진 / 톱니바퀴 / 사이클 / 눈덩이
참고서: 첫걸음 / 완벽가이드 / 핵심정리 / 기초/중급/고급 / 첫 설계 / 첫 셋업
챕터 제목 라벨: "발굴 —", "1단계:", "STEP 1." 같은 분리 형식
의문문 문어체: "왜 ~는 ~하지 못하는가" 식의 한 PART에 1개까지만
콜론(:) — 단 한 번도 쓰지 마라
숫자 중복 금지: 전체 목차에서 같은 숫자(예: 3시간 + 3개월) 두 번 등장 금지
숫자 표기 - 부호 금지: "-12%" "+30%" 같은 부호 사용 X. "12% 폭락에도", "30% 상승하면" 식으로

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 출력 형식 (이 형식 외 어떤 텍스트도 출력 금지)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[책 시그니처 컨셉]
○○○ + (시스템/구조/공식/알고리즘) | 한 줄 설명

PART 1. [13~18자, 좌절 + 통계 충격]
- [통계 충격 또는 결정적 함정형]
- [인지부조화/통념 박살형]
- [패턴 - 위 8개 중 다른 것]
- [컨셉명 첫 등장: 자연스러운 도입]

PART 2. [13~18자, 통념 박살 + 뇌과학 권위]
- [패턴]
- [패턴 - 다른 것]
- [패턴 - 또 다른 것]
- [클리프행어: PART 3의 '첫 사건'을 예고하는 떡밥]

PART 3. [13~18자, 첫 사건/결정적 순간]
- [패턴]
- [임박감/손해회피형]
- [컨셉명 등장 가능: "○○○를 처음 적용한 사람들이 가장 먼저 느낀 변화" — PART 4에 넣을 거면 여기는 일반 패턴]
- [클리프행어: PART 4의 '정체성 전환'을 예고하는 떡밥]

PART 4. [13~18자, 정체성 전환 선언]
- [컨셉명 등장 가능: "○○○가 본격 작동하는 4가지 신호" — PART 3에 안 넣었다면 여기에]
- [뇌과학 권위형 또는 통계형]
- [패턴]
- [클리프행어: PART 5의 '도약/보상'을 예고하는 떡밥]

PART 5. [13~18자, 도약/이중 보상]
- [호기심 갭형]
- [이중/삼중 보상형]
- [패턴 - 또 다른 것]
- [컨셉명 세 번째 등장: 인그룹 사회증명]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔍 출력 전 자가 점검 (반드시 통과)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

출력 직전에 5가지 모두 통과해야 한다. 하나라도 No면 다시 써라.

체크 1. 5개 챕터 제목만 이어 읽었을 때 "X였던 사람이 Y로 바뀐다"는 변화 서사가 또렷한가?
체크 2. 통념 박살(인지부조화) 패턴이 5개 챕터+20개 소제목 안에 3개 이상 박혀있는가?
체크 3. 구체적 숫자(시간/금액/비율)가 8개 이상 등장하는가?
체크 4. 평이한 표현("~의 방법", "~의 모든 것", "~의 중요성", "효과적인", "성공적인")이 0개인가?
체크 5. 5초간 훑어본 가상 독자가 "이건 안 사면 손해"라고 느낄 만한 손실회피 트리거가 챕터 제목 5개 중 2개 이상에 있는가?
체크 6. 호기심 갭(결과만 보이고 방법/이유는 숨김)이 20개 소제목 중 14개 이상에 박혀 있는가? — "결과만 보이는데 본문을 사야 알 수 있는 한 줄"이 절반 이상이어야 결제 전환됨.
체크 7. 컨셉명 글자 수를 단정하는 표현("세 글자", "네 글자", "단 한 단어")이 단 하나도 없는가? — 하나라도 있으면 즉시 전체 다시 쓰기.
체크 8. 5개 챕터 제목을 "그래서/하지만/그 순간/마침내"로 이어 읽으면 한 사람의 변화 스토리가 한 문단으로 자연스럽게 읽히는가? — 끊기면 챕터 제목부터 다시.
체크 9. PART 2~4의 마지막 소제목이 각각 '다른 방식'으로 다음 PART를 궁금하게 만드는 떡밥인가?
체크 10. 25줄 전부 감정 온도 테스트를 통과하는가? "아, 그렇구나"로 끝나는 설명문이 한 줄이라도 있으면 그 줄만 다시 써라.
체크 11. 존댓말("~습니다", "제가")이 0건인가? 챕터 제목에 "~하게 된다"가 0건인가?
체크 12. 구체적 수단(도구/플랫폼/행동)이 보이는 줄이 8줄 이상인가? 빈 떡밥("여기서 끝이 아니었다" 류)이 0건인가?

목차만 출력. 콜론 금지. 매 소제목 다른 패턴. 명사·동사 엔딩 섞기. 어법 어색한 표현 즉시 폐기. 각 PART는 정확히 소제목 4개. 컨셉명 글자수 단정 금지. 자가점검 결과는 출력하지 말 것."""
    return ask_ai(prompt, 0.85, ensure_quality=True)


def generate_content_premium(subtopic, chapter, questions, answers, topic, persona):
    """자기계발 후킹 + 권석천 칼럼 깊이의 몰입형 본문"""
    import random

    qa_pairs = ""
    for i, (q, a) in enumerate(zip(questions, answers), 1):
        if a.strip():
            qa_pairs += f"\n질문{i}: {q}\n답변{i}: {a}\n"

    # 다양한 시작 스타일
    hook_styles = [
        "도발적 질문으로 시작 (예: '왜 99%는 여기서 실패할까요?')",
        "충격적 고백으로 시작 (예: '저도 2년간 완전히 틀리고 있었습니다.')",
        "반전 사실로 시작 (예: '사실 정반대였습니다.')",
        "구체적 숫자로 시작 (예: '정확히 23일 만에 달라졌습니다.')",
        "생생한 에피소드로 시작 (예: '그날 밤 컴퓨터 앞에서 깨달았습니다.')",
        "단호한 선언으로 시작 (예: '핵심부터 말씀드리겠습니다.')",
        "대화체로 시작 (예: '\"진짜요?\" 처음 들었을 때 저도 의심했습니다.')",
        "before/after로 시작 (예: '1년 전의 저는 완전히 다른 상황이었습니다.')",
        "상식 뒤집기로 시작 (예: '열심히 하면 된다? 틀렸습니다.')",
        "사건 장면으로 시작 (예: '2023년 3월의 일이었습니다.')",
        "인용으로 시작 (예: '한 후배가 이런 말을 했습니다.')",
    ]
    current_hook = random.choice(hook_styles)

    # 표 포함 여부 (40% 확률 — 모든 장이 아닌 필요한 장에만)
    include_table = random.random() < 0.4

    prompt = f"""당신은 한국 자기계발 베스트셀러 작가입니다. '{subtopic}'에 대해 첫 문장으로 독자를 멈춰 세우고, 권석천 기자 칼럼처럼 정확한 디테일과 통찰로 끝까지 끌고 가는 본문을 씁니다.

[주제]: {topic}
[챕터]: {chapter}
[참고 내용]
{qa_pairs}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 첫 문장이 전부다
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

반드시 이 스타일로 시작:
👉 {current_hook}

(일반론·교훈으로 시작하면 즉시 폐기. 사건/숫자/대사/의문 중 하나로만 시작)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ 본문 톤: 자기계발 후킹 + 권석천 칼럼 깊이
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[문체]
- 합쇼체 기본 ("~입니다", "~합니다") + 가끔 구어체 ("~거든요", "~더라고요")
- ⚠️ 모든 문장은 존댓말로 종결. 반말(~했다, ~이다, ~한다) 혼용 절대 금지 — 한 문장도 예외 없이
- 현재 시제로 장면을 그리듯
- 짧은 문장과 긴 문장을 교차해 리듬
- 추상보다 구체. "많은 사람" → "월급 280만원짜리 7년차 회사원"

[권석천식 깊이]
- 사실 → 분석 → 통찰 순서로 전개
- 가설 검증식 흐름: "왜 그럴까. 이유를 되짚어봤습니다"
- 사회적 맥락이나 통계, 책/논문 인용을 자연스럽게 섞기
- 결론 강요 X, 독자가 스스로 깨닫게 단서를 깔기

[자기계발식 후킹]
- 첫 문장: 사건/숫자/대사/의문
- 본문 중간: 작은 반전 1회 ("그런데 진짜 흥미로운 건 그 다음이었습니다")
- 마지막 문장: 발견의 결과로서의 통찰 한 줄

[문단 - 가독성 필수 규칙]
- 한 문단은 2~3문장만. 절대 4문장을 넘기지 마라
- 문단과 문단 사이는 반드시 빈 줄로 구분
- 긴 설명은 여러 개의 짧은 문단으로 쪼갤 것

[쉽게 쓰기 - 전문성과 이해도를 동시에 (필수)]
- 내용은 논문 수준으로 정확하게, 문장은 중학생이 한 번에 이해하게
- 어려운 개념은 일상 비유 1개로 풀어라 (월급, 장보기, 통장, 운동 같은 생활 소재.
  단, 마법·황금열쇠 같은 유치한 비유와 층·탑·계단을 "쌓는" 건축 비유는 금지)
- 전문 용어는 처음 등장할 때 바로 한 줄로 뜻풀이
- 비문 금지: 모든 문장을 소리 내어 읽었을 때 자연스러운 한국어여야 한다. 꼬인 구문("~도 아니었습니다. ~부터였습니다" 류) 금지
- 한 문장에 추상 명사 2개 이상 금지

[강조 표시 - 필수]
- 독자가 밑줄 칠 핵심 문장·구절 2~3곳을 **굵게** 표시 (마크다운 ** 사용)
- 남발 금지: 문단당 최대 1곳, 한 번에 한 문장 이내. 통찰·반전·핵심 숫자에만
- 이 글의 단 하나의 핵심 통찰 문장을 골라, 줄 시작에 ★ 를 붙여 단독 문단(앞뒤 빈 줄)으로 1회 배치
  (책에서 괘선 인용구로 디자인되어 들어가니, 짧고 강한 한 문장으로. 예: ★ 시간은 만드는 것이 아니라 발견하는 것입니다)
- 본문 끝에 '핵심 정리'/'요약' 같은 제목·섹션 금지 — 파트 끝에 핵심 정리 박스가 자동으로 들어간다

[표 포함: {'예' if include_table else '아니오'}]
{'''
📊 본문 흐름상 자연스러운 위치에 표 1개 필수 (마크다운 형식):
- 비교(기존 vs 새로운), 단계별 정리, 수치 데이터 중 내용에 맞는 것
- 형식 예시 (앞뒤로 빈 줄 필수):

| 구분 | 기존 방식 | 새로운 방식 |
| 시간 | 하루 3시간 | 하루 40분 |
| 결과 | 월 50만원 | 월 340만원 |

- 셀 내용은 15자 이내로 간결하게, 3~5행이 적당
''' if include_table else '- 이번 글에는 표를 넣지 말 것 (순수 텍스트만)'}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 구체성을 끝까지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- 숫자: "많이" X → "월 340만원, 정확히 47일" O
- 사례: 수강생/지인/현장 (이름은 가끔만)
- 실제 도구명: 네이버, 카카오, 노션, 구글 시트 등

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ 자청, 역행자, 자의식 해체, 유전자 역행, 원시인, 추월차선, 아토믹 해빗 (특정 작가/책 고유 표현)
❌ 유치한 비유: 마법, 황금열쇠, 나침반, 로켓, 눈덩이, 톱니바퀴, 파이프라인
❌ AI 어휘: 중요합니다, 따라서, 결론적으로, ~를 통해, 다양한, 효과적인, 진정한
❌ 직접 호칭: 여러분, 당신, 독자님
❌ 형식: 1. 2. 첫째, 글머리 기호, 이모지 (단, **굵게** 강조와 지시된 표는 허용)
❌ 같은 이름 반복 (민준, 지수가 계속 X)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📏 분량: 1800~2200자
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

'{subtopic}' 본문 작성.
- 시작: {current_hook}
- 권석천 칼럼처럼 사실 → 분석 → 통찰 순서로 인과 추적
- 마지막 한 줄에 발견된 통찰 하나
- {'비교표 1개 포함' if include_table else '순수 텍스트만'}

⛔ 절대 금지: 본문 첫 줄에 소제목('{subtopic}')을 다시 쓰지 마라.
   소제목은 위에 이미 표시되므로, 본문은 곧장 첫 후킹 문장으로 시작한다.
   ❌ 잘못: "{subtopic}\\n\\n그날 새벽 두 시였습니다..." (소제목 반복)
   ✅ 올바름: "그날 새벽 두 시였습니다..." (바로 본문 시작)"""
    return ask_ai(prompt, 0.75, ensure_quality=True)


def format_content_html(content):
    """본문을 HTML 형식으로 변환 (강조 표시 적용)"""
    if not content:
        return ""
    # 「」 → 주황색 볼드
    formatted = re.sub(r'「([^」]+)」', r'<b style="color:#e67e22;">\1</b>', content)
    # ★ → 주황색 볼드 문장
    formatted = re.sub(r'★\s*(.+?)(?=\n|$)', r'<p style="color:#e67e22;font-weight:700;margin:20px 0;font-size:17px;">★ \1</p>', formatted)
    # 문단 구분 (빈 줄) → 문단 간격
    formatted = formatted.replace('\n\n', '</p><p style="color:#000000;margin:25px 0;line-height:2.0;font-size:17px;">')
    # 단일 줄바꿈 제거 (문단 내 연결)
    formatted = formatted.replace('\n', ' ')
    formatted = f'<p style="color:#000000;margin:25px 0;line-height:2.0;font-size:17px;">{formatted}</p>'
    # 빈 <p> 태그 정리
    formatted = re.sub(r'<p[^>]*>\s*</p>', '', formatted)
    return formatted


def generate_questions(subtopic, chapter, topic):
    prompt = f"""'{topic}' 전자책 '{chapter}' 챕터의 '{subtopic}' 작성용 질문 3개:

Q1: [질문]
Q2: [질문]
Q3: [질문]"""
    return ask_ai(prompt, 0.7)


# ==========================================
# 메인 UI
# ==========================================
# 비디오 배경 헤더
# ebook.py와 같은 폴더의 title_bg.mp4 사용 (GitHub 저장소에 함께 푸시)
video_path = str(Path(__file__).parent / "title_bg.mp4")
header_video_b64 = get_video_base64_cached(video_path)

if header_video_b64:
    st.markdown(f"""
    <style>
    .writey-brandbar {{
        display:flex; align-items:center; justify-content:space-between;
        padding:15px 24px; border-bottom:0.5px solid rgba(255,255,255,0.06);
        margin-bottom:30px;
    }}
    .writey-wordmark {{ font-size:18px; font-weight:500; letter-spacing:0.18em; color:#F5F3EF; }}
    .writey-cashtag {{ font-size:10px; letter-spacing:0.28em; color:#7A776F; margin-left:10px; }}
    .writey-author {{ font-size:11px; letter-spacing:0.1em; color:#7A776F; }}
    .writey-title {{
        font-family: 'S-CoreDream','Pretendard',sans-serif !important;
        font-size: 60px !important;
        font-weight: 300 !important;
        letter-spacing: 0.16em;
        text-indent: 0.16em;
        margin: 0;
        background: linear-gradient(180deg,#FFFFFF 0%,#F4ECD4 52%,#D8B45E 130%);
        -webkit-background-clip: text !important;
        background-clip: text !important;
        -webkit-text-fill-color: transparent !important;
        filter: drop-shadow(0 2px 14px rgba(201,162,75,0.28));
    }}
    .writey-eyebrow {{
        display:flex; align-items:center; justify-content:center; gap:16px;
        margin-bottom: 22px;
    }}
    .writey-eyebrow .we-text {{
        color: #E8Cf8C !important;
        font-size: 11px !important;
        letter-spacing: 0.45em !important;
        text-indent: 0.45em;
        font-weight: 500;
    }}
    .writey-eyebrow .we-line-l {{ width:30px; height:1px; background:linear-gradient(90deg,transparent,#C9A24B); }}
    .writey-eyebrow .we-line-r {{ width:30px; height:1px; background:linear-gradient(90deg,#C9A24B,transparent); }}
    .writey-divider {{
        display:flex; align-items:center; justify-content:center; gap:10px; margin:24px auto 0;
    }}
    .writey-divider .wd-line-l {{ width:42px; height:1px; background:linear-gradient(90deg,transparent,#C9A24B); }}
    .writey-divider .wd-line-r {{ width:42px; height:1px; background:linear-gradient(90deg,#C9A24B,transparent); }}
    .writey-divider .wd-diamond {{ width:5px; height:5px; background:#C9A24B; transform:rotate(45deg); box-shadow:0 0 8px rgba(201,162,75,0.7); }}
    .writey-tagline {{
        color: #8A8780 !important;
        font-size: 15px !important;
        margin-top: 18px;
        font-weight: 300;
        letter-spacing: 0.03em;
    }}
    </style>
    <div class="writey-brandbar">
        <div><span class="writey-wordmark">WRITEY</span><span class="writey-cashtag">CASHMAKER</span></div>
        <span class="writey-author">남현우 작가</span>
    </div>
    <div style="position:relative;border-radius:16px;overflow:hidden;margin-bottom:35px;border:0.5px solid rgba(201,162,75,0.18);">
        <video autoplay muted loop playsinline style="width:100%;height:340px;object-fit:cover;filter:brightness(0.3) saturate(1.05);">
            <source src="data:video/mp4;base64,{header_video_b64}" type="video/mp4">
        </video>
        <div style="position:absolute;top:0;left:0;right:0;bottom:0;background:linear-gradient(180deg, rgba(11,11,13,0.4) 0%, rgba(11,11,13,0.7) 100%);"></div>
        <div style="position:absolute;top:0;left:0;right:0;bottom:0;display:flex;flex-direction:column;justify-content:center;align-items:center;text-align:center;">
            <div class="writey-eyebrow"><span class="we-line-l"></span><span class="we-text">PREMIUM EBOOK MAKER</span><span class="we-line-r"></span></div>
            <h1 class="writey-title">WRITEY</h1>
            <div class="writey-divider"><span class="wd-line-l"></span><span class="wd-diamond"></span><span class="wd-line-r"></span></div>
            <p class="writey-tagline">6개의 질문에 답하면 AI가 목차부터 본문까지 완성합니다</p>
        </div>
    </div>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
    <div style="display:flex;align-items:center;justify-content:space-between;padding:15px 24px;border-bottom:0.5px solid rgba(255,255,255,0.06);margin-bottom:30px;">
        <div><span style="font-size:18px;font-weight:500;letter-spacing:0.18em;color:#F5F3EF;">WRITEY</span><span style="font-size:10px;letter-spacing:0.28em;color:#7A776F;margin-left:10px;">CASHMAKER</span></div>
        <span style="font-size:11px;letter-spacing:0.1em;color:#7A776F;">남현우 작가</span>
    </div>
    <div style="position:relative;text-align:center;padding:66px 20px 54px;margin-bottom:30px;background:radial-gradient(ellipse at 50% 0%, rgba(201,162,75,0.12) 0%, rgba(11,11,13,0) 62%), #0B0B0D;border-radius:18px;border:1px solid rgba(201,162,75,0.28);box-shadow:inset 0 0 70px rgba(201,162,75,0.05), 0 22px 60px rgba(0,0,0,0.45);overflow:hidden;">
        <div style="position:absolute;top:0;left:50%;transform:translateX(-50%);width:170px;height:1px;background:linear-gradient(90deg,transparent,rgba(201,162,75,0.75),transparent);"></div>
        <div style="display:flex;align-items:center;justify-content:center;gap:16px;margin-bottom:24px;">
            <span style="width:30px;height:1px;background:linear-gradient(90deg,transparent,#C9A24B);"></span>
            <span style="font-size:11px;letter-spacing:0.45em;color:#C9A24B;font-weight:500;text-indent:0.45em;">PREMIUM EBOOK MAKER</span>
            <span style="width:30px;height:1px;background:linear-gradient(90deg,#C9A24B,transparent);"></span>
        </div>
        <h1 style="font-family:'S-CoreDream','Pretendard',sans-serif;font-size:58px;font-weight:300;letter-spacing:0.16em;margin:0;text-indent:0.16em;background:linear-gradient(180deg,#FFFFFF 0%,#F0E6CC 52%,#C9A24B 128%);-webkit-background-clip:text;background-clip:text;-webkit-text-fill-color:transparent;">WRITEY</h1>
        <div style="display:flex;align-items:center;justify-content:center;gap:10px;margin:24px auto 0;">
            <span style="width:42px;height:1px;background:linear-gradient(90deg,transparent,#C9A24B);"></span>
            <span style="width:5px;height:5px;background:#C9A24B;transform:rotate(45deg);box-shadow:0 0 8px rgba(201,162,75,0.6);"></span>
            <span style="width:42px;height:1px;background:linear-gradient(90deg,#C9A24B,transparent);"></span>
        </div>
        <p style="color:#9A968C;font-size:15px;margin-top:20px;font-weight:300;letter-spacing:0.04em;">6개의 질문에 답하면 AI가 목차부터 본문까지 완성합니다</p>
    </div>
    """, unsafe_allow_html=True)

# ==========================================
# 인터뷰 모드 (interview_completed가 False일 때)
# ==========================================
if not st.session_state.get('interview_completed', False):
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">INTERVIEW</span>
        <h2>나만의 전자책 만들기</h2>
        <p>몇 가지 질문에 답하면 AI가 전자책을 완성해드립니다</p>
    </div>
    """, unsafe_allow_html=True)

    # 인터뷰 진행 상태
    if 'interview_step' not in st.session_state:
        st.session_state['interview_step'] = 1

    step = st.session_state['interview_step']
    total_steps = 6

    # 진행률 표시
    st.progress(step / total_steps)
    st.caption(f"질문 {step} / {total_steps}")

    st.markdown("---")

    # 인터뷰 데이터 임시 저장
    if 'temp_interview' not in st.session_state:
        st.session_state['temp_interview'] = {}

    # ========== STEP 1: 기본 정보 ==========
    if step == 1:
        st.markdown("""
        <div style="background:rgba(255,255,255,0.015);padding:30px 34px;border-radius:14px;border:1px solid rgba(201,162,75,0.16);border-left:2px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--cream);margin:0 0 10px 0;font-family:'Playfair Display','Noto Serif KR',serif;font-weight:500;">먼저 당신에 대해 알려주세요</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">전자책의 저자로서 기본 정보를 입력해주세요</p>
        </div>
        """, unsafe_allow_html=True)

        with st.form(key="step1_form"):
            author_name = st.text_input("저자명 (필명 가능)", value=st.session_state['temp_interview'].get('author_name', ''), placeholder="예: 김성장, 머니메이커 등")
            field = st.text_input("당신의 전문 분야는?", value=st.session_state['temp_interview'].get('field', ''), placeholder="예: 주식투자, 블로그 수익화, 다이어트, 영어회화 등")
            author_job = st.text_input("현재 직업/상황 (선택 — 프롤로그·에필로그에 '내 이야기'로 들어갑니다)", value=st.session_state['temp_interview'].get('author_job', ''), placeholder="예: 자영업 8년 차, 프리랜서 디자이너, 두 아이 키우는 주부, 전업 투자자")

            exp_options = ["선택하세요", "1년 미만", "1~2년", "3~5년", "5~10년", "10년 이상"]
            saved_exp = st.session_state['temp_interview'].get('experience_years', '선택하세요')
            exp_index = exp_options.index(saved_exp) if saved_exp in exp_options else 0
            experience = st.selectbox("이 분야 경험은?", exp_options, index=exp_index)

            col1, col2 = st.columns([1, 1])
            with col2:
                submitted = st.form_submit_button("다음 →", use_container_width=True, type="primary")

            if submitted:
                if not author_name.strip() or not field.strip() or experience == "선택하세요":
                    st.error("모든 항목을 입력해주세요")
                else:
                    st.session_state['temp_interview']['author_name'] = author_name.strip()
                    st.session_state['temp_interview']['field'] = field.strip()
                    st.session_state['temp_interview']['author_job'] = author_job.strip()
                    st.session_state['temp_interview']['experience_years'] = experience
                    st.session_state['interview_step'] = 2
                    st.rerun()

    # ========== STEP 2: 주제와 노하우 ==========
    elif step == 2:
        st.markdown("""
        <div style="background:rgba(255,255,255,0.015);padding:30px 34px;border-radius:14px;border:1px solid rgba(201,162,75,0.16);border-left:2px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--cream);margin:0 0 10px 0;font-family:'Playfair Display','Noto Serif KR',serif;font-weight:500;">어떤 내용을 담을까요?</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">당신만의 핵심 노하우를 알려주세요</p>
        </div>
        """, unsafe_allow_html=True)

        col_prev, col_next = st.columns([1, 1])
        with col_prev:
            if st.button("← 이전", key="interview_prev_2", use_container_width=True):
                st.session_state['interview_step'] = 1
                st.rerun()

        with st.form(key="step2_form"):
            topic = st.text_input("전자책 주제", value=st.session_state['temp_interview'].get('topic', ''), placeholder="예: 월 100만원 배당 투자, 하루 1시간 블로그로 월 300 벌기")
            core_method = st.text_area("당신만의 핵심 방법/노하우는?", value=st.session_state['temp_interview'].get('core_method', ''), height=120, placeholder="예: 저는 고배당 ETF를 활용해서 안정적으로 수익을 내는 방법을 알려드립니다. 핵심은 분산투자와 복리의 마법입니다...")

            col1, col2 = st.columns([1, 1])
            with col2:
                submitted = st.form_submit_button("다음 →", use_container_width=True, type="primary")

            if submitted:
                if not topic.strip() or not core_method.strip():
                    st.error("모든 항목을 입력해주세요")
                else:
                    st.session_state['temp_interview']['topic'] = topic.strip()
                    st.session_state['temp_interview']['core_method'] = core_method.strip()
                    st.session_state['interview_step'] = 3
                    st.rerun()

    # ========== STEP 3: 타겟 독자 (AI 추천) ==========
    elif step == 3:
        st.markdown("""
        <div style="background:rgba(255,255,255,0.015);padding:30px 34px;border-radius:14px;border:1px solid rgba(201,162,75,0.16);border-left:2px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--cream);margin:0 0 10px 0;font-family:'Playfair Display','Noto Serif KR',serif;font-weight:500;">누구를 위한 책인가요?</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">AI가 시장 데이터를 분석해 최적의 타겟을 추천해드립니다</p>
        </div>
        """, unsafe_allow_html=True)

        topic = st.session_state['temp_interview'].get('topic', '')

        # AI 타겟 분석 (캐시)
        if 'ai_target_suggestions' not in st.session_state or st.session_state.get('ai_target_topic') != topic:
            if st.button("🔍 AI 타겟 분석 시작", key="analyze_target", use_container_width=True, type="primary"):
                with st.spinner("시장 데이터 분석 중..."):
                    result = suggest_targets(topic)
                    parsed = parse_json(result)
                    if parsed and parsed.get('personas'):
                        st.session_state['ai_target_suggestions'] = parsed['personas']
                        st.session_state['ai_target_topic'] = topic
                        st.rerun()
                    else:
                        st.error("분석 실패. 다시 시도해주세요.")

        # AI 추천 결과 표시
        if st.session_state.get('ai_target_suggestions'):
            st.markdown("### AI 추천 타겟")
            personas = st.session_state['ai_target_suggestions']

            selected_idx = st.session_state.get('selected_target_idx', 0)

            for idx, persona in enumerate(personas[:3]):
                is_selected = (idx == selected_idx)
                border_color = "var(--gold)" if is_selected else "var(--line)"
                bg_color = "rgba(201,162,75,0.1)" if is_selected else "rgba(20,20,20,0.5)"

                pain_list = persona.get('pain_points', [])[:3]
                pains_text = " / ".join(pain_list) if pain_list else "고민 분석 중..."

                st.markdown(f"""
                <div style="background:{bg_color};border:1px solid {border_color};border-radius:10px;padding:15px;margin-bottom:10px;">
                    <div style="font-weight:bold;color:var(--gold);margin-bottom:5px;">{persona.get('name', '타겟')}</div>
                    <div style="font-size:13px;color:var(--text2);margin-bottom:8px;">{persona.get('demographics', '')}</div>
                    <div style="font-size:12px;color:var(--text);opacity:0.8;">💭 {pains_text}</div>
                </div>
                """, unsafe_allow_html=True)

                if st.button(f"✓ 이 타겟 선택" if not is_selected else "✓ 선택됨", key=f"select_target_{idx}", use_container_width=True, disabled=is_selected):
                    st.session_state['selected_target_idx'] = idx
                    st.rerun()

            st.markdown("---")

            # 선택된 타겟 정보 자동 입력
            selected_persona = personas[selected_idx] if selected_idx < len(personas) else personas[0]
            default_reader = f"{selected_persona.get('name', '')} ({selected_persona.get('demographics', '')})"
            default_problem = " ".join(selected_persona.get('pain_points', [])[:3])

            st.markdown("##### 선택된 타겟 정보 (수정 가능)")
            target_reader = st.text_input("타겟 독자", value=st.session_state['temp_interview'].get('target_reader', '') or default_reader, key="target_reader_input")
            target_problem = st.text_area("이 독자들의 가장 큰 고민/문제", value=st.session_state['temp_interview'].get('target_problem', '') or default_problem, height=80, key="target_problem_input")

            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("← 이전", key="interview_prev_3", use_container_width=True):
                    st.session_state['interview_step'] = 2
                    st.rerun()
            with col2:
                if st.button("다음 →", key="interview_next_3", use_container_width=True, type="primary"):
                    if not target_reader or not target_problem:
                        st.error("타겟 독자와 고민을 입력해주세요")
                    else:
                        st.session_state['temp_interview']['target_reader'] = target_reader
                        st.session_state['temp_interview']['target_problem'] = target_problem
                        st.session_state['interview_step'] = 4
                        st.rerun()
        else:
            # AI 분석 전 직접 입력 옵션
            st.markdown("---")
            st.markdown("##### 또는 직접 입력")
            target_reader = st.text_input("타겟 독자", value=st.session_state['temp_interview'].get('target_reader', ''), placeholder="예: 30대 직장인, 투자 초보자")
            target_problem = st.text_area("이 독자들의 가장 큰 고민/문제는?", value=st.session_state['temp_interview'].get('target_problem', ''), height=80, placeholder="예: 월급만으로는 부족하고, 어디서부터 시작해야 할지 모르겠다...")

            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("← 이전", key="interview_prev_3_manual", use_container_width=True):
                    st.session_state['interview_step'] = 2
                    st.rerun()
            with col2:
                if st.button("다음 →", key="interview_next_3_manual", use_container_width=True, type="primary"):
                    if not target_reader or not target_problem:
                        st.error("타겟 독자와 고민을 입력해주세요")
                    else:
                        st.session_state['temp_interview']['target_reader'] = target_reader
                        st.session_state['temp_interview']['target_problem'] = target_problem
                        st.session_state['interview_step'] = 4
                        st.rerun()

    # ========== STEP 4: 스토리 & 경력 ==========
    elif step == 4:
        st.markdown("""
        <div style="background:rgba(255,255,255,0.015);padding:30px 34px;border-radius:14px;border:1px solid rgba(201,162,75,0.16);border-left:2px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--cream);margin:0 0 10px 0;font-family:'Playfair Display','Noto Serif KR',serif;font-weight:500;">당신의 이야기를 들려주세요</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">독자들이 공감할 수 있는 진솔한 경험담과 경력</p>
        </div>
        """, unsafe_allow_html=True)

        struggle_story = st.text_area("처음 시작할 때 겪었던 어려움/실패는?", value=st.session_state['temp_interview'].get('struggle_story', ''), height=100, placeholder="예: 처음에는 주식으로 500만원을 잃었습니다. 유튜브 정보만 믿고 투자했다가 큰 손실을 봤죠...")
        breakthrough = st.text_area("어떻게 극복하고 성과를 냈나요?", value=st.session_state['temp_interview'].get('breakthrough', ''), height=100, placeholder="예: 그 후 기본서 10권을 정독하고, 나만의 원칙을 세웠습니다. 1년 후 손실을 모두 만회하고 수익을 내기 시작했습니다...")

        st.markdown("---")
        st.markdown("##### 📌 작가 경력/경험 (선택)")
        author_career = st.text_area("관련 경력이나 자격, 성과가 있다면?", value=st.session_state['temp_interview'].get('author_career', ''), height=100, placeholder="예: 금융회사 7년 근무, 투자 관련 유튜브 구독자 5만명, 월 수익 3천만원 달성, CFA 자격증 보유, 강의 경력 3년...")

        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("← 이전", key="interview_prev_4", use_container_width=True):
                st.session_state['interview_step'] = 3
                st.rerun()
        with col2:
            if st.button("다음 →", key="interview_next_4", use_container_width=True, type="primary"):
                if not struggle_story or not breakthrough:
                    st.error("어려움/실패와 극복 스토리는 필수입니다")
                else:
                    st.session_state['temp_interview']['struggle_story'] = struggle_story
                    st.session_state['temp_interview']['breakthrough'] = breakthrough
                    st.session_state['temp_interview']['author_career'] = author_career
                    st.session_state['interview_step'] = 5
                    st.rerun()

    # ========== STEP 5: 마무리 ==========
    elif step == 5:
        st.markdown("""
        <div style="background:rgba(255,255,255,0.015);padding:30px 34px;border-radius:14px;border:1px solid rgba(201,162,75,0.16);border-left:2px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--cream);margin:0 0 10px 0;font-family:'Playfair Display','Noto Serif KR',serif;font-weight:500;">마지막으로!</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">독자에게 전하고 싶은 메시지</p>
        </div>
        """, unsafe_allow_html=True)

        why_write = st.text_area("왜 이 책을 쓰려고 하나요?", value=st.session_state['temp_interview'].get('why_write', ''), height=80, placeholder="예: 저처럼 헤매는 사람들이 시행착오 없이 바로 성과를 낼 수 있도록 도와주고 싶습니다...")
        final_message = st.text_area("독자에게 마지막으로 전하고 싶은 말", value=st.session_state['temp_interview'].get('final_message', ''), height=80, placeholder="예: 누구나 할 수 있습니다. 포기하지 않으면 반드시 성공합니다...")

        # 입력 내용 미리보기
        st.markdown("---")
        st.markdown("### 입력 내용 확인")

        preview_data = st.session_state['temp_interview']
        st.markdown(f"""
        <div style="background:rgba(20,20,20,0.8);padding:20px;border-radius:10px;border:1px solid var(--line);">
            <p><b>저자:</b> {preview_data.get('author_name', '')}</p>
            <p><b>분야:</b> {preview_data.get('field', '')} ({preview_data.get('experience_years', '')})</p>
            <p><b>주제:</b> {preview_data.get('topic', '')}</p>
            <p><b>타겟:</b> {preview_data.get('target_reader', '')}</p>
        </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("← 이전", key="interview_prev_5", use_container_width=True):
                st.session_state['interview_step'] = 4
                st.rerun()
        with col2:
            if st.button("목차 생성하기", key="interview_generate_outline", use_container_width=True, type="primary"):
                if not get_api_key():
                    st.error("사이드바에서 API 키를 먼저 입력해주세요")
                elif not why_write or not final_message:
                    st.error("모든 항목을 입력해주세요")
                else:
                    st.session_state['temp_interview']['why_write'] = why_write
                    st.session_state['temp_interview']['final_message'] = final_message

                    # 목차만 먼저 생성
                    progress_box = st.empty()
                    interview_data = st.session_state['temp_interview']
                    success = generate_outline_only(interview_data, progress_box)

                    if success:
                        import time
                        time.sleep(1)
                        st.session_state['interview_step'] = 6  # 목차 확인 단계로 이동
                        st.rerun()

    # ========== STEP 6: 목차 확인 및 본문 생성 ==========
    elif step == 6:
        st.markdown("""
        <div style="background:rgba(255,255,255,0.015);padding:30px 34px;border-radius:14px;border:1px solid rgba(201,162,75,0.16);border-left:2px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--cream);margin:0 0 10px 0;font-family:'Playfair Display','Noto Serif KR',serif;font-weight:500;">목차 확인 및 수정</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">생성된 목차를 확인하고, 직접 수정하거나 AI로 재생성할 수 있습니다</p>
        </div>
        """, unsafe_allow_html=True)

        # 제목 표시
        book_title = st.session_state.get('book_title', '')
        subtitle = st.session_state.get('subtitle', '')
        book_concept = st.session_state.get('book_concept', '')

        if book_title:
            st.markdown(f"""
            <div style="background:rgba(30,30,30,0.9);padding:25px;border-radius:15px;border:0.5px solid var(--gold);margin-bottom:20px;text-align:center;">
                <h2 style="color:var(--gold);margin:0 0 10px 0;font-size:32px;">{book_title}</h2>
                <p style="color:var(--text2);margin:0;font-size:18px;">{subtitle}</p>
            </div>
            """, unsafe_allow_html=True)

        # ── 제목 후보 선택 (클릭하면 제목·부제·표지 카피가 함께 교체됨) ──
        _tc = st.session_state.get('title_candidates') or []
        if len(_tc) >= 2:
            with st.expander("✏️ 다른 제목 후보 보기 / 바꾸기", expanded=False):
                _tc_labels = [f"{c.get('title', '')}  —  {c.get('subtitle', '')}" for c in _tc]
                _cur_i = next((i for i, c in enumerate(_tc) if c.get('title') == book_title), 0)
                _sel_label = st.radio("제목 후보", _tc_labels, index=_cur_i, key="title_candidate_radio", label_visibility="collapsed")
                _sel_c = _tc[_tc_labels.index(_sel_label)]
                if _sel_c.get('title') and _sel_c.get('title') != book_title:
                    st.session_state['book_title'] = _sel_c.get('title')
                    st.session_state['subtitle'] = _sel_c.get('subtitle', '')
                    st.session_state['cover_eyebrow'] = (_sel_c.get('cover_eyebrow') or '').strip()
                    st.session_state['cover_belt_copy'] = (_sel_c.get('belt_copy') or '').strip()
                    st.rerun()

        # 컨셉 표시 — 마크다운 기호·표·체크마크를 걷어내고 깔끔한 줄글로
        if book_concept:
            _bc = book_concept
            _bc = re.sub(r'\|[^\n]*\|', '', _bc)                  # 표 행 제거
            _bc = re.sub(r'-{3,}', '\n', _bc)                     # 구분선 제거 (줄 중간 포함)
            _bc = re.sub(r'[#*`>✅❌━]+', '', _bc)                # 마크다운/체크 기호 제거
            _bc = re.sub(r'\[발음/중복[^\]]*\][\s\S]*$', '', _bc)  # 자체 검증 결과 블록 제거
            _bc = re.sub(r'\[목차 활용[^\]]*\][\s\S]*?(?=\[|$)', '', _bc)  # 목차 예시 블록 제거
            _bc = re.sub(r'\n{3,}', '\n\n', _bc).strip()
            with st.expander("💡 이 책의 고유 컨셉 보기", expanded=False):
                st.markdown(f"""
                <div style="background:rgba(201,162,75,0.1);padding:20px;border-radius:10px;border-left:3px solid var(--gold);line-height:1.8;">
                    {html.escape(_bc).replace(chr(10), '<br>')}
                </div>
                """, unsafe_allow_html=True)

        st.markdown("---")

        # 목차 표시 및 편집
        outline = st.session_state.get('outline', [])
        chapters = st.session_state.get('chapters', {})

        # ── 목차 수정사항 선반영 ──
        # 입력창에서 바뀐 챕터/소제목 이름을 렌더링 '전에' 적용한다.
        # (이전 방식은 렌더링 중 st.rerun()을 호출해서, 목차 수정 직후
        #  '본문 생성하기' 클릭이 무시되고 공백 입력 시 무한 깜박임이 발생했음)
        for i in range(len(outline)):
            ch_old = outline[i]
            w = st.session_state.get(f"ch_edit_{i}")
            if isinstance(w, str):
                new_name = w.strip()
                if new_name and new_name != ch_old and new_name not in chapters and ch_old in chapters:
                    outline[i] = new_name
                    chapters[new_name] = chapters.pop(ch_old)
            ch_now = outline[i]
            ch_entry = chapters.get(ch_now)
            if not ch_entry:
                continue
            subs = ch_entry.get('subtopics', [])
            sd = ch_entry.setdefault('subtopic_data', {})
            for j in range(len(subs)):
                wv = st.session_state.get(f"sub_edit_{i}_{j}")
                if isinstance(wv, str):
                    ns = wv.strip()
                    if ns and ns != subs[j]:
                        sd[ns] = sd.pop(subs[j], {'questions': [], 'answers': [], 'content': ''})
                        subs[j] = ns

        if outline:
            st.markdown("### 목차 구성")
            st.caption("각 챕터와 소제목을 직접 수정하거나, 🔄 버튼으로 AI가 새로 생성합니다")

            for i, ch in enumerate(outline):
                ch_data = chapters.get(ch, {})
                subtopics = ch_data.get('subtopics', [])

                # 챕터 헤더
                st.markdown(f"""
                <div style="background:linear-gradient(90deg, rgba(201,162,75,0.2) 0%, rgba(30,30,30,0.9) 100%);
                            padding:15px 20px;border-radius:10px;margin:20px 0 10px 0;
                            border-left:4px solid var(--gold);">
                    <span style="color:var(--gold);font-weight:bold;font-size:18px;">PART {i+1}</span>
                </div>
                """, unsafe_allow_html=True)

                # 챕터 제목 편집 + 챕터 소제목 전체 재생성 버튼
                ch_col1, ch_col2 = st.columns([5.2, 0.8])
                with ch_col1:
                    new_ch_name = st.text_input(
                        f"챕터 {i+1} 제목",
                        value=ch,
                        key=f"ch_edit_{i}",
                        label_visibility="collapsed"
                    )
                with ch_col2:
                    if st.button("🔄", key=f"regen_ch_{i}", help="이 챕터의 소제목 4개를 AI로 새로 생성", use_container_width=True):
                        with st.spinner("소제목 재생성 중..."):
                            new_subs = regenerate_chapter_subtopics(ch, i)
                        if new_subs:
                            ch_data['subtopics'] = new_subs
                            ch_data['subtopic_data'] = {s: {'questions': [], 'answers': [], 'content': ''} for s in new_subs}
                            # 입력창에 남은 이전 소제목 상태 제거 (되돌림 방지)
                            for _j in range(10):
                                st.session_state.pop(f"sub_edit_{i}_{_j}", None)
                            st.rerun()
                        else:
                            st.error("재생성 실패. 다시 시도해주세요.")

                # 소제목들
                for j, sub in enumerate(subtopics):
                    col1, col2, col3 = st.columns([0.4, 4.8, 0.8])
                    with col1:
                        st.markdown(f"<div style='color:var(--text2);padding-top:8px;'>•</div>", unsafe_allow_html=True)
                    with col2:
                        new_sub = st.text_input(
                            f"소제목 {j+1}",
                            value=sub,
                            key=f"sub_edit_{i}_{j}",
                            label_visibility="collapsed"
                        )
                    with col3:
                        if st.button("🔄", key=f"regen_sub_{i}_{j}", help="이 소제목만 AI로 새로 생성", use_container_width=True):
                            with st.spinner("소제목 재생성 중..."):
                                new_one = regenerate_single_subtopic(ch, j, subtopics)
                            if new_one:
                                sd = ch_data.setdefault('subtopic_data', {})
                                sd[new_one] = sd.pop(sub, {'questions': [], 'answers': [], 'content': ''})
                                sd[new_one]['content'] = ''  # 제목이 바뀌었으니 본문은 새로 생성
                                ch_data['subtopics'][j] = new_one
                                st.session_state.pop(f"sub_edit_{i}_{j}", None)
                                st.rerun()
                            else:
                                st.error("재생성 실패. 다시 시도해주세요.")

            st.markdown("---")

        # 하단 버튼
        progress_box = st.empty()  # 진행 상황을 전체 폭으로 표시
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("← 이전", key="interview_prev_6", use_container_width=True):
                st.session_state['interview_step'] = 5
                st.rerun()
        with col2:
            if st.button("본문 생성하기", key="generate_body", use_container_width=True, type="primary"):
                interview_data = st.session_state.get('interview_data', st.session_state.get('temp_interview', {}))
                success = generate_body_from_outline(interview_data, progress_box)

                if success:
                    import time
                    time.sleep(1)
                    st.session_state['current_page'] = 7  # 최종 출력 페이지로 이동
                    st.rerun()

    st.stop()

# ==========================================
# 여기서부터 기존 페이지 로직 (인터뷰 완료 후)
# ==========================================

# 페이지 네비게이션 (간소화: 4단계)
simple_pages = ["주제", "목차", "본문", "완성"]
page_mapping = [0, 4, 5, 7]  # 실제 페이지 인덱스
current = st.session_state['current_page']

# 현재 페이지가 간소화된 네비게이션의 어디에 해당하는지
def get_simple_index(current_page):
    if current_page <= 0:
        return 0
    elif current_page <= 4:
        return 1
    elif current_page <= 5:
        return 2
    else:
        return 3

simple_current = get_simple_index(current)

# 프리미엄 네비게이션 바 (4단계)
st.markdown('<div class="premium-nav-container">', unsafe_allow_html=True)
cols = st.columns(4)
for i, (col, page) in enumerate(zip(cols, simple_pages)):
    with col:
        if i == simple_current:
            st.markdown(f'<div class="nav-item active">{i+1}. {page}</div>', unsafe_allow_html=True)
        else:
            if st.button(f"{i+1}. {page}", key=f"nav_{i}", use_container_width=True):
                st.session_state['current_page'] = page_mapping[i]
                st.rerun()
st.markdown('</div>', unsafe_allow_html=True)

st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

# API 키가 없으면 환영 화면 표시
if not get_api_key():
    st.markdown("""
    <div style="background:linear-gradient(135deg, rgba(201,162,75,0.2) 0%, rgba(30,30,30,0.98) 100%);
                border:0.5px solid rgba(201,162,75,0.4);border-radius:25px;padding:50px 40px;text-align:center;margin:20px 0;">
        <p style="font-size:60px;margin:0 0 20px 0;">👋</p>
        <h2 style="color:var(--gold);font-size:42px;margin-bottom:20px;font-weight:700;">환영합니다!</h2>
        <p style="color:var(--text);font-size:24px;margin-bottom:10px;line-height:1.8;">
            AI가 전자책을 대신 써주는 프로그램입니다
        </p>
        <p style="color:var(--text2);font-size:20px;">
            처음 한 번만 설정하면 바로 사용할 수 있어요
        </p>
    </div>
    """, unsafe_allow_html=True)

    # 큰 안내 박스
    st.markdown("""
    <div style="background:#1a1a2e;border:3px solid #e74c3c;padding:30px;border-radius:20px;margin:30px 0;">
        <p style="font-size:28px;margin:0;line-height:1.6;color:#fff;text-align:center;">
            🔑 <b style="color:#e74c3c;">첫 번째 할 일</b><br><br>
            <span style="font-size:24px;">👈 왼쪽에 <span style="color:#C9A24B;font-weight:700;">"API 키"</span>를 넣어야 해요</span>
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <p style="text-align:center;font-size:32px;color:var(--gold);margin:40px 0 30px 0;font-weight:700;">
        📖 딱 3단계만 하면 끝!
    </p>
    """, unsafe_allow_html=True)

    # STEP 1 - Anthropic 가입
    st.markdown("""
    <div style="background:linear-gradient(135deg, #7c3aed 0%, #5b21b6 100%);padding:30px;border-radius:20px;margin:20px 0;">
        <p style="font-size:28px;margin:0 0 15px 0;color:#fff;font-weight:800;text-align:center;">
            1️⃣ Anthropic 회원가입
        </p>
        <div style="background:rgba(255,255,255,0.15);padding:15px;border-radius:10px;">
            <p style="font-size:16px;margin:0;color:#fff;line-height:1.8;">
                1. 아래 버튼 클릭<br>
                2. <b>"Sign up"</b> 클릭<br>
                3. Google 계정으로 가입 (가장 쉬움)
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.link_button("👆 Anthropic 가입하기", "https://console.anthropic.com/", use_container_width=True, type="primary")

    st.markdown("<p style='height:20px;'></p>", unsafe_allow_html=True)

    # STEP 2 - 결제 등록
    st.markdown("""
    <div style="background:linear-gradient(135deg, #e74c3c 0%, #c0392b 100%);padding:30px;border-radius:20px;margin:20px 0;">
        <p style="font-size:28px;margin:0 0 15px 0;color:#fff;font-weight:800;text-align:center;">
            2️⃣ 결제 수단 & 크레딧 충전
        </p>
        <div style="background:rgba(255,255,255,0.15);padding:15px;border-radius:10px;">
            <p style="font-size:16px;margin:0;color:#fff;line-height:1.8;">
                1. 로그인 후 왼쪽 메뉴 <b>"Settings"</b> 클릭<br>
                2. <b>"Billing"</b> 클릭<br>
                3. <b>"Add payment method"</b>로 카드 등록<br>
                4. <b>"Add credits"</b>로 $5~10 충전
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.link_button("👆 Billing 페이지 바로가기", "https://console.anthropic.com/settings/billing", use_container_width=True, type="primary")

    st.markdown("<p style='height:20px;'></p>", unsafe_allow_html=True)

    # STEP 3 - 키 받기
    st.markdown("""
    <div style="background:linear-gradient(135deg, #27ae60 0%, #2ecc71 100%);padding:30px;border-radius:20px;margin:20px 0;">
        <p style="font-size:28px;margin:0 0 15px 0;color:#fff;font-weight:800;text-align:center;">
            3️⃣ API 키 발급
        </p>
        <div style="background:rgba(255,255,255,0.15);padding:15px;border-radius:10px;">
            <p style="font-size:16px;margin:0;color:#fff;line-height:1.8;">
                1. 왼쪽 메뉴 <b>"API Keys"</b> 클릭<br>
                2. <b>"Create Key"</b> 버튼 클릭<br>
                3. 이름 입력 (예: ebook) → <b>"Create Key"</b><br>
                4. 생성된 키 <b>복사</b> (sk-ant-api03-... 형식)
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.link_button("👆 API Keys 페이지 바로가기", "https://console.anthropic.com/settings/keys", use_container_width=True, type="primary")

    # 마무리
    st.markdown("""
    <div style="background:linear-gradient(135deg, #C9A24B 0%, #A8852F 100%);padding:30px;border-radius:20px;margin:40px 0;text-align:center;">
        <p style="font-size:28px;margin:0 0 10px 0;color:#000;font-weight:800;">
            👈 복사한 키를 왼쪽 사이드바에 붙여넣기
        </p>
        <p style="font-size:16px;margin:0;color:#000;">
            💰 비용: 전자책 1권 약 200~500원 (Claude Sonnet 4 기준)
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.warning("⚠️ API 키는 생성 시 한 번만 보여줍니다. 꼭 복사해두세요!")

    st.markdown("---")

    # 도움말
    st.markdown("""
    <p style="text-align:center;font-size:20px;color:var(--text2);margin:20px 0;">
        😕 어려우시면 유튜브 영상을 보세요
    </p>
    """, unsafe_allow_html=True)

    st.link_button("📺 Claude API 키 발급 방법 (유튜브)", "https://www.youtube.com/results?search_query=anthropic+claude+api+key+발급", use_container_width=True)

    st.markdown("""
    <div style="background:rgba(100,100,100,0.2);padding:20px;border-radius:15px;margin:30px 0;text-align:center;">
        <p style="font-size:18px;margin:0;color:var(--text2);">
            💡 <b>팁:</b> 키는 한 번만 넣으면 저장돼요. 다음부터는 바로 시작!
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.stop()  # API 키 없으면 여기서 멈춤

# ==========================================
# PAGE 0: 주제 & 시장분석
# ==========================================
if current == 0:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 01</span>
        <h2>주제 선정 & 시장 분석</h2>
        <p>AI가 전자책의 성공 가능성을 분석합니다</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1], gap="large")

    with col1:
        st.markdown("### 주제 입력")

        topic = st.text_input("어떤 주제로 전자책을 쓸까요?", value=st.session_state['topic'], placeholder="예: 주식 배당으로 월 100만원", key="p0_topic", label_visibility="collapsed")
        if topic != st.session_state['topic']:
            st.session_state['topic'] = topic
            st.session_state['score_details'] = None

        # 빠른 제작 버튼 (자동 모드)
        st.markdown("""
        <div style="background:linear-gradient(135deg, #C9A24B 0%, #A8852F 100%);padding:20px;border-radius:15px;margin:20px 0;text-align:center;">
            <p style="font-size:14px;margin:0 0 5px 0;color:#000;opacity:0.8;">⚡ 클릭 한 번으로</p>
            <p style="font-size:20px;margin:0;color:#000;font-weight:800;">목차 + 본문 자동 완성</p>
        </div>
        """, unsafe_allow_html=True)

        if st.button("🚀 빠른 제작 시작", use_container_width=True, key="p0_auto", type="primary"):
            if not topic:
                st.error("주제를 먼저 입력해주세요")
            elif not get_api_key():
                st.error("사이드바에서 API 키를 입력해주세요")
            else:
                progress_box = st.empty()
                success = auto_generate_all(topic, progress_box)
                if success:
                    import time
                    time.sleep(1)
                    st.session_state['current_page'] = 5  # 본문 페이지로 이동
                    st.rerun()

        st.markdown("---")
        st.caption("또는 시장 분석부터 단계별로 진행:")

        if st.button("📊 시장 분석 먼저 하기", use_container_width=True, key="p0_analyze"):
            if not topic:
                st.error("주제를 입력해주세요")
            elif not get_api_key():
                st.error("사이드바에서 API 키를 입력해주세요")
            else:
                with st.spinner("AI가 시장을 분석하고 있습니다..."):
                    result = analyze_market_deep(topic)
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['score_details'] = parsed
                        st.rerun()

    with col2:
        if st.session_state.get('score_details'):
            d = st.session_state['score_details']
            score = d.get('total_score', 0)
            verdict = d.get('verdict', '')
            v_class = "verdict-go" if "추천" in verdict else ("verdict-wait" if "보류" in verdict else "verdict-no")

            st.markdown(f"""
            <div class="score-card">
                <div class="score-number">{score}</div>
                <div style="font-size:14px;color:var(--text-dim);margin-top:8px;">종합 점수</div>
                <div style="margin-top:24px;"><span class="{v_class}">{verdict}</span></div>
            </div>
            """, unsafe_allow_html=True)

            st.markdown(f"""
            <div class="result-card" style="margin-top:20px;">
                <div style="font-size:13px;color:var(--text-dim);margin-bottom:8px;">AI 분석 요약</div>
                <div style="font-size:15px;color:var(--text-bright);line-height:1.7;">{d.get('verdict_reason', '')}</div>
            </div>
            """, unsafe_allow_html=True)

            sd = d.get('search_data', {})
            if sd:
                st.markdown(f"""
                <div class="data-card" style="margin-top:16px;">
                    <b>검색 데이터</b><br><br>
                    <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px;">
                        <div>• 네이버: <b>{sd.get('naver_monthly', 'N/A')}</b></div>
                        <div>• 구글: <b>{sd.get('google_monthly', 'N/A')}</b></div>
                        <div>• 블로그: <b>{sd.get('naver_blog_posts', 'N/A')}</b></div>
                        <div>• 유튜브: <b>{sd.get('youtube_videos', 'N/A')}</b></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

            ms = d.get('market_size', {})
            comp = d.get('competition', {})

            c1, c2 = st.columns(2)
            with c1:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{ms.get("level", "")}</div><div class="stat-label">시장 규모 ({ms.get("score", 0)}점)</div></div>', unsafe_allow_html=True)
            with c2:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{comp.get("level", "")}</div><div class="stat-label">경쟁 강도 ({comp.get("score", 0)}점)</div></div>', unsafe_allow_html=True)

            if comp.get('your_opportunity'):
                st.success(f"**차별화 기회:** {comp.get('your_opportunity', '')}")

            # 경쟁 도서 검색 - 주제 키워드로 직접 검색
            current_topic = st.session_state.get('topic', '')
            if current_topic:
                st.markdown("""
                <div style="margin-top:35px;">
                    <div style="display:flex;align-items:center;gap:12px;margin-bottom:25px;">
                        <div style="width:50px;height:50px;background:linear-gradient(135deg, var(--gold) 0%, var(--gold-dark) 100%);border-radius:12px;display:flex;align-items:center;justify-content:center;">
                            <span style="font-size:26px;">🔍</span>
                        </div>
                        <div>
                            <h4 style="color:var(--gold);margin:0;font-size:22px;font-weight:600;">경쟁 도서 직접 확인하기</h4>
                            <p style="color:var(--text2);margin:4px 0 0 0;font-size:14px;">각 플랫폼에서 이 주제의 책들을 살펴보세요</p>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

                # 플랫폼별 검색 URL 매핑
                platforms = [
                    {
                        'name': '크몽 전자책',
                        'icon': '📘',
                        'url': 'https://kmong.com/search?c=ebook&q=',
                        'desc': '전자책/PDF 마켓',
                        'gradient': 'linear-gradient(135deg, #667eea 0%, #764ba2 100%)'
                    },
                    {
                        'name': '리디북스',
                        'icon': '📗',
                        'url': 'https://ridibooks.com/search?q=',
                        'desc': '국내 최대 전자책',
                        'gradient': 'linear-gradient(135deg, #11998e 0%, #38ef7d 100%)'
                    },
                    {
                        'name': 'YES24',
                        'icon': '📙',
                        'url': 'https://www.yes24.com/Product/Search?domain=BOOK&query=',
                        'desc': '종합 서점',
                        'gradient': 'linear-gradient(135deg, #f093fb 0%, #f5576c 100%)'
                    },
                    {
                        'name': '교보문고',
                        'icon': '📕',
                        'url': 'https://search.kyobobook.co.kr/search?keyword=',
                        'desc': '국내 대표 서점',
                        'gradient': 'linear-gradient(135deg, #4facfe 0%, #00f2fe 100%)'
                    },
                    {
                        'name': '클래스101',
                        'icon': '🎓',
                        'url': 'https://class101.net/search?query=',
                        'desc': '온라인 클래스',
                        'gradient': 'linear-gradient(135deg, #fa709a 0%, #fee140 100%)'
                    },
                    {
                        'name': '탈잉',
                        'icon': '👨‍🏫',
                        'url': 'https://taling.me/search?query=',
                        'desc': '재능 마켓',
                        'gradient': 'linear-gradient(135deg, #a8edea 0%, #fed6e3 100%)'
                    }
                ]

                search_query = urllib.parse.quote(current_topic)

                cols = st.columns(3)
                for idx, platform in enumerate(platforms):
                    with cols[idx % 3]:
                        search_url = platform['url'] + search_query
                        st.markdown(f"""
                        <a href="{search_url}" target="_blank" style="text-decoration:none;display:block;margin-bottom:15px;">
                            <div style="background:rgba(25,25,25,0.9);border:1px solid rgba(201,162,75,0.3);border-radius:16px;overflow:hidden;transition:all 0.3s ease;">
                                <div style="height:80px;background:{platform['gradient']};display:flex;align-items:center;justify-content:center;">
                                    <span style="font-size:40px;">{platform['icon']}</span>
                                </div>
                                <div style="padding:18px;text-align:center;">
                                    <div style="font-size:17px;color:var(--text);font-weight:700;margin-bottom:6px;">
                                        {platform['name']}
                                    </div>
                                    <div style="font-size:13px;color:var(--text2);margin-bottom:12px;">
                                        {platform['desc']}
                                    </div>
                                    <div style="background:linear-gradient(135deg, var(--gold) 0%, var(--gold-dark) 100%);color:var(--dark);padding:10px 16px;border-radius:8px;font-size:13px;font-weight:700;">
                                        🔍 "{current_topic[:15]}{'...' if len(current_topic) > 15 else ''}" 검색
                                    </div>
                                </div>
                            </div>
                        </a>
                        """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="empty-state">
                <div class="empty-state-text">
                    주제를 입력하고 <b>AI 시장 분석</b>을 시작하세요<br>
                    검색량, 경쟁 강도, 수익 가능성을 분석합니다
                </div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 2, 1])
    with c2:
        if st.button("다음 단계로 타겟 설정", key="p0_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 1: 타겟 & 컨셉
# ==========================================
elif current == 1:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 02</span>
        <h2>타겟 설정 & 제목 생성</h2>
        <p>구매할 사람을 정하고 끌리는 제목을 만듭니다</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 타겟 선정")

        if st.button("AI 타겟 추천", key="p1_target"):
            if st.session_state['topic'] and get_api_key():
                with st.spinner("분석 중..."):
                    result = suggest_targets(st.session_state['topic'])
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['suggested_targets'] = parsed
                        st.rerun()

        if st.session_state.get('suggested_targets'):
            personas = st.session_state['suggested_targets'].get('personas', [])[:3]

            for i, p in enumerate(personas):
                target_name = p.get('name', '')
                target_demo = p.get('demographics', '')
                target_needs = p.get('needs', '')
                target_pains = p.get('pain_points', [])

                st.markdown(f"""<div class="data-card">
                    <b>{html.escape(str(target_name))}</b><br>
                    <small>{html.escape(str(target_demo))}</small><br>
                    <small style="color:var(--gold);">{html.escape(str(target_needs))}</small>
                </div>""", unsafe_allow_html=True)

                if st.button(f"이 타겟 선택", key=f"sel_target_{i}", use_container_width=True):
                    selected_target = f"{target_name} - {target_demo}"
                    st.session_state['target_persona'] = selected_target
                    st.session_state['p1_persona'] = selected_target
                    st.session_state['pain_points'] = ", ".join(target_pains[:5])
                    st.session_state['suggested_targets'] = None
                    st.rerun()

        st.markdown("---")
        st.markdown("### 선택된 타겟")
        persona = st.text_area("타겟:", value=st.session_state.get('target_persona', ''), height=60, key="p1_persona", placeholder="AI 추천에서 선택하거나 직접 입력")
        st.session_state['target_persona'] = persona

        if st.button("고민 심층 분석", key="p1_analyze", use_container_width=True):
            if not persona:
                st.error("타겟을 먼저 입력해주세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("심층 분석 중..."):
                    r = analyze_pains_deep(st.session_state['topic'], persona)
                    parsed = parse_json(r)
                    if parsed:
                        st.session_state['analyzed_pains'] = parsed
                        surface = parsed.get('surface_pains', {}).get('pains', [])
                        hidden = parsed.get('hidden_pains', {}).get('pains', [])
                        st.session_state['pain_points'] = ", ".join((surface + hidden)[:6])
                        st.rerun()
                    else:
                        st.error("분석 실패. 다시 시도해주세요")

        if st.session_state.get('analyzed_pains'):
            p = st.session_state['analyzed_pains']
            st.markdown("**표면적 고민**")
            for pain in p.get('surface_pains', {}).get('pains', []):
                st.write(f"• {pain}")
            st.markdown("**숨겨진 진짜 고민**")
            for pain in p.get('hidden_pains', {}).get('pains', []):
                st.write(f"• {pain}")
            if p.get('marketing_hook'):
                st.info(f"**마케팅 훅:** {p.get('marketing_hook', '')}")

    with col2:
        st.markdown("### 베스트셀러급 제목 생성")

        # 선택된 제목이 있으면 상단에 확정 표시
        if st.session_state.get('book_title'):
            st.markdown(f"""
            <div style="background:linear-gradient(135deg,#10b981,#059669);padding:16px 20px;border-radius:12px;margin-bottom:20px;">
                <div style="color:white;font-size:12px;margin-bottom:6px;">✓ 확정된 제목</div>
                <div style="color:white;font-size:20px;font-weight:700;">{html.escape(st.session_state.get('book_title', ''))}</div>
                <div style="color:rgba(255,255,255,0.85);font-size:14px;margin-top:4px;">{html.escape(st.session_state.get('subtitle', ''))}</div>
            </div>
            """, unsafe_allow_html=True)

        pain_points = st.text_area("독자의 고민:", value=st.session_state['pain_points'], height=60, key="p1_pains")
        st.session_state['pain_points'] = pain_points

        if st.button("베스트셀러 제목 생성", key="p1_title"):
            if st.session_state['topic']:
                with st.spinner("베스트셀러 패턴 분석 중..."):
                    r = generate_titles_bestseller(st.session_state['topic'], st.session_state['target_persona'], st.session_state['pain_points'])
                    parsed = parse_json(r)
                    if parsed:
                        st.session_state['generated_titles'] = parsed
                        st.rerun()

        if st.session_state.get('generated_titles'):
            titles_list = st.session_state['generated_titles'].get('titles', [])[:5]
            for i, t in enumerate(titles_list):
                title_val = t.get('title', '')
                subtitle_val = t.get('subtitle', '')
                concept_val = t.get('concept', '')

                st.markdown(f"""
                <div class="title-card">
                    <div class="title-main">{html.escape(title_val)}</div>
                    <div class="title-sub">{html.escape(subtitle_val)}</div>
                    <div style="font-size:11px;color:var(--gold);margin-top:12px;letter-spacing:2px;">{html.escape(concept_val)}</div>
                </div>
                """, unsafe_allow_html=True)

                if st.button(f"✓ 이 제목으로 확정", key=f"sel_title_{i}", use_container_width=True):
                    st.session_state['book_title'] = title_val
                    st.session_state['subtitle'] = subtitle_val
                    st.toast(f"'{title_val}' 제목이 확정되었습니다!")
                    st.rerun()

        # 직접 입력 옵션
        st.markdown("---")
        st.markdown("#### 또는 직접 입력")
        manual_title = st.text_input("제목 입력", key="manual_title_v3")
        manual_subtitle = st.text_input("부제 입력", key="manual_subtitle_v3")
        if st.button("✓ 직접 입력한 제목으로 확정", key="manual_confirm_v3", use_container_width=True):
            if manual_title:
                st.session_state['book_title'] = manual_title
                st.session_state['subtitle'] = manual_subtitle if manual_subtitle else ''
                st.toast(f"'{manual_title}' 제목이 확정되었습니다!")
                st.rerun()

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p1_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 경쟁분석", key="p1_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 2: 경쟁도서 분석
# ==========================================
elif current == 2:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 03</span>
        <h2>경쟁 도서 분석</h2>
        <p>기존 도서의 부정 리뷰를 분석해서 숨은 니즈를 찾습니다</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 부정 리뷰 분석")

        if st.button("경쟁 도서 분석하기", use_container_width=True, key="p2_analyze"):
            if not st.session_state['topic']:
                st.error("주제를 먼저 입력해주세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("경쟁 도서 분석 중..."):
                    result = analyze_competitor_reviews(st.session_state['topic'])
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['review_analysis'] = parsed
                        concepts = parsed.get('concept_suggestions', [])
                        st.session_state['market_gaps'] = [c.get('concept', '') for c in concepts]
                        st.rerun()

        if st.session_state.get('review_analysis'):
            a = st.session_state['review_analysis']
            scope = a.get('analysis_scope', {})
            col_s1, col_s2 = st.columns(2)
            with col_s1:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{scope.get("books_analyzed", "N/A")}</div><div class="stat-label">분석 도서</div></div>', unsafe_allow_html=True)
            with col_s2:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{scope.get("negative_reviews", "N/A")}</div><div class="stat-label">부정 리뷰</div></div>', unsafe_allow_html=True)

    with col2:
        st.markdown("### 분석 결과")

        if st.session_state.get('review_analysis'):
            a = st.session_state['review_analysis']

            patterns = a.get('negative_patterns', [])
            if patterns:
                st.markdown("#### 독자 불만 패턴")
                for i, p in enumerate(patterns[:3], 1):
                    st.markdown(f"""<div class="data-card">
                        <b>{i}. {p.get('pattern', '')} ({p.get('frequency', '')})</b>
                    </div>""", unsafe_allow_html=True)
                    for rev in p.get('example_reviews', []):
                        st.caption(f'"{rev}"')
                    st.info(f"**숨겨진 니즈:** {p.get('hidden_need', '')}")
                    st.success(f"**해결책:** {p.get('solution', '')}")

            concepts = a.get('concept_suggestions', [])
            if concepts:
                st.markdown("#### 차별화 컨셉")
                for c in concepts[:2]:
                    st.markdown(f"""
                    <div class="info-card">
                        <b>「{html.escape(c.get('concept', ''))}」</b><br>
                        <span style="color:rgba(255,255,255,0.7);">{html.escape(c.get('why_works', ''))}</span>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.markdown('<div style="text-align:center;padding:60px;background:rgba(255,255,255,0.03);border-radius:16px;border:1px solid rgba(201,162,75,0.15);"><p style="color:rgba(255,255,255,0.5);">분석 버튼을 눌러주세요</p></div>', unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p2_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 학습", key="p2_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 3: 학습 & 리서치
# ==========================================
elif current == 3:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 04</span>
        <h2>학습 & 리서치</h2>
        <p>베스트셀러 분석, 트렌드 파악, 핵심 인사이트를 수집합니다</p>
    </div>
    """, unsafe_allow_html=True)

    # 탭으로 구분
    tab1, tab2, tab3 = st.tabs(["레퍼런스 추천", "트렌드 분석", "경쟁서 분석"])

    # ========== 탭1: 레퍼런스 추천 & 아이디어 ==========
    with tab1:
        topic = st.session_state.get('topic', '')

        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown("### 레퍼런스 자동 추천")
            st.markdown('<p style="color:var(--text2);font-size:13px;">주제에 맞는 참고 자료를 AI가 자동으로 추천합니다</p>', unsafe_allow_html=True)

            if not topic:
                st.warning("먼저 시장분석 페이지에서 주제를 입력해주세요")
            else:
                st.markdown(f'<p style="color:var(--accent);font-size:14px;margin:10px 0;">현재 주제: <b>{html.escape(topic)}</b></p>', unsafe_allow_html=True)

                ref_category = st.selectbox("추천 카테고리", ["베스트셀러 도서", "핵심 개념/이론", "성공 사례", "전문가 인사이트"], key="ref_cat")

                if st.button("레퍼런스 추천받기", use_container_width=True, key="auto_ref_btn"):
                    if not get_api_key():
                        st.error("API 키를 입력해주세요")
                    else:
                        with st.spinner("관련 레퍼런스 심층 분석 중..."):
                            prompt = f"""'{topic}' 주제로 전자책을 쓰려고 합니다.
'{ref_category}' 카테고리에서 참고할 만한 자료 3개를 추천해주세요.

중요: 마치 이 책/자료를 직접 읽은 것처럼 아주 상세하게 설명해주세요.

각 추천 자료에 대해 다음을 포함해주세요:
1. 제목과 저자
2. 책/자료의 핵심 메시지 (10문장 이상으로 상세히)
3. 주요 챕터/섹션별 핵심 내용
4. 저자의 핵심 주장과 근거
5. 실제 사례나 스토리
6. 전자책에 활용할 수 있는 구체적 인사이트

중요: 책의 모든 주요 챕터를 빠짐없이 요약해주세요. 일부만 하지 말고 전체 목차를 다 포함해주세요.

JSON 형식으로 응답:
{{
    "recommendations": [
        {{
            "title": "자료 제목",
            "author": "저자/출처",
            "core_message": "이 책의 핵심 메시지와 주장을 10문장 이상으로 상세하게 설명",
            "chapters": [
                {{"name": "1장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "2장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "3장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "4장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "5장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "6장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "7장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "8장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}}
            ],
            "key_arguments": ["저자의 핵심 주장1과 근거", "핵심 주장2와 근거", "핵심 주장3과 근거"],
            "real_examples": ["책에 나온 실제 사례/스토리 1", "사례 2", "사례 3"],
            "key_insights": ["전자책에 활용할 인사이트 1", "인사이트 2", "인사이트 3", "인사이트 4", "인사이트 5"],
            "application": "내 전자책에 구체적으로 활용하는 방법 (3문장 이상)"
        }}
    ]
}}"""
                            result = ask_ai(prompt, 0.8)
                            parsed = parse_json(result)
                            if parsed and parsed.get('recommendations'):
                                st.session_state['recommended_refs'] = parsed['recommendations']
                                st.rerun()
                            else:
                                st.error("추천 생성 실패")

                # 추천된 레퍼런스 표시
                if st.session_state.get('recommended_refs'):
                    st.markdown("---")
                    st.markdown("#### 추천 레퍼런스")

                    for i, ref in enumerate(st.session_state['recommended_refs']):
                        st.markdown(f"""<div class="data-card">
                            <b>{html.escape(str(ref.get('title', '')))}</b>
                            <br><small style="color:var(--text2);">{html.escape(str(ref.get('author', '')))}</small>
                        </div>""", unsafe_allow_html=True)

                        # 핵심 메시지
                        if ref.get('core_message'):
                            st.markdown("**핵심 메시지**")
                            st.markdown(f'<p style="color:#e0e0e0;line-height:1.8;">{html.escape(str(ref.get("core_message", "")))}</p>', unsafe_allow_html=True)

                        # 챕터 요약
                        if ref.get('chapters'):
                            st.markdown("**챕터별 요약**")
                            for ch in ref.get('chapters', []):
                                st.markdown(f"""<div style="background:rgba(255,255,255,0.05);padding:12px 16px;margin:8px 0;border-left:3px solid var(--gold);">
                                    <b style="color:var(--gold);">{html.escape(str(ch.get('name', '')))}</b><br>
                                    <span style="color:#e0e0e0;">{html.escape(str(ch.get('summary', '')))}</span>
                                </div>""", unsafe_allow_html=True)

                        # 핵심 주장
                        if ref.get('key_arguments'):
                            st.markdown("**저자의 핵심 주장**")
                            for arg in ref.get('key_arguments', []):
                                st.info(arg)

                        # 실제 사례
                        if ref.get('real_examples'):
                            st.markdown("**실제 사례**")
                            for ex in ref.get('real_examples', []):
                                st.markdown(f'<p style="color:#e0e0e0;">- {html.escape(str(ex))}</p>', unsafe_allow_html=True)

                        # 활용 인사이트
                        if ref.get('key_insights'):
                            st.markdown("**활용 인사이트**")
                            for insight in ref.get('key_insights', []):
                                st.success(insight)

                        # 적용 방법
                        if ref.get('application'):
                            st.markdown("**내 책에 적용하는 방법**")
                            st.markdown(f'<p style="color:#e0e0e0;line-height:1.8;">{html.escape(str(ref.get("application", "")))}</p>', unsafe_allow_html=True)

                        if st.button("이 자료 저장하기", key=f"save_ref_{i}", use_container_width=True):
                            ref_item = {
                                'title': ref.get('title', ''),
                                'source': ref.get('author', ''),
                                'core_message': ref.get('core_message', ''),
                                'chapters': ref.get('chapters', []),
                                'key_arguments': ref.get('key_arguments', []),
                                'real_examples': ref.get('real_examples', []),
                                'key_insights': ref.get('key_insights', []),
                                'application': ref.get('application', ''),
                                'type': 'recommended',
                                'added_at': datetime.now().strftime('%Y-%m-%d %H:%M')
                            }
                            st.session_state['knowledge_hub'].append(ref_item)
                            st.success("저장 완료")
                            st.rerun()

                        st.markdown("---")

        with col2:
            st.markdown("### 저장된 자료 & 아이디어 도출")
            hub = st.session_state.get('knowledge_hub', [])

            if hub:
                st.caption(f"총 {len(hub)}개 자료 저장됨")

                for i, item in enumerate(hub):
                    title = item.get('title', item.get('main_topic', item.get('source', f'자료 {i+1}')))

                    st.markdown(f"""<div class="data-card">
                        <b>{html.escape(str(title))}</b><br>
                        <small>{html.escape(str(item.get('source', '')))} | {item.get('added_at', '')}</small>
                    </div>""", unsafe_allow_html=True)

                    # 핵심 메시지 전체 표시
                    if item.get('core_message'):
                        st.write(item['core_message'])

                    # 핵심 인사이트 표시
                    if item.get('key_insights'):
                        st.markdown("**핵심 인사이트:**")
                        for insight in item.get('key_insights', [])[:3]:
                            st.success(insight)

                    # 적용 방법 표시
                    if item.get('application'):
                        st.info(f"적용법: {item['application']}")

                    col_a, col_b = st.columns([1, 1])
                    with col_a:
                        if st.button("상세보기", key=f"view_ref_{i}"):
                            st.session_state[f'show_detail_{i}'] = not st.session_state.get(f'show_detail_{i}', False)
                            st.rerun()
                    with col_b:
                        if st.button("삭제", key=f"del_ref_{i}"):
                            st.session_state['knowledge_hub'].pop(i)
                            st.rerun()

                    # 상세 보기 토글
                    if st.session_state.get(f'show_detail_{i}', False):
                        if item.get('chapters'):
                            st.markdown("**챕터 요약:**")
                            for ch in item.get('chapters', []):
                                st.markdown(f"""<div style="background:rgba(255,255,255,0.05);padding:10px 14px;margin:6px 0;border-left:2px solid var(--gold);">
                                    <b style="color:var(--gold);">{html.escape(str(ch.get('name', '')))}</b><br>
                                    <span style="color:#e0e0e0;font-size:14px;">{html.escape(str(ch.get('summary', '')))}</span>
                                </div>""", unsafe_allow_html=True)
                        if item.get('key_arguments'):
                            st.markdown("**핵심 주장:**")
                            for arg in item.get('key_arguments', []):
                                st.markdown(f'<p style="color:#e0e0e0;">- {html.escape(str(arg))}</p>', unsafe_allow_html=True)
                        if item.get('real_examples'):
                            st.markdown("**실제 사례:**")
                            for ex in item.get('real_examples', []):
                                st.markdown(f'<p style="color:#e0e0e0;">- {html.escape(str(ex))}</p>', unsafe_allow_html=True)

                    st.markdown("---")

                st.markdown("---")
                st.markdown("#### 아이디어 도출")
                st.markdown('<p style="color:var(--text2);font-size:13px;">수집된 자료를 바탕으로 전자책 아이디어를 생성합니다</p>', unsafe_allow_html=True)

                if st.button("아이디어 생성하기", use_container_width=True, key="ideate_btn"):
                    if not get_api_key():
                        st.error("API 키를 입력해주세요")
                    else:
                        with st.spinner("아이디어 생성 중..."):
                            hub_summary = ""
                            for item in hub[:5]:
                                hub_summary += f"\n[{item.get('title', '')}]\n"
                                if item.get('core_message'):
                                    hub_summary += f"핵심 메시지: {item.get('core_message', '')}\n"
                                if item.get('key_insights'):
                                    hub_summary += f"인사이트: {', '.join(item.get('key_insights', []))}\n"
                                if item.get('application'):
                                    hub_summary += f"적용법: {item.get('application', '')}\n"

                            prompt = f"""다음 수집된 자료들을 철저히 분석하여 '{topic}' 주제의 전자책 아이디어를 도출해주세요:

수집된 자료:
{hub_summary}

위 자료들의 공통점, 차이점, 빈틈을 분석하고 다음을 포함해서 아이디어를 생성해주세요:
1. 기존 책들과 확실히 다른 차별화된 콘셉트
2. 독자의 문제를 해결하는 독특한 관점
3. 구체적인 목차/콘텐츠 구성 아이디어
4. 타겟 독자에게 강하게 어필할 포인트

JSON 형식으로 응답:
{{
    "main_concept": "핵심 콘셉트 한 문장 (경쟁작과 어떻게 다른지 명확히)",
    "unique_angles": ["독특한 관점 1 (왜 이 관점이 효과적인지 설명)", "관점 2", "관점 3"],
    "content_ideas": ["챕터 아이디어 1", "챕터 아이디어 2", "챕터 아이디어 3", "챕터 아이디어 4", "챕터 아이디어 5"],
    "appeal_points": ["어필 포인트 1", "포인트 2", "포인트 3"],
    "title_suggestions": ["제목 제안 1 (부제 포함)", "제목 제안 2 (부제 포함)", "제목 제안 3 (부제 포함)"],
    "differentiation": "경쟁작 대비 구체적인 차별화 전략 (3문장 이상)"
}}"""
                            result = ask_ai(prompt, 0.9)
                            parsed = parse_json(result)
                            if parsed:
                                st.session_state['generated_ideas'] = parsed
                                st.rerun()
                            else:
                                st.error("아이디어 생성 실패")

                # 생성된 아이디어 표시
                if st.session_state.get('generated_ideas'):
                    ideas = st.session_state['generated_ideas']

                    st.markdown(f"""<div class="summary-hub">
                        <b>핵심 콘셉트</b><br>
                        {html.escape(str(ideas.get('main_concept', '')))}
                    </div>""", unsafe_allow_html=True)

                    if ideas.get('unique_angles'):
                        st.markdown("**독특한 관점**")
                        for angle in ideas.get('unique_angles', []):
                            st.info(angle)

                    if ideas.get('title_suggestions'):
                        st.markdown("**제목 제안**")
                        for title in ideas.get('title_suggestions', []):
                            st.success(title)

                    if ideas.get('content_ideas'):
                        st.markdown("**콘텐츠 아이디어**")
                        for idea in ideas.get('content_ideas', []):
                            st.write(f"- {idea}")

                    if ideas.get('differentiation'):
                        st.markdown(f"""<div class="data-card">
                            <b>차별화 전략</b><br>
                            <small>{html.escape(str(ideas.get('differentiation', '')))}</small>
                        </div>""", unsafe_allow_html=True)
            else:
                st.markdown('<div style="text-align:center;padding:60px 20px;background:var(--card);border:1px solid var(--line);"><p style="color:var(--text2);">추천받은 레퍼런스를 저장하면<br>아이디어를 도출할 수 있습니다</p></div>', unsafe_allow_html=True)

    # ========== 탭2: 트렌드 분석 ==========
    with tab2:
        st.markdown("### 시장 트렌드 분석")
        st.markdown('<p style="color:var(--text2);">현재 인기 있는 전자책 트렌드와 키워드를 파악합니다</p>', unsafe_allow_html=True)

        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown("#### 트렌드 키워드 분석")
            trend_topic = st.text_input("분석할 분야", key="trend_topic", placeholder="예: 재테크, 자기계발, 다이어트...")

            if st.button("트렌드 분석", use_container_width=True, key="trend_btn"):
                if not trend_topic:
                    st.error("분야를 입력하세요")
                elif not get_api_key():
                    st.error("API 키를 입력해주세요")
                else:
                    with st.spinner("트렌드 분석 중..."):
                        prompt = f"""'{trend_topic}' 분야의 전자책 시장 트렌드를 분석해주세요.

JSON 형식으로 응답:
{{
    "hot_keywords": ["인기 키워드 1", "키워드 2", "키워드 3", "키워드 4", "키워드 5"],
    "rising_topics": ["떠오르는 주제 1", "주제 2", "주제 3"],
    "reader_needs": ["독자가 원하는 것 1", "원하는 것 2", "원하는 것 3"],
    "content_gaps": ["시장에서 부족한 콘텐츠 1", "부족한 콘텐츠 2"],
    "recommended_angles": ["추천 접근 방식 1", "접근 방식 2", "접근 방식 3"],
    "avoid": ["피해야 할 것 1", "피해야 할 것 2"]
}}"""
                        result = ask_ai(prompt, 0.8)
                        parsed = parse_json(result)
                        if parsed:
                            st.session_state['trend_analysis'] = parsed
                            st.rerun()

            if st.session_state.get('trend_analysis'):
                ta = st.session_state['trend_analysis']
                if ta.get('hot_keywords'):
                    st.write("**인기 키워드**")
                    st.write(" | ".join(ta.get('hot_keywords', [])))
                if ta.get('rising_topics'):
                    st.write("**떠오르는 주제**")
                    for t in ta.get('rising_topics', []):
                        st.write(f"- {t}")

        with col2:
            st.markdown("#### 독자 니즈")
            if st.session_state.get('trend_analysis'):
                ta = st.session_state['trend_analysis']
                if ta.get('reader_needs'):
                    st.write("**독자가 원하는 것**")
                    for n in ta.get('reader_needs', []):
                        st.info(n)
                if ta.get('content_gaps'):
                    st.write("**시장 빈틈**")
                    for g in ta.get('content_gaps', []):
                        st.success(g)
                if ta.get('recommended_angles'):
                    st.write("**추천 접근법**")
                    for r in ta.get('recommended_angles', []):
                        st.write(f"- {r}")
            else:
                st.markdown('<div style="text-align:center;padding:60px;background:var(--card);border:1px solid var(--line);"><p style="color:var(--text2);">분야를 입력하고<br>트렌드 분석을 시작하세요</p></div>', unsafe_allow_html=True)

    # ========== 탭3: 경쟁서 분석 ==========
    with tab3:
        st.markdown("### 경쟁 도서 분석")
        st.markdown('<p style="color:var(--text2);">경쟁 전자책의 목차, 리뷰, 강점을 분석합니다</p>', unsafe_allow_html=True)

        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown("#### 경쟁서 정보 입력")
            comp_title = st.text_input("책 제목", key="comp_title", placeholder="예: 돈의 심리학")
            comp_toc = st.text_area("목차 (복사/붙여넣기)", height=150, key="comp_toc", placeholder="1장. 제목\n2장. 제목\n...")
            comp_reviews = st.text_area("대표 리뷰 (선택)", height=100, key="comp_reviews", placeholder="인상적인 리뷰를 붙여넣으세요...")

            if st.button("경쟁서 분석", use_container_width=True, key="comp_btn"):
                if not comp_title or not comp_toc:
                    st.error("제목과 목차를 입력하세요")
                elif not get_api_key():
                    st.error("API 키를 입력해주세요")
                else:
                    with st.spinner("분석 중..."):
                        prompt = f"""다음 경쟁 도서를 분석해주세요:

제목: {comp_title}
목차:
{comp_toc}

리뷰: {comp_reviews if comp_reviews else '없음'}

JSON 형식으로 응답:
{{
    "book_summary": "이 책의 핵심 콘셉트",
    "target_audience": "예상 타겟 독자",
    "strengths": ["강점 1", "강점 2", "강점 3"],
    "weaknesses": ["약점/빈틈 1", "약점 2"],
    "unique_selling_point": "이 책만의 차별점",
    "improvement_opportunities": ["내 책에서 더 잘할 수 있는 것 1", "기회 2", "기회 3"],
    "key_chapters": ["핵심 챕터 1", "챕터 2"],
    "content_structure": "콘텐츠 구성 방식"
}}"""
                        result = ask_ai(prompt, 0.7)
                        parsed = parse_json(result)
                        if parsed:
                            parsed['title'] = comp_title
                            parsed['added_at'] = datetime.now().strftime('%Y-%m-%d %H:%M')
                            if 'competitor_analysis' not in st.session_state:
                                st.session_state['competitor_analysis'] = []
                            st.session_state['competitor_analysis'].append(parsed)
                            st.success("분석 완료")
                            st.rerun()

        with col2:
            st.markdown("#### 분석 결과")
            comps = st.session_state.get('competitor_analysis', [])

            if comps:
                for i, comp in enumerate(comps):
                    st.markdown(f"""<div class="data-card">
                        <b>{html.escape(str(comp.get('title', f'경쟁서 {i+1}')))}</b>
                        <br><small>{comp.get('added_at', '')}</small>
                    </div>""", unsafe_allow_html=True)
                    st.caption(comp.get('book_summary', ''))

                    if comp.get('improvement_opportunities'):
                        for o in comp.get('improvement_opportunities', [])[:2]:
                            st.success(f"차별화: {o}")

                    if st.button("삭제", key=f"del_comp_{i}"):
                        st.session_state['competitor_analysis'].pop(i)
                        st.rerun()
            else:
                st.markdown('<div style="text-align:center;padding:60px;background:var(--card);border:1px solid var(--line);"><p style="color:var(--text2);">경쟁 도서 정보를 입력하고<br>분석해보세요</p></div>', unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p3_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음", key="p3_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 4: 목차 설계
# ==========================================
elif current == 4:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 05</span>
        <h2>목차 설계</h2>
        <p>독자의 호기심을 자극하는 목차를 만듭니다</p>
    </div>
    """, unsafe_allow_html=True)

    if st.session_state.get('market_gaps'):
        st.success(f"{len(st.session_state['market_gaps'])}개 차별화 포인트 반영")

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 목차 생성")

        st.markdown("""
        <div class="info-card">
            <b>🔥 목차 작성 팁</b><br><br>
            • 설명하지 말고 <b>궁금하게</b><br>
            • 구체적 <b>숫자 + 결과</b> 보여주기<br>
            • <b>실패담/고백</b>으로 공감 얻기<br>
            • "99%가 모르는" <b>비밀</b> 암시<br>
            • <b>반전</b>이 있을 것 같은 느낌<br><br>
            <span style="color:var(--gold);">❌ "시간관리의 중요성"</span><br>
            <span style="color:#50c878;">✓ "20대에 이걸 몰라서 5년 날렸다"</span>
        </div>
        """, unsafe_allow_html=True)

        if st.button("목차 생성하기", use_container_width=True, key="p4_outline_btn"):
            if not st.session_state.get('topic'):
                st.error("주제를 입력하세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("목차 생성 중..."):
                    result = generate_outline(
                        st.session_state['topic'],
                        st.session_state.get('target_persona', ''),
                        st.session_state.get('pain_points', ''),
                        st.session_state.get('market_gaps', [])
                    )

                    if result:
                        lines = result.split('\n')
                        chapters = []
                        current_ch = None
                        subtopics = {}

                        for line in lines:
                            orig_line = line
                            line = line.strip()
                            if not line:
                                continue

                            # 마크다운 정리 (먼저 정리한 후 검사)
                            clean_line = re.sub(r'^[#\*\s]+', '', line).strip()
                            clean_line = clean_line.replace('**', '').replace('*', '').strip()

                            # PART 또는 챕터 형식 인식 (더 유연하게)
                            is_chapter = False

                            # PART 형식 (다양한 변형)
                            if re.search(r'PART\s*\d+', clean_line, re.IGNORECASE):
                                is_chapter = True
                            # 파트 형식 (한글)
                            elif re.search(r'파트\s*\d+', clean_line):
                                is_chapter = True
                            # Chapter 형식
                            elif re.search(r'(Chapter|챕터)\s*\d+', clean_line, re.IGNORECASE):
                                is_chapter = True
                            # "1. 제목" 형식
                            elif re.match(r'^\d+[\.\)]\s*.+', clean_line) and not clean_line.startswith('-'):
                                is_chapter = True
                            # 숫자로 시작하는 제목 (예: "1 첫번째 파트")
                            elif re.match(r'^\d+\s+[가-힣A-Za-z]', clean_line):
                                is_chapter = True

                            if is_chapter:
                                name = clean_line
                                if name and len(name) > 3:
                                    current_ch = name
                                    chapters.append(current_ch)
                                    subtopics[current_ch] = []

                            # 소제목 - 다양한 형식 지원
                            elif current_ch:
                                is_subtopic = False
                                st_name = ""

                                # "-" 또는 "•" 또는 "·" 로 시작
                                if re.match(r'^\s*[\-\•\·]\s*', line):
                                    is_subtopic = True
                                    st_name = re.sub(r'^[\s\-\•\·]+', '', line).strip()
                                # 들여쓰기 된 내용
                                elif orig_line.startswith('  ') or orig_line.startswith('\t'):
                                    if not any(x in line.upper() for x in ['PART', 'CHAPTER', '파트']):
                                        is_subtopic = True
                                        st_name = line.strip().lstrip('-•· ')
                                # "  1)" 또는 "  a)" 형식
                                elif re.match(r'^\s+[\da-z][\)\.]', orig_line):
                                    is_subtopic = True
                                    st_name = re.sub(r'^[\s\da-z\)\.\-]+', '', line).strip()

                                if is_subtopic:
                                    st_name = st_name.replace('**', '').replace('*', '').replace('#', '').strip()
                                    # 소제목이 충분히 길고 유효한 경우만 추가
                                    if st_name and len(st_name) > 3 and not re.match(r'^(PART|파트|Chapter|챕터)', st_name, re.IGNORECASE):
                                        subtopics[current_ch].append(st_name)

                        if chapters:
                            st.session_state['outline'] = chapters
                            st.session_state['chapters'] = {}
                            for ch in chapters:
                                st.session_state['chapters'][ch] = {
                                    'subtopics': subtopics.get(ch, []),
                                    'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in subtopics.get(ch, [])}
                                }
                            st.success(f"{len(chapters)}개 챕터 생성!")
                            st.rerun()
                        else:
                            st.error("목차 생성 실패. 다시 시도해주세요.")
                    else:
                        st.error("AI 응답 없음. 다시 시도해주세요.")

    with col2:
        st.markdown("### 현재 목차")

        if st.session_state.get('outline'):
            # 수정 모드 토글
            if 'edit_outline_mode' not in st.session_state:
                st.session_state['edit_outline_mode'] = False

            col_view, col_edit = st.columns([1, 1])
            with col_view:
                if st.button("👁 보기 모드", use_container_width=True, disabled=not st.session_state['edit_outline_mode']):
                    st.session_state['edit_outline_mode'] = False
                    st.rerun()
            with col_edit:
                if st.button("✏️ 수정 모드", use_container_width=True, disabled=st.session_state['edit_outline_mode']):
                    st.session_state['edit_outline_mode'] = True
                    st.rerun()

            st.markdown("---")

            if st.session_state['edit_outline_mode']:
                # 수정 모드
                st.markdown('<p style="color:var(--gold);font-size:14px;">📝 제목을 직접 수정할 수 있습니다</p>', unsafe_allow_html=True)

                updated_outline = []
                updated_chapters = {}

                for ch_idx, ch in enumerate(st.session_state['outline']):
                    # 챕터 제목 수정
                    new_ch_title = st.text_input(
                        f"PART {ch_idx + 1}",
                        value=ch,
                        key=f"edit_ch_{ch_idx}"
                    )
                    updated_outline.append(new_ch_title)
                    updated_chapters[new_ch_title] = {'subtopics': [], 'subtopic_data': {}}

                    # 소제목 수정
                    subtopics = st.session_state['chapters'].get(ch, {}).get('subtopics', [])
                    new_subtopics = []
                    for st_idx, st_name in enumerate(subtopics):
                        new_st = st.text_input(
                            f"  └ 소제목 {st_idx + 1}",
                            value=st_name,
                            key=f"edit_st_{ch_idx}_{st_idx}",
                            label_visibility="collapsed"
                        )
                        if new_st.strip():
                            new_subtopics.append(new_st)
                            # 기존 데이터 유지
                            old_data = st.session_state['chapters'].get(ch, {}).get('subtopic_data', {}).get(st_name, {'questions': [], 'answers': [], 'content': ''})
                            updated_chapters[new_ch_title]['subtopic_data'][new_st] = old_data

                    updated_chapters[new_ch_title]['subtopics'] = new_subtopics
                    st.markdown("---")

                # 저장 버튼
                if st.button("💾 수정 내용 저장", use_container_width=True, type="primary"):
                    st.session_state['outline'] = updated_outline
                    st.session_state['chapters'] = updated_chapters
                    st.session_state['edit_outline_mode'] = False
                    st.success("목차가 수정되었습니다!")
                    st.rerun()

            else:
                # 보기 모드 - 예쁘게 표시
                for ch_idx, ch in enumerate(st.session_state['outline']):
                    st.markdown(f"""
                    <div style="background:linear-gradient(135deg, rgba(201,162,75,0.15) 0%, rgba(201,162,75,0.05) 100%);
                                padding:16px 20px;border-radius:12px;margin-bottom:8px;border-left:4px solid var(--gold);">
                        <span style="color:var(--gold);font-size:13px;font-weight:600;">PART {ch_idx + 1}</span>
                        <p style="color:var(--text);font-size:17px;font-weight:600;margin:8px 0 0 0;">{ch}</p>
                    </div>
                    """, unsafe_allow_html=True)

                    subtopics = st.session_state['chapters'].get(ch, {}).get('subtopics', [])
                    for st_idx, st_name in enumerate(subtopics):
                        st.markdown(f"""
                        <div style="padding:10px 20px 10px 35px;color:var(--text);font-size:15px;">
                            <span style="color:var(--gold);margin-right:8px;">•</span>{st_name}
                        </div>
                        """, unsafe_allow_html=True)

                    st.markdown("<div style='margin-bottom:20px;'></div>", unsafe_allow_html=True)

        else:
            st.markdown('<div style="text-align:center;padding:60px;background:rgba(255,255,255,0.03);border-radius:16px;border:1px solid rgba(201,162,75,0.15);"><p style="color:rgba(255,255,255,0.5);">목차를 생성해주세요</p></div>', unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p4_prev", use_container_width=True):
            if st.session_state.get('interview_completed'):
                st.session_state['current_page'] = 0  # 인터뷰 사용자는 주제 페이지로
            else:
                go_prev()
            st.rerun()
    with c3:
        if st.button("다음 본문", key="p4_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 5: 본문 작성
# ==========================================
elif current == 5:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 06</span>
        <h2>본문 작성</h2>
        <p>AI가 각 챕터의 콘텐츠를 작성합니다</p>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.get('outline'):
        st.warning("먼저 목차를 설계하세요")
    else:
        col_sel1, col_sel2 = st.columns([1, 1])
        with col_sel1:
            selected_ch = st.selectbox("챕터", st.session_state['outline'], key="p5_chapter")

        # 선택된 챕터가 있고 chapters에 존재하는지 확인
        if selected_ch and selected_ch in st.session_state.get('chapters', {}):
            ch_data = st.session_state['chapters'][selected_ch]
            subtopics_list = ch_data.get('subtopics', [])

            # 소제목이 있는 경우에만 선택박스 표시
            selected_st = None
            if subtopics_list:
                with col_sel2:
                    selected_st = st.selectbox("소제목", subtopics_list, key="p5_subtopic")

            # 진행률 표시
            completed = sum(1 for s in subtopics_list if ch_data.get('subtopic_data', {}).get(s, {}).get('content'))
            total = len(subtopics_list)
            if total > 0:
                st.progress(completed / total)
                st.caption(f"{completed}/{total} 완료")

            # 소제목이 선택된 경우에만 편집 UI 표시
            if selected_st:
                # subtopic_data 초기화 확인
                if 'subtopic_data' not in ch_data:
                    ch_data['subtopic_data'] = {}
                if selected_st not in ch_data['subtopic_data']:
                    ch_data['subtopic_data'][selected_st] = {'questions': [], 'answers': [], 'content': ''}

                st_data = ch_data['subtopic_data'][selected_st]

                col1, col2 = st.columns([1, 1])

                # 버튼 키를 위한 고유 식별자
                st_key = f"{selected_ch}_{selected_st}".replace(" ", "_")

                with col1:
                    st.markdown("### 인터뷰")
                    if st.button("질문 생성", key=f"gen_q_{st_key}"):
                        if not get_api_key():
                            st.error("API 키를 입력해주세요")
                        else:
                            with st.spinner("생성 중..."):
                                q_text = generate_questions(selected_st, selected_ch, st.session_state['topic'])
                                if q_text:
                                    questions = re.findall(r'Q\d+:\s*(.+)', q_text)
                                    if not questions:
                                        questions = [q.strip() for q in q_text.split('\n') if '?' in q][:3]
                                    if questions:
                                        st_data['questions'] = questions
                                        st_data['answers'] = [''] * len(questions)
                                        st.rerun()
                                    else:
                                        st.error("질문 생성에 실패했습니다")

                    if st_data.get('questions'):
                        for i, q in enumerate(st_data['questions']):
                            st.markdown(f"**Q{i+1}.** {q}")
                            # answers 리스트 크기 확인
                            while len(st_data.get('answers', [])) <= i:
                                st_data['answers'].append('')
                            st_data['answers'][i] = st.text_area(f"A{i+1}", value=st_data['answers'][i], height=80, key=f"ans_{st_key}_{i}", label_visibility="collapsed")

                with col2:
                    st.markdown("### 본문")
                    has_ans = st_data.get('questions') and any(a.strip() for a in st_data.get('answers', []))

                    if has_ans:
                        if st.button("본문 생성", key=f"gen_content_{st_key}", use_container_width=True, type="primary"):
                            if not get_api_key():
                                st.error("API 키를 입력해주세요")
                            else:
                                with st.spinner("본문 작성 중... (1~2분 소요)"):
                                    content = generate_content_premium(selected_st, selected_ch, st_data['questions'], st_data['answers'], st.session_state['topic'], st.session_state['target_persona'])
                                    if content:
                                        st_data['content'] = content
                                        st.success("본문 생성 완료!")
                                        st.rerun()
                                    else:
                                        st.error("본문 생성에 실패했습니다. 다시 시도해주세요.")
                    else:
                        st.info("왼쪽에서 질문에 답변을 입력하면 본문을 생성할 수 있습니다")

                    # 본문 표시
                    current_content = st_data.get('content', '')
                    if current_content:
                        # HTML 형식으로 변환하여 표시
                        formatted_html = format_content_html(current_content)
                        st.markdown(f"""
                        <style>
                        .content-preview-box {{
                            background:#ffffff !important;
                            padding:25px 30px;
                            border-radius:12px;
                            border:1px solid rgba(201,162,75,0.3);
                            margin:15px 0;
                            font-family:'S-CoreDream', sans-serif !important;
                            font-size:17px;
                            max-height:500px;
                            overflow-y:auto;
                        }}
                        .content-preview-box,
                        .content-preview-box p,
                        .content-preview-box span,
                        .content-preview-box div {{
                            color:#000000 !important;
                            -webkit-text-fill-color:#000000 !important;
                        }}
                        .content-preview-box b[style*="color:#e67e22"],
                        .content-preview-box p[style*="color:#e67e22"] {{
                            color:#e67e22 !important;
                            -webkit-text-fill-color:#e67e22 !important;
                        }}
                        </style>
                        <div class="content-preview-box">
                            {formatted_html}
                        </div>
                        """, unsafe_allow_html=True)
                        st.caption(f"📝 {len(current_content.replace(' ', '').replace(chr(10), '')):,}자")

                        # 이미지 추가 기능
                        st.markdown("---")
                        st.markdown("**📷 이미지 추가**")
                        uploaded_img = st.file_uploader("이미지 업로드", type=['png', 'jpg', 'jpeg'], key=f"img_{st_key}", label_visibility="collapsed")
                        if uploaded_img:
                            # 이미지 저장
                            if 'images' not in st_data:
                                st_data['images'] = []
                            img_b64 = base64.b64encode(uploaded_img.read()).decode()
                            st_data['images'].append({'name': uploaded_img.name, 'data': img_b64})
                            st.success(f"이미지 '{uploaded_img.name}' 추가됨!")
                            st.rerun()

                        # 추가된 이미지 표시
                        if st_data.get('images'):
                            st.caption(f"추가된 이미지: {len(st_data['images'])}개")
                            for idx, img in enumerate(st_data['images']):
                                col_img, col_del = st.columns([4, 1])
                                with col_img:
                                    st.image(f"data:image/png;base64,{img['data']}", caption=img['name'], width=200)
                                with col_del:
                                    if st.button("삭제", key=f"del_img_{st_key}_{idx}"):
                                        st_data['images'].pop(idx)
                                        st.rerun()

                        # 수정 기능
                        st.markdown("---")
                        with st.expander("✏️ 본문 직접 수정"):
                            st.caption("「중요단어」 → 주황색 강조 | ★ 문장 → 핵심 강조")
                            edited = st.text_area("본문 편집", value=current_content, height=400, key=f"content_{st_key}", label_visibility="collapsed")
                            if edited != current_content:
                                st_data['content'] = edited
                                st.rerun()
                    else:
                        st.markdown('<div style="text-align:center;padding:80px 20px;background:rgba(255,255,255,0.03);border-radius:12px;border:1px dashed rgba(201,162,75,0.3);"><p style="color:var(--text2);font-size:16px;">본문이 아직 없습니다<br>질문에 답변 후 "본문 생성" 버튼을 누르세요</p></div>', unsafe_allow_html=True)
            else:
                st.info("이 챕터에는 소제목이 없습니다. 목차를 다시 생성해주세요.")

        st.markdown("---")
        st.markdown("### 전체 본문")
        full_content = get_full_content()
        if full_content:
            char_count = len(full_content.replace(' ', '').replace('\n', ''))
            st.success(f"총 {char_count:,}자 | 약 {char_count//500}페이지")

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p5_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 출력", key="p5_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 6: 표지 디자인
# ==========================================
elif current == 6:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 07</span>
        <h2>표지 디자인</h2>
        <p>전문 디자인 툴로 고품질 표지를 만드세요</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 표지 정보 정리")

        # 이전 페이지에서 설정한 제목/부제 자동 연동
        saved_title = st.session_state.get('book_title', '')
        saved_subtitle = st.session_state.get('subtitle', '')

        cover_title = st.text_input("표지 제목", value=saved_title, key="cover_title", placeholder="예: 돈의 속성")
        cover_subtitle = st.text_input("부제목", value=saved_subtitle, key="cover_subtitle", placeholder="예: 당신이 모르는 부의 법칙")
        cover_author = st.text_input("저자명", key="cover_author", placeholder="예: 홍길동")

        st.markdown("---")
        st.markdown("### AI 표지 스타일 추천")

        if st.button("내 주제에 맞는 표지 스타일 추천받기", use_container_width=True, key="ai_cover_suggest"):
            topic = st.session_state.get('topic', '')
            if not topic:
                st.error("시장분석 페이지에서 주제를 먼저 입력해주세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("베스트셀러 표지 분석 중..."):
                    prompt = f"""'{topic}' 주제의 전자책 표지 디자인을 추천해주세요.

이 분야의 실제 베스트셀러 책 표지를 분석해서 추천해주세요.

JSON 형식으로 응답:
{{
    "recommended_style": "추천 스타일명",
    "color_scheme": "추천 색상 조합 (예: 검정 배경 + 금색 텍스트)",
    "design_concept": "디자인 콘셉트 설명 (2문장)",
    "typography_tip": "타이포그래피 팁 (폰트 스타일, 크기 등)",
    "reference_books": ["참고할 베스트셀러 표지 1", "표지 2", "표지 3"],
    "canva_search_keyword": "Canva에서 검색할 키워드 (영문)"
}}"""
                    result = ask_ai(prompt, 0.7)
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['cover_suggestion'] = parsed
                        st.rerun()

        if st.session_state.get('cover_suggestion'):
            sug = st.session_state['cover_suggestion']
            st.markdown(f"""<div class="data-card">
                <b>추천 스타일: {html.escape(str(sug.get('recommended_style', '')))}</b><br>
                <small>색상: {html.escape(str(sug.get('color_scheme', '')))}</small>
            </div>""", unsafe_allow_html=True)
            st.write(sug.get('design_concept', ''))
            if sug.get('typography_tip'):
                st.info(f"💡 타이포그래피 팁: {sug.get('typography_tip', '')}")
            if sug.get('reference_books'):
                st.markdown("**참고 베스트셀러:**")
                for book in sug.get('reference_books', []):
                    st.caption(f"- {book}")
            if sug.get('canva_search_keyword'):
                st.session_state['canva_keyword'] = sug.get('canva_search_keyword', '')

    with col2:
        st.markdown("### 표지 미리보기")

        _label_to_id = {v: k for k, v in COVER_TEMPLATES.items()}
        _labels = list(_label_to_id.keys())
        # 주제에 맞는 아키타입 자동 추천 — 책 주제/제목이 바뀌면 추천도 갱신
        # (이전 책에서 고른 템플릿이 남아 새 주제와 무관한 표지가 나오는 문제 방지)
        _cover_auto_key = f"{st.session_state.get('topic', '')}|{cover_title or ''}"
        if ("cover_template_choice" not in st.session_state
                or st.session_state.get('_cover_auto_key') != _cover_auto_key):
            _auto = pick_cover_template(st.session_state.get('topic', ''), cover_title or '')
            _auto_label = COVER_TEMPLATES.get(_auto)
            if _auto_label in _labels:
                st.session_state["cover_template_choice"] = _auto_label
            st.session_state['_cover_auto_key'] = _cover_auto_key

        # ── 10가지 디자인을 내 책 제목으로 한눈에 보기 ──
        with st.expander("🎨 10가지 디자인 모두 보기 (내 제목으로 미리보기)", expanded=True):
            _gal_items = list(COVER_TEMPLATES.items())
            for _row_start in range(0, len(_gal_items), 5):
                _gal_cols = st.columns(5)
                for _gc, (_gid, _glabel) in zip(_gal_cols, _gal_items[_row_start:_row_start + 5]):
                    with _gc:
                        _gsvg = build_cover_svg(_gid, cover_title or "제목", cover_subtitle, cover_author)
                        st.markdown(
                            f'<div style="border-radius:4px;overflow:hidden;'
                            f'box-shadow:0 6px 18px rgba(0,0,0,0.35);">{_gsvg}</div>'
                            f'<p style="text-align:center;font-size:11px;color:#999;margin:6px 0 2px;">{html.escape(_glabel.split(" — ")[0])}</p>',
                            unsafe_allow_html=True,
                        )

        _choice = st.radio(
            "표지 스타일 선택 (주제에 맞춰 자동 추천됨)",
            _labels,
            key="cover_template_choice",
        )
        _tmpl = _label_to_id[_choice]

        _svg = build_cover_svg(
            _tmpl,
            cover_title or "제목을 입력하세요",
            cover_subtitle,
            cover_author,
        )
        st.markdown(
            '<div style="max-width:320px;margin:12px auto 18px;border-radius:8px;'
            'overflow:hidden;box-shadow:0 24px 70px rgba(0,0,0,0.55);">'
            f'{_svg}</div>',
            unsafe_allow_html=True,
        )
        st.download_button(
            "표지 다운로드 (SVG · 고해상도 벡터)",
            _svg,
            file_name=f"{(cover_title or 'cover')}_cover.svg",
            mime="image/svg+xml",
            use_container_width=True,
            key="cover_svg_dl",
        )
        st.caption(
            "SVG는 무손실 벡터 파일입니다. 브라우저에서 열어 캡처하거나, "
            "Canva·Figma·미리캔버스에 올려 PNG/JPG로 내보낼 수 있습니다."
        )

        st.markdown("---")
        st.markdown("### Canva로 표지 만들기")

        st.markdown("""
        <div class="data-card">
            <p style="font-size:16px;margin-bottom:15px;">
                <b>Canva</b>는 전문 디자이너 수준의 표지를 무료로 만들 수 있는 온라인 툴입니다.
            </p>
            <p style="color:var(--text2);font-size:14px;">
                ✓ 수천 개의 프로 템플릿<br>
                ✓ 드래그 앤 드롭 편집<br>
                ✓ 무료 이미지/아이콘<br>
                ✓ 한글 폰트 지원
            </p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("---")

        # Canva 검색 키워드 설정
        canva_keyword = st.session_state.get('canva_keyword', 'book cover')

        # Canva 책 표지 템플릿 링크
        canva_url = f"https://www.canva.com/templates/?query={canva_keyword}%20book%20cover"

        st.markdown(f"""
        <a href="{canva_url}" target="_blank" style="
            display:block;
            background:linear-gradient(135deg,#7c3aed,#6366f1);
            color:white;
            padding:18px 24px;
            border-radius:12px;
            text-decoration:none;
            text-align:center;
            font-size:18px;
            font-weight:600;
            margin-bottom:15px;
            transition:transform 0.2s;
        ">
            🎨 Canva에서 표지 만들기
        </a>
        """, unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### 사용 방법")
        st.markdown("""
        1. **Canva 열기** - 위 버튼 클릭 (무료 가입)
        2. **템플릿 선택** - 마음에 드는 디자인 클릭
        3. **텍스트 수정** - 제목, 부제목, 저자명 입력
        4. **다운로드** - PNG 또는 PDF로 저장
        """)

        st.markdown("---")
        st.markdown("### 추천 검색어")

        search_keywords = [
            "ebook cover", "book cover minimalist",
            "book cover gold", "book cover business",
            "korean book cover", "self help book cover"
        ]

        cols = st.columns(2)
        for i, kw in enumerate(search_keywords):
            with cols[i % 2]:
                if st.button(kw, key=f"canva_kw_{i}", use_container_width=True):
                    st.session_state['canva_keyword'] = kw
                    st.rerun()

        st.markdown("---")

        # 복사할 텍스트
        if cover_title or cover_subtitle or cover_author:
            st.markdown("### 복사할 텍스트")
            copy_text = f"제목: {cover_title}\n부제목: {cover_subtitle}\n저자: {cover_author}"
            st.code(copy_text, language=None)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p6_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음", key="p6_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 7: 최종 출력
# ==========================================
elif current == 7:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 08</span>
        <h2>최종 출력</h2>
        <p>완성된 전자책을 다운로드하세요</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1.5, 1])

    with col1:
        st.markdown("### 다운로드")

        final_title = st.text_input("제목", value=st.session_state.get('book_title', ''), key="p6_title")
        final_subtitle = st.text_input("부제", value=st.session_state.get('subtitle', ''), key="p6_subtitle")

        full = f"{final_title}\n{final_subtitle}\n\n{'='*50}\n\n"
        for ch in st.session_state.get('outline', []):
            if ch in st.session_state.get('chapters', {}):
                ch_data = st.session_state['chapters'][ch]
                ch_content = ""
                for s in ch_data.get('subtopics', []):
                    c = ch_data.get('subtopic_data', {}).get(s, {}).get('content', '')
                    if c:
                        ch_content += f"\n\n【{s}】\n\n{clean_content(c)}"
                if ch_content:
                    full += f"\n\n{ch}\n{'-'*40}{ch_content}\n"

        st.markdown("**미리보기**")
        st.text_area("전체 내용", value=full, height=300, disabled=True, key="p7_preview")

        # 저자명 가져오기
        author_name = st.session_state.get('author_name', '') or st.session_state.get('interview_data', {}).get('author_name', '')

        # 다운로드 버튼 3개
        st.markdown("### 내보내기")

        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button("TXT", full, file_name=f"{final_title or 'ebook'}.txt", use_container_width=True, key="p7_txt")
        with c2:
            # HTML 내보내기 - 특수문자 이스케이프 처리
            escaped_title = html.escape(final_title)
            escaped_content = html.escape(full).replace('\n', '<br>')
            html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{escaped_title}</title>
    <style>
        body {{
            max-width: 800px;
            margin: 0 auto;
            padding: 60px 40px;
            font-family: 'Pretendard', -apple-system, BlinkMacSystemFont, sans-serif;
            line-height: 2;
            background: #fafafa;
            color: #333;
        }}
        h1 {{ font-size: 32px; color: #1a1a2e; margin-bottom: 10px; }}
        h2 {{ font-size: 14px; color: #888; font-weight: normal; }}
    </style>
</head>
<body>
{escaped_content}
</body>
</html>"""
            st.download_button("HTML", html_content, file_name=f"{final_title or 'ebook'}.html", use_container_width=True, key="p7_html")

        with c3:
            # DOCX 다운로드
            if DOCX_AVAILABLE:
                docx_data, docx_error = create_ebook_docx(
                    final_title,
                    final_subtitle,
                    author_name,
                    st.session_state.get('chapters', {}),
                    st.session_state.get('outline', []),
                    st.session_state.get('interview_data', {})
                )
                if docx_data:
                    st.download_button(
                        "WORD",
                        docx_data,
                        file_name=f"{final_title or 'ebook'}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="p7_docx"
                    )
                else:
                    st.button("WORD", disabled=True, use_container_width=True, key="p7_docx_disabled")
                    if docx_error:
                        st.caption(f"⚠️ {docx_error[:30]}")
            else:
                st.button("WORD", disabled=True, use_container_width=True, key="p7_docx_na")
                st.caption("pip install python-docx")

        total = len(full.replace(' ', '').replace('\n', ''))
        if total > 0:
            # Word(A5) 실제 구조 기반 페이지 추정:
            # 표지1 + 판권1 + 프롤로그2 + 차례2 + 에필로그2 = 8, 챕터 오프너 각 1,
            # 소제목은 각각 새 페이지에서 시작 (A5·명조 10.5pt·행간 1.9 ≈ 700자/페이지)
            import math as _math
            _pages = 8
            for _ch in st.session_state.get('outline', []):
                _cd = st.session_state.get('chapters', {}).get(_ch, {})
                _pages += 1
                for _s in _cd.get('subtopics', []):
                    _c = _cd.get('subtopic_data', {}).get(_s, {}).get('content', '')
                    if _c:
                        _pages += max(1, _math.ceil(len(_c) / 700))
            st.success(f"총 {total:,}자 | Word 기준 약 {_pages}페이지")

    with col2:
        st.markdown("### 현황")
        total_st = sum(len(ch.get('subtopics', [])) for ch in st.session_state.get('chapters', {}).values())
        done = sum(1 for ch in st.session_state.get('chapters', {}).values() for s in ch.get('subtopic_data', {}).values() if s.get('content'))

        if total_st > 0:
            st.progress(done / total_st)
            st.write(f"**완료:** {done}/{total_st}")

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 2, 1])
    with c1:
        if st.button("← 본문 수정", key="p7_prev", use_container_width=True):
            st.session_state['current_page'] = 5  # 본문 편집 페이지로
            st.rerun()
    with c3:
        if st.button("표지 디자인", key="p7_cover", use_container_width=True):
            st.session_state['current_page'] = 6
            st.rerun()


st.markdown("""
<div style="
    text-align: center;
    padding: 30px 20px;
    margin-top: 50px;
    border-top: 1px solid rgba(201,162,75,0.3);
    color: #ffffff !important;
    font-size: 16px;
    letter-spacing: 2px;
    background: rgba(0,0,0,0.3);
">
    <div style="font-family:'Playfair Display',serif;font-size:13px;letter-spacing:0.5em;color:#C9A24B;text-indent:0.5em;">W R I T E Y</div>
    <div style="font-size:10px;letter-spacing:0.3em;color:#7A776F;margin-top:10px;text-indent:0.3em;">CASHMAKER · 남현우 작가</div>
</div>
""", unsafe_allow_html=True)
